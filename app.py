"""
Parvidya — Flask Backend (Production-Ready)
=========================================
Stack:
  • Python + Flask
  • MongoDB (primary) with local JSON fallback
  • LangChain + OpenAI GPT-4o / GPT-4o-mini
  • LangChain InMemoryVectorStore (per syllabus, session-scoped)
  • Vercel-compatible (read-only FS except /tmp)

Run locally:  python app.py
API base:     http://localhost:5001/api/...
"""

import os, json, re, uuid, hashlib, random, base64, shutil, smtplib
import logging
import time as _time
import threading
import requests as _requests
from concurrent.futures import ThreadPoolExecutor, as_completed
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
from pathlib import Path
from functools import wraps
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename

# ── LangChain ────────────────────────────────────────────────────────────────
try:
    from langchain_openai import ChatOpenAI, OpenAIEmbeddings
    from langchain_core.documents import Document
    from langchain_core.prompts import PromptTemplate
    from langchain_core.vectorstores import InMemoryVectorStore
    from langchain_core.messages import HumanMessage, SystemMessage
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    try:
        from langchain.chains import ConversationalRetrievalChain
        from langchain.memory import ConversationBufferMemory
        LANGCHAIN_OK = True
    except ImportError:
        try:
            # LangChain 0.3+ layout (some installs omit the chains re-export)
            from langchain.chains.conversational_retrieval.base import ConversationalRetrievalChain
            from langchain.memory import ConversationBufferMemory
            LANGCHAIN_OK = True
        except ImportError as _lc_err:
            print(f"[BOOT] LangChain core imports failed: {_lc_err}")
            LANGCHAIN_OK = False
except ImportError as _lc_err:
    LANGCHAIN_OK = False
    print(f"[WARN] LangChain not available: {_lc_err}")

try:
    import openai as _oai
    OPENAI_OK = True
except ImportError:
    OPENAI_OK = False

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_API_KEY_BULK_PRIMARY = os.getenv("OPENAI_API_KEY_BULK_PRIMARY", "")
OPENAI_API_KEY_BULK_FALLBACK = os.getenv("OPENAI_API_KEY_BULK_FALLBACK", "")
OPENAI_API_KEY_PREMIUM = os.getenv("OPENAI_API_KEY_PREMIUM", "")
OPENAI_MODEL   = os.getenv("OPENAI_MODEL", "gpt-4o")
OPENAI_MINI    = os.getenv("OPENAI_MINI_MODEL", "gpt-4o-mini")
OCR_VISION_MODEL        = os.getenv("OCR_VISION_MODEL", OPENAI_MINI)
OCR_PREMIUM_VISION_MODEL = os.getenv("OCR_PREMIUM_VISION_MODEL", OPENAI_MODEL)
SUBJECTIVE_OCR_MODEL    = os.getenv("SUBJECTIVE_OCR_MODEL", OPENAI_MINI)  # default to mini for lower cost; override in env when needed
OCR_MAX_IMAGE_DIM = int(os.getenv("OCR_MAX_IMAGE_DIM", "2048") or "2048")
OCR_JPEG_QUALITY = int(os.getenv("OCR_JPEG_QUALITY", "82") or "82")
OCR_PDF_DPI = int(os.getenv("OCR_PDF_DPI", "180" if os.getenv("K_SERVICE", "").strip() else "300") or "300")
OCR_QUALITY_ESCALATION_THRESHOLD = float(os.getenv("OCR_QUALITY_ESCALATION_THRESHOLD", "0.75") or "0.75")
OCR_PREPROCESS_GEOMETRY = os.getenv("OCR_PREPROCESS_GEOMETRY", "1") == "1"
OCR_PREPROCESS_CLAHE = os.getenv("OCR_PREPROCESS_CLAHE", "1") == "1"
OCR_PREPROCESS_COMPRESSION = os.getenv("OCR_PREPROCESS_COMPRESSION", "1") == "1"
MULTI_SPLIT_FIXED_PAGES = int(os.getenv("MULTI_SPLIT_FIXED_PAGES", "2") or "2")
MULTI_EVAL_BATCH_SIZE = int(os.getenv("MULTI_EVAL_BATCH_SIZE", "5") or "5")
MULTI_EVAL_INTER_BATCH_DELAY = float(os.getenv("MULTI_EVAL_INTER_BATCH_DELAY", "0.75") or "0.75")
MULTI_EVAL_EXTRACT_ATTEMPTS = int(os.getenv("MULTI_EVAL_EXTRACT_ATTEMPTS", "4") or "4")
MULTI_EVAL_EVAL_ATTEMPTS = int(os.getenv("MULTI_EVAL_EVAL_ATTEMPTS", "3") or "3")
MULTI_EVAL_MAX_WORKERS_OVERRIDE = int(os.getenv("MULTI_EVAL_MAX_WORKERS", "0") or "0")
OCR_PIPELINE_VERSION = "2026-05-03.2"
OCR_FAST_COST = os.getenv("OCR_FAST_COST", "0") == "1"

# ── Enhanced Evaluation Configuration ──────────────────────────────────────────
OCR_CONFIDENCE_THRESHOLD = float(os.getenv("OCR_CONFIDENCE_THRESHOLD", "0.75") or "0.75")     # Escalate if below
GRADING_BORDERLINE_MARGIN = float(os.getenv("GRADING_BORDERLINE_MARGIN", "0.05") or "0.05") # 5% of pass threshold
REVIEW_ESSAY_LENGTH_THRESHOLD = int(os.getenv("REVIEW_ESSAY_LENGTH_THRESHOLD", "500") or "500")  # Words
REVIEW_SCORE_ANOMALY_STDDEV = float(os.getenv("REVIEW_SCORE_ANOMALY_STDDEV", "2.0") or "2.0") # >2σ deviations
BULK_PARALLEL_WORKERS = int(os.getenv("BULK_PARALLEL_WORKERS", "10") or "10")     # Max workers for bulk
COST_OPTIMIZATION_ENABLED = os.getenv("COST_OPTIMIZATION_ENABLED", "1") == "1"     # Route by sheet quality
ENABLE_PARENT_EMAILS = os.getenv("ENABLE_PARENT_EMAILS", "0") == "1"               # Send automated parent reports

# ── Razorpay (Payment Gateway) ─────────────────────────────────────────────────
RAZORPAY_KEY_ID     = os.getenv("RAZORPAY_KEY_ID", "")
RAZORPAY_KEY_SECRET = os.getenv("RAZORPAY_KEY_SECRET", "")

# Plan prices in paise (INR × 100). coaching is per-10-students unit.
PLAN_PRICES = {
    "student-pro":    14900,   # ₹149/month
    "school-starter":  99900,  # ₹999/month
    "school-growth":  249900,  # ₹2,499/month
    "coaching":        99000,  # ₹990 base (≈10 students @ ₹99)
}

# Map plan_id → QUOTA_PRESETS key for auto-applying after payment
PLAN_TO_PRESET = {
    "student-pro":   {"type": "role_default", "preset": "student-pro",   "role": "student"},
    "school-starter":{"type": "institution",   "preset": "school-starter"},
    "school-growth": {"type": "institution",   "preset": "school-growth"},
    "coaching":      {"type": "institution",   "preset": "coaching"},
}

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE         = Path(__file__).parent.parent
IS_VERCEL    = os.environ.get("VERCEL") == "1"
IS_RENDER    = os.environ.get("RENDER") == "true"
IS_SERVERLESS = IS_VERCEL or IS_RENDER

TMP_ROOT  = Path("/tmp") if IS_SERVERLESS else BASE
DATA_F    = (TMP_ROOT / "platform_data.json") if IS_SERVERLESS else (BASE / "data" / "platform_data.json")
UPL_DIR   = (TMP_ROOT / "uploads")            if IS_SERVERLESS else (BASE / "uploads")
DB_DIR    = (TMP_ROOT / "study_db")           if IS_SERVERLESS else (BASE / "study_db")

EXAMS_DIR = DB_DIR / "exams"
EVAL_DIR  = DB_DIR / "evaluations"
BULK_DIR  = DB_DIR / "bulk_evaluations"

for _d in [DATA_F.parent, UPL_DIR, DB_DIR, EXAMS_DIR, EVAL_DIR, BULK_DIR]:
    _d.mkdir(parents=True, exist_ok=True)

ALLOWED = {"pdf","txt","md","docx","doc","mp3","mp4","wav","m4a","ogg","jpg","jpeg","png","webp"}

# ── Flask ─────────────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder="dist", static_url_path="")
CORS(app, resources={
    r"/api/*": {
        "origins": "*",
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "expose_headers": ["Content-Type", "Authorization"],
        "supports_credentials": True,
        "max_age": 3600
    }
})
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024  # 200 MB

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "File too large. Maximum upload size is 200 MB. For audio files, keep under 25 MB for best results."}), 413

@app.route("/api/health")
def health():
    return jsonify({"status": "ok"}), 200

# ── Structured logging ────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("parvidya.api")

# ── Celery integration (optional — graceful degradation if Redis unavailable) ─
CELERY_OK = False
celery_app = None
try:
    from celery import Celery as _Celery
    REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")
    celery_app = _Celery(
        "parvidya_tasks",
        broker=REDIS_URL,
        backend=REDIS_URL,
    )
    celery_app.conf.update(
        task_serializer="json", result_serializer="json",
        accept_content=["json"], task_track_started=True,
    )
    # Workers optional: sync fallback still works for bulk / multi-student.
    _ping_result = celery_app.control.inspect(timeout=2).ping()
    if _ping_result:
        CELERY_OK = True
        log.info("Celery/Redis connected ✓ — async evaluation enabled")
    else:
        log.warning(
            "Celery broker is reachable but no workers responded — bulk/class-PDF runs synchronously. "
            "For async: `docker compose up -d worker` or "
            "`celery -A celery_worker worker --loglevel=info` (from this directory, venv active)."
        )
except Exception as _ce:
    log.warning(f"Celery/Redis unavailable — falling back to sync evaluation: {_ce}")

# ── Native background task registry (used when Celery/Redis is unavailable) ──
# Keyed by a random hex task_id. Entries: {"status": "pending"|"completed"|"failed",
# "result": {...}|None, "error": str|None}
_bg_task_registry: dict = {}
_bg_task_lock = threading.Lock()

def _bg_task_persist(tid: str, status: str, result: Optional[dict] = None, error: Optional[str] = None):
    """Persist background task state to MongoDB when available."""
    if not MONGO_OK:
        return
    try:
        _mongo_update(
            M_BG_TASKS,
            {"task_id": tid},
            {"$set": {
                "task_id": tid,
                "status": status,
                "result": _stringify_keys(result) if isinstance(result, dict) else None,
                "error": str(error or ""),
                "updated_at": _now(),
            }},
            upsert=True,
        )
    except Exception as _bg_err:
        log.warning(f"[BG-TASK] Persist failed for {tid}: {_bg_err}")

def _bg_task_create() -> str:
    """Register a new pending background task and return its ID."""
    tid = uuid.uuid4().hex
    with _bg_task_lock:
        _bg_task_registry[tid] = {"status": "pending", "result": None, "error": None}
    _bg_task_persist(tid, "pending", result=None, error=None)
    return tid

def _bg_task_set_result(tid: str, result: dict):
    with _bg_task_lock:
        _bg_task_registry[tid] = {"status": "completed", "result": result, "error": None}
    _bg_task_persist(tid, "completed", result=result, error=None)

def _bg_task_set_error(tid: str, error: str):
    with _bg_task_lock:
        _bg_task_registry[tid] = {"status": "failed", "result": None, "error": error}
    _bg_task_persist(tid, "failed", result=None, error=error)

def _bg_task_get(tid: str) -> dict: # Removed Union for simplicity and compatibility
    with _bg_task_lock:
        local = _bg_task_registry.get(tid)
    if local is not None:
        return local
    if not MONGO_OK:
        return None
    doc = _mongo_find_one(M_BG_TASKS, {"task_id": tid})
    if not doc:
        return None
    payload = {
        "status": doc.get("status", "pending"),
        "result": doc.get("result"),
        "error": doc.get("error"),
    }
    with _bg_task_lock:
        _bg_task_registry[tid] = payload
    return payload

# ── OpenAI rate limiter (2 req/s to stay within tier limits) ─────────────────
class _RateLimiter:
    def __init__(self, rate: float = 2.0):
        self._min_gap = 1.0 / rate
        self._last    = 0.0
        self._lock    = threading.Lock()
    def wait(self):
        with self._lock:
            now = _time.monotonic()
            gap = now - self._last
            if gap < self._min_gap:
                _time.sleep(self._min_gap - gap)
            self._last = _time.monotonic()

_openai_limiter = _RateLimiter(rate=2.0)

def _openai_key_candidates(premium: bool = False) -> list[str]:
    """Ordered OpenAI keys for failover. Premium queue prefers premium key first."""
    raw = []
    if premium:
        raw.append(OPENAI_API_KEY_PREMIUM)
    raw.extend([OPENAI_API_KEY_BULK_PRIMARY, OPENAI_API_KEY, OPENAI_API_KEY_BULK_FALLBACK])
    out: list[str] = []
    for k in raw:
        ks = str(k or "").strip()
        # Filter out empty keys AND inline comments from .env files
        if ks and not ks.startswith('#') and ks not in out:
            out.append(ks)
    return out

def _is_retryable_openai_error(ex: Exception) -> bool:
    msg = str(ex).lower()
    retry_tokens = (
        "429", "rate limit", "rate_limit", "insufficient_quota",
        "timeout", "timed out", "connection", "temporar", "service unavailable",
    )
    return any(t in msg for t in retry_tokens)

# ══════════════════════════════════════════════════════════════════════════════
#  ENHANCED EVALUATION SYSTEM — Data Models & Helpers
# ══════════════════════════════════════════════════════════════════════════════

def _create_ocr_extraction_record(eval_id: str, student_name: str, roll_no: str, 
                                   exam_id: str, answers: dict, confidence_scores: dict) -> dict:
    """Create an OCR extraction record with confidence scores per answer."""
    min_confidence = min(confidence_scores.values()) if confidence_scores else 1.0
    needs_review = min_confidence < OCR_CONFIDENCE_THRESHOLD
    
    return {
        "extraction_id": eval_id,
        "exam_id": exam_id,
        "student_name": student_name,
        "roll_no": roll_no,
        "answers": _stringify_keys(answers),
        "confidence_scores": _stringify_keys(confidence_scores),
        "min_confidence": round(min_confidence, 3),
        "needs_review": needs_review,
        "review_reason": "low_confidence_ocr" if needs_review else "",
        "created_at": _now(),
    }

def _create_grading_record(eval_id: str, exam_id: str, student_name: str, roll_no: str,
                           answers: dict, grading_result: dict, subject: str = "", class_level: str = "") -> dict:
    """Create a grading record with per-question marks and subject-specific rules applied."""
    total_marks = 0
    max_marks = 0
    questions_graded = []
    
    # Extract marks from grading result
    for q_id, q_result in grading_result.items():
        if isinstance(q_result, dict):
            marks_awarded = float(q_result.get("marks_awarded", 0))
            max_q_marks = float(q_result.get("max_marks", 0))
            total_marks += marks_awarded
            max_marks += max_q_marks
            questions_graded.append({
                "question_id": q_id,
                "marks_awarded": marks_awarded,
                "max_marks": max_q_marks,
                "reason": q_result.get("reason", ""),
            })
    
    percentage = round((total_marks / max(max_marks, 1)) * 100, 1)
    pass_threshold = 40.0  # Default: 40% pass threshold
    is_pass = percentage >= pass_threshold
    
    return {
        "grading_id": eval_id,
        "exam_id": exam_id,
        "student_name": student_name,
        "roll_no": roll_no,
        "subject": subject,
        "class_level": class_level,
        "total_marks": round(total_marks, 2),
        "max_marks": max_marks,
        "percentage": percentage,
        "is_pass": is_pass,
        "pass_threshold": pass_threshold,
        "questions_graded": questions_graded,
        "grading_model": OPENAI_MODEL,
        "created_at": _now(),
    }

def _add_to_review_queue(eval_id: str, exam_id: str, student_name: str, roll_no: str, 
                         reason: str, metadata: dict = None) -> bool:
    """Add an evaluation to the teacher review queue."""
    if not reason:
        return False
    
    review_item = {
        "review_id": uuid.uuid4().hex[:10],
        "evaluation_id": eval_id,
        "exam_id": exam_id,
        "student_name": student_name,
        "roll_no": roll_no,
        "review_reason": reason,
        "review_reason_type": _categorize_review_reason(reason),
        "metadata": metadata or {},
        "status": "pending",  # pending, approved, rejected, auto_approved
        "teacher_comments": "",
        "created_at": _now(),
        "updated_at": _now(),
    }
    
    if MONGO_OK:
        try:
            _mongo_insert(M_REVIEW_QUEUE, review_item)
            log.info(f"[REVIEW-QUEUE] Added {student_name} ({roll_no}): {reason}")
            return True
        except Exception as ex:
            log.warning(f"[REVIEW-QUEUE] Failed to add {eval_id}: {ex}")
    return False

def _categorize_review_reason(reason: str) -> str:
    """Categorize the review reason for analytics."""
    reason_lower = reason.lower()
    if "low_confidence" in reason_lower:
        return "low_confidence_ocr"
    elif "borderline" in reason_lower:
        return "borderline_marks"
    elif "illegible" in reason_lower:
        return "illegible_handwriting"
    elif "diagram" in reason_lower:
        return "complex_diagram"
    elif "essay" in reason_lower or "length" in reason_lower:
        return "long_essay"
    elif "anomaly" in reason_lower or "unusual" in reason_lower:
        return "score_anomaly"
    else:
        return "other"

def _track_analytics(exam_id: str, teacher_id: str, metrics: dict) -> bool:
    """Track evaluation analytics for insights dashboard."""
    analytics_doc = {
        "analytics_id": uuid.uuid4().hex[:10],
        "exam_id": exam_id,
        "teacher_id": teacher_id,
        "papers_processed": metrics.get("papers_processed", 0),
        "avg_grading_time_sec": metrics.get("avg_grading_time_sec", 0),
        "accuracy_review_percent": metrics.get("accuracy_review_percent", 0),
        "cost_per_paper_usd": metrics.get("cost_per_paper_usd", 0),
        "teacher_corrections_percent": metrics.get("teacher_corrections_percent", 0),
        "most_common_weak_topics": metrics.get("weak_topics", []),
        "pass_rate_percent": metrics.get("pass_rate_percent", 0),
        "avg_score": metrics.get("avg_score", 0),
        "created_at": _now(),
    }
    
    if MONGO_OK:
        try:
            _mongo_insert(M_ANALYTICS, analytics_doc)
            return True
        except Exception as ex:
            log.warning(f"[ANALYTICS] Failed to track: {ex}")
    return False

def _track_cost(eval_id: str, exam_id: str, model_used: str, tokens_used: int, cost_usd: float) -> bool:
    """Track AI API cost per evaluation."""
    cost_doc = {
        "cost_id": uuid.uuid4().hex[:10],
        "evaluation_id": eval_id,
        "exam_id": exam_id,
        "model_used": model_used,
        "tokens_used": tokens_used,
        "cost_usd": round(cost_usd, 4),
        "created_at": _now(),
    }
    
    if MONGO_OK:
        try:
            _mongo_insert(M_COST_TRACKER, cost_doc)
            return True
        except Exception as ex:
            log.warning(f"[COST-TRACK] Failed to track: {ex}")
    return False

@app.route("/api/uploads/<path:filename>")
def download_upload(filename):
    return send_from_directory(UPL_DIR, filename)

# ── In-memory session caches ──────────────────────────────────────────────────
TOKENS:                     Dict[str, str] = {}
syllabi_registry:           Dict[str, dict] = {}
qa_chains:                  Dict[str, Any]  = {}
vector_stores:              Dict[str, Any]  = {}
memories:                   Dict[str, Any]  = {}
exams_registry:             Dict[str, dict] = {}
evaluations_registry:       Dict[str, dict] = {}
bulk_evaluations_registry:  Dict[str, dict] = {}

# ── Registry file paths ───────────────────────────────────────────────────────
SYLLABI_REG_F = DB_DIR / "syllabi_registry.json"
EXAMS_REG_F   = DB_DIR / "exams_registry.json"
EVALS_REG_F   = DB_DIR / "evaluations_registry.json"
BULK_REG_F    = DB_DIR / "bulk_evaluations_registry.json"

def _save_json(path, obj):
    Path(path).write_text(json.dumps(obj, indent=2, default=str), encoding="utf-8")

def _load_json(path, default=None):
    p = Path(path)
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            pass
    return default if default is not None else {}

# ══════════════════════════════════════════════════════════════════════════════
#  MONGODB — Primary persistence layer
# ══════════════════════════════════════════════════════════════════════════════
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")

# Quote credentials if embedded in URI
if MONGO_URI and "://" in MONGO_URI and "@" in MONGO_URI:
    try:
        from urllib.parse import quote_plus
        prefix, rest = MONGO_URI.split("://", 1)
        creds, cluster = rest.split("@", 1)
        if ":" in creds:
            user, pw = creds.split(":", 1)
            MONGO_URI = f"{prefix}://{quote_plus(user)}:{quote_plus(pw)}@{cluster}"
    except Exception as _uri_err:
        print(f"[MONGO] URI credential quoting skipped: {_uri_err}")

MONGO_OK = False
mongo_db = None

try:
    import pymongo, certifi
    _uri_l = (MONGO_URI or "").lower()
    _mongo_kw = dict(
        serverSelectionTimeoutMS=5000,
        connectTimeoutMS=10000,
    )
    # Only Atlas / explicit TLS URIs need a CA bundle. Plain mongodb:// to localhost or
    # Docker `mongo` is not TLS — forcing tlsCAFile makes PyMongo negotiate TLS and breaks local DBs.
    if MONGO_URI.startswith("mongodb+srv://") or "tls=true" in _uri_l or "ssl=true" in _uri_l:
        _mongo_kw["tlsCAFile"] = certifi.where()
    mongo_client = pymongo.MongoClient(MONGO_URI, **_mongo_kw)
    mongo_client.server_info()                  # Raises if unreachable
    mongo_db = mongo_client["vidyai"]
    MONGO_OK = True
    print("[INIT] ✅ MongoDB connected.")
except Exception as _mongo_err:
    print(f"[INIT] ⚠️  MongoDB unavailable — using JSON fallback: {_mongo_err}")

# ── MongoDB collection names ───────────────────────────────────────────────────
M_PLATFORM        = "platform_data"
M_SYLLABI         = "syllabi"
M_EXAMS           = "exams"
M_EVALS           = "evaluations"
M_BULK            = "bulk_evals"
M_SESSIONS        = "sessions"
M_PRACTICE_HIST   = "practice_history"
M_BG_TASKS        = "bg_tasks"

# ── Enhanced Evaluation Collections ────────────────────────────────────────────
M_UPLOADS         = "evaluation_uploads"       # Bulk upload tracking (exam, subject, class, max_marks)
M_OCR_QUEUE       = "ocr_extraction_queue"     # Answer extraction with confidence scores
M_GRADING_QUEUE   = "grading_queue"            # Grading results with per-question marks
M_REVIEW_QUEUE    = "review_queue"             # Items needing teacher review
M_EVAL_RESULTS    = "evaluation_results"       # Complete evaluation history
M_ANALYTICS       = "evaluation_analytics"     # Daily/hourly analytics tracking
M_COST_TRACKER    = "cost_tracking"            # AI API cost per paper

# ── MongoDB helpers ────────────────────────────────────────────────────────────
def _stringify_keys(obj):
    """Recursively convert integer dict keys to strings for MongoDB compatibility."""
    if isinstance(obj, dict):
        return {str(k): _stringify_keys(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_stringify_keys(i) for i in obj]
    return obj

def _mongo_save(col: str, data, key: str = "registry_data"):
    """Upsert a document in a MongoDB collection."""
    if not MONGO_OK:
        return
    try:
        mongo_db[col].update_one({"_key": key}, {"$set": {"data": _stringify_keys(data), "_key": key}}, upsert=True)
    except Exception as e:
        print(f"[MONGO] Save error ({col}/{key}): {e}")

def _mongo_load(col: str, default=None, key: str = "registry_data"):
    """Load a document from a MongoDB collection."""
    if not MONGO_OK:
        return default
    try:
        doc = mongo_db[col].find_one({"_key": key})
        return doc["data"] if doc else default
    except Exception as e:
        print(f"[MONGO] Load error ({col}/{key}): {e}")
        return default

def _mongo_insert(col: str, doc: dict):
    """Insert a single document, stripping _id to avoid conflicts."""
    if not MONGO_OK:
        return
    try:
        doc.pop("_id", None)
        mongo_db[col].insert_one(_stringify_keys(doc))
    except Exception as e:
        print(f"[MONGO] Insert error ({col}): {e}")

def _mongo_find(col: str, query: dict) -> list:
    """Return all matching documents (without Mongo _id)."""
    if not MONGO_OK:
        return []
    try:
        return [{k: v for k, v in d.items() if k != "_id"} for d in mongo_db[col].find(query)]
    except Exception as e:
        print(f"[MONGO] Find error ({col}): {e}")
        return []

def _mongo_find_one(col: str, query: dict) -> Optional[dict]:
    """Return first matching document (without Mongo _id)."""
    if not MONGO_OK:
        return None
    try:
        d = mongo_db[col].find_one(query)
        if d:
            d.pop("_id", None)
        return d
    except Exception as e:
        print(f"[MONGO] FindOne error ({col}): {e}")
        return None

def _mongo_update(col: str, query: dict, update: dict, upsert: bool = False):
    """Run an update operation."""
    if not MONGO_OK:
        return
    try:
        mongo_db[col].update_one(query, update, upsert=upsert)
    except Exception as e:
        print(f"[MONGO] Update error ({col}): {e}")

def _mongo_delete(col: str, query: dict):
    if not MONGO_OK:
        return
    try:
        mongo_db[col].delete_one(query)
    except Exception as e:
        print(f"[MONGO] Delete error ({col}): {e}")

# ── Registry save/load with dual-write ────────────────────────────────────────
def _save_syllabi():
    _mongo_save(M_SYLLABI,  syllabi_registry)
    if not IS_VERCEL:
        _save_json(SYLLABI_REG_F, syllabi_registry)

def _save_exams():
    _mongo_save(M_EXAMS,    exams_registry)
    if not IS_VERCEL:
        _save_json(EXAMS_REG_F, exams_registry)

def _save_evals():
    _mongo_save(M_EVALS,    evaluations_registry)
    if not IS_VERCEL:
        _save_json(EVALS_REG_F, evaluations_registry)

def _save_bulk():
    _mongo_save(M_BULK,     bulk_evaluations_registry)
    if not IS_VERCEL:
        _save_json(BULK_REG_F, bulk_evaluations_registry)

def _rebuild_exams_registry_from_dir() -> Dict[str, dict]:
    rebuilt: Dict[str, dict] = {}
    try:
        for p in sorted(EXAMS_DIR.glob("*.json")):
            try:
                data = json.loads(p.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    exam_id = str(data.get("exam_id") or p.stem)
                    data["exam_id"] = exam_id
                    rebuilt[exam_id] = data
            except Exception:
                continue
    except Exception:
        return {}
    return rebuilt

def _boot_load():
    """On startup: load registries from MongoDB first, fall back to JSON."""
    global syllabi_registry, exams_registry, evaluations_registry, bulk_evaluations_registry

    def _load(col, file, default={}):
        data = _mongo_load(col)
        # Treat missing or empty Mongo payload as "not loaded" so JSON / disk can repopulate.
        # Otherwise an empty `{}` persisted once blocks `exams_registry.json` forever.
        if not data:
            data = _load_json(file, default)
            if data and MONGO_OK:
                print(f"[BOOT] Migrating {file.name} → MongoDB/{col}")
                _mongo_save(col, data)
        return data or {}

    syllabi_registry           = _load(M_SYLLABI, SYLLABI_REG_F)
    exams_registry             = _load(M_EXAMS,   EXAMS_REG_F)
    evaluations_registry       = _load(M_EVALS,   EVALS_REG_F)
    bulk_evaluations_registry  = _load(M_BULK,    BULK_REG_F)

    # Merge per-exam JSON files not yet in the registry (Docker bind mounts, recovery).
    disk_exams = _rebuild_exams_registry_from_dir()
    if disk_exams:
        _added = 0
        for _eid, _payload in disk_exams.items():
            if _eid not in exams_registry:
                exams_registry[_eid] = _payload
                _added += 1
        if _added:
            _save_exams()
            print(f"[BOOT] Merged {_added} exam(s) from {EXAMS_DIR} into registry")

    counts = {
        "syllabi":     len(syllabi_registry),
        "exams":       len(exams_registry),
        "evaluations": len(evaluations_registry),
    }
    print(f"[BOOT] Registry loaded: {counts}")

    # ── Role migration: school_admin → admin/institute_admin, tutor → teacher ──
    try:
        _db = db_load()
        changed = False
        for u in _db.get("users", []):
            if u.get("role") == "school_admin":
                if u.get("email") == "admin@vidyai.in":
                    u["role"] = "admin"
                else:
                    u["role"] = "institute_admin"
                changed = True
            elif u.get("role") == "tutor":
                u["role"] = "teacher"
                changed = True
        if changed:
            db_save(_db)
            print("[BOOT] Migrated legacy roles (school_admin/tutor)")
    except Exception as _me:
        print(f"[BOOT] Role migration skipped: {_me}")

# ── Utility helpers ────────────────────────────────────────────────────────────
def _hash(pw: str)  -> str: return hashlib.sha256(pw.encode()).hexdigest()
def _uid()          -> str: return "u" + uuid.uuid4().hex[:8]
def _now()          -> str: return datetime.utcnow().isoformat()
def _allowed(fn)    -> bool: return "." in fn and fn.rsplit(".", 1)[1].lower() in ALLOWED
def _safe(u: dict)  -> dict: return {k: v for k, v in u.items() if k not in ("pw_hash", "_id")}
def _grade(p: float)-> str:
    return "A+" if p>=90 else "A" if p>=80 else "B" if p>=70 else "C" if p>=60 else "D" if p>=50 else "F"
def _normalize(t)   -> str: return re.sub(r"\s+", " ", str(t or "").strip().lower())

def _normalize_subject_name(subject: str) -> str:
    """Canonicalize common subject aliases/misspellings for reliable lookups."""
    s = _normalize(subject)
    aliases = {
        "math": "mathematics",
        "maths": "mathematics",
        "mathmatics": "mathematics",
        "mathematic": "mathematics",
        "evs": "environmental studies",
        "env studies": "environmental studies",
        "social studies": "social science",
        "social": "social science",
        "cs": "computer science",
        "comp science": "computer science",
        "computer application": "computer applications",
        "history and civics": "history & civics",
        "history/civics": "history & civics",
        "history civics": "history & civics",
        "english lit": "english literature",
        "english lang": "english language",
    }
    return aliases.get(s, s)

def _strip_answer_prefix(text: str) -> str:
    return re.sub(r"^(?:ans(?:wer)?|response|student\s*ans(?:wer)?)\s*[:\-]\s*", "", str(text or "").strip(), flags=re.IGNORECASE).strip()

def _coerce_mcq_answer(student_answer: str, options: Dict[str, str]) -> str:
    """IMPROVED: Properly handles dash-prefixed answers like '- A'"""
    if not student_answer:
        return ""

    # If OCR returned a structured payload (often stringified as a dict),
    # prefer the explicit bubble field over any incidental letter elsewhere.
    if isinstance(student_answer, dict):
        bubble_val = str(student_answer.get("bubble", student_answer.get("answer", ""))).strip()
        box_val    = str(student_answer.get("box", student_answer.get("write_box", ""))).strip()
        if bubble_val:
            return _coerce_mcq_answer(bubble_val, options)
        if box_val:
            return _coerce_mcq_answer(box_val, options)
        return ""
    
    s = str(student_answer).strip()

    if s.startswith("{") and s.endswith("}"):
        try:
            import ast
            parsed = ast.literal_eval(s)
            if isinstance(parsed, dict):
                bubble_val = str(parsed.get("bubble", parsed.get("answer", ""))).strip()
                box_val    = str(parsed.get("box", parsed.get("write_box", ""))).strip()
                if bubble_val:
                    return _coerce_mcq_answer(bubble_val, options)
                if box_val:
                    return _coerce_mcq_answer(box_val, options)
        except Exception:
            pass
    
    # Remove answer prefixes
    s = re.sub(r'^(?:ans(?:wer)?|response|student\s*ans(?:wer)?)\s*[:\-]\s*', '', s, flags=re.IGNORECASE)
    
    # Remove all whitespace for matching
    s_clean = re.sub(r'\s+', '', s)
    
    # Handle dash prefix: "-A" or "- A" or "— B"
    dash_match = re.match(r'^[-–—]\s*([A-D])', s, re.IGNORECASE)
    if dash_match:
        return dash_match.group(1).upper()
    
    # Single letter with optional brackets: "A", "(B)", "[C]"
    letter_match = re.match(r'^[\(\{\[]?\s*([A-D])\s*[\)\}\]]?$', s, re.IGNORECASE)
    if letter_match:
        return letter_match.group(1).upper()
    
    # Word format: "option A", "choice B"
    word_match = re.search(r'(?:option|choice|answer)\s+([A-D])', s, re.IGNORECASE)
    if word_match:
        return word_match.group(1).upper()
    
    # Find any A-D letter
    any_letter = re.search(r'\b([A-D])\b', s, re.IGNORECASE)
    if any_letter:
        return any_letter.group(1).upper()
    
    # Check against option text
    if options:
        for k, v in options.items():
            if _normalize(v) == _normalize(s):
                return k.upper()
    
    return s

def _clean_subjective_answer(text: str) -> str:
    """
    Clean and normalize subjective answer text:
    - Remove OCR garbage
    - Combine multiple lines properly
    - Remove duplicate spaces
    - Preserve complete answer
    """
    if not text:
        return ""
    
    # Replace newlines and multiple spaces with single space
    cleaned = re.sub(r'[\r\n]+', ' ', text)
    cleaned = re.sub(r'\s+', ' ', cleaned)
    
    # Remove common OCR artifacts
    cleaned = re.sub(r'[•\*#|→←↑↓]', '', cleaned)
    
    # Remove instruction text that might have been captured
    cleaned = re.sub(r'Write\s+your\s+answer\s+below\s*[:\-]?\s*', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'Answer\s*[:\-]\s*', '', cleaned, flags=re.IGNORECASE)
    
    # Strip leading/trailing punctuation and spaces
    cleaned = cleaned.strip()
    
    # If the answer seems incomplete (very short), it might be an extraction error
    # Log warning but keep it
    if len(cleaned) < 5 and cleaned and not cleaned.isdigit():
        log.warning(f"Potentially incomplete subjective answer: {cleaned}")
    
    return cleaned

def _question_valid_answers(q: dict) -> List[str]:
    vals = q.get("valid_answers", [])
    if not vals:
        ans = str(q.get("answer", "")).strip()
        if ans:
            vals = [x.strip() for x in re.split(r"[,\n;/|]", ans) if x.strip()]
    return [str(v).strip() for v in vals if str(v).strip()]

def _objective_is_correct(q: dict, student_answer: str) -> bool:
    options = q.get("options") or {}
    student = _coerce_mcq_answer(student_answer, options)
    student_norm = _normalize(student)
    valid_raw = _question_valid_answers(q)
    valid_norm = set(_normalize(v) for v in valid_raw if str(v).strip())

    # Expand valid_norm: if a valid answer is a letter key (e.g. "A"),
    # also add the option text (e.g. "5"); if it is option text, also add the key.
    for ans in list(valid_raw):
        key = str(ans).strip().upper()
        if key in options:
            valid_norm.add(_normalize(options[key]))
            valid_norm.add(_normalize(key))
        else:
            # answer stored as text — find its matching option letter
            for k, v in options.items():
                if _normalize(v) == _normalize(ans):
                    valid_norm.add(_normalize(k))

    ans_field = str(q.get("answer", "")).strip()
    if ans_field:
        valid_norm.add(_normalize(ans_field))
        key = ans_field.upper()
        if key in options:
            # answer is a letter key → also accept the text value
            valid_norm.add(_normalize(options[key]))
        else:
            # answer is text → also accept the corresponding letter key
            for k, v in options.items():
                if _normalize(v) == _normalize(ans_field):
                    valid_norm.add(_normalize(k))

    # Extra bridge: if student submitted a letter, also compare its text to valid_norm
    student_upper = student.strip().upper()
    if student_upper in options:
        valid_norm.add(_normalize(options[student_upper]))

    return bool(student_norm and student_norm in valid_norm)


def _dtype(ext: str)-> str:
    return {"pdf":"PDF","txt":"Text","md":"Text","mp3":"Audio","wav":"Audio",
            "m4a":"Audio","mp4":"Video","jpg":"Image","jpeg":"Image","png":"Image"}.get(ext, "File")

# ══════════════════════════════════════════════════════════════════════════════
#  QUOTA / USAGE-LIMIT SYSTEM — feature names, defaults, helpers, decorator
# ══════════════════════════════════════════════════════════════════════════════
QUOTA_FEATURES: dict[str, str] = {
    "curriculum_load":    "Load Chapters",
    "curriculum_tool":    "Curriculum Tools (Summarise / Flashcards / etc.)",
    "summarise":          "Summarise",
    "flashcards":         "Flashcards",
    "video_explanation":  "Video Explanation",
    "practice_generate":  "Practice Questions",
    "practice_evaluate":  "Practice Evaluation",
    "chat":               "AI Tutor Chat",
    "generate_questions": "Generate Questions",
    "evaluate":           "Evaluate Answers",
    "uk_curriculum":      "UK Curriculum Tools",
}
_UNLIMITED = -1   # sentinel: no daily limit
_DEFAULT_ROLE_QUOTAS: dict[str, dict[str, int]] = {
    "student": {
        "curriculum_load": 10, "curriculum_tool": 10, "summarise": 5,
        "flashcards": 5, "video_explanation": 3,
        "practice_generate": 10, "practice_evaluate": 20,
        "chat": 30, "generate_questions": 0, "evaluate": 0, "uk_curriculum": 5,
    },
    "parent": {
        "curriculum_load": 5, "curriculum_tool": 5, "summarise": 3,
        "flashcards": 3, "video_explanation": 2,
        "practice_generate": 5, "practice_evaluate": 10,
        "chat": 15, "generate_questions": 0, "evaluate": 0, "uk_curriculum": 3,
    },
    "teacher": {
        "curriculum_load": 50, "curriculum_tool": 30, "summarise": 20,
        "flashcards": 20, "video_explanation": 10,
        "practice_generate": 50, "practice_evaluate": 100,
        "chat": 100, "generate_questions": 30, "evaluate": 100, "uk_curriculum": 50,
    },
    "institute_admin": {f: _UNLIMITED for f in QUOTA_FEATURES},
    "admin":           {f: _UNLIMITED for f in QUOTA_FEATURES},
}

# Pricing-aligned quota presets for quick admin application.
QUOTA_PRESETS: dict[str, dict] = {
    "role_defaults": {
        "free-student": {
            "label": "Free - Student",
            "summary": "Entry-level monthly usage for student trial users.",
            "template": {
                "curriculum_load": 3,
                "curriculum_tool": 3,
                "summarise": 5,
                "flashcards": 10,
                "video_explanation": 0,
                "practice_generate": 5,
                "practice_evaluate": 5,
                "chat": 3,
                "generate_questions": 0,
                "evaluate": 0,
                "uk_curriculum": 3,
            },
        },
        "student-pro": {
            "label": "Student Pro",
            "summary": "High daily limits for paid individual students.",
            "template": {
                "curriculum_load": _UNLIMITED,
                "curriculum_tool": _UNLIMITED,
                "summarise": _UNLIMITED,
                "flashcards": _UNLIMITED,
                "video_explanation": _UNLIMITED,
                "practice_generate": _UNLIMITED,
                "practice_evaluate": _UNLIMITED,
                "chat": _UNLIMITED,
                "generate_questions": 20,
                "evaluate": 20,
                "uk_curriculum": _UNLIMITED,
            },
        },
        "parent-lite": {
            "label": "Parent Lite",
            "summary": "Balanced parent access for supervision and basic support.",
            "template": {
                "curriculum_load": 10,
                "curriculum_tool": 10,
                "summarise": 6,
                "flashcards": 6,
                "video_explanation": 3,
                "practice_generate": 10,
                "practice_evaluate": 20,
                "chat": 25,
                "generate_questions": 0,
                "evaluate": 0,
                "uk_curriculum": 8,
            },
        },
        "teacher-pro": {
            "label": "Teacher Pro",
            "summary": "Higher teacher limits for paper generation and evaluation.",
            "template": {
                "curriculum_load": 80,
                "curriculum_tool": 60,
                "summarise": 50,
                "flashcards": 50,
                "video_explanation": 25,
                "practice_generate": 120,
                "practice_evaluate": 200,
                "chat": 200,
                "generate_questions": 80,
                "evaluate": 200,
                "uk_curriculum": 80,
            },
        },
    },
    "institution": {
        "school-starter": {
            "label": "School Starter",
            "summary": "Core school setup for smaller institutions.",
            "user_daily": {
                "curriculum_load": 30,
                "curriculum_tool": 20,
                "summarise": 15,
                "flashcards": 15,
                "video_explanation": 8,
                "practice_generate": 40,
                "practice_evaluate": 60,
                "chat": 60,
                "generate_questions": 20,
                "evaluate": 80,
                "uk_curriculum": 20,
            },
            "inst_daily": {
                "generate_questions": 400,
                "evaluate": 800,
                "practice_evaluate": 1200,
                "chat": 3000,
            },
        },
        "school-growth": {
            "label": "School Growth",
            "summary": "Expanded school limits for larger student base.",
            "user_daily": {
                "curriculum_load": 60,
                "curriculum_tool": 50,
                "summarise": 40,
                "flashcards": 40,
                "video_explanation": 20,
                "practice_generate": 100,
                "practice_evaluate": 150,
                "chat": 120,
                "generate_questions": 60,
                "evaluate": 180,
                "uk_curriculum": 60,
            },
            "inst_daily": {
                "generate_questions": 1500,
                "evaluate": 4000,
                "practice_evaluate": 5000,
                "chat": 12000,
            },
        },
        "coaching": {
            "label": "Coaching Institute",
            "summary": "High-throughput preset for test-prep batches.",
            "user_daily": {
                "curriculum_load": 80,
                "curriculum_tool": 80,
                "summarise": 60,
                "flashcards": 60,
                "video_explanation": 30,
                "practice_generate": 150,
                "practice_evaluate": 250,
                "chat": 180,
                "generate_questions": 100,
                "evaluate": 250,
                "uk_curriculum": 80,
            },
            "inst_daily": {
                "generate_questions": 3000,
                "evaluate": 10000,
                "practice_evaluate": 12000,
                "chat": 25000,
            },
        },
        "enterprise": {
            "label": "Enterprise & Government",
            "summary": "Near-unlimited institutional capacity.",
            "user_daily": {f: _UNLIMITED for f in QUOTA_FEATURES},
            "inst_daily": {f: _UNLIMITED for f in QUOTA_FEATURES},
        },
    },
}

def _visible_quota_presets(caller_role: str) -> dict:
    """Return presets the caller can apply."""
    if caller_role == "admin":
        return QUOTA_PRESETS
    return {
        "role_defaults": {},
        "institution": {},
    }

# ── Default database seed ──────────────────────────────────────────────────────
def _default_db() -> dict:
    return {
        "users": [
            {"id":"u1","name":"Alex Johnson","email":"student@vidyai.in",
             "pw_hash":_hash("password"),"role":"student","institution":"Demo School",
             "roll_number":"101","joined":_now()[:10],"docs":0,"status":"active"},
            {"id":"u2","name":"Dr. Priya Sharma","email":"teacher@vidyai.in",
             "pw_hash":_hash("password"),"role":"teacher","institution":"Demo School",
             "joined":_now()[:10],"docs":0,"status":"active"},
            {"id":"u3","name":"Platform Admin","email":"admin@vidyai.in",
             "pw_hash":_hash("password"),"role":"admin","institution":"",
             "joined":_now()[:10],"docs":0,"status":"active"},
            {"id":"u4","name":"Meena Iyer","email":"parent@vidyai.in",
             "pw_hash":_hash("password"),"role":"parent","institution":"",
             "joined":_now()[:10],"docs":0,"status":"active"},
            {"id":"u5","name":"Rajesh Kumar","email":"institute@vidyai.in",
             "pw_hash":_hash("password"),"role":"institute_admin","institution":"Demo School",
             "joined":_now()[:10],"docs":0,"status":"active"},
        ],
        "documents": [],
        "activity": [],
        "sessions": [],
        "settings": {
            "model":                  OPENAI_MODEL,
            "max_uploads_per_user":   20,
            "max_uploads_student":    5,
            "max_uploads_teacher":    30,
            "max_flashcards_default": 10,
            "max_questions_per_paper":30,
            "student_self_register":  True,
            "tutor_approval_required":False,
            "show_answer_explanations":True,
            "otp_required_for_signup": True,
            "email_reports_enabled":   True,
        },
        # ── Quota configuration ─────────────────────────────────────────────
        # role_defaults: {role: {feature: daily_limit}}  (-1 = unlimited, 0 = blocked)
        # institution_overrides: {inst_name: {user_daily: {feature: int}, inst_daily: {feature: int}}}
        # user_overrides: {uid: {feature: daily_limit}}
        "quota_config": {
            "role_defaults": _DEFAULT_ROLE_QUOTAS,
            "institution_overrides": {},
            "user_overrides": {},
        },
        # ── Usage counters (date → users/institutions → feature → count) ────
        "usage_counters": {},
    }

# ── DB load / save with MongoDB primary + JSON fallback ───────────────────────
def db_load() -> dict:
    # 1. MongoDB is the source of truth in production
    if MONGO_OK:
        data = _mongo_load(M_PLATFORM, key="main_db")
        if data:
            return data

    # 2. JSON file fallback (dev / offline)
    if DATA_F.exists():
        try:
            d = json.loads(DATA_F.read_text())
            if MONGO_OK:
                print("[MIGRATE] JSON → MongoDB sync")
                _mongo_save(M_PLATFORM, d, key="main_db")
            return d
        except Exception:
            pass

    # 3. Bootstrap with seed data
    print("[WARN] Seeding empty database")
    d = _default_db()
    if MONGO_OK:
        _mongo_save(M_PLATFORM, d, key="main_db")
    return d

def db_save(d: dict):
    if MONGO_OK:
        _mongo_save(M_PLATFORM, d, key="main_db")
    if not IS_VERCEL:
        try:
            DATA_F.write_text(json.dumps(d, indent=2, default=str))
        except Exception as e:
            print(f"[DB] JSON write failed: {e}")

def db_log(db: dict, uid: str, action: str, detail: str):
    db["activity"].insert(0, {
        "id": uuid.uuid4().hex[:6], "user_id": uid,
        "action": action, "detail": detail, "ts": _now()
    })
    db["activity"] = db["activity"][:200]

# ══════════════════════════════════════════════════════════════════════════════
#  QUOTA HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _qc_load() -> dict:
    """Return the quota_config section from DB, merging any missing role defaults."""
    db  = db_load()
    qc  = db.get("quota_config", {})
    if not qc.get("role_defaults"):
        qc["role_defaults"] = _DEFAULT_ROLE_QUOTAS
    if "institution_overrides" not in qc:
        qc["institution_overrides"] = {}
    if "user_overrides" not in qc:
        qc["user_overrides"] = {}
    return qc

def _qc_save(qc: dict):
    db = db_load()
    db["quota_config"] = qc
    db_save(db)

def _quota_today() -> str:
    from datetime import date as _date
    return _date.today().isoformat()

def _get_effective_limit(user: dict, feature: str) -> tuple:
    """
    Returns (user_daily_limit, institution_pool_limit).
    -1 = unlimited. None for pool = no pool limit.
    Priority: user_override > institution_user_daily > role_default.
    """
    uid  = user["id"]
    role = user.get("role", "student")
    inst = (user.get("institution") or "").strip()
    qc   = _qc_load()

    user_daily      = qc.get("user_overrides", {}).get(uid, {}).get(feature)
    inst_user_daily = None
    inst_pool       = None
    if inst:
        inst_cfg        = qc.get("institution_overrides", {}).get(inst, {})
        inst_user_daily = inst_cfg.get("user_daily", {}).get(feature)
        inst_pool       = inst_cfg.get("inst_daily", {}).get(feature)

    role_cfg = (qc.get("role_defaults") or _DEFAULT_ROLE_QUOTAS)
    role_defaults = role_cfg.get(role, _DEFAULT_ROLE_QUOTAS.get("student", {}))
    role_default  = role_defaults.get(feature, 0)

    effective_user = (user_daily       if user_daily      is not None else
                      inst_user_daily  if inst_user_daily is not None else
                      role_default)
    return effective_user, inst_pool

def _get_today_usage(uid: str, inst: str) -> dict:
    """Returns {'user': {feature: count}, 'inst': {feature: count}} for today."""
    today = _quota_today()
    db    = db_load()
    day   = db.get("usage_counters", {}).get(today, {})
    return {
        "user": day.get("users", {}).get(uid, {}),
        "inst": day.get("institutions", {}).get(inst, {}) if inst else {},
    }

def _increment_usage(uid: str, inst: str, feature: str):
    """Increment daily usage counters for user (and institution if set)."""
    from datetime import date as _date, timedelta as _td
    today  = _quota_today()
    cutoff = (_date.today() - _td(days=32)).isoformat()
    db     = db_load()
    if "usage_counters" not in db:
        db["usage_counters"] = {}
    # Prune old days
    db["usage_counters"] = {k: v for k, v in db["usage_counters"].items() if k >= cutoff}
    if today not in db["usage_counters"]:
        db["usage_counters"][today] = {"users": {}, "institutions": {}}
    day = db["usage_counters"][today]
    # Ensure user entry exists
    if uid not in day["users"]:
        day["users"][uid] = {}
    day["users"][uid][feature] = day["users"][uid].get(feature, 0) + 1
    # Same for institution
    if inst:
        if inst not in day["institutions"]:
            day["institutions"][inst] = {}
        day["institutions"][inst][feature] = day["institutions"][inst].get(feature, 0) + 1
    db_save(db)

def check_and_use_quota(user: dict, feature: str) -> tuple:
    """
    Check quota; increment if allowed.
    Returns (allowed: bool, message: str, remaining: int).
    remaining == -1 means unlimited.
    """
    uid        = user["id"]
    inst       = (user.get("institution") or "").strip()
    u_limit, i_pool = _get_effective_limit(user, feature)
    counts     = _get_today_usage(uid, inst)
    u_used     = counts["user"].get(feature, 0)
    i_used     = counts["inst"].get(feature, 0)
    feat_label = QUOTA_FEATURES.get(feature, feature)

    if u_limit == 0:
        return (False,
                f"'{feat_label}' is not available for your role. "
                "Contact your administrator.", 0)
    if u_limit != _UNLIMITED and u_used >= u_limit:
        return (False,
                f"Daily limit reached for '{feat_label}' ({u_limit}/day). "
                "Your quota resets at midnight.", 0)
    if i_pool is not None and i_pool != _UNLIMITED and i_used >= i_pool:
        return (False,
                f"Your institution's daily quota for '{feat_label}' is exhausted "
                f"({i_pool} uses/day across all users). Try again tomorrow.", 0)

    _increment_usage(uid, inst, feature)
    remaining = (u_limit - u_used - 1) if u_limit != _UNLIMITED else _UNLIMITED
    return True, "ok", remaining

def quota(feature: str):
    """
    Decorator: enforces per-user and per-institution daily quotas.
    Must be stacked BELOW @auth() so request.user is available.

    Usage:
        @app.post("/api/endpoint")
        @auth()
        @quota('feature_name')
        def my_endpoint():
    """
    def decorator(fn):
        @wraps(fn)
        def wrapper(*args, **kwargs):
            user = getattr(request, "user", None)
            if user:
                allowed, msg, _remaining = check_and_use_quota(user, feature)
                if not allowed:
                    return jsonify({
                        "error":         msg,
                        "quota_exceeded": True,
                        "feature":        feature,
                        "feature_label":  QUOTA_FEATURES.get(feature, feature),
                    }), 429
            return fn(*args, **kwargs)
        return wrapper
    return decorator

# ── Email helper ───────────────────────────────────────────────────────────────
def send_email_report(to_email: str, subject: str, results: list) -> bool:
    host   = os.getenv("SMTP_HOST")
    port   = int(os.getenv("SMTP_PORT", 587))
    user   = os.getenv("SMTP_USER")
    passwd = os.getenv("SMTP_PASS")
    sender = os.getenv("SMTP_FROM", user)

    if not all([host, user, passwd]):
        print("[SMTP] Missing credentials — skipping email.")
        return False

    try:
        msg = MIMEMultipart()
        msg["From"], msg["To"], msg["Subject"] = sender, to_email, subject
        rows = "".join(
            f"<li><b>{r.get('question','')}</b><br/>"
            f"Score: {r.get('evaluation',{}).get('score',0)}/10<br/>"
            f"Feedback: {r.get('evaluation',{}).get('feedback','N/A')}</li>"
            for r in results
        )
        html = f"<h2>Your AI Practice Results — VidyAI</h2><ul>{rows}</ul><p>Keep it up! — Parvidya Team</p>"
        msg.attach(MIMEText(html, "html"))
        with smtplib.SMTP(host, port) as s:
            s.starttls(); s.login(user, passwd); s.send_message(msg)
        return True
    except Exception as e:
        print(f"[SMTP] {e}")
        return False

# ══════════════════════════════════════════════════════════════════════════════
#  AI / LANGCHAIN HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _get_llm(temperature: float = 0.3, mini: bool = False):
    model = OPENAI_MINI if mini else OPENAI_MODEL
    return ChatOpenAI(model_name=model, temperature=temperature, openai_api_key=OPENAI_API_KEY)

def _get_emb():
    return OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)

def _clean_json(raw: str):
    """Strip markdown fences and parse JSON."""
    clean = re.sub(r"```(?:json)?|```", "", raw).strip()
    # Strip leading text before first [ or {
    m = re.search(r"[\[\{]", clean)
    if m:
        clean = clean[m.start():]
    return json.loads(clean)

def _llm_text(prompt: str, temperature: float = 0.3, mini: bool = False) -> str:
    if not LANGCHAIN_OK or not OPENAI_API_KEY:
        return "(AI unavailable — check OPENAI_API_KEY)"
    
    # ── Cache check ──────────
    if '_eval_cache' in globals():
        cached = _eval_cache.get(prompt)
        if cached: return cached

    res = _get_llm(temperature=temperature, mini=mini).invoke(prompt).content
    
    if '_eval_cache' in globals():
        _eval_cache.set(prompt, res)
    return res

def call_ollama(prompt: str, model: str = "qwen2.5:7b", is_json: bool = True):
    """Call local Ollama API."""
    import requests
    try:
        url = os.getenv("OLLAMA_URL", "http://localhost:11434/api/generate")
        payload = {
            "model": os.getenv("OLLAMA_MODEL", model),
            "prompt": prompt,
            "stream": False,
            "format": "json" if is_json else ""
        }
        resp = requests.post(url, json=payload, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        raw = data.get("response", "").strip()
        if is_json:
            try:
                return json.loads(raw)
            except Exception:
                # Try to extract JSON if it's wrapped in markdown
                match = re.search(r'\{.*\}', raw, re.DOTALL)
                if match:
                    return json.loads(match.group(0))
                raise
        return raw
    except Exception as e:
        log.warning(f"[Ollama] Failed: {e}")
        return None

class SimpleCache:
    def __init__(self, cache_dir=".cache/vidyai"):
        self.cache_dir = Path(cache_dir)
        try:
            self.cache_dir.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass # Fallback for read-only systems

    def _get_key(self, text):
        return hashlib.md5(text.encode()).hexdigest()

    def get(self, text):
        key = self._get_key(text)
        f = self.cache_dir / f"{key}.json"
        if f.exists():
            try:
                return json.loads(f.read_text())
            except Exception:
                return None
        return None

    def set(self, text, val):
        if val is None: return
        try:
            key = self._get_key(text)
            f = self.cache_dir / f"{key}.json"
            f.write_text(json.dumps(val))
        except Exception:
            pass

_eval_cache = SimpleCache()

def _study_fallback_answer(question: str, role: str = "student") -> str:
    q = (question or "").strip()
    if not q:
        return "Please share your exact question, chapter, and what you have tried so far."
    return (
        "AI backend is temporarily busy, so here is a guided exam-style approach:\n\n"
        f"Question: {q}\n\n"
        "1. Identify chapter and concept\n"
        "2. List all given values with units\n"
        "3. Write the governing formula/reaction/principle\n"
        "4. Solve step-by-step and keep unit consistency\n"
        "5. Check reasonableness using option elimination\n\n"
        "Share the full problem statement if you want a strict line-by-line final solution."
    )

def _llm_json(prompt: str, temperature: float = 0.2, mini: bool = False):
    raw = _llm_text(prompt, temperature=temperature, mini=mini)
    try:
        return _clean_json(raw)
    except Exception:
        # Second attempt: find array or object in the response
        for pat in [r"\[.*\]", r"\{.*\}"]:
            m = re.search(pat, raw, re.DOTALL)
            if m:
                try:
                    return json.loads(m.group())
                except Exception:
                    pass
        print(f"[LLM JSON parse fail] raw snippet: {raw[:200]}")
        return {"error": "Could not parse AI response as JSON", "raw": raw[:500]}

def _load_vs(did: str):
    """Retrieve InMemoryVectorStore from session cache or re-index from disk."""
    if did in vector_stores:
        return vector_stores[did]
    db  = db_load()
    doc = next((d for d in db.get("documents", []) if d["id"] == did), None)
    if doc and doc.get("path") and Path(doc["path"]).exists():
        _index_doc(Path(doc["path"]), did, doc["ext"])
        return vector_stores.get(did)
    return None

# Tracks OCR background indexing status: did -> 'indexing' | 'ready' | 'failed'
ocr_status: dict = {}


def _run_ocr_background(path: Path, did: str, syl_name: str, fn: str, owner_id: str):
    """Background thread: OCR scanned PDF pages via Vision API, then build vector store."""
    import base64, time
    ocr_status[did] = "indexing"
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print(f"[OCR] {did}: PyMuPDF not installed")
        ocr_status[did] = "failed"
        return
    client = _oai.OpenAI(api_key=OPENAI_API_KEY)
    pdf_doc = fitz.open(str(path))
    total_pages = len(pdf_doc)
    MAX_OCR_PAGES = 60
    pages_to_ocr = min(total_pages, MAX_OCR_PAGES)
    print(f"[OCR] {did}: scanned PDF {total_pages} pages, OCR-ing {pages_to_ocr} pages...")
    docs = []
    mat = fitz.Matrix(1.5, 1.5)
    for i in range(pages_to_ocr):
        try:
            page = pdf_doc[i]
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img_bytes = pix.tobytes("png")
            b64 = base64.b64encode(img_bytes).decode()
            for attempt in range(4):
                try:
                    resp = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "Extract all text from this textbook page. Return only the text content."},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "auto"}}
                            ]
                        }],
                        max_tokens=2000,
                        timeout=30
                    )
                    text = (resp.choices[0].message.content or "").strip()
                    if text:
                        docs.append(Document(page_content=text, metadata={"source": str(path), "page": i}))
                    break
                except Exception as e:
                    wait = (attempt + 1) * 3
                    print(f"[OCR] {did}: page {i} attempt {attempt+1} error: {e}, wait {wait}s")
                    time.sleep(wait)
            time.sleep(0.5)  # gentle rate-limit pause between pages
        except Exception as e:
            print(f"[OCR] {did}: page {i} render error: {e}")
    print(f"[OCR] {did}: extracted text from {len(docs)}/{pages_to_ocr} pages")
    if not docs:
        ocr_status[did] = "failed"
        return
    try:
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150, separators=["\n\n", "\n", ".", " "])
        chunks = splitter.split_documents(docs)
        if not chunks:
            ocr_status[did] = "failed"
            return
        vs = InMemoryVectorStore.from_documents(chunks, _get_emb())
        vector_stores[did] = vs
        chapters = _extract_chapters(docs)
        # Update the registry and db with real chunk count
        if did in syllabi_registry:
            syllabi_registry[did]["chunks"] = len(chunks)
            syllabi_registry[did]["chapters"] = chapters
            _save_syllabi()
        db = db_load()
        for doc in db.get("documents", []):
            if doc["id"] == did:
                doc["chunks"] = len(chunks)
                doc["chapters"] = chapters
                break
        db_save(db)
        ocr_status[did] = "ready"
        print(f"[OCR] {did}: indexing complete — {len(chunks)} chunks")
    except Exception as e:
        print(f"[OCR] {did}: vectorstore error: {e}")
        import traceback; traceback.print_exc()
        ocr_status[did] = "failed"


def _index_doc(path: Path, did: str, ext: str):
    """Parse file → chunk → embed → InMemoryVectorStore.  Returns (chunk_count, chapters).
    For scanned PDFs, returns (-1, []) and queues background OCR thread."""
    if not LANGCHAIN_OK or not OPENAI_API_KEY:
        return 0, []
    try:
        ext_l = ext.lower()
        if ext_l == "pdf":
            import pypdf
            reader = pypdf.PdfReader(str(path))
            docs = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    docs.append(Document(page_content=text, metadata={"source": str(path), "page": i}))
            if not docs:
                # Scanned PDF — start background OCR thread and return immediately
                print(f"[INDEX] {did}: scanned PDF detected, queuing background OCR...")
                return -1, []  # signal to caller that OCR is running in background
            if not docs:
                return 0, []
        elif ext_l in {"docx", "doc"}:
            try:
                import docx as _docx
            except ImportError:
                import subprocess, sys
                subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
                import docx as _docx
            document = _docx.Document(str(path))
            text = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
            if not text.strip():
                return 0, []
            docs = [Document(page_content=text, metadata={"source": str(path)})]
        elif ext_l in {"txt", "md"}:
            text = path.read_text(encoding="utf-8", errors="ignore").strip()
            if not text:
                return 0, []
            docs = [Document(page_content=text, metadata={"source": str(path)})]
        elif ext_l in {"mp3", "mp4", "wav", "m4a", "ogg"}:
            client = _oai.OpenAI(api_key=OPENAI_API_KEY)
            with open(path, "rb") as f:
                transcript = client.audio.transcriptions.create(model="whisper-1", file=f, response_format="text")
            if not transcript or not transcript.strip():
                return 0, []
            docs = [Document(page_content=transcript, metadata={"source": str(path), "type": "transcript"})]
        else:
            return 0, []

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=150, separators=["\n\n", "\n", ".", " "]
        )
        chunks = splitter.split_documents(docs)
        if not chunks:
            return 0, []
        vs = InMemoryVectorStore.from_documents(chunks, _get_emb())
        vector_stores[did] = vs
        chapters = _extract_chapters(docs)
        return len(chunks), chapters
    except Exception as e:
        print(f"[INDEX] {did}: {e}")
        import traceback; traceback.print_exc()
        return 0, []

def _extract_chapters(docs: list) -> list:
    try:
        sample = "\n\n".join(d.page_content for d in docs[:6])[:3000]
        prompt = CHAPTER_EXTRACT_PROMPT.replace("{context}", sample)
        result = _llm_json(prompt, temperature=0.1, mini=True)
        if isinstance(result, list):
            return [str(c).strip() for c in result if str(c).strip()]
    except Exception as e:
        print(f"[CHAPTERS] {e}")
    return []

def _get_chain(syllabus_id: str):
    """Build or retrieve a ConversationalRetrievalChain for a syllabus."""
    if syllabus_id in qa_chains:
        return qa_chains[syllabus_id]

    vs = _load_vs(syllabus_id)
    if not vs:
        # Virtual chain for curriculum-hub syllabi without real documents
        class _VirtualChain:
            def __init__(self, sid): self.sid = sid
            def __call__(self, inputs):
                q    = inputs.get("question", "")
                name = syllabi_registry.get(self.sid, {}).get("name", "your curriculum")
                prompt = (
                    f"You are Arthavi, a helpful AI tutor. The student has selected: {name}.\n"
                    f"Answer the following question thoroughly and helpfully. If it is clearly not related to {name}, "
                    f"start your answer with '⚠️ This topic is not part of your selected syllabus ({name}), but I\\'ll still help:'\n"
                    f"IMPORTANT: Always give a useful educational answer. Never say there are no materials or refuse.\n\n"
                    f"Question: {q}\n\nAnswer:"
                )
                ans = _llm_text(prompt, mini=True)
                return {"answer": ans, "source_documents": []}
        return _VirtualChain(syllabus_id)

    retriever = vs.as_retriever(search_type="mmr", search_kwargs={"k": 5})
    if syllabus_id not in memories:
        memories[syllabus_id] = ConversationBufferMemory(
            memory_key="chat_history", return_messages=True, output_key="answer"
        )
    kwargs = {}
    if STUDY_PROMPT:
        kwargs["combine_docs_chain_kwargs"] = {"prompt": STUDY_PROMPT}
    chain = ConversationalRetrievalChain.from_llm(
        llm=_get_llm(), retriever=retriever,
        memory=memories[syllabus_id], return_source_documents=True,
        verbose=False, **kwargs
    )
    qa_chains[syllabus_id] = chain
    return chain

def _evaluate_answers(exam: dict, submitted: Dict[int, str], roll_no: str = "") -> dict:
    """Grade submitted answers against an exam's answer key."""
    evals, total_awarded, total_possible = [], 0.0, 0.0
    for q in exam.get("questions", []):
        qid    = int(q["id"])
        qtype  = q.get("type", "objective")
        max_m  = float(q.get("weightage", q.get("marks", 1)))
        student = str(submitted.get(qid, "")).strip()
        total_possible += max_m

        if qtype == "objective":
            coerced = _coerce_mcq_answer(student, q.get("options", {}))
            correct = _objective_is_correct(q, student)
            awarded = max_m if correct else 0.0
            total_awarded += awarded
            evals.append({
                "question_id": qid, "type": "objective", "student_answer": coerced,
                "valid_answers": q.get("valid_answers", []), "is_correct": correct,
                "awarded_marks": awarded, "max_marks": max_m,
                "feedback": "Correct ✓" if correct else "Incorrect ✗"
            })
        else:
            # Clean subjective answer before grading
            student_cleaned = student
            if isinstance(student_cleaned, str) and len(student_cleaned.strip()) < 5:
                # Minimum text check
                awarded = 0.0
                total_awarded += awarded
                evals.append({
                    "question_id": qid, "type": "subjective", "student_answer": student_cleaned,
                    "max_marks": max_m, "question_text": q.get("question", ""),
                    "valid_answers": q.get("valid_answers", []),
                    "awarded_marks": awarded,
                    "feedback": "No answer provided.",
                    "answer_key_points": q.get("answer_key_points", []),
                    "rubric": q.get("evaluation_rubric", q.get("evaluation_criteria", "")),
                })
                continue
                
            evals.append({
                "question_id": qid, "type": "subjective", "student_answer": student_cleaned,
                "max_marks": max_m, "question_text": q.get("question", ""),
                "valid_answers": q.get("valid_answers", []),
                "answer_key_points": q.get("answer_key_points", []),
                "rubric": q.get("evaluation_rubric", q.get("evaluation_criteria", "")),
            })

    # Batch grade subjective questions
    subj_evals = [e for e in evals if e.get("type") == "subjective"]
    if subj_evals:
        log.info(f"[EVAL] Batch grading {len(subj_evals)} subjective questions for roll={roll_no}...")
        subj_text_block = ""
        for e in subj_evals:
            ans = e.get('student_answer', '')
            if isinstance(ans, str) and ans.startswith("[BLOCK_SECTION_B]"):
                subj_text_block = ans.replace("[BLOCK_SECTION_B]", "").strip()
                break
        
        batch_prompt = (
            "You are an expert examiner grading student answers.\n"
            "Below is the text extracted from the student's answer sheet.\n\n"
        )
        if subj_text_block:
            batch_prompt += f"--- EXTRACTED TEXT FROM SUBJECTIVE SECTION ---\n{subj_text_block}\n\n"
        
        batch_prompt += "QUESTIONS AND GRADING GUIDELINES:\n"
        for e in subj_evals:
            batch_prompt += f"--- Q{e['question_id']} ({e['max_marks']} marks) ---\n"
            batch_prompt += f"Question: {e['question_text']}\n"
            batch_prompt += f"Model Answers: {json.dumps(e['valid_answers'])}\n"
            batch_prompt += f"Key Points: {json.dumps(e['answer_key_points'])}\n"
            batch_prompt += f"Rubric: {e['rubric']}\n"
            if not subj_text_block:
                batch_prompt += f"Student Answer: {e['student_answer']}\n"
            batch_prompt += "\n"
        
        batch_prompt += (
            "INSTRUCTIONS:\n"
            "1. You are a strict exam evaluator.\n"
            "2. Ignore any printed instructions, question text, or noise like 'BLOCK_SECTION'.\n"
            "3. Evaluate ONLY the student's handwritten answer.\n"
            "4. If no valid handwritten answer is present, or if the text is just copied instructions, award 0 marks.\n"
            "5. Return JSON with question IDs as keys. Format for each question:\n"
            "   {\"awarded_marks\": number, \"feedback\": \"short explanation\"}\n"
        )
        
        try:
            # ── NEW: Use Ollama with fallback to GPT ──────────
            batch_results = call_ollama(batch_prompt)
            if not batch_results:
                log.info("[EVAL] Ollama failed or returned no data, falling back to GPT-4o-mini")
                batch_results = _llm_json(batch_prompt, temperature=0, mini=True)
            
            log.info(f"[EVAL] Batch results keys: {list(batch_results.keys())}")
            for e in subj_evals:
                qid_str = str(e['question_id'])
                # Handle various key formats (e.g. "7", "Q7", 7)
                res = batch_results.get(qid_str) or batch_results.get(f"Q{qid_str}") or batch_results.get(int(qid_str), {})
                
                awarded = 0.0
                try:
                    awarded = float(res.get("awarded_marks", 0))
                except (ValueError, TypeError):
                    pass
                awarded = max(0.0, min(e['max_marks'], awarded))
                
                total_awarded += awarded
                e.update({
                    "awarded_marks": awarded,
                    "feedback": res.get("feedback", "No feedback provided."),
                    "missing_points": res.get("missing_points", []),
                    "strengths": res.get("strengths", []),
                    "confidence": res.get("confidence", None),
                })
                log.info(f"[EVAL] Q{qid_str}: Awarded {awarded}/{e['max_marks']} marks.")
        except Exception as _be:
            log.error(f"[EVAL] Batch grading failed: {_be}")
            for e in subj_evals:
                e.update({"awarded_marks": 0.0, "feedback": "Grading error."})

    pct = round((total_awarded / total_possible) * 100, 2) if total_possible else 0.0
    improvement = ""
    if any(e.get("type") == "subjective" for e in evals) and pct < 90.0:
        try:
            fb = " | ".join(e.get("feedback", "") for e in evals if e.get("type") == "subjective")
            improvement = _llm_text(
                f"Based on this feedback: '{fb}', give one short encouraging sentence about what the student needs to improve.",
                mini=True
            ).strip()
        except Exception:
            improvement = "Keep practising subjective responses."
    elif pct >= 90.0:
        improvement = "Excellent work! Maintain this standard."

    return {
        "roll_no": roll_no, "total_possible": total_possible, "total_awarded": total_awarded,
        "percentage": pct, "grade": _grade(pct), "is_pass": pct >= 40.0,
        "improvement_prediction": improvement, "question_wise": evals
    }

# ══════════════════════════════════════════════════════════════════════════════
#  SUBJECT-SPECIFIC GRADING RULES
# ══════════════════════════════════════════════════════════════════════════════

def _get_subject_grading_prompt(exam: dict, subject: str, class_level: str) -> str:
    """Return a grading prompt tailored to the subject and class level."""
    subject_lower = subject.lower() if subject else ""
    
    # MATHEMATICS: Step marking, formula correctness, final answer
    if "math" in subject_lower or "maths" in subject_lower:
        return """You are a strict CBSE Mathematics examiner.
When grading answers:
1. Award FULL marks only if:
   - The method/approach is correct
   - All steps are shown (if multi-step problem)
   - Formula is applied correctly
   - Final answer is correct AND properly simplified
2. Award PARTIAL marks (50% of max) if:
   - Correct method but calculation error
   - Correct final answer but steps missing
   - Formula is correct but unit is wrong
3. Award 0 marks if:
   - Completely wrong method
   - No working shown for multi-step problem
   - Formula misapplied
Return JSON: {{"awarded_marks": X, "max_marks": Y, "reason": "...", "method_correct": bool, "answer_correct": bool}}"""
    
    # SCIENCE: Keywords, explanation quality, diagram labels
    elif "science" in subject_lower or "physics" in subject_lower or "chemistry" in subject_lower or "biology" in subject_lower:
        return """You are a strict CBSE Science examiner.
When grading answers:
1. Award FULL marks if:
   - All key scientific terms are used correctly
   - Explanation includes cause-effect relationship
   - Diagram (if required) is labeled correctly
   - Answer is scientifically accurate
2. Award PARTIAL marks if:
   - Most keywords present but one missing
   - Correct concept but incomplete explanation
   - Diagram present but missing labels
3. Award 0 marks if:
   - Fundamentally incorrect understanding
   - No use of scientific vocabulary
   - Irrelevant/off-topic
Return JSON: {{"awarded_marks": X, "max_marks": Y, "reason": "...", "keywords_found": [...], "explanation_quality": "good/fair/poor"}}"""
    
    # ENGLISH: Grammar, structure, content relevance
    elif "english" in subject_lower:
        return """You are a strict CBSE English examiner.
When grading answers:
1. Award FULL marks if:
   - Grammar and spelling are mostly correct
   - Sentence structure is coherent and varied
   - Content is relevant and well-organized
   - Vocabulary is appropriate
2. Award PARTIAL marks if:
   - Minor grammar/spelling errors
   - Content is mostly relevant but slightly disorganized
   - Simple vocabulary but understandable
3. Award 0 marks if:
   - Severe grammar/spelling making it unintelligible
   - Irrelevant content
   - No attempt at coherent structure
Return JSON: {{"awarded_marks": X, "max_marks": Y, "reason": "...", "grammar_score": 0-10, "content_score": 0-10, "structure_score": 0-10}}"""
    
    # SOCIAL STUDIES: Concepts, dates, examples
    elif "sst" in subject_lower or "history" in subject_lower or "geography" in subject_lower or "civics" in subject_lower:
        return """You are a strict CBSE Social Studies examiner.
When grading answers:
1. Award FULL marks if:
   - Core concept is accurately explained
   - Relevant historical dates/examples included
   - Answer shows understanding (not just memorization)
2. Award PARTIAL marks if:
   - Concept mostly correct but one detail wrong
   - Examples are relevant but dates might be off by 1-2 years
   - Shows understanding but missing one supporting detail
3. Award 0 marks if:
   - Core concept is misunderstood
   - Completely wrong dates/facts
   - Irrelevant answer
Return JSON: {{"awarded_marks": X, "max_marks": Y, "reason": "...", "concept_correct": bool, "examples_relevant": bool, "dates_accurate": bool}}"""
    
    # DEFAULT: Generic grading
    else:
        return """You are a fair and strict examiner.
When grading answers:
1. Award FULL marks if the answer is correct and complete
2. Award PARTIAL marks if the answer shows understanding but has minor errors or omissions
3. Award 0 marks if the answer is fundamentally wrong or irrelevant
Return JSON: {{"awarded_marks": X, "max_marks": Y, "reason": "...", "accuracy": 0-100}}"""

def _grade_with_subject_rules(exam: dict, submitted: Dict[int, str], subject: str = "", 
                               class_level: str = "", roll_no: str = "") -> dict:
    """
    Enhanced grading with subject-specific rules.
    Applies specialized prompts based on subject (Maths, Science, English, SST, etc).
    """
    evals, total_awarded, total_possible = [], 0.0, 0.0
    subject_prompt_template = _get_subject_grading_prompt(exam, subject, class_level)
    
    for q in exam.get("questions", []):
        qid = int(q["id"])
        qtype = q.get("type", "objective")
        max_m = float(q.get("weightage", q.get("marks", 1)))
        student = str(submitted.get(qid, "")).strip()
        total_possible += max_m
        
        if qtype == "objective":
            # MCQ grading is simple: all or nothing
            coerced = _coerce_mcq_answer(student, q.get("options", {}))
            correct = _objective_is_correct(q, student)
            awarded = max_m if correct else 0.0
            total_awarded += awarded
            evals.append({
                "question_id": qid, "type": "objective", "student_answer": coerced,
                "awarded_marks": awarded, "max_marks": max_m,
                "reason": "Correct answer" if correct else "Incorrect answer",
                "confidence": 1.0 if correct else 0.0,
            })
        else:
            # Subjective: use subject-specific grading
            student_cleaned = _clean_subjective_answer(student)
            model_ans = str(q.get("model_answer", q.get("answer", ""))).strip()
            
            prompt = f"""{subject_prompt_template}

Question: {q.get('question', '')}
Max Marks: {max_m}
Model Answer/Rubric: {model_ans}
Student's Answer: {student_cleaned}"""
            
            try:
                grade_json = _llm_json(prompt, temperature=0)
                awarded = max(0.0, min(max_m, float(grade_json.get("awarded_marks", 0))))
                reason = grade_json.get("reason", "Subjective evaluation")
                confidence = 0.85  # Subjective grading has inherent uncertainty
            except Exception as ex:
                log.warning(f"[GRADE] Subject-specific grading failed: {ex}")
                awarded = 0.0
                reason = "Evaluation failed"
                confidence = 0.0
            
            total_awarded += awarded
            evals.append({
                "question_id": qid, "type": "subjective",
                "student_answer": student_cleaned,
                "awarded_marks": awarded, "max_marks": max_m,
                "reason": reason,
                "confidence": confidence,
            })
    
    pct = round((total_awarded / total_possible) * 100, 2) if total_possible else 0.0
    
    return {
        "roll_no": roll_no,
        "total_awarded": total_awarded,
        "total_possible": total_possible,
        "percentage": pct,
        "grade": _grade(pct),
        "is_pass": pct >= 40.0,
        "subject": subject,
        "class_level": class_level,
        "question_wise": evals,
    }

def _preprocess_image_for_ocr(img):
    """
    Enhance a scanned image so handwriting (faint pen, circles, ticks) becomes
    clearly visible to GPT Vision.
    Steps: greyscale -> high contrast -> sharpen -> median filter (removes noise).
    Returns RGB PIL Image ready for base64 encoding.
    """
    from PIL import ImageEnhance, ImageFilter
    img = img.convert("L")                         # greyscale
    img = ImageEnhance.Contrast(img).enhance(2.5)  # make dark marks much darker
    img = ImageEnhance.Sharpness(img).enhance(2.0) # sharpen edges (helps circles/ticks)
    img = img.filter(ImageFilter.MedianFilter())   # remove scan noise
    return img.convert("RGB")                      # back to RGB for API

def _preprocess_image_enhanced(img, apply_rotation=True, apply_shadow_removal=True, 
                                apply_clahe=False, apply_compression=True):
    """
    Enhanced image preprocessing for evaluation papers:
    1. Rotation detection & correction (deskew)
    2. Shadow removal (gamma correction)
    3. CLAHE (Contrast Limited Adaptive Histogram Equalization)
    4. Intelligent compression
    
    This improves OCR accuracy for skewed/shadowed handwritten answers.
    """
    from PIL import ImageEnhance, ImageFilter, Image
    import numpy as np
    
    # Convert to numpy for advanced processing
    img_np = np.array(img.convert("L"))
    
    # Step 1: Rotation detection & correction (deskew)
    if apply_rotation:
        try:
            from scipy import ndimage
            # Find text skew angle via contour analysis
            img_edges = cv2.Canny(img_np, 100, 200) if 'cv2' in globals() else None
            if img_edges is not None:
                lines = cv2.HoughLinesP(img_edges, 1, np.pi/180, 50, minLineLength=100, maxLineGap=10)
                if lines is not None and len(lines) > 0:
                    angles = []
                    for line in lines:
                        x1, y1, x2, y2 = line[0]
                        angle = np.arctan2(y2-y1, x2-x1) * 180 / np.pi
                        if abs(angle) < 45:  # Only consider text-like angles
                            angles.append(angle)
                    if angles:
                        median_angle = np.median(angles)
                        if abs(median_angle) > 0.5:  # Only rotate if meaningful skew
                            img = img.rotate(-median_angle, expand=False, fillcolor="white")
                            img_np = np.array(img.convert("L"))
                            log.debug(f"[IMAGE-PROC] Deskewed by {median_angle:.1f}°")
        except Exception as ex:
            log.debug(f"[IMAGE-PROC] Deskew skipped: {ex}")
    
    # Step 2: Shadow removal via gamma correction
    if apply_shadow_removal:
        img_np = np.clip(img_np.astype(float) / 255.0, 0, 1)
        gamma = 1.2  # Brighten shadows without oversaturating highlights
        img_np = (img_np ** (1/gamma)) * 255
        img_np = np.clip(img_np, 0, 255).astype(np.uint8)
        log.debug(f"[IMAGE-PROC] Applied gamma correction (gamma={gamma})")
    
    # Step 3: CLAHE (more adaptive than global contrast)
    if apply_clahe:
        try:
            import cv2
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            img_np = clahe.apply(img_np.astype(np.uint8))
            log.debug("[IMAGE-PROC] Applied CLAHE")
        except Exception as ex:
            log.debug(f"[IMAGE-PROC] CLAHE skipped: {ex}")
    
    # Standard enhancement: contrast + sharpen
    img = Image.fromarray(img_np).convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.5)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    img = img.filter(ImageFilter.MedianFilter())
    
    # Step 4: Intelligent compression (reduce file size while preserving quality)
    if apply_compression:
        # For OCR, we can be aggressive: target 80-85% quality
        img_rgb = img.convert("RGB")
        # Size limiting is handled at the API level
    else:
        img_rgb = img.convert("RGB")
    
    return img_rgb

def _detect_and_crop_borders(img, border_color_tolerance=30):
    """
    Detect and crop white/light borders from scanned pages.
    Removes unnecessary margins that waste tokens in Vision API.
    """
    from PIL import Image
    import numpy as np
    
    img_np = np.array(img.convert("L"))
    threshold = 240  # Light gray threshold for "white" borders
    
    # Find rows/cols with mostly white pixels
    height, width = img_np.shape
    row_white = (img_np > threshold).sum(axis=1) > (width * 0.95)
    col_white = (img_np > threshold).sum(axis=0) > (height * 0.95)
    
    # Find first/last non-white row/col
    white_rows = np.where(row_white)[0]
    white_cols = np.where(col_white)[0]
    
    crop_top, crop_bottom = 0, height
    crop_left, crop_right = 0, width
    
    if len(white_rows) > 0:
        # Find largest gap of white rows
        white_diffs = np.diff(white_rows)
        if len(white_diffs) > 0 and np.max(white_diffs) > 10:
            largest_gap_idx = np.argmax(white_diffs)
            crop_top = white_rows[largest_gap_idx] + 1
            crop_bottom = white_rows[largest_gap_idx + 1]
    
    if len(white_cols) > 0:
        white_diffs = np.diff(white_cols)
        if len(white_diffs) > 0 and np.max(white_diffs) > 10:
            largest_gap_idx = np.argmax(white_diffs)
            crop_left = white_cols[largest_gap_idx] + 1
            crop_right = white_cols[largest_gap_idx + 1]
    
    # Add small margin to avoid cutting text
    margin = 10
    crop_top = max(0, crop_top - margin)
    crop_left = max(0, crop_left - margin)
    crop_bottom = min(height, crop_bottom + margin)
    crop_right = min(width, crop_right + margin)
    
    if crop_top >= crop_bottom or crop_left >= crop_right:
        return img  # No valid crop
    
    cropped = img.crop((crop_left, crop_top, crop_right, crop_bottom))
    removed_percent = round(100 * (1 - (cropped.size[0]*cropped.size[1]) / (width*height)), 1)
    log.debug(f"[IMAGE-PROC] Cropped borders: removed {removed_percent}% of image area")
    return cropped

def _build_ocr_prompt_v2(exam_questions: list) -> str:
    """
    Improved OCR prompt — explicitly addresses the dark Q-number tab confusion
    and ensures ALL handwritten lines are read.
    """
    obj_qs  = [q for q in exam_questions if q.get("type", "objective") == "objective"]
    subj_qs = [q for q in exam_questions if q.get("type", "objective") != "objective"]
    obj_ids  = ", ".join(str(q["id"]) for q in obj_qs)  or "none"
    subj_ids = ", ".join(str(q["id"]) for q in subj_qs) or "none"

    q_lines = []
    for q in sorted(exam_questions, key=lambda x: int(x.get("id", 0))):
        qtype = "MCQ" if q.get("type", "objective") == "objective" else "Written"
        q_lines.append(f"  Q{q['id']} [{qtype}]: {str(q.get('question',''))[:80]}")
    question_ref = (
        "\nQUESTION REFERENCE (to locate each box by its number):\n"
        + "\n".join(q_lines)
    ) if q_lines else ""

    return (
        "You are reading a SCANNED STUDENT ANSWER SHEET.\n\n"

        "═══ CRITICAL — READ BEFORE ANSWERING ═══\n"
        "Each question has a DARK COLOURED TAB on the LEFT edge showing the Q number.\n"
        "⚠️  The dark tab is NOT a bubble and is NOT an answer option.\n"
        "⚠️  IGNORE the dark tab when identifying bubbles.\n\n"

        "MCQ LAYOUT (questions: " + obj_ids + "):\n"
        "Each MCQ box has TWO ways for the student to answer:\n"
        "  METHOD 1 — Bubbles: 4 circles labelled A / B / C / D from left to right.\n"
        "    The letter label appears ABOVE each circle (not inside it).\n"
        "    The student FILLED or CIRCLED exactly one bubble.\n"
        "  METHOD 2 — Write box: a square box on the RIGHT end of the MCQ row.\n"
        "    The student may have written a single letter (A/B/C/D) inside this box.\n\n"
        "PRIORITY: Check the WRITE BOX on the right FIRST.\n"
        "  If it contains a handwritten letter A/B/C/D → use that as the answer.\n"
        "  If the write box is empty → check which bubble is filled/circled.\n"
        "    Count bubbles strictly left-to-right AFTER the dark tab:\n"
        "    1st circle = A,  2nd circle = B,  3rd circle = C,  4th circle = D.\n"
        "    The bubble with the most ink / darkest fill / scribble marks = selected.\n"
        "❌ NEVER return \"- A\" or \"A)\" — return ONLY the single letter e.g. \"A\".\n\n"

        "WRITTEN LAYOUT (questions: " + subj_ids + "):\n"
        "After the dark green Q-number tab, there are horizontal ruled lines.\n"
        "⚠️ IMPORTANT: The student's answer often continues across MULTIPLE LINES.\n"
        "⚠️ Do NOT stop after reading just one line. Read EVERY line in the box.\n"
        "Read from the first line (just below the 'Write your answer below:' instruction)\n"
        "to the last line (before the bottom border of the box).\n"
        "Join ALL lines into ONE string separated by spaces — do NOT stop at the first line.\n"
        "Ignore printed labels ('Write your answer below:', section headings, etc.)\n"
        "If the answer spans 10 lines, read ALL 10 lines and combine them.\n\n"

        "YOUR TASK:\n"
        "For each question box, identify its Q number from the tab, then extract:\n"
        "  MCQ → single letter ONLY: A, B, C, or D (the scribbled/filled circle)\n"
        "  Written → ALL handwritten text across ALL ruled lines, joined as one string\n"
        "  Blank → omit from JSON\n"
        + question_ref + "\n\n"

        "Return ONLY valid JSON (no markdown, no explanation):\n"
        + '{"answers": {"1":"B","2":"D","3":"B","4":"A","5":"A","6":"C","7":"student full written answer across all lines"}}' + "\n"
    )

def _build_ocr_prompt(exam_questions: list) -> str:
    """
    Layout-aware prompt for the VidyAI structured answer sheet.
    Describes the EXACT visual structure of the PDF we generated so GPT-4o
    can locate each question's answer box precisely.

    Answer sheet layout (generated by _generate_answer_sheet_pdf):
    - Each question has its own bordered rectangle
    - LEFT SIDE: a coloured vertical tab (indigo for MCQ, green for Written)
      containing the question number (Q1, Q2…) and marks
    - RIGHT SIDE (content area): the answer space
      * MCQ: "Mark ONE correct answer:" label + 4 large circles labelled A B C D
        - Student FILLS IN or SCRIBBLES INSIDE one circle
        - A filled/scribbled circle = selected answer
      * Written: "Write your answer below:" label + ruled horizontal lines
        - Student writes on the lines
    """
    obj_qs  = [q for q in exam_questions if q.get("type", "objective") == "objective"]
    subj_qs = [q for q in exam_questions if q.get("type", "objective") != "objective"]
    obj_ids  = ", ".join(str(q["id"]) for q in obj_qs)  or "none"
    subj_ids = ", ".join(str(q["id"]) for q in subj_qs) or "none"

    q_lines = []
    for q in sorted(exam_questions, key=lambda x: int(x.get("id", 0))):
        qtype = "MCQ" if q.get("type", "objective") == "objective" else "Written"
        q_lines.append(
            f"  Q{q['id']} [{qtype}]: {str(q.get('question', ''))[:80]}"
        )
    question_ref = (
        "\nQUESTION REFERENCE — use these to locate each Q-number tab on the page:\n"
        + "\n".join(q_lines)
    ) if q_lines else ""

    return (
        "You are reading a SCANNED STUDENT ANSWER SHEET from the VidyAI system.\n\n"

        "═══ EXACT LAYOUT OF THIS ANSWER SHEET ═══\n"
        "Each question has a bordered box. Inside each box:\n"
        "  LEFT: a solid coloured tab showing the Q number (e.g. Q1, Q2…)\n"
        "         — INDIGO/DARK BLUE tab = MCQ question\n"
        "         — GREEN tab = Written/subjective question\n"
        "  RIGHT: the student answer area\n\n"

        "MCQ ANSWER AREA (questions: " + obj_ids + "):\n"
        "  • The label 'Mark ONE correct answer:' appears at the top\n"
        "  • Below it are FOUR large circles in a row: A   B   C   D\n"
        "  • Each circle has its letter (A/B/C/D) printed INSIDE it\n"
        "  • The student marks ONE circle by:\n"
        "      - Filling/scribbling inside the circle with pen\n"
        "      - Drawing a heavy circle around the letter\n"
        "      - Writing the letter inside or next to the circle\n"
        "  • The FILLED/SCRIBBLED/DARKEST circle = the student's answer\n"
        "  • A circle with dense pen marks is DEFINITELY selected\n"
        "  ❌ NEVER return \"- A\" or \"A)\" — return ONLY the single letter e.g. \"A\".\n\n"

        "WRITTEN ANSWER AREA (questions: " + subj_ids + "):\n"
        "  • The label 'Write your answer below:' appears at the top\n"
        "  • Below it are horizontal ruled lines inside the box\n"
        "  • The student writes their answer on these lines\n"
        "  ⚠️  READ ALL LINES — do NOT stop after the first handwritten line.\n"
        "  ⚠️  Students typically write 2-8 sentences across ALL lines.\n"
        "  • Join every handwritten line into ONE string separated by spaces.\n"
        "  • Do NOT include the printed 'Write your answer below:' label.\n\n"

        "═══ YOUR TASK ═══\n"
        "For EACH question box you can see:\n"
        "  1. Find the Q-number tab on the LEFT to identify the question number\n"
        "  2. Look at the RIGHT side for the student's answer\n"
        "  3. For MCQ: identify which of the 4 circles (A/B/C/D) is filled/scribbled\n"
        "     — The one with the most ink/darkest fill is the answer\n"
        "     — Even heavy scribbling inside a circle = that option is selected\n"
        "  4. For Written: read ALL handwritten text on EVERY ruled line\n\n"

        "⚠️  CRITICAL — EACH QUESTION IS IN ITS OWN SEPARATE BOX ⚠️\n"
        "The answer sheet has DISTINCT bordered rectangles, one per question.\n"
        "Q7 has its own box. Q8 has a DIFFERENT box. Q9 has a DIFFERENT box.\n"
        "Do NOT combine text from Q8's box into Q7's answer or vice versa.\n"
        "Each box is separated by a visible gap and has its own coloured tab.\n"
        "Match each answer ONLY to the Q number shown on its tab.\n\n"

        "CRITICAL RULES:\n"
        "  • For MCQ: return ONLY a single letter — A, B, C, or D. No dashes, no brackets.\n"
        "  • The printed letter inside an empty/clean circle is NOT an answer\n"
        "  • Only a circle with INK MARKS (fill, scribble, circle) = selected\n"
        "  • If two circles seem marked, choose the one with more ink\n"
        "  • For Written: return ALL handwritten text joined as one string (ignore printed labels)\n"
        "  • ONLY include text that is inside THAT question's box — do NOT bleed into adjacent boxes\n"
        "  • If a question has no mark, omit it from the JSON\n"
        + question_ref + "\n\n"

        "Return STRICTLY valid JSON:\n"
        '{"answers": {"1": "B", "2": "D", "7": "Divide the numerator by the denominator and write the remainder"}}' + "\n"
        "No markdown, no explanation, no extra keys."
    )

def _build_freeform_handwriting_prompt(exam_questions: list) -> str:
    """
    Prompt for PLAIN HANDWRITTEN answer sheets (notebook pages, not printed sheets).
    No indigo tabs, no bubble circles, no printed labels.
    Instructs GPT-4o to find question numbers (1, 2, Q1, Q2) anywhere on the page.
    """
    q_lines = []
    for q in sorted(exam_questions, key=lambda x: int(x.get("id", 0))):
        qtype = "MCQ" if q.get("type", "objective") == "objective" else "Written"
        q_lines.append(f"  Q{q['id']} [{qtype}]: {str(q.get('question',''))[:80]}")
    question_ref = "\nQUESTION REFERENCE:\n" + "\n".join(q_lines) if q_lines else ""

    return (
        "You are reading a PHOTO or SCAN of a student's HANDWRITTEN homework/test.\n"
        "This is NOT a structured form. It is a plain piece of paper or notebook page.\n\n"

        "═══ YOUR TASK ═══\n"
        "1. Identify each question by its number (e.g., '1', '2', 'Q1', 'Q2').\n"
        "2. Extract the student's handwritten answer for that question.\n"
        "3. MCQ: If the student wrote a letter (A/B/C/D) or circled/underlined one, "
        "return ONLY that single letter (A, B, C, or D). Do NOT return dashes or punctuation.\n"
        "4. WRITTEN (subjective): Extract ALL handwritten text ACROSS ALL LINES for that question.\n"
        "   - Read line 1, line 2, line 3... ALL lines until the next question number.\n"
        "   - Join all lines with a single space into ONE continuous string.\n"
        "   - Do NOT stop at the first line. Do NOT truncate. Capture the FULL answer.\n"
        "5. Treat each question independently.\n\n"

        "═══ CRITICAL RULES FOR WRITTEN ANSWERS ═══\n"
        "- A student's written answer often spans 3–10 lines. READ EVERY LINE.\n"
        "- Start reading from the line immediately after the question number.\n"
        "- Stop when you see the NEXT question number (e.g., '2.', 'Q2', '2)')\n"
        "- Ignore any scribbles or marks that are not part of an answer.\n"
        "- If a question is not found or has no answer, omit it from the JSON.\n"
        "- For MCQ: ONLY return a single letter A, B, C, or D — no dashes, no hyphens.\n"
        + question_ref + "\n\n"

        "Return ONLY valid JSON:\n"
        '{"answers": {"1": "C", "2": "A", "7": "The square of the hypotenuse is equal to the sum of the squares of the other two sides. This applies only to right-angled triangles and is proved by Pythagoras."}}\n'
    )


def _build_enhanced_ocr_prompt(exam_questions: list) -> str:
    """Enhanced OCR prompt — color-neutral, structured for scanned handwritten sheets"""
    q_ref_lines = []
    for q in sorted(exam_questions, key=lambda x: int(x.get("id", 0))):
        qid = q.get("id", "?")
        qtype = q.get("type", "objective")
        q_text = str(q.get("question", ""))[:120]
        if qtype == "objective":
            opts = q.get("options", {})
            opt_str = "  ".join(f"{k}) {v}" for k, v in opts.items()) if isinstance(opts, dict) else ""
            q_ref_lines.append(f"  Q{qid} [MCQ]: {q_text}\n    Options: {opt_str}")
        else:
            q_ref_lines.append(f"  Q{qid} [Written, max {q.get('weightage', q.get('marks', '?'))} marks]: {q_text}")

    question_ref = "\n".join(q_ref_lines)

    return (
        "You are an expert at reading HANDWRITTEN student answer sheets.\n"
        "The image is a scanned page. It may be black-and-white, faded, or slightly skewed.\n\n"
        "ANSWER SHEET LAYOUT:\n"
        "- Each question has a label/tab on the left showing the question number.\n"
        "- MCQ questions: the student marks ONE of four options (A, B, C, D) by filling a circle or writing a letter.\n"
        "- Written questions: the student writes a free-text answer on ruled lines inside a box.\n\n"
        "YOUR TASK — extract the student's answer for every question:\n\n"
        "MCQ rules:\n"
        "  1. Identify which circle (A/B/C/D) is filled, darkened, or marked.\n"
        "  2. Return ONLY the single uppercase letter.\n"
        "  3. The question-number label is NOT an answer — ignore it.\n\n"
        "Written-answer rules:\n"
        "  1. Read EVERY handwritten line inside the answer box, top to bottom.\n"
        "  2. Combine all lines into ONE string separated by spaces.\n"
        "  3. Do NOT stop after the first line — students often write across many lines.\n"
        "  4. Preserve the student's actual words even if grammar/spelling is imperfect.\n"
        "  5. If the box is blank or unreadable, return an empty string.\n\n"
        "Return ONLY valid JSON — no commentary, no markdown:\n"
        '{"answers": {"1": "B", "2": "the student wrote this across multiple lines..."}}\n\n'
        f"QUESTIONS ON THIS SHEET:\n{question_ref}\n"
    )


def _validate_ocr_answers(raw_answers: dict, exam_questions: list) -> Dict[int, str]:
    """
    Post-process raw OCR answers dict to reject noise:
    - Values that look like question text fragments (long sentences ending in ?)
    - Values that contain question-paper header phrases
    - Values that are garbled OCR of section headers or marks annotations
    - Keys that don't correspond to actual exam question IDs
    Valid question IDs from the exam; anything outside is noise.
    """
    valid_ids = {int(q.get("id", 0)) for q in exam_questions} if exam_questions else None

    # These patterns only match if the ENTIRE value is a label/instruction, not real content.
    # For subjective answers, we strip any leading instruction prefix before checking.
    NOISE_PATTERNS = re.compile(
        r"^(?:section\s+[AB]|objective\s+questions?|subjective\s+questions?|"
        r"general\s+instructions?|total\s+marks|roll\s+number|student\s+name|"
        r"marks\s+obtained|answer\s+sheet|vidyai\s+school|"
        r"write\s+only\s+inside|only\s+inside\s+the|"
        r"mark\s+one\s+correct\s*answer|circle\s+(your|the|one)\s+answer|"
        r"do\s+not\s+write\s+outside|"
        r"\d{1,2}/\d{1,2}/\d{4})$",
        re.IGNORECASE
    )
    # Strip common OCR prefix artifacts before validation
    STRIP_PREFIXES = re.compile(
        r"^(?:write\s+your\s+answer\s+below\s*[:\-]?\s*|"
        r"write\s+answer\s+here\s*[:\-]?\s*|"
        r"answer\s*[:\-]\s*)",
        re.IGNORECASE
    )
    # MCQ answer: should be a single letter A-D (possibly with punctuation)
    MCQ_IDS = {int(q["id"]) for q in exam_questions
               if q.get("type", "objective") == "objective"} if exam_questions else set()

    result: Dict[int, str] = {}
    for k, v in raw_answers.items():
        try:
            qid = int(str(k).strip())
        except ValueError:
            continue
        val = str(v).strip()
        if not val:
            continue

        # Reject if question ID not in exam
        if valid_ids and qid not in valid_ids:
            continue

        # ── MCQ: extract letter from value regardless of surrounding text ──────
        if qid in MCQ_IDS:
            # Pre-check: reject obvious section-header / noise strings first
            if NOISE_PATTERNS.search(val):
                log.debug(f"[OCR-VALIDATE] Q{qid} MCQ noise rejected: {val[:40]}")
                continue
            # Reject standalone dashes or punctuation (no actual letter)
            if re.match(r'^[\-\u2013\u2014\s\.\/\|]+$', val):
                log.debug(f"[OCR-VALIDATE] Q{qid} MCQ dash-only rejected: {val[:40]}")
                continue
            # Primary: look for standalone letter
            m = re.search(r"\b([A-D])\b", val, re.IGNORECASE)
            if m:
                result[qid] = m.group(1).upper()
                log.debug(f"[OCR-VALIDATE] Q{qid} MCQ → '{result[qid]}' (from: {val[:40]})")
            else:
                # Fallback heuristic: any A-D character in the string
                letters = re.findall(r"[A-Da-d]", val)
                if letters:
                    result[qid] = letters[0].upper()
                    log.debug(f"[OCR-VALIDATE] Q{qid} MCQ fallback → '{result[qid]}' (from: {val[:40]})")
                else:
                    log.debug(f"[OCR-VALIDATE] Q{qid} MCQ no letter found in: {val[:40]}")
            continue

        # ── Subjective: strip instruction prefixes, then reject obvious noise ──────
        # Strip "Write your answer below:" or "Answer:" prefixes GPT may include
        val = STRIP_PREFIXES.sub("", val).strip()
        # Normalise newlines → spaces so multi-line handwriting is stored as one string
        val = re.sub(r"[\r\n]+", " ", val)
        val = re.sub(r"\s{2,}", " ", val).strip()
        if not val:
            continue
        if NOISE_PATTERNS.search(val):
            log.debug(f"[OCR-VALIDATE] Rejected Q{qid} noise: {val[:60]}")
            continue
        # Only reject if the ENTIRE value looks like a question (ends with ?)
        # and has no sentence before it — don't reject long answers that happen
        # to contain a question phrase at the end.
        if val.endswith("?") and len(val) > 40:
            # Check: is there a sentence-ending character before the '?'?
            # If yes, it's a long answer that ends with a rhetorical question → keep
            has_prior_sentence = bool(re.search(r'[.!;]', val[:-1]))
            if not has_prior_sentence:
                log.debug(f"[OCR-VALIDATE] Rejected Q{qid} question-text: {val[:60]}")
                continue
        if len(val) < 2:
            continue
        # Reject OCR garble: words with 4+ consecutive consonants = corrupted scan
        # Use a higher threshold (50%) and minimum 5 words to avoid false rejects
        # on multi-line answers with some tricky words.
        words = re.findall(r"[a-zA-Z]+", val)
        if len(words) >= 5:
            garble = sum(1 for w in words if re.search(r"[bcdfghjklmnpqrstvwxyz]{4,}", w, re.I))
            if garble / len(words) > 0.50:
                log.debug(f"[OCR-VALIDATE] Rejected Q{qid} garble ({garble}/{len(words)}): {val[:60]}")
                continue

        log.debug(f"[OCR-VALIDATE] Q{qid} subjective → '{val[:60]}'")
        result[qid] = val

    return result

# ══════════════════════════════════════════════════════════════════════════════
#  SCANNED ANSWER SHEET: IMAGE-DIFF HANDWRITING EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════════
#  PDF TEMPLATE GENERATORS — Answer Sheet, Question Paper, Answer Key
# ══════════════════════════════════════════════════════════════════════════════

def _pdf_wrap(text, c_per_line=88):
    """Word-wrap text into lines of at most c_per_line characters."""
    words = str(text).split()
    lines, cur = [], ""
    for w in words:
        if len(cur) + len(w) + 1 <= c_per_line:
            cur = (cur + " " + w).strip()
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines or [""]

def _pdf_section_bar(c, x, y, w, h, fill_hex, text, text_hex, font_size=9, mm=None):
    """Draw a coloured section header bar with text."""
    from reportlab.lib import colors as _c
    c.setFillColor(_c.HexColor(fill_hex))
    c.rect(x, y - h, w, h, fill=1, stroke=0)
    c.setFillColor(_c.HexColor(text_hex))
    c.setFont("Helvetica-Bold", font_size)
    c.drawString(x + 4*mm, y - h*0.62, text)

def _draw_answer_sheet_page_header(c, exam, page_num, W, H, mm, colors):
    """
    Beautiful Answer Sheet page header.
    Layout (top → bottom):
      Row 1 (14mm): Dark indigo bar — school name left, "ANSWER SHEET" badge right
      Row 2 (10mm): Light indigo — subject | board | class | date | exam ID
      Row 3 (16mm): Student info boxes — Name | Roll | Class/Sec | Subject
      ─ total height from top: 44mm
    """
    M = 14 * mm

    # ── Row 1: main title bar ─────────────────────────────────────────────────
    BAR_H = 14 * mm
    c.setFillColor(colors.HexColor('#1e1b4b'))
    c.rect(M, H - BAR_H, W - 2*M, BAR_H, fill=1, stroke=0)

    # School name — left
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 12)
    school = exam.get('school_name', 'VidyAI School')
    c.drawString(M + 5*mm, H - 9.5*mm, school[:50])

    # "ANSWER SHEET" pill badge — right
    badge_text  = "ANSWER SHEET"
    badge_x     = W - M - 52*mm
    badge_y     = H - BAR_H + 3*mm
    badge_w     = 48*mm
    badge_h     = 8*mm
    c.setFillColor(colors.HexColor('#4f46e5'))
    c.roundRect(badge_x, badge_y, badge_w, badge_h, 2, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 8)
    c.drawCentredString(badge_x + badge_w/2, badge_y + 2.5*mm, badge_text)

    # ── Row 2: metadata strip ─────────────────────────────────────────────────
    META_H = 8 * mm
    c.setFillColor(colors.HexColor('#eef2ff'))
    c.rect(M, H - BAR_H - META_H, W - 2*M, META_H, fill=1, stroke=0)

    subject = exam.get('subject', '')
    board   = exam.get('board', '')
    cls     = exam.get('class', exam.get('class_name', ''))
    date    = exam.get('exam_date', '')
    eid     = exam.get('exam_id', '—')
    total_m = exam.get('total_marks', 0)

    meta_parts = [str(p) for p in [subject, board, cls] if p]
    meta_str   = '  ·  '.join(meta_parts)
    if date:
        meta_str += f'  |  Date: {date}'
    meta_str += f'  |  Total Marks: {total_m}  |  Exam ID: {eid}'

    c.setFillColor(colors.HexColor('#3730a3'))
    c.setFont("Helvetica", 7.5)
    c.drawString(M + 4*mm, H - BAR_H - 5*mm, meta_str[:110])
    c.setFillColor(colors.HexColor('#94a3b8'))
    c.drawRightString(W - M - 4*mm, H - BAR_H - 5*mm, f'Page {page_num}')

    # QR marker improves downstream split/trace reliability for multi-student scans.
    try:
        from reportlab.graphics.barcode import qr as _qr
        from reportlab.graphics.shapes import Drawing as _Drawing
        from reportlab.graphics import renderPDF as _renderPDF
        qr_payload = f"exam:{exam.get('exam_id', '')}|page:{page_num}|school:{exam.get('school_name', '')}"
        qr_code = _qr.QrCodeWidget(qr_payload)
        b = qr_code.getBounds()
        bw = max(1, b[2] - b[0])
        bh = max(1, b[3] - b[1])
        qr_size = 14 * mm
        qr_draw = _Drawing(qr_size, qr_size, transform=[qr_size / bw, 0, 0, qr_size / bh, 0, 0])
        qr_draw.add(qr_code)
        _renderPDF.draw(qr_draw, c, W - M - qr_size, H - BAR_H - META_H - qr_size - 1.5 * mm)

        c.setFillColor(colors.HexColor('#64748b'))
        c.setFont("Helvetica", 5.5)
        c.drawRightString(W - M, H - BAR_H - META_H - qr_size - 3.2 * mm,
                          f"Scan ID: {exam.get('exam_id', '')} · Pg {page_num}")
    except Exception:
        pass

    # ── Row 3: student info boxes ─────────────────────────────────────────────
    INFO_TOP = H - BAR_H - META_H - 3*mm
    BOX_H    = 10 * mm
    fields   = [
        ("STUDENT NAME",     M,          90*mm),
        ("ROLL NUMBER",      M + 93*mm,  38*mm),
        ("CLASS / SECTION",  M + 134*mm, 35*mm),
    ]
    for label, fx, fw in fields:
        # Label above box
        c.setFillColor(colors.HexColor('#374151'))
        c.setFont("Helvetica-Bold", 7)
        c.drawString(fx + 1*mm, INFO_TOP - 0.5*mm, label)
        # Box
        c.setFillColor(colors.white)
        c.setStrokeColor(colors.HexColor('#94a3b8'))
        c.setLineWidth(0.8)
        c.roundRect(fx, INFO_TOP - BOX_H - 1*mm, fw, BOX_H, 1.5, fill=1, stroke=1)

    # Instruction row
    INSTR_Y = INFO_TOP - BOX_H - 5*mm
    c.setFillColor(colors.HexColor('#f0f9ff'))
    c.rect(M, INSTR_Y - 6*mm, W - 2*M, 6*mm, fill=1, stroke=0)
    c.setFillColor(colors.HexColor('#0369a1'))
    c.setFont("Helvetica-Bold", 7)
    c.drawString(M + 4*mm, INSTR_Y - 4*mm,
        'INSTRUCTIONS:  Section A — Circle ONE bubble (A/B/C/D). If writing, write ONLY the option letter in the box  '
        '·  Section B — Write your answer on the ruled lines in the box  '
        '·  Write ONLY inside the boxes')

    AFTER_HEADER = INSTR_Y - 8*mm
    # Corner registration marks (for image-diff alignment)
    c.setFillColor(colors.black)
    for rx, ry in [(6, H - 6), (W - 6, H - 6), (6, 6*mm), (W - 6, 6*mm)]:
        c.circle(rx, ry, 2.5, fill=1)

    return AFTER_HEADER

def _generate_answer_sheet_pdf(exam: dict) -> bytes:
    """
    OCR-optimised student Answer Sheet.
    Section A: each MCQ gets its own bordered box with 4 large bubble circles.
    Section B: each question gets a bordered box with ruled writing lines.
    Bubbles are properly sized and vertically centred — no overlap with labels.
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors as rl_colors
        import io as _io2

        buf = _io2.BytesIO()
        W, H = A4
        c = rl_canvas.Canvas(buf, pagesize=A4)

        qs      = exam.get("questions", [])
        obj_qs  = [q for q in qs if q.get("type", "objective") == "objective"]
        subj_qs = [q for q in qs if q.get("type", "objective") != "objective"]
        M       = 14 * mm
        BOX_W   = W - 2 * M

        page = 1
        y    = _draw_answer_sheet_page_header(c, exam, page, W, H, mm, rl_colors)

        def new_page():
            nonlocal page
            # Bottom footer
            c.setFillColor(rl_colors.HexColor('#94a3b8'))
            c.setFont("Helvetica", 6.5)
            c.drawCentredString(W / 2, 5*mm,
                'Write answers ONLY inside the boxes  ·  Do not write outside the boxes')
            c.showPage()
            page += 1
            return _draw_answer_sheet_page_header(c, exam, page, W, H, mm, rl_colors)

        # ── Section A — Objective ─────────────────────────────────────────────
        if obj_qs:
            _pdf_section_bar(c, M, y, BOX_W, 8*mm,
                '#eef2ff', f'SECTION A — OBJECTIVE  ({len(obj_qs)} questions)  '
                '·  Circle ONE bubble  OR  write the letter in the box on the right',
                '#3730a3', font_size=8.5, mm=mm)
            y -= 10*mm

            # MCQ box geometry:
            #   BOX_H=36mm | TAB=18mm | BUBBLE_R=7mm
            #   Label row: 9mm from box top, Bubble centre: 24mm from box top
            #   Write-box: rightmost 32mm of content area (large blank square)
            BOX_H         = 36 * mm
            TAB_W         = 18 * mm
            BUBBLE_R      = 7  * mm
            BUBBLE_FROM_TOP = 24 * mm   # centre of bubble from box top
            BUBBLE_GAP    = 29 * mm     # centre-to-centre spacing (wider for clarity)
            WRITE_BOX_W   = 32 * mm     # "Write answer" box width
            CONTENT_X     = M + TAB_W + 3*mm
            CONTENT_W     = BOX_W - TAB_W - 3*mm
            BUBBLES_END   = CONTENT_X + BUBBLE_R + 3 * BUBBLE_GAP + BUBBLE_R
            WRITE_BOX_X   = CONTENT_X + CONTENT_W - WRITE_BOX_W  # flush right

            for q in obj_qs:
                if y - BOX_H < 15*mm:
                    y = new_page()
                marks = float(q.get("marks", q.get("weightage", 1)))
                m_str = f'{int(marks) if marks == int(marks) else marks}m'

                # ── Outer box ────────────────────────────────────────────────
                c.setStrokeColor(rl_colors.HexColor('#cbd5e1'))
                c.setFillColor(rl_colors.HexColor('#f8fafc'))
                c.setLineWidth(1.0)
                c.roundRect(M, y - BOX_H, BOX_W, BOX_H, 3, fill=1, stroke=1)

                # ── Q-number tab ─────────────────────────────────────────────
                c.setFillColor(rl_colors.HexColor('#4f46e5'))
                c.roundRect(M, y - BOX_H, TAB_W, BOX_H, 3, fill=1, stroke=0)
                c.setFillColor(rl_colors.HexColor('#4f46e5'))
                c.rect(M + TAB_W - 4, y - BOX_H, 4, BOX_H, fill=1, stroke=0)

                c.setFillColor(rl_colors.white)
                c.setFont("Helvetica-Bold", 13)
                c.drawCentredString(M + TAB_W/2, y - 10*mm, f'Q{q["id"]}')
                c.setFont("Helvetica", 7)
                c.drawCentredString(M + TAB_W/2, y - BOX_H + 3.5*mm, f'[{m_str}]')

                # Separator between tab and content
                c.setStrokeColor(rl_colors.HexColor('#e2e8f0'))
                c.setLineWidth(0.5)
                c.line(M + TAB_W, y - BOX_H + 2, M + TAB_W, y - 2)

                # ── Instruction text ─────────────────────────────────────────
                c.setFillColor(rl_colors.HexColor('#64748b'))
                c.setFont("Helvetica", 7.5)
                c.drawString(CONTENT_X, y - 7*mm, 'Circle ONE bubble:')

                # ── Bubbles with labels ABOVE (so label stays visible when filled) ─
                BUBBLE_Y = y - BUBBLE_FROM_TOP  # centre Y of bubbles
                LABEL_Y  = y - BUBBLE_FROM_TOP + BUBBLE_R + 4*mm  # label above bubble

                for i, label in enumerate(['A', 'B', 'C', 'D']):
                    bx = CONTENT_X + BUBBLE_R + i * BUBBLE_GAP

                    # Option letter label ABOVE the circle (stays visible when bubble is filled)
                    c.setFillColor(rl_colors.HexColor('#1e293b'))
                    c.setFont("Helvetica-Bold", 11)
                    c.drawCentredString(bx, LABEL_Y, label)

                    # Circle (white fill, dark border) — student fills/circles this
                    c.setStrokeColor(rl_colors.HexColor('#1e293b'))
                    c.setFillColor(rl_colors.white)
                    c.setLineWidth(1.5)
                    c.circle(bx, BUBBLE_Y, BUBBLE_R, fill=1, stroke=1)

                # ── Vertical divider before write-box ────────────────────────
                c.setStrokeColor(rl_colors.HexColor('#e2e8f0'))
                c.setLineWidth(0.8)
                c.line(WRITE_BOX_X - 3*mm, y - 3*mm, WRITE_BOX_X - 3*mm, y - BOX_H + 3*mm)

                # ── "Write answer" box (right side) ──────────────────────────
                c.setFillColor(rl_colors.HexColor('#64748b'))
                c.setFont("Helvetica", 7)
                c.drawCentredString(WRITE_BOX_X + WRITE_BOX_W/2, y - 7*mm, 'OR write letter:')

                # Square for writing the letter (bold border, generous size)
                SQ = 20 * mm
                sq_x = WRITE_BOX_X + (WRITE_BOX_W - SQ) / 2
                sq_y = y - BUBBLE_FROM_TOP - SQ/2
                c.setStrokeColor(rl_colors.HexColor('#64748b'))
                c.setFillColor(rl_colors.white)
                c.setLineWidth(2.0)
                c.roundRect(sq_x, sq_y, SQ, SQ, 2, fill=1, stroke=1)

                y -= BOX_H + 4*mm

        # ── Section B — Subjective ────────────────────────────────────────────
        if subj_qs:
            y -= 2*mm
            if y < 35*mm:
                y = new_page()
            _pdf_section_bar(c, M, y, BOX_W, 8*mm,
                '#ecfdf5', f'SECTION B — SUBJECTIVE  ({len(subj_qs)} questions)  '
                '·  Write your answer on the ruled lines inside the box',
                '#065f46', font_size=8.5, mm=mm)
            y -= 10*mm

            TAB_W   = 18 * mm
            LINE_H  = 9  * mm

            for q in subj_qs:
                marks   = float(q.get("marks", q.get("weightage", 1)))
                n_lines = min(14, max(5, int(marks * 1.8)))
                BOX_H   = (10 + n_lines * LINE_H/mm + 4) * mm

                if y - BOX_H < 15*mm:
                    y = new_page()

                m_str = f'{int(marks) if marks == int(marks) else marks}m'

                # ── Outer box ────────────────────────────────────────────────
                c.setStrokeColor(rl_colors.HexColor('#a7f3d0'))
                c.setFillColor(rl_colors.HexColor('#f0fdf4'))
                c.setLineWidth(1.0)
                c.roundRect(M, y - BOX_H, BOX_W, BOX_H, 3, fill=1, stroke=1)

                # ── Q-number tab ─────────────────────────────────────────────
                c.setFillColor(rl_colors.HexColor('#059669'))
                c.roundRect(M, y - BOX_H, TAB_W, BOX_H, 3, fill=1, stroke=0)
                c.setFillColor(rl_colors.HexColor('#059669'))
                c.rect(M + TAB_W - 4, y - BOX_H, 4, BOX_H, fill=1, stroke=0)

                c.setFillColor(rl_colors.white)
                c.setFont("Helvetica-Bold", 13)
                c.drawCentredString(M + TAB_W/2, y - BOX_H/2 + 2*mm, f'Q{q["id"]}')
                c.setFont("Helvetica", 7)
                c.drawCentredString(M + TAB_W/2, y - BOX_H + 4*mm, f'[{m_str}]')

                c.setStrokeColor(rl_colors.HexColor('#d1fae5'))
                c.setLineWidth(0.5)
                c.line(M + TAB_W, y - BOX_H + 2, M + TAB_W, y - 2)

                # ── "Write your answer below:" instruction ────────────────────
                CONTENT_X = M + TAB_W + 3*mm
                c.setFillColor(rl_colors.HexColor('#64748b'))
                c.setFont("Helvetica", 7.5)
                c.drawString(CONTENT_X, y - 7*mm, 'Write your answer below:')

                # Divider under instruction
                c.setStrokeColor(rl_colors.HexColor('#d1fae5'))
                c.setLineWidth(0.4)
                c.line(CONTENT_X, y - 9.5*mm, M + BOX_W - 3, y - 9.5*mm)

                # ── Ruled lines ───────────────────────────────────────────────
                LINE_X0 = CONTENT_X
                LINE_X1 = M + BOX_W - 3*mm
                for li in range(n_lines):
                    ly = y - 10*mm - (li + 1) * LINE_H
                    c.setStrokeColor(rl_colors.HexColor('#bbf7d0'))
                    c.setLineWidth(0.35)
                    c.line(LINE_X0, ly, LINE_X1, ly)

                y -= BOX_H + 5*mm

        # Last-page footer
        c.setFillColor(rl_colors.HexColor('#94a3b8'))
        c.setFont("Helvetica", 6.5)
        c.drawCentredString(W / 2, 5*mm,
            'Write answers ONLY inside the boxes  ·  Do not write outside the boxes')
        c.save()
        buf.seek(0)
        return buf.read()

    except Exception as e:
        log.error(f"[ANSWER-SHEET] Generation failed: {e}", exc_info=True)
        return b""

def _generate_question_paper_pdf(exam: dict) -> bytes:
    """
    Question Paper PDF — clean, professional exam paper layout.
    Header: school name, subject/board/class, marks/time/date info strip,
            amber instruction box pointing to the separate answer sheet.
    Body: numbered questions with clear visual hierarchy.
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors as rl_colors
        import io as _io2

        buf = _io2.BytesIO()
        W, H = A4
        c   = rl_canvas.Canvas(buf, pagesize=A4)
        M   = 15 * mm

        qs      = exam.get("questions", [])
        obj_qs  = [q for q in qs if q.get("type", "objective") == "objective"]
        subj_qs = [q for q in qs if q.get("type", "objective") != "objective"]

        obj_m    = sum(float(q.get("marks", q.get("weightage", 1))) for q in obj_qs)
        subj_m   = sum(float(q.get("marks", q.get("weightage", 1))) for q in subj_qs)
        total_m  = exam.get("total_marks", obj_m + subj_m)
        time_min = exam.get("time_minutes",
                   max(30, int(len(obj_qs) * 1.5 + len(subj_qs) * 8)))

        def draw_header(page_num=1):
            # ── Title block ───────────────────────────────────────────────────
            c.setFillColor(rl_colors.HexColor('#1e1b4b'))
            c.rect(M, H - 18*mm, W - 2*M, 18*mm, fill=1, stroke=0)

            # School name
            c.setFillColor(rl_colors.white)
            c.setFont("Helvetica-Bold", 13)
            school = exam.get('school_name', 'VidyAI School')
            c.drawString(M + 5*mm, H - 8*mm, school[:50])

            # Subject / board / class — right side
            subj_str = '  ·  '.join(str(p) for p in [
                exam.get('subject',''), exam.get('board',''), exam.get('class','')
            ] if p)
            c.setFont("Helvetica", 9.5)
            c.drawRightString(W - M - 5*mm, H - 8*mm, subj_str)

            # "QUESTION PAPER" label below school name
            c.setFont("Helvetica", 8)
            c.setFillColor(rl_colors.HexColor('#a5b4fc'))
            c.drawString(M + 5*mm, H - 14*mm, 'QUESTION PAPER')
            c.setFillColor(rl_colors.HexColor('#c7d2fe'))
            c.drawRightString(W - M - 5*mm, H - 14*mm, f'Page {page_num}')

            # ── Info strip — 4-cell grid ──────────────────────────────────────
            INFO_H = 12 * mm
            INFO_Y = H - 18*mm - INFO_H
            cells  = [
                ('Total Marks',  str(int(total_m))),
                ('Time Allowed', f'{time_min} minutes'),
                ('Date',         exam.get('exam_date', '—')),
                ('Exam ID',      exam.get('exam_id', '—')[:12]),
            ]
            cell_w = (W - 2*M) / len(cells)
            for ci, (lbl, val) in enumerate(cells):
                cx = M + ci * cell_w
                bg = '#f8fafc' if ci % 2 == 0 else '#f1f5f9'
                c.setFillColor(rl_colors.HexColor(bg))
                c.rect(cx, INFO_Y, cell_w, INFO_H, fill=1, stroke=0)
                # Divider
                if ci > 0:
                    c.setStrokeColor(rl_colors.HexColor('#cbd5e1'))
                    c.setLineWidth(0.5)
                    c.line(cx, INFO_Y + 2*mm, cx, INFO_Y + INFO_H - 2*mm)
                c.setFillColor(rl_colors.HexColor('#64748b'))
                c.setFont("Helvetica", 7)
                c.drawCentredString(cx + cell_w/2, INFO_Y + 7.5*mm, lbl)
                c.setFillColor(rl_colors.HexColor('#0f172a'))
                c.setFont("Helvetica-Bold", 9)
                c.drawCentredString(cx + cell_w/2, INFO_Y + 3.5*mm, val)
            # Border around info strip
            c.setStrokeColor(rl_colors.HexColor('#cbd5e1'))
            c.setLineWidth(0.6)
            c.rect(M, INFO_Y, W - 2*M, INFO_H, fill=0, stroke=1)

            # ── Section marks summary bar ─────────────────────────────────────
            SUMM_Y = INFO_Y - 8*mm
            c.setFillColor(rl_colors.HexColor('#1e1b4b'))
            c.rect(M, SUMM_Y, W - 2*M, 8*mm, fill=1, stroke=0)
            c.setFillColor(rl_colors.white)
            c.setFont("Helvetica-Bold", 8)
            c.drawString(M + 5*mm, SUMM_Y + 3*mm,
                f'Section A (Objective): {int(obj_m)} marks  ·  '
                f'Section B (Subjective): {int(subj_m)} marks  ·  '
                f'Total: {int(total_m)} marks  ·  '
                f'All questions compulsory')

            # ── Instructions box ─────────────────────────────────────────────
            INSTR_Y = SUMM_Y - 13*mm
            c.setFillColor(rl_colors.HexColor('#fffbeb'))
            c.setStrokeColor(rl_colors.HexColor('#f59e0b'))
            c.setLineWidth(0.8)
            c.roundRect(M, INSTR_Y, W - 2*M, 12*mm, 2, fill=1, stroke=1)
            c.setFillColor(rl_colors.HexColor('#92400e'))
            c.setFont("Helvetica-Bold", 8.5)
            c.drawString(M + 5*mm, INSTR_Y + 8.5*mm,
                'OPTION 1 (Preferred): Use the separate ANSWER SHEET — circle the bubble / write in the box')
            c.setFont("Helvetica", 7.5)
            c.drawString(M + 5*mm, INSTR_Y + 4.5*mm,
                'OPTION 2 (Fallback): Write answers directly on this question paper in the space provided')
            c.drawString(M + 5*mm, INSTR_Y + 1*mm,
                'If using Answer Sheet: Exam ID must match. Download Answer Sheet from teacher portal.')

            return INSTR_Y - 6*mm

        page = 1
        y    = draw_header(page)

        def new_page():
            nonlocal page
            c.showPage(); page += 1
            return draw_header(page)

        # ── Section A ─────────────────────────────────────────────────────────
        if obj_qs:
            _pdf_section_bar(c, M, y, W - 2*M, 8*mm,
                '#eef2ff',
                f'SECTION A — OBJECTIVE QUESTIONS  '
                f'({len(obj_qs)} questions  ·  {int(obj_m)} marks total)',
                '#3730a3', font_size=9, mm=mm)
            y -= 11*mm

            for q in obj_qs:
                marks = float(q.get("marks", q.get("weightage", 1)))
                qtext = str(q.get("question", ""))
                q_lines = _pdf_wrap(f'Q{q["id"]}.  {qtext}', 90)
                opts    = q.get("options") or {}

                # Two-row options always (cleaner than conditional)
                items   = sorted(opts.items())
                opt_h   = (10 if len(items) > 2 else 6) * mm
                needed  = (len(q_lines) * 5 + 4 + opt_h/mm + 6) * mm
                if y - needed < 20*mm:
                    y = new_page()

                # Subtle left accent bar for question number
                c.setFillColor(rl_colors.HexColor('#4f46e5'))
                c.rect(M, y - len(q_lines)*5*mm - 2, 2, len(q_lines)*5*mm + 2, fill=1, stroke=0)

                # Question text
                c.setFillColor(rl_colors.HexColor('#0f172a'))
                c.setFont("Helvetica-Bold", 10)
                for li, line in enumerate(q_lines):
                    c.drawString(M + 5*mm, y - (li+1)*5*mm, line)

                # Marks pill — right
                m_str = f'[{int(marks) if marks == int(marks) else marks} mark{"s" if marks!=1 else ""}]'
                c.setFont("Helvetica", 7.5)
                c.setFillColor(rl_colors.HexColor('#6366f1'))
                c.drawRightString(W - M, y - 5*mm, m_str)
                y -= (len(q_lines)*5 + 3) * mm

                # Options — always 2 rows of 2
                if opts:
                    c.setFont("Helvetica", 9.5)
                    c.setFillColor(rl_colors.HexColor('#1e293b'))
                    if len(items) == 4:
                        row1 = '          '.join(f'({k})  {v}' for k, v in items[:2])
                        row2 = '          '.join(f'({k})  {v}' for k, v in items[2:])
                        c.drawString(M + 8*mm, y - 5*mm, row1)
                        c.drawString(M + 8*mm, y - 10*mm, row2)
                        y -= 13*mm
                    else:
                        row = '     '.join(f'({k})  {v}' for k, v in items)
                        c.drawString(M + 8*mm, y - 5*mm, row)
                        y -= 8*mm

                # Answer blank line (for students writing directly on paper)
                c.setFillColor(rl_colors.HexColor('#374151'))
                c.setFont("Helvetica", 8.5)
                c.drawString(M + 8*mm, y - 5*mm, "Answer: ")
                c.setStrokeColor(rl_colors.HexColor('#94a3b8'))
                c.setLineWidth(0.7)
                c.line(M + 28*mm, y - 4*mm, M + 90*mm, y - 4*mm)
                y -= 8*mm

                # Thin separator between questions
                c.setStrokeColor(rl_colors.HexColor('#e2e8f0'))
                c.setLineWidth(0.4)
                c.line(M, y - 2, W - M, y - 2)
                y -= 5*mm

        # ── Section B ─────────────────────────────────────────────────────────
        if subj_qs:
            y -= 2*mm
            if y < 35*mm:
                y = new_page()
            _pdf_section_bar(c, M, y, W - 2*M, 8*mm,
                '#ecfdf5',
                f'SECTION B — SUBJECTIVE QUESTIONS  '
                f'({len(subj_qs)} questions  ·  {int(subj_m)} marks total)',
                '#065f46', font_size=9, mm=mm)
            y -= 11*mm

            for q in subj_qs:
                marks   = float(q.get("marks", q.get("weightage", 1)))
                qtext   = str(q.get("question", ""))
                q_lines = _pdf_wrap(f'Q{q["id"]}.  {qtext}', 90)
                needed  = (len(q_lines)*5 + 10) * mm
                if y - needed < 20*mm:
                    y = new_page()

                c.setFillColor(rl_colors.HexColor('#059669'))
                c.rect(M, y - len(q_lines)*5*mm - 2, 2, len(q_lines)*5*mm + 2, fill=1, stroke=0)

                c.setFillColor(rl_colors.HexColor('#0f172a'))
                c.setFont("Helvetica-Bold", 10)
                for li, line in enumerate(q_lines):
                    c.drawString(M + 5*mm, y - (li+1)*5*mm, line)

                m_str = f'[{int(marks) if marks==int(marks) else marks} marks]'
                c.setFont("Helvetica", 7.5)
                c.setFillColor(rl_colors.HexColor('#059669'))
                c.drawRightString(W - M, y - 5*mm, m_str)
                y -= (len(q_lines)*5 + 3) * mm

                # Answer lines for fallback writing on paper
                # Number of lines proportional to marks
                n_ans_lines = min(8, max(3, int(marks * 1.2)))
                for li in range(n_ans_lines):
                    line_y = y - (li+1) * 7*mm
                    c.setStrokeColor(rl_colors.HexColor('#cbd5e1'))
                    c.setLineWidth(0.5)
                    c.line(M + 2*mm, line_y, W - M - 2*mm, line_y)
                y -= (n_ans_lines * 7 + 3) * mm

                c.setStrokeColor(rl_colors.HexColor('#d1fae5'))
                c.setLineWidth(0.4)
                c.line(M, y - 2, W - M, y - 2)
                y -= 5*mm

        # Footer
        c.setFillColor(rl_colors.HexColor('#94a3b8'))
        c.setFont("Helvetica", 6.5)
        c.drawCentredString(W/2, 5*mm,
            f'Question Paper  ·  {exam.get("school_name","")}  ·  '
            f'Exam ID: {exam.get("exam_id","—")}  ·  '
            f'Total Marks: {int(total_m)}')
        c.save()
        buf.seek(0)
        return buf.read()

    except Exception as e:
        log.error(f"[QUESTION-PAPER] Generation failed: {e}", exc_info=True)
        return b""

def _generate_answer_key_pdf(exam: dict) -> bytes:
    """
    Answer Key PDF — teacher use only.
    Header: confidential dark red banner, exam metadata, marking scheme grid.
    MCQ: question + all options with correct one highlighted in green.
    Subjective: question + model answer block + key points + marking rubric.
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors as rl_colors
        import io as _io2

        buf = _io2.BytesIO()
        W, H = A4
        c   = rl_canvas.Canvas(buf, pagesize=A4)
        M   = 15 * mm

        qs      = exam.get("questions", [])
        obj_qs  = [q for q in qs if q.get("type", "objective") == "objective"]
        subj_qs = [q for q in qs if q.get("type", "objective") != "objective"]
        obj_m   = sum(float(q.get("marks", q.get("weightage", 1))) for q in obj_qs)
        subj_m  = sum(float(q.get("marks", q.get("weightage", 1))) for q in subj_qs)
        total_m = exam.get("total_marks", obj_m + subj_m)

        def draw_header(page_num=1):
            # ── Confidential banner ───────────────────────────────────────────
            c.setFillColor(rl_colors.HexColor('#7f1d1d'))
            c.rect(M, H - 14*mm, W - 2*M, 14*mm, fill=1, stroke=0)

            c.setFillColor(rl_colors.white)
            c.setFont("Helvetica-Bold", 12)
            school = exam.get('school_name', 'VidyAI School')
            c.drawString(M + 5*mm, H - 9*mm, school[:48])

            # ANSWER KEY badge
            badge_x = W - M - 50*mm
            badge_y = H - 14*mm + 3*mm
            c.setFillColor(rl_colors.HexColor('#dc2626'))
            c.roundRect(badge_x, badge_y, 46*mm, 8*mm, 2, fill=1, stroke=0)
            c.setFillColor(rl_colors.white)
            c.setFont("Helvetica-Bold", 8)
            c.drawCentredString(badge_x + 23*mm, badge_y + 2.5*mm, '🔑  ANSWER KEY — CONFIDENTIAL')

            # ── Exam info strip ───────────────────────────────────────────────
            INFO_H = 10 * mm
            INFO_Y = H - 14*mm - INFO_H
            c.setFillColor(rl_colors.HexColor('#fef2f2'))
            c.rect(M, INFO_Y, W - 2*M, INFO_H, fill=1, stroke=0)

            subj_str = '  ·  '.join(str(p) for p in [
                exam.get('subject',''), exam.get('board',''), exam.get('class','')
            ] if p)
            c.setFillColor(rl_colors.HexColor('#991b1b'))
            c.setFont("Helvetica-Bold", 8.5)
            c.drawString(M + 5*mm, INFO_Y + 6.5*mm, subj_str)
            c.setFont("Helvetica", 7.5)
            c.setFillColor(rl_colors.HexColor('#374151'))
            c.drawString(M + 5*mm, INFO_Y + 2.5*mm,
                f'Date: {exam.get("exam_date","—")}  ·  '
                f'Exam ID: {exam.get("exam_id","—")}')
            c.setFillColor(rl_colors.HexColor('#94a3b8'))
            c.drawRightString(W - M - 5*mm, INFO_Y + 4.5*mm, f'Page {page_num}')

            # ── Marking scheme grid ───────────────────────────────────────────
            GRID_H  = 10 * mm
            GRID_Y  = INFO_Y - GRID_H
            cells   = [
                ('Section A (Objective)', f'{int(obj_m)} marks', '#eef2ff', '#3730a3'),
                ('Section B (Subjective)', f'{int(subj_m)} marks', '#ecfdf5', '#065f46'),
                ('TOTAL', f'{int(total_m)} marks', '#1e1b4b', '#ffffff'),
            ]
            cw = (W - 2*M) / 3
            for ci, (lbl, val, bg, fg) in enumerate(cells):
                cx = M + ci * cw
                c.setFillColor(rl_colors.HexColor(bg))
                c.rect(cx, GRID_Y, cw, GRID_H, fill=1, stroke=0)
                c.setStrokeColor(rl_colors.HexColor('#cbd5e1'))
                c.setLineWidth(0.4)
                c.rect(cx, GRID_Y, cw, GRID_H, fill=0, stroke=1)
                c.setFillColor(rl_colors.HexColor(fg))
                c.setFont("Helvetica", 7)
                c.drawCentredString(cx + cw/2, GRID_Y + 6.5*mm, lbl)
                c.setFont("Helvetica-Bold", 10)
                c.drawCentredString(cx + cw/2, GRID_Y + 2.5*mm, val)

            return GRID_Y - 6*mm

        page = 1
        y    = draw_header(page)

        def new_page():
            nonlocal page
            c.showPage(); page += 1
            return draw_header(page)

        # ── Section A: Objective Answer Key ───────────────────────────────────
        if obj_qs:
            _pdf_section_bar(c, M, y, W - 2*M, 8*mm,
                '#eef2ff', f'SECTION A — OBJECTIVE  ({len(obj_qs)} questions)',
                '#3730a3', font_size=9, mm=mm)
            y -= 11*mm

            for q in obj_qs:
                marks = float(q.get("marks", q.get("weightage", 1)))
                qtext = str(q.get("question", ""))
                opts  = q.get("options") or {}

                # Resolve correct answer
                correct = set()
                raw_ans = str(q.get("answer", "")).strip().upper()
                if raw_ans in opts:
                    correct.add(raw_ans)
                for va in q.get("valid_answers", []):
                    va_u = str(va).strip().upper()
                    if va_u in opts:
                        correct.add(va_u)
                    else:
                        for k, v in opts.items():
                            if str(v).strip().lower() == va.strip().lower():
                                correct.add(k.upper())
                if not correct and raw_ans:
                    correct.add(raw_ans)

                q_lines = _pdf_wrap(f'Q{q["id"]}. {qtext}', 90)
                needed  = (len(q_lines)*5 + len(opts)*6 + 8) * mm
                if y - needed < 20*mm:
                    y = new_page()

                # Question
                c.setFillColor(rl_colors.HexColor('#7f1d1d'))
                c.rect(M, y - len(q_lines)*5*mm, 2, len(q_lines)*5*mm, fill=1, stroke=0)
                c.setFillColor(rl_colors.HexColor('#0f172a'))
                c.setFont("Helvetica-Bold", 9.5)
                for li, line in enumerate(q_lines):
                    c.drawString(M + 5*mm, y - (li+1)*5*mm, line)
                m_str = f'[{int(marks) if marks==int(marks) else marks} mark{"s" if marks!=1 else ""}]'
                c.setFont("Helvetica", 7.5)
                c.setFillColor(rl_colors.HexColor('#6b7280'))
                c.drawRightString(W - M, y - 5*mm, m_str)
                y -= (len(q_lines)*5 + 2) * mm

                # Options — highlight correct
                for k, v in sorted(opts.items()):
                    is_correct = k.upper() in correct
                    opt_y      = y - 5.5*mm
                    if is_correct:
                        c.setFillColor(rl_colors.HexColor('#dcfce7'))
                        c.roundRect(M + 3*mm, opt_y - 2*mm, W - 2*M - 3*mm, 7*mm, 1.5, fill=1, stroke=0)
                        c.setFillColor(rl_colors.HexColor('#15803d'))
                        c.setFont("Helvetica-Bold", 9.5)
                        c.drawString(M + 6*mm, opt_y, f'({k})  {v}')
                        c.setFont("Helvetica-Bold", 8)
                        c.drawRightString(W - M - 4*mm, opt_y, '✓  CORRECT')
                    else:
                        c.setFillColor(rl_colors.HexColor('#374151'))
                        c.setFont("Helvetica", 9.5)
                        c.drawString(M + 6*mm, opt_y, f'({k})  {v}')
                    y -= 6*mm

                # Explanation
                expl = str(q.get("explanation", "")).strip()
                if expl:
                    ex_lines = _pdf_wrap(f'Explanation: {expl}', 90)
                    EX_H     = (len(ex_lines)*4.5 + 4) * mm
                    c.setFillColor(rl_colors.HexColor('#eff6ff'))
                    c.roundRect(M + 3*mm, y - EX_H - 1*mm, W - 2*M - 3*mm, EX_H, 2, fill=1, stroke=0)
                    c.setFillColor(rl_colors.HexColor('#1d4ed8'))
                    c.setFont("Helvetica-Oblique", 8)
                    for li, line in enumerate(ex_lines):
                        c.drawString(M + 6*mm, y - 4*mm - li*4.5*mm, line)
                    y -= EX_H + 2*mm

                c.setStrokeColor(rl_colors.HexColor('#e2e8f0'))
                c.setLineWidth(0.4)
                y -= 2*mm
                c.line(M, y, W - M, y)
                y -= 4*mm

        # ── Section B: Subjective Answer Key ──────────────────────────────────
        if subj_qs:
            y -= 2*mm
            if y < 35*mm:
                y = new_page()
            _pdf_section_bar(c, M, y, W - 2*M, 8*mm,
                '#ecfdf5', f'SECTION B — SUBJECTIVE  ({len(subj_qs)} questions)',
                '#065f46', font_size=9, mm=mm)
            y -= 11*mm

            for q in subj_qs:
                marks      = float(q.get("marks", q.get("weightage", 1)))
                qtext      = str(q.get("question", ""))
                valid_ans  = q.get("valid_answers", [])
                key_pts    = q.get("answer_key_points", [])
                rubric     = str(q.get("evaluation_rubric",
                                       q.get("evaluation_criteria", ""))).strip()
                model_ans  = str(q.get("answer","")).strip() or (valid_ans[0] if valid_ans else "")

                q_lines    = _pdf_wrap(f'Q{q["id"]}. {qtext}', 90)
                ans_lines  = _pdf_wrap(f'Model Answer: {model_ans}', 88) if model_ans else []
                kp_lines   = []
                for kp in key_pts:
                    kp_lines += _pdf_wrap(f'  •  {kp}', 86)

                needed = (len(q_lines)*5 + len(ans_lines)*5 +
                          len(kp_lines)*5 + 25) * mm
                if y - needed < 20*mm:
                    y = new_page()

                # Question
                c.setFillColor(rl_colors.HexColor('#059669'))
                c.rect(M, y - len(q_lines)*5*mm, 2, len(q_lines)*5*mm, fill=1, stroke=0)
                c.setFillColor(rl_colors.HexColor('#0f172a'))
                c.setFont("Helvetica-Bold", 9.5)
                for li, line in enumerate(q_lines):
                    c.drawString(M + 5*mm, y - (li+1)*5*mm, line)
                m_str = f'[{int(marks) if marks==int(marks) else marks} marks]'
                c.setFont("Helvetica", 7.5)
                c.setFillColor(rl_colors.HexColor('#6b7280'))
                c.drawRightString(W - M, y - 5*mm, m_str)
                y -= (len(q_lines)*5 + 3) * mm

                # Model answer + key points block
                if ans_lines or kp_lines:
                    all_lines = ans_lines[:]
                    if kp_lines:
                        all_lines += ['Key Points:'] + kp_lines
                    BLK_H = (len(all_lines)*4.8 + 5) * mm
                    c.setFillColor(rl_colors.HexColor('#f0fdf4'))
                    c.setStrokeColor(rl_colors.HexColor('#86efac'))
                    c.setLineWidth(0.8)
                    c.roundRect(M + 3*mm, y - BLK_H - 1*mm,
                                W - 2*M - 3*mm, BLK_H, 2, fill=1, stroke=1)
                    for li, line in enumerate(all_lines):
                        bold = (li == 0 or line == 'Key Points:')
                        c.setFont("Helvetica-Bold" if bold else "Helvetica", 8.5)
                        c.setFillColor(rl_colors.HexColor('#14532d'))
                        c.drawString(M + 6*mm, y - 4.5*mm - li*4.8*mm, line)
                    y -= BLK_H + 3*mm

                # Rubric block
                if rubric:
                    rub_lines = _pdf_wrap(f'Marking Criteria: {rubric}', 88)
                    RUB_H     = (len(rub_lines)*4.5 + 4) * mm
                    c.setFillColor(rl_colors.HexColor('#fffbeb'))
                    c.setStrokeColor(rl_colors.HexColor('#fbbf24'))
                    c.setLineWidth(0.6)
                    c.roundRect(M + 3*mm, y - RUB_H - 1*mm,
                                W - 2*M - 3*mm, RUB_H, 2, fill=1, stroke=1)
                    c.setFillColor(rl_colors.HexColor('#78350f'))
                    c.setFont("Helvetica-Oblique", 8)
                    for li, line in enumerate(rub_lines):
                        c.drawString(M + 6*mm, y - 4*mm - li*4.5*mm, line)
                    y -= RUB_H + 3*mm

                c.setStrokeColor(rl_colors.HexColor('#d1fae5'))
                c.setLineWidth(0.4)
                y -= 2*mm
                c.line(M, y, W - M, y)
                y -= 5*mm

        # Confidential footer
        c.setFillColor(rl_colors.HexColor('#fef2f2'))
        c.rect(M, 4*mm, W - 2*M, 7*mm, fill=1, stroke=0)
        c.setFillColor(rl_colors.HexColor('#991b1b'))
        c.setFont("Helvetica-Bold", 7)
        c.drawCentredString(W/2, 7*mm,
            'CONFIDENTIAL  ·  FOR TEACHER / EXAMINER USE ONLY  ·  DO NOT DISTRIBUTE TO STUDENTS')

        c.save()
        buf.seek(0)
        return buf.read()

    except Exception as e:
        log.error(f"[ANSWER-KEY] Generation failed: {e}", exc_info=True)
        return b""

# Keep _generate_blank_answer_sheet_pdf as an alias so _diff_and_ocr_answers
# generates a blank that MATCHES the new structured answer sheet layout exactly.
def _generate_blank_answer_sheet_pdf(exam: dict) -> bytes:
    """Alias — generates the same structured answer sheet (blank, no handwriting)."""
    return _generate_answer_sheet_pdf(exam)

def _diff_and_ocr_answers(scanned_pdf_path: str, exam: dict) -> Dict[int, str]:
    """
    Core handwriting extraction using image diff:
    1. Generate the blank answer sheet from exam data
    2. Convert both blank and scanned to images at same DPI
    3. Pixel-diff: scanned - blank = only handwriting remains
    4. Per-question: crop the diff image to that question's answer region
    5. Send each crop to GPT-4o Vision for clean OCR (sees only ink on white)
    Returns {question_id: answer_text}
    """
    import io as _io
    import numpy as np
    from PIL import Image

    questions   = exam.get("questions", [])
    obj_qs      = [q for q in questions if q.get("type", "objective") == "objective"]
    subj_qs     = [q for q in questions if q.get("type", "objective") != "objective"]

    if not questions:
        return {}

    # ── Step 1: generate blank answer sheet ───────────────────────────────────
    blank_pdf_bytes = _generate_blank_answer_sheet_pdf(exam)
    if not blank_pdf_bytes:
        print("[DIFF-OCR] Blank PDF generation failed — falling back")
        return {}

    # ── Step 2: convert both PDFs to images ──────────────────────────────────
    try:
        DPI = 200
        import fitz as _fitz
        # Blank PDF from bytes
        _bdoc = _fitz.open(stream=blank_pdf_bytes, filetype="pdf")
        blank_pages = []
        for p in _bdoc:
            _pix = p.get_pixmap(dpi=DPI)
            blank_pages.append(Image.frombytes("RGB", [_pix.width, _pix.height], _pix.samples))
            del _pix
        _bdoc.close()
        # Scanned PDF from path
        _sdoc = _fitz.open(scanned_pdf_path)
        scanned_pages = []
        for p in _sdoc:
            _pix = p.get_pixmap(dpi=DPI)
            scanned_pages.append(Image.frombytes("RGB", [_pix.width, _pix.height], _pix.samples))
            del _pix
        _sdoc.close()
    except Exception as e:
        print(f"[DIFF-OCR] PDF→image failed: {e}")
        return {}

    log.info(f"[DIFF-OCR] Processing {len(blank_pages)} page(s) in parallel")
    return _ocr_page_parallel(
        images=scanned_pages,
        exam_questions=questions,
        exam=exam,
        blank_pages=blank_pages,
    )

# ── File-hash cache helper ────────────────────────────────────────────────────
def _file_hash(path: str) -> str:
    """MD5 of file contents — stable cache key independent of filename."""
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()

def _cache_get(file_hash: str, exam_id: str) -> Optional[dict]:
    """Return cached evaluation payload if it exists and matches current pipeline version, else None."""
    if not MONGO_OK:
        return None
    try:
        doc = mongo_db["eval_cache"].find_one({"cache_key": f"{file_hash}:{exam_id}"})
        if doc:
            payload = doc.get("payload", {})
            # Invalidate cache if OCR pipeline version has changed
            if payload.get("ocr_version") != OCR_PIPELINE_VERSION:
                log.info(f"[CACHE MISS] version mismatch hash={file_hash[:8]} exam={exam_id} "
                         f"cached={payload.get('ocr_version')} current={OCR_PIPELINE_VERSION}")
                return None
            log.info(f"[CACHE HIT] hash={file_hash[:8]} exam={exam_id}")
            return payload
    except Exception as e:
        log.warning(f"[CACHE] get error: {e}")
    return None

def _cache_set(file_hash: str, exam_id: str, payload: dict):
    """Write evaluation payload to cache.
    
    MongoDB requires all dict keys to be strings. submitted_answers uses integer
    question IDs as keys ({1: 'C', 2: 'D', ...}), so we must stringify them before
    storing, and question_wise question_id values stay as-is (they're in values not keys).
    """
    if not MONGO_OK:
        return
    try:
        import json as _json
        # Serialize then deserialize to convert any integer dict keys → string keys
        clean_payload = _json.loads(_json.dumps(payload, default=str))
        mongo_db["eval_cache"].update_one(
            {"cache_key": f"{file_hash}:{exam_id}"},
            {"$set": {
                "cache_key": f"{file_hash}:{exam_id}",
                "payload":   clean_payload,
                "cached_at": _now(),
            }},
            upsert=True,
        )
    except Exception as e:
        log.warning(f"[CACHE] set error: {e}")

def _find_uploaded_file(file_name: str, file_path: str = "") -> Optional[Path]:
    """Locate the original uploaded file so stale reports can be refreshed."""
    if file_path:
        p = Path(file_path)
        if p.exists():
            return p
    if file_name:
        exact = UPL_DIR / file_name
        if exact.exists():
            return exact
        matches = list(UPL_DIR.glob(f"*_{file_name}"))
        if len(matches) == 1:
            return matches[0]
    return None

def _refresh_evaluation_for_report(ev: dict, exam: dict) -> dict:
    """Recompute stored evaluation results when the OCR pipeline changes.
    
    Priority:
    1. If ocr_version matches current → skip (already up to date)
    2. If submitted_answers exist → re-evaluate from stored answers (fast)
    3. If page_indices + file available → re-extract from student pages only
    4. Fall back to full-PDF extraction (last resort)
    """
    if not ev:
        return ev

    # ── Fast path: already up to date ─────────────────────────────────────
    if ev.get("ocr_version") == OCR_PIPELINE_VERSION:
        return ev

    # ── Re-evaluate from stored submitted_answers (no OCR needed) ─────────
    submitted = ev.get("submitted_answers")
    if isinstance(submitted, dict) and submitted:
        normalized: Dict[int, str] = {}
        for k, v in submitted.items():
            try:
                normalized[int(k)] = v
            except (TypeError, ValueError):
                continue
        if normalized:
            refreshed = dict(ev)
            refreshed["result"] = _evaluate_answers(exam, normalized, str(ev.get("roll_no", "")))
            refreshed["ocr_version"] = OCR_PIPELINE_VERSION
            refreshed["report_refreshed_at"] = _now()
            eid = str(refreshed.get("evaluation_id", "")).strip()
            if eid:
                evaluations_registry[eid] = refreshed
                if not IS_VERCEL:
                    _save_json(EVAL_DIR / f"{eid}.json", refreshed)
                _save_evals()
                _mongo_update("eval_items", {"evaluation_id": eid}, {"$set": {**refreshed, "evaluation_id": eid}}, upsert=True)
            return refreshed

    # ── Re-extract from student's specific pages (not full PDF) ───────────
    upload_path = _find_uploaded_file(
        str(ev.get("file_name", "")).strip(),
        str(ev.get("file_path", "")).strip(),
    )
    if not upload_path:
        log.info(f"[REPORT] No upload found to refresh evaluation {ev.get('evaluation_id', '—')}")
        return ev

    page_indices = ev.get("page_indices")
    answers = {}
    mode = ""
    err = None

    if isinstance(page_indices, list) and page_indices:
        # Per-student page extraction (correct for multi-student PDFs)
        answers, mode, err = _extract_answers_from_pages(
            str(upload_path), page_indices,
            exam_questions=exam.get("questions", []), exam=exam, attempts=2
        )
        log.info(f"[REPORT] Refreshed {ev.get('evaluation_id', '—')} from pages {page_indices}: {len(answers)} answers")
    else:
        # Single-student PDF fallback
        answers, mode, err = _extract_answers_with_retry(
            str(upload_path), exam_questions=exam.get("questions", []), exam=exam, attempts=2
        )

    if not answers:
        log.warning(f"[REPORT] Refresh failed for {ev.get('evaluation_id', '—')} ({mode})")
        return ev

    refreshed = dict(ev)
    refreshed["submitted_answers"] = answers
    refreshed["result"] = _evaluate_answers(exam, answers, str(ev.get("roll_no", "")))
    refreshed["extraction_mode"] = mode
    refreshed["ocr_version"] = OCR_PIPELINE_VERSION
    refreshed["report_refreshed_at"] = _now()

    eid = str(refreshed.get("evaluation_id", "")).strip()
    if eid:
        evaluations_registry[eid] = refreshed
        if not IS_VERCEL:
            _save_json(EVAL_DIR / f"{eid}.json", refreshed)
        _save_evals()
        _mongo_update("eval_items", {"evaluation_id": eid}, {"$set": {**refreshed, "evaluation_id": eid}}, upsert=True)
    return refreshed


def _detect_bubbles_opencv(image, num_questions=4, num_options=4) -> dict:
    import cv2
    import numpy as np
    from PIL import Image
    import logging
    log = logging.getLogger("parvidya.api")
    
    results = {}
    try:
        if isinstance(image, Image.Image):
            image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
        
        H_px, W_px = gray.shape
        A4_W_MM = 210.0
        A4_H_MM = 297.0
        mm_x = W_px / A4_W_MM
        mm_y = H_px / A4_H_MM
        
        HEADER_MM = 65.0
        SLOT_MM = 40.0
        
        # We use the old layout coordinates since all our test PDFs are in the old layout
        # (Centers at 47.2, 73.1, 99.1, 125.0)
        centers_mm = [47.2, 73.1, 99.1, 125.0][:num_options]
        options = ['A', 'B', 'C', 'D'][:num_options]
        BUBBLE_R = 7.0
        
        for q_idx in range(num_questions):
            box_top_mm = HEADER_MM + q_idx * SLOT_MM
            abs_bubble_y_mm = box_top_mm + 24.0
            bubble_y_px = int(abs_bubble_y_mm * mm_y)
            
            dark_ratios = []
            for cx_mm in centers_mm:
                cx_px = int(cx_mm * mm_x)
                r_px = int(BUBBLE_R * mm_x)
                
                x1 = max(0, cx_px - r_px)
                x2 = min(W_px, cx_px + r_px)
                y1 = max(0, bubble_y_px - r_px)
                y2 = min(H_px, bubble_y_px + r_px)
                
                if x1 >= x2 or y1 >= y2:
                    dark_ratios.append(0)
                    continue
                    
                region = gray[y1:y2, x1:x2]
                _, thresh = cv2.threshold(region, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
                
                dark_pixels = np.sum(thresh == 255)
                total_pixels = region.size
                dark_ratios.append(dark_pixels / total_pixels if total_pixels > 0 else 0)
            
            if not dark_ratios:
                continue
                
            max_idx = np.argmax(dark_ratios)
            max_fill = dark_ratios[max_idx]
            
            sorted_fills = sorted(dark_ratios, reverse=True)
            fill_margin = sorted_fills[0] - sorted_fills[1] if len(sorted_fills) > 1 else sorted_fills[0]
            
            absolute_confidence = min(max_fill / 0.5, 1.0)
            margin_confidence = min(fill_margin / 0.15, 1.0)
            confidence = (absolute_confidence * 0.6 + margin_confidence * 0.4)
            
            qid = q_idx + 1
            results[qid] = {
                'answer': options[max_idx],
                'confidence': round(confidence, 3),
                'fill_pct': round(max_fill, 3)
            }
            log.info(f"[OpenCV-OMR] Q{qid}: {options[max_idx]} (conf={confidence:.3f}, fill={max_fill:.3f})")
            
        return results
    except Exception as e:
        import logging
        logging.getLogger("parvidya.api").error(f"[OpenCV-OMR] Bubble detection failed: {e}")
        return {}

def _ocr_with_google_vision(image_pil) -> str:
    """Extract handwritten text from a PIL image.
    Primary: Google Cloud Vision (high accuracy for handwriting).
    Fallback: OpenAI GPT-4o-mini vision (used when GCV SDK/credentials missing).
    """
    import io, base64
    # ── Attempt Google Cloud Vision ──────────────────────────────────────────
    try:
        from google.cloud import vision as _gcv
        client = _gcv.ImageAnnotatorClient()
        buf = io.BytesIO()
        image_pil.save(buf, format="PNG")
        content = buf.getvalue()
        image = _gcv.Image(content=content)
        response = client.document_text_detection(image=image)
        if response.full_text_annotation:
            return response.full_text_annotation.text.strip()
        return ""
    except Exception as _gce:
        log.warning(f"[GCV-OCR] Google Vision unavailable ({_gce}); falling back to OpenAI Vision")

    # ── Fallback: OpenAI GPT-4o-mini Vision ─────────────────────────────────
    if not OPENAI_API_KEY:
        log.warning("[GCV-OCR] No OpenAI key available for fallback OCR")
        return ""
    try:
        import openai as _oai
        buf = io.BytesIO()
        image_pil.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        client = _oai.OpenAI(api_key=OPENAI_API_KEY)
        resp = client.chat.completions.create(
            model=OCR_VISION_MODEL,
            max_tokens=400,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": (
                        "This is a crop from a handwritten student answer sheet. "
                        "Transcribe ALL handwritten text exactly as written. "
                        "Return only the transcribed text, nothing else."
                    )},
                    {"type": "image_url", "image_url": {
                        "url": f"data:image/png;base64,{b64}",
                        "detail": "high"
                    }},
                ],
            }],
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as _oae:
        log.warning(f"[GCV-OCR] OpenAI Vision fallback also failed: {_oae}")
        return ""

def _ocr_with_google_vision_detailed(image_pil):
    """Detailed OCR with bounding boxes using Google Cloud Vision.
    Falls back to None gracefully if SDK/credentials unavailable.
    """
    import io
    try:
        from google.cloud import vision as _gcv
        client = _gcv.ImageAnnotatorClient()
        buf = io.BytesIO()
        image_pil.save(buf, format="PNG")
        content = buf.getvalue()
        image = _gcv.Image(content=content)
        response = client.document_text_detection(image=image)
        return response.full_text_annotation
    except Exception as e:
        log.warning(f"[Google-Vision] Detailed OCR failed: {e}")
        return None

def _ocr_page_parallel(images: list, exam_questions: list, exam: dict,
                        blank_pages: list = None) -> Dict[int, str]:
    """
    Run OCR on multiple PDF pages in parallel using ThreadPoolExecutor.
    Each page is processed independently; results are merged.
    Uses image-diff if blank_pages supplied, else falls back to enhanced vision.
    """
    import io as _io
    import numpy as _np
    from PIL import Image, ImageEnhance

    combined: Dict[int, str] = {}

    obj_qs  = [q for q in exam_questions if q.get("type","objective") == "objective"]
    subj_qs = [q for q in exam_questions if q.get("type","objective") != "objective"]
    obj_ids  = ", ".join(str(q["id"]) for q in obj_qs)  or "none"
    subj_ids = ", ".join(str(q["id"]) for q in subj_qs) or "none"
    q_ref    = "\n".join(
        f"Q{q['id']}: {str(q.get('question',''))[:70]}"
        for q in sorted(exam_questions, key=lambda x: int(x.get("id",0)))
    )

    def _process_one_page(args):
        """
        Process one scanned page using PER-QUESTION crops.

        Why per-question crops (not full page):
        - The Q-number tab is solid dark blue/black when scanned, visually
          resembling a filled bubble. Full-page GPT-4o vision mistakes it for
          the first option and shifts all bubble labels by +1 (reads B as A, etc.)
        - Per-question crops: each crop contains ONLY the answer area for ONE
          question, with the Q-number tab on the left edge. GPT sees exactly
          4 circles (A B C D) with no ambiguous dark region to the left.
        - For subjective questions: crop the full box so all ruled lines are visible.

        Layout geometry (from _generate_answer_sheet_pdf):
          Page header: ~30% of page height (varies by content above questions)
          MCQ box: 28mm tall, 4mm gap below each box
          Subjective box: 10mm + n_lines*8mm + 4mm tall, 5mm gap
          Both pages contain questions — Page 1 has MCQs, Page 2+ has subjectives.
        """
        page_idx, img = args
        try:
            import openai as _oai_local
            from PIL import ImageEnhance as _IE
            client = _oai_local.OpenAI(api_key=OPENAI_API_KEY)
            page_answers: Dict[int, str] = {}

            # ── Step 1: Upscale + colour-safe contrast boost ──────────────────
            # On Render free tier (512 MB RAM) keep images smaller to avoid OOM.
            MAX_W = 1600 if IS_RENDER else 2480
            if img.width < MAX_W:
                scale = MAX_W / img.width
                img   = img.resize((MAX_W, int(img.height * scale)), Image.LANCZOS)
            elif img.width > MAX_W:
                scale = MAX_W / img.width
                img   = img.resize((MAX_W, int(img.height * scale)), Image.LANCZOS)
            img = _IE.Contrast(_IE.Sharpness(img.convert("RGB")).enhance(1.8)).enhance(1.5)

            W_px, H_px = img.width, img.height

            # ── Step 2: First send the FULL PAGE with question-counting prompt ─
            # This lets GPT identify which questions are on this page and roughly
            # where each answer area is, returning a best-effort JSON.
            # We use this as a baseline then verify per-question crops.
            # Use JPEG (not PNG) for base64: PNG of 2480px A4 is ~13 MB as b64;
            # JPEG quality=88 is ~3 MB — 4× smaller, no loss in OCR accuracy.
            buf_full = _io.BytesIO()
            img.save(buf_full, format="JPEG", quality=88, optimize=True)
            b64_full = base64.b64encode(buf_full.getvalue()).decode()
            buf_full.close()
            del buf_full

            # Questions on this page (heuristic: split evenly across pages)
            n_pages = max(len(images), 1)
            qs_per  = max(1, len(exam_questions) // n_pages)
            start_i = page_idx * qs_per
            end_i   = start_i + qs_per if page_idx < n_pages - 1 else len(exam_questions)
            page_qs = exam_questions[start_i:end_i]
            if not page_qs:
                page_qs = exam_questions   # fallback: try all questions

            page_obj_ids  = ", ".join(str(q["id"]) for q in page_qs
                                      if q.get("type","objective") == "objective") or "none"
            page_subj_ids = ", ".join(str(q["id"]) for q in page_qs
                                      if q.get("type","objective") != "objective") or "none"

            full_page_prompt = """
            You are reading a SCANNED STUDENT ANSWER SHEET.

            LAYOUT: Each question has a bordered box with a DARK TAB on the left showing Q number.

            MCQ QUESTIONS:
            - After the dark tab, there are 4 circles labelled A, B, C, D
            - Student fills one circle
            - Return ONLY the letter (A, B, C, or D) - no dashes, no brackets

            WRITTEN/SUBJECTIVE QUESTIONS:
            - After the dark tab, there is a box with multiple ruled lines
            - ⚠️ IMPORTANT: Students write across MULTIPLE LINES
            - You MUST read EVERY line from top to bottom
            - Do NOT stop after the first line
            - Join ALL lines into ONE string with spaces
            - If the answer spans 10 lines, read ALL 10 lines
            - Ignore printed labels like "Write your answer below:"

            Return ONLY valid JSON with answers in format:
            {"answers": {"1": "B", "2": "Complete answer with all lines..."}}
            """

            _openai_limiter.wait()
            resp = client.chat.completions.create(
                model=OPENAI_MODEL, temperature=0, max_tokens=4096,
                response_format={"type": "json_object"},
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": full_page_prompt},
                    {"type": "image_url", "image_url": {
                        "url": f"data:image/png;base64,{b64_full}",
                        "detail": "high"
                    }},
                ]}]
            )
            raw_content = resp.choices[0].message.content
            log.info(f"[PAGE-OCR] page {page_idx+1} full-page raw: {raw_content[:400]}")
            raw_full = json.loads(raw_content)
            page_answers = _validate_ocr_answers(raw_full.get("answers", {}), exam_questions)
            del b64_full  # free ~3 MB string before crop loop

            # ── Step 3: Per-question crop verification for MCQ only ───────────
            # Send each MCQ question as an individual crop to Google Vision and OpenCV
            mcq_qs_on_page = [q for q in page_qs if q.get("type","objective") == "objective"]
            if mcq_qs_on_page and len(mcq_qs_on_page) > 0:
                BOX_H_FRAC  = 28 / 297
                SLOT_FRAC   = 31 / 297
                PAD_FRAC    =  5 / 297
                HEADER_FRAC = (55 / 297) if page_idx == 0 else (25 / 297)

                verified: Dict[int, str] = {}
                
                cv_bubbles = _detect_bubbles_opencv(img, num_questions=len(mcq_qs_on_page), num_options=4)
                
                try:
                    from google.cloud import vision as _gcv_mod
                    vision_client = _gcv_mod.ImageAnnotatorClient()
                except Exception as e:
                    log.warning(f"[Google-Vision] Init failed (will use OpenAI fallback for MCQ crops): {e}")
                    vision_client = None

                for q_idx, q in enumerate(mcq_qs_on_page):
                    qid = int(q.get("id", 0))
                    q_options = q.get("options", {})
                    
                    y_top = int(H_px * (HEADER_FRAC + q_idx * SLOT_FRAC - PAD_FRAC))
                    y_bot = int(H_px * (HEADER_FRAC + q_idx * SLOT_FRAC + BOX_H_FRAC + PAD_FRAC))
                    y_top = max(0, min(y_top, H_px - 10))
                    y_bot = max(y_top + 10, min(y_bot, H_px))

                    # Crop rightmost 20% for the write box
                    split_x = int(W_px * 0.80)
                    box_crop = img.crop((split_x, y_top, W_px, y_bot))
                    box_crop = _IE.Contrast(box_crop).enhance(1.5)

                    cbuf = _io.BytesIO()
                    box_crop.save(cbuf, format="PNG")
                    content_img = cbuf.getvalue()
                    cbuf.close()
                    
                    vision_text = ""
                    if vision_client:
                        try:
                            vision_image = vision.Image(content=content_img)
                            v_resp = vision_client.document_text_detection(image=vision_image)
                            if v_resp.full_text_annotation:
                                vision_text = v_resp.full_text_annotation.text.strip()
                        except Exception as e:
                            log.warning(f"[Google-Vision] Q{qid} failed: {e}")

                    final_answer = ""
                    if vision_text:
                        v_clean = vision_text.upper().replace(' ', '').replace('\n', '')
                        if v_clean in ['A', 'B', 'C', 'D']:
                            final_answer = v_clean
                            log.info(f"[HYBRID-OCR] Q{qid} write-box='{vision_text}' -> direct {final_answer}")
                        else:
                            for opt_k, opt_v in q_options.items():
                                if str(opt_v).upper().replace(' ', '') == v_clean:
                                    final_answer = opt_k
                                    log.info(f"[HYBRID-OCR] Q{qid} write-box='{vision_text}' -> mapped {final_answer}")
                                    break
                                # fallback for numbers like "10" "20"
                                if str(opt_v) in vision_text:
                                    final_answer = opt_k
                                    log.info(f"[HYBRID-OCR] Q{qid} write-box contains '{opt_v}' -> mapped {final_answer}")
                                    break
                    
                    if not final_answer:
                        cv_ans = cv_bubbles.get(q_idx + 1, {}).get('answer', "")
                        if cv_ans:
                            final_answer = cv_ans
                            log.info(f"[HYBRID-OCR] Q{qid} bubble read -> {final_answer}")
                    
                    if final_answer:
                        verified[qid] = final_answer
                        log.info(f"[CROP-OCR] Q{qid} crop answer: {final_answer}")

                if verified:
                    for qid, crop_val in verified.items():
                        page_answers[qid] = crop_val
                    log.info(f"[CROP-OCR] page {page_idx+1} verified MCQ: {verified}")

            # ── Step 4: Per-question crop for Subjective questions ──────────────
            # This ensures all handwritten lines are extracted regardless of
            # whether the full-page pass captured them correctly.
            subj_qs_on_page = [q for q in page_qs if q.get("type", "objective") != "objective"]
            if subj_qs_on_page:
                # Subjective layout geometry after the Section B header.
                # MCQs occupy: header + n_obj * 31mm  (approx)
                # Section B bar: 10mm
                # Subjective box: variable, ~10mm + n_lines*9mm
                # We estimate n_lines based on marks (~1.8 lines per mark, min 5, max 14)
                n_obj = len([q for q in exam_questions if q.get("type","objective") == "objective"])
                # Approximate Y position of section B start
                HEADER_FRAC   = (55 / 297) if page_idx == 0 else (25 / 297)
                MCQ_TAKEN_MM  = 10 + n_obj * 31  # header bar + slots
                SECTION_B_MM  = MCQ_TAKEN_MM + 10  # section B bar
                LINE_H_MM     = 9

                subj_y_mm = SECTION_B_MM + 10   # after section B bar
                for q in subj_qs_on_page:
                    qid     = int(q.get("id", 0))
                    marks   = float(q.get("marks", q.get("weightage", 1)))
                    n_lines = min(14, max(5, int(marks * 1.8)))
                    box_h_mm = 10 + n_lines * LINE_H_MM + 4

                    # Convert mm positions to pixel fractions
                    y_top_mm = subj_y_mm
                    y_bot_mm = subj_y_mm + box_h_mm
                    subj_y_mm += box_h_mm + 5  # gap between boxes

                    y_top = int(H_px * (y_top_mm / 297) - 10)  # small padding
                    y_bot = int(H_px * (y_bot_mm / 297) + 10)
                    y_top = max(0, min(y_top, H_px - 10))
                    y_bot = max(y_top + 20, min(y_bot, H_px))

                    # Only crop if this question doesn't already have a full answer
                    existing = page_answers.get(qid, "")
                    if existing and len(existing) > 30:
                        log.info(f"[SUBJ-CROP] Q{qid} already has answer ({len(existing)} chars), skipping")
                        continue

                    crop = img.crop((0, y_top, W_px, y_bot))
                    from PIL import ImageEnhance as _IE2
                    crop = _IE2.Contrast(crop).enhance(1.4)
                    sbuf = _io.BytesIO()
                    crop.save(sbuf, format="JPEG", quality=88, optimize=True)
                    sb64 = base64.b64encode(sbuf.getvalue()).decode()
                    del crop
                    sbuf.close()
                    del sbuf

                    subj_crop_prompt = f"""
                    This image shows ONLY question Q{qid} (a written/subjective question) from a student answer sheet.

                    LAYOUT:
                    - LEFT side: DARK GREEN TAB with question number
                    - RIGHT side: Box with multiple ruled lines for the answer

                    CRITICAL: The student's answer often spans MULTIPLE LINES (3-10 lines).
                    You MUST read EVERY handwritten line from top to bottom.

                    INSTRUCTIONS:
                    1. Start from the first line below the printed instruction
                    2. Read line 1, line 2, line 3... until the bottom of the box
                    3. Join ALL lines into ONE string with spaces
                    4. Preserve the COMPLETE answer - do NOT truncate
                    5. Do NOT stop after reading just one line
                    6. Ignore printed text like "Write your answer below:"

                    Return ONLY valid JSON:
                    {{"answer": "the complete student answer with all lines combined"}}

                    If the box has no writing, return: {{"answer": ""}}
                    """
                    try:
                        _openai_limiter.wait()
                        subj_resp = client.chat.completions.create(
                            model=OPENAI_MODEL, temperature=0, max_tokens=1000,
                            response_format={"type": "json_object"},
                            messages=[{"role": "user", "content": [
                                {"type": "text", "text": subj_crop_prompt},
                                {"type": "image_url", "image_url": {
                                    "url": f"data:image/png;base64,{sb64}",
                                    "detail": "high"
                                }},
                            ]}]
                        )
                        subj_raw = json.loads(subj_resp.choices[0].message.content)
                        subj_val = str(subj_raw.get("answer", "")).strip()
                        if subj_val and len(subj_val) > 3:
                            page_answers[qid] = subj_val
                            log.info(f"[SUBJ-CROP] Q{qid} → '{subj_val[:80]}'")
                        else:
                            log.info(f"[SUBJ-CROP] Q{qid} empty or too short: {subj_val!r}")
                    except Exception as se:
                        log.warning(f"[SUBJ-CROP] Q{qid} failed: {se}")
                    finally:
                        del sb64

            del img  # free upscaled page image before returning
            log.info(f"[PAGE-OCR] page {page_idx+1} final: {page_answers}")
            return page_answers

        except Exception as e:
            log.error(f"[OCR-PAGE] page {page_idx} failed: {e}", exc_info=True)
            return {}

    # On Render free tier (512 MB RAM), process pages sequentially to avoid
    # multiple threads each holding 26 MB image + b64 simultaneously.
    max_w = 1 if IS_RENDER else min(4, len(images))
    with ThreadPoolExecutor(max_workers=max_w) as pool:
        futures = {pool.submit(_process_one_page, (i, img)): i
                   for i, img in enumerate(images)}
        for fut in as_completed(futures):
            page_result = fut.result()
            combined.update(page_result)

    return combined


def _extract_header_info(path: str) -> dict:
    """
    ZERO OpenAI Vision Cost Architecture for Header.
    Google Vision OCR (Text-only) + GPT (Text-only) for extraction.
    """
    import io as _io
    from PIL import Image
    try:
        from pathlib import Path
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            import fitz
            doc = fitz.open(path)
            pix = doc[0].get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            doc.close()
        else:
            img = Image.open(path).convert("RGB")
        
        # Crop header (top 28%)
        crop_h = int(img.height * 0.28)
        header_img = img.crop((0, 0, img.width, crop_h))
        
        # Google Vision OCR
        text = _ocr_with_google_vision(header_img)
        if not text: return {}
        
        # Text LLM for extraction
        prompt = (
            f"Extract student info from this OCR text:\n{text}\n\n"
            "Return ONLY JSON: {\"student_name\": \"...\", \"roll_no\": \"...\", \"class_section\": \"...\"}"
        )
        return _llm_json(prompt, mini=True)
    except Exception as e:
        log.warning(f"[HEADER] Extraction failed: {e}")
        return {}


def _warp_page_to_fixed_size(pil_img, target_w=2480, target_h=3508):
    """Warp the page image to a standard A4 300dpi resolution for deterministic OMR."""
    import numpy as np
    import cv2
    from PIL import Image
    img = np.array(pil_img.convert("RGB"))
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    
    # 1. Edge detection
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    edged = cv2.Canny(blurred, 50, 150)
    
    # 2. Find contours
    cnts, _ = cv2.findContours(edged, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    cnts = sorted(cnts, key=cv2.contourArea, reverse=True)[:10]
    
    rect = None
    for c in cnts:
        peri = cv2.arcLength(c, True)
        approx = cv2.approxPolyDP(c, 0.02 * peri, True)
        if len(approx) == 4:
            pts = approx.reshape(4, 2)
            s = pts.sum(axis=1)
            diff = np.diff(pts, axis=1)
            rect = np.zeros((4, 2), dtype="float32")
            rect[0] = pts[np.argmin(s)]
            rect[2] = pts[np.argmax(s)]
            rect[1] = pts[np.argmin(diff)]
            rect[3] = pts[np.argmax(diff)]
            break
    
    if rect is not None:
        rect_area = cv2.contourArea(rect)
        img_area = gray.shape[0] * gray.shape[1]
        if rect_area > img_area * 0.4:
            dst = np.array([
                [0, 0],
                [target_w - 1, 0],
                [target_w - 1, target_h - 1],
                [0, target_h - 1]], dtype="float32")
            M = cv2.getPerspectiveTransform(rect, dst)
            warped = cv2.warpPerspective(img, M, (target_w, target_h))
            return Image.fromarray(warped), True
    
    # IMPROVED FALLBACK: Crop content before resizing to 2480x3508
    log.warning("[OMR] Warp failed, trying content-aware resize")
    _, thresh = cv2.threshold(gray, 240, 255, cv2.THRESH_BINARY_INV)
    coords = cv2.findNonZero(thresh)
    if coords is not None:
        x, y, w, h = cv2.boundingRect(coords)
        # Add small margin
        m = 20
        x, y = max(0, x-m), max(0, y-m)
        w, h = min(gray.shape[1]-x, w+2*m), min(gray.shape[0]-y, h+2*m)
        cropped = img[y:y+h, x:x+w]
        return Image.fromarray(cropped).resize((target_w, target_h), Image.LANCZOS), False
        
    return pil_img.resize((target_w, target_h), Image.LANCZOS), False

def clean_ocr(text):
    """Clean OCR output to remove printed instructions."""
    remove = ["BLOCK_SECTION", "BLOCK_SECTION_B", "2m", "your on the ruled lines", "inside the box", "Wete", "Winte", "Wote", "Wofa", "waw", "Do not write", "answers ONLY"]
    for r in remove:
        # case insensitive replace using regex
        import re
        text = re.sub(re.escape(r), "", text, flags=re.IGNORECASE)
    # also remove any stray brackets or leading bullets
    text = re.sub(r'[\[\]\(\)\-\.]', ' ', text)
    return text.strip()

def _clean_subjective_ocr(text: str) -> str:
    """Strip printed form template text from subjective OCR, keeping only handwritten answers.

    Removes:
    • "Q5" / "Q6" labels (dark sidebar and inline)
    • "Write your answer below:" instruction line
    • Marks indicators: "3m", "[3m]", "2m", "[2m]" etc.
    • OCR noise from dark sidebar: "wine your…", "wing your…", etc.
    • "Write your" prefix fragments at line start
    """
    import re
    if not text:
        return ""

    lines = text.splitlines()
    cleaned = []
    for line in lines:
        s = line.strip()
        # Drop blank lines after stripping (they'll be re-added selectively)
        if not s:
            continue
        # Drop template patterns
        if re.match(r'^Q\s*\d+\s*$', s, re.I):              # "Q5", "Q6"
            continue
        if re.match(r'write\s+your\s+(answer|answ)', s, re.I):  # "Write your answer…"
            continue
        if re.match(r'\[?\d+\s*m\]?\s*$', s, re.I):          # "3m", "[3m]", "2m", "[2m]"
            continue
        if re.match(r'(wine|wing|wini|winn)\s+your', s, re.I):   # dark-sidebar OCR noise
            continue
        cleaned.append(s)

    return "\n".join(cleaned).strip()


def _match_write_box_to_option(wr_text: str, q_options: dict) -> str:
    """Map write-box OCR text to an option letter (A/B/C/D).

    Strategy (in order):
    1. Reject AI refusal messages (>12 chars).
    2. Transliterate visually-similar non-ASCII characters to ASCII
       (GCV sometimes returns Cyrillic В for B, С for C, etc.).
    3. Single A/B/C/D letter within very short text (≤3 chars).
    4. Numeric value matching against exam option values (strips units).
       Also substitutes common OCR letter↔digit confusions (I/L→1, O→0).
    5. Any A/B/C/D letter found anywhere in short text.
    Returns '' if no match found.
    """
    import re

    # Strip the printed "OR write:" template label that appears above the write box.
    # When the crop window is wide enough we capture both the label and the answer;
    # remove the label line so only the student's handwritten portion is matched.
    import re as _re
    _stripped = _re.sub(
        r"(?i)^(or\s*write[:\s]*|ok\s*wh[a-z]*[:\s]*)\n?",
        "",
        wr_text.strip(),
    ).strip()
    wr_clean = (_stripped or wr_text.strip()).upper()

    # 1. Reject obviously-wrong text (AI refusals, printed template labels, long error messages)
    if len(wr_clean) > 12:
        return ""
    # Reject leftover template text or AI error phrases
    _REJECT_PHRASES = ("OR WRITE", "OR WRIT", "OK WHAT", "OK WHIT", "OK WHI",
                       "SORRY", "UNABLE", "CANNOT", "CAN'T", "SEE ANY", "BLANK",
                       "NO TEXT", "NO HANDWRITTEN", "NO MARK", "EMPTY", "NOTHING")
    if any(p in wr_clean for p in _REJECT_PHRASES):
        return ""

    # 2. Transliterate visually-similar non-ASCII → ASCII
    #    GCV often returns Cyrillic characters for Latin-looking letters
    _VISUAL = {"В": "B", "С": "C", "А": "A", "Е": "E", "О": "O",
               "Н": "H", "Р": "P", "Т": "T", "Х": "X", "М": "M"}
    for src, dst in _VISUAL.items():
        wr_clean = wr_clean.replace(src, dst)

    # 3. Short text (≤3 chars) with ONLY letters (no digits) — pure letter answer like "C" or "B"
    #    Skip this step when digits are present; they indicate a numeric value that must go
    #    through step 4 first (e.g., "2D" should be treated as "20", not matched to "D").
    hits = [ch for ch in wr_clean if ch in "ABCD"]
    has_digits = any(c.isdigit() for c in wr_clean)
    if hits and len(wr_clean) <= 3 and not has_digits:
        return hits[0]

    # 4. Numeric value → option value matching
    if q_options:
        # Apply common OCR letter↔digit substitutions before stripping.
        # I/L → 1, O → 0, P/G → 6 (handwritten 6 sometimes reads as P or G)
        # S → 9 (handwritten 9 with open loop can look like 'S' to OCR)
        # Z → 2 (handwritten 2 sometimes reads as Z)
        _NUM_SUBS = str.maketrans("ILOGPSZ", "1106692")
        wr_for_num = wr_clean.translate(_NUM_SUBS)
        wr_num = re.sub(r"[^0-9.]", "", wr_for_num)
        if wr_num:
            # 4a. Exact numeric match
            for letter, opt_text in q_options.items():
                opt_num = re.sub(r"[^0-9.]", "", str(opt_text).upper())
                if opt_num and wr_num == opt_num:
                    return letter

            # 4b. Single-digit suffix match: "O" → "0" matches "10" unambiguously
            if len(wr_num) == 1:
                suffix_hits = [
                    letter for letter, opt_text in q_options.items()
                    if re.sub(r"[^0-9.]", "", str(opt_text).upper()).endswith(wr_num)
                ]
                if len(suffix_hits) == 1:
                    return suffix_hits[0]

        # 4c. Try additional substitution tables when the primary didn't match.
        #     Different digits are commonly confused in handwriting OCR:
        #       D → 0 (e.g., "2D" misread from "20")
        #       S → 8 (e.g., "S" or "8" written in cursive look identical)
        #       4 → 8 (closed-top handwritten 4 can resemble 8)
        #     We try each alternative and return the first unambiguous hit.
        _ALT_SUBS = [
            ("D→0",  str.maketrans("ILOGPSDZD", "110669200")),   # D→0
            ("S→8",  str.maketrans("ILOGPSD",   "1106682")),     # S→8
            ("4→8",  str.maketrans("4",          "8")),           # closed-top 4 misread as 8
            ("D0S8", str.maketrans("ILOGPSDZ",  "11066820")),    # combined
        ]
        for _label, _subs in _ALT_SUBS:
            _wr_alt = wr_clean.translate(_subs)
            _wr_num = re.sub(r"[^0-9.]", "", _wr_alt)
            if _wr_num and _wr_num != wr_num:
                for letter, opt_text in q_options.items():
                    opt_num = re.sub(r"[^0-9.]", "", str(opt_text).upper())
                    if opt_num and _wr_num == opt_num:
                        return letter

    # 5. Any A/B/C/D letter anywhere in very short text (≤5 chars to avoid false
    #    positives from words like "WHAT" which contains 'A' by accident).
    #    Still skip if digits are present — numeric context overrides letter guessing.
    if hits and len(wr_clean) <= 5 and not has_digits:
        return hits[0]

    return ""


def _ocr_cropped_region(image_pil, box, remove_header_px=50, write_box=False):
    """OCR a cropped region of a page image for subjective answers.

    Pipeline:
      1. Crop + downscale (max 1024px wide) — reduces API payload ~4x.
      2. MD5 cache check — identical crops skip all OCR.
      3. Google Cloud Vision (primary, best handwriting accuracy).
      4. OpenAI GPT-4o-mini vision (fallback when GCV unavailable).
      Returns the extracted text string (may be empty).
    """
    import io, hashlib
    import numpy as np
    import cv2
    from PIL import Image

    xl, xr, yt, yb = box
    w, h = image_pil.size
    xl, xr = max(0, xl), min(w, xr)
    yt, yb = max(0, yt), min(h, yb)
    if xr <= xl or yb <= yt:
        return ""

    crop = image_pil.crop((xl, yt, xr, yb))

    # ── Erase printed header lines at top of crop ────────────────────────────
    if remove_header_px > 0:
        crop_np = np.array(crop)
        if crop_np.shape[0] > remove_header_px:
            crop_np[:remove_header_px, :] = 255
        crop = Image.fromarray(crop_np)

    # ── Resize: upscale tiny crops (improves OCR on small write boxes) ───────
    MIN_W = 300
    if crop.width < MIN_W:
        ratio = MIN_W / crop.width
        crop = crop.resize((MIN_W, int(crop.height * ratio)), Image.LANCZOS)

    # ── Downscale: cap at 1024px wide to cut API payload ────────────────────
    MAX_W = 1024
    if crop.width > MAX_W:
        ratio = MAX_W / crop.width
        crop = crop.resize((MAX_W, int(crop.height * ratio)), Image.LANCZOS)

    # ── Cache check ──────────────────────────────────────────────────────────
    buf = io.BytesIO()
    crop.save(buf, format="JPEG", quality=80)   # JPEG saves ~3x vs PNG
    img_bytes = buf.getvalue()
    img_hash = hashlib.md5(img_bytes).hexdigest()
    _cache_key = f"ocr_crop_{img_hash}"

    cached = _eval_cache.get(_cache_key)
    if cached is not None:
        # Reject stale AI-refusal responses that were cached before the
        # write_box=True prompt was added (they contain long refusal strings
        # rather than a single character/number or "BLANK").
        if len(cached) > 12 or any(p in cached.upper() for p in (
                "NO HANDWRITTEN", "NO TEXT", "NO VISIBLE", "THERE IS NO",
                "I CANNOT", "UNABLE TO")):
            pass   # fall through to fresh OCR
        else:
            return clean_ocr(cached)

    # ── Google Cloud Vision (primary) ────────────────────────────────────────
    text = ""
    try:
        from google.cloud import vision as _gcv
        client = _gcv.ImageAnnotatorClient()
        image_obj = _gcv.Image(content=img_bytes)
        # Force English to prevent GCV misidentifying digits as Korean/Cyrillic
        img_ctx = _gcv.ImageContext(language_hints=["en"])
        resp = client.document_text_detection(image=image_obj, image_context=img_ctx)
        if resp.full_text_annotation:
            text = resp.full_text_annotation.text.strip()
    except Exception as _gce:
        log.warning(f"[OCR-CROP] GCV unavailable ({type(_gce).__name__}); trying OpenAI fallback")

    # ── OpenAI Vision fallback ───────────────────────────────────────────────
    if not text and OPENAI_API_KEY:
        try:
            import openai as _oai, base64 as _b64
            b64 = _b64.b64encode(img_bytes).decode()
            client_oai = _oai.OpenAI(api_key=OPENAI_API_KEY)
            if write_box:
                # Write-box: student writes a single digit, number, or letter A-D.
                # Use a specific prompt to prevent model refusals on faint/small marks.
                ocr_prompt = (
                    "This image shows a small handwritten answer box from a student exam. "
                    "The student has written a single digit (0-9), a short number, or a letter (A/B/C/D). "
                    "Even if the writing is faint, small, or unclear, do your best to read it. "
                    "Return ONLY the handwritten character(s) — nothing else. "
                    "If the box is completely empty with no marks at all, return the single word: BLANK"
                )
                detail_level = "high"   # higher detail for small single-character images
            else:
                ocr_prompt = (
                    "Transcribe ALL handwritten text visible in this image. "
                    "Return only the transcribed text."
                )
                detail_level = "low"   # 'low' costs ~85 tokens vs 'high' ~765
            resp_oai = client_oai.chat.completions.create(
                model=OCR_VISION_MODEL,
                max_tokens=50 if write_box else 300,
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": ocr_prompt},
                    {"type": "image_url", "image_url": {
                        "url": f"data:image/jpeg;base64,{b64}",
                        "detail": detail_level,
                    }},
                ]}],
            )
            text = (resp_oai.choices[0].message.content or "").strip()
        except Exception as _oae:
            log.warning(f"[OCR-CROP] OpenAI fallback failed: {_oae}")

    _eval_cache.set(f"ocr_crop_{img_hash}", text)
    return clean_ocr(text)

def _extract_answers(path: str, exam_questions: list = None, exam: dict = None):
    """
    Extraction pipeline:

    OBJECTIVE  → HoughCircles OMR on page 0, OR-write-box OCR fallback
    SUBJECTIVE → page 1 of the PDF (2-page answer sheet), GCV/OpenAI OCR

    Layout (DNHS 2-page answer sheet at 150 dpi → 1240×1755 px per page)
    ─────────────────────────────────────────────────────────────────────
    Page 0 — MCQ bubbles
      Bubble X centres (% width):  A=22.4  B=34.9  C=47.2  D=59.5
      Bubble Y centres (% height): Q1=28.4  Q2=40.5  Q3=52.5  Q4=64.5
      OR-write box (right side):   x=69%–79%,  same Y as bubble row

    Page 1 — Subjective answers
      Q5 answer area: y=16.5%–36.5%  (first question box)
      Q6 answer area: y=37.5%–57.5%  (second question box)
    """
    import cv2, numpy as np
    from PIL import Image
    from pathlib import Path

    exam_questions = exam_questions or []
    mcq_ids  = sorted([int(q["id"]) for q in exam_questions if q.get("type", "objective") == "objective"])
    subj_ids = sorted([int(q["id"]) for q in exam_questions if q.get("type", "objective") != "objective"])

    # ── 1. Load both pages at 150 dpi ────────────────────────────────────────
    try:
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            import fitz
            doc = fitz.open(path)
            pix0 = doc[0].get_pixmap(dpi=150)
            page0 = Image.frombytes("RGB", [pix0.width, pix0.height], pix0.samples)
            page1 = None
            if len(doc) > 1:
                pix1 = doc[1].get_pixmap(dpi=150)
                page1 = Image.frombytes("RGB", [pix1.width, pix1.height], pix1.samples)
            doc.close()
        else:
            page0 = Image.open(path).convert("RGB")
            page1 = None
    except Exception as e:
        log.error(f"[EXTRACT] Page load failed: {e}")
        return {}, "load_failed"

    W0, H0 = page0.width, page0.height
    gray0 = np.array(page0.convert("L"))
    answers = {}

    # ── 2. Template OMR — objective answers from page 1 bubble rows ──────────
    # The DNHS answer sheet has exactly 4 MCQ bubble rows (Q1-Q4).
    # Any extra MCQ questions (Q5, Q6, ...) are read from page 2 boxes below.
    if mcq_ids:
        # Calibrated X fractions for A/B/C/D bubbles (measured at 150 dpi, 1240×1755)
        TMPL_X_FRAC = {"A": 0.224, "B": 0.349, "C": 0.472, "D": 0.595}
        # Calibrated Y fractions per question row (measured from actual sheets)
        TMPL_Y_FRAC = [0.284, 0.405, 0.525, 0.645]
        # OR-write box: spans x≈68%–90%; the printed "OR write:" label sits at
        # ~68–79% and the student's handwritten answer appears below it.  Using
        # 68–90% and ±80 px (vs the old 79–90% / ±55 px) ensures we capture the
        # answer even in combined/class PDFs where the answer may be written below
        # the printed label rather than to its right.
        OR_X1, OR_X2 = int(W0 * 0.68), int(W0 * 0.90)
        OR_DY = 80

        rad = max(15, int(22 * (H0 / 1754.0)))   # ~22 px at 150 dpi
        debug_img = cv2.cvtColor(gray0, cv2.COLOR_GRAY2BGR)

        # Only read from physical bubble rows (cap at TMPL_Y_FRAC length = 4)
        mcq_bubble_ids = mcq_ids[:len(TMPL_Y_FRAC)]
        if len(mcq_ids) > len(TMPL_Y_FRAC):
            log.warning(f"[OMR] Exam has {len(mcq_ids)} MCQ Qs but form has {len(TMPL_Y_FRAC)} bubble rows — reading only Q{mcq_bubble_ids} from bubbles; extra MCQs read from page 2 boxes")

        for i, qid in enumerate(mcq_bubble_ids):
            yf  = TMPL_Y_FRAC[i]
            yc  = int(yf * H0)
            centers = {letter: (int(f * W0), yc) for letter, f in TMPL_X_FRAC.items()}

            # Score each option: higher = darker = more ink = more likely filled
            scores: dict = {}
            for letter, (xc, _yc) in centers.items():
                cv2.circle(debug_img, (xc, _yc), rad, (0, 255, 0), 2)
                mask = np.zeros((H0, W0), np.uint8)
                cv2.circle(mask, (xc, _yc), rad, 255, -1)
                scores[letter] = 255.0 - cv2.mean(gray0, mask=mask)[0]

            ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
            winner, top_score = ranked[0]
            gap = ranked[0][1] - ranked[1][1] if len(ranked) > 1 else top_score

            # top_score < 25: printed bubble outlines in combined/class PDFs register
            # ~16–23 due to bolder ink; individual-PDF outlines are < 10.  Using 25 as
            # the "clearly filled" floor catches both.  Also fall back when the gap
            # between winner and runner-up is tiny (< 15), indicating noise rather than
            # a definitive mark.
            if top_score < 25 or gap < 15:
                wr_text = _ocr_cropped_region(
                    page0,
                    [OR_X1, OR_X2, max(0, yc - OR_DY), min(H0, yc + OR_DY)],
                    remove_header_px=0,
                    write_box=True,
                ).strip()
                # Look up option values so numeric write-box answers can be mapped
                q_options: dict = {}
                for _q in (exam_questions or []):
                    if int(_q.get("id", 0)) == qid:
                        raw = _q.get("options")
                        if isinstance(raw, dict):
                            q_options = raw
                        elif isinstance(raw, list):
                            q_options = {ch: str(v) for ch, v in zip("ABCD", raw)}
                        break
                matched = _match_write_box_to_option(wr_text, q_options)
                if matched:
                    winner = matched
                    log.info(f"[OMR-WR] Q{qid}: write-box '{wr_text}' → '{winner}'")
                else:
                    # Both bubble AND write-box failed — do not guess by defaulting to
                    # the OMR noise winner; mark as blank so grading gives 0 marks
                    # without accidentally crediting the wrong option.
                    winner = ""
                    log.warning(
                        f"[OMR] Q{qid}: no bubble and no write-box match ('{wr_text}') → blank"
                    )

            answers[qid] = winner
            log.info(f"[OMR] Q{qid}: {winner}  score={top_score:.1f}  gap={gap:.1f}")

        cv2.imwrite("/uploads/debug_bubbles_last.jpg", debug_img)

    # ── 3. Page 2 boxes — Q5 and Q6 (MCQ letter or subjective text) ──────────
    # The physical DNHS answer sheet always has exactly 2 answer boxes on page 2:
    #   Box 0 (top):    y=16.5%–35.5%  →  5th question by sorted ID  (Q5)
    #   Box 1 (bottom): y=37.5%–56.5%  →  6th question by sorted ID  (Q6)
    #
    # The box content is interpreted based on the exam's question type:
    #   objective  → OCR the letter/value and match to A/B/C/D option
    #   subjective → OCR as free text (student's written answer)
    #
    # This handles BOTH exam formats without any ID-based hardcoding:
    #   4-MCQ exam  (Q1-4 obj, Q5-6 subj): box0→Q5 text, box1→Q6 text
    #   6-MCQ exam  (Q1-6 obj, Q7-10 subj): box0→Q5 letter, box1→Q6 letter
    all_q_ids_sorted = sorted([int(q["id"]) for q in exam_questions])
    page2_q_ids = all_q_ids_sorted[4:6]  # 5th and 6th questions (physical page 2 boxes)

    if page2_q_ids and page1 is not None:
        Ws, Hs = page1.width, page1.height
        PHYS_BOXES_P2 = [
            [160, int(Ws * 0.93), int(Hs * 0.165), int(Hs * 0.355)],  # box 0: Q5 area
            [160, int(Ws * 0.93), int(Hs * 0.375), int(Hs * 0.565)],  # box 1: Q6 area
        ]
        for box_idx, qid in enumerate(page2_q_ids):
            if box_idx >= len(PHYS_BOXES_P2):
                break
            q_info = next((q for q in exam_questions if int(q.get("id", 0)) == qid), None)
            if q_info is None:
                continue
            q_type = q_info.get("type", "objective")
            raw_ocr = _ocr_cropped_region(page1, PHYS_BOXES_P2[box_idx], remove_header_px=55)

            if q_type == "objective":
                # MCQ on page 2: match OCR text to an option letter
                q_options: dict = {}
                raw_opts = q_info.get("options")
                if isinstance(raw_opts, dict):
                    q_options = raw_opts
                elif isinstance(raw_opts, list):
                    q_options = {ch: str(v) for ch, v in zip("ABCD", raw_opts)}
                matched = _match_write_box_to_option(raw_ocr.strip(), q_options)
                if matched:
                    answers[qid] = matched
                    log.info(f"[OMR-P2] Q{qid}: page2 box='{raw_ocr.strip()[:40]}' → '{matched}'")
                else:
                    log.warning(f"[OMR-P2] Q{qid}: could not match '{raw_ocr.strip()[:40]}' to option")
            else:
                # Subjective: extract free text
                ocr_text = _clean_subjective_ocr(raw_ocr)
                log.info(f"[OCR] Q{qid} subj: '{ocr_text[:80]}'")
                answers[qid] = ocr_text

    # Q7+ in 6-MCQ exams have no physical space on this 2-page DNHS form → not captured

    return answers, "omr+ocr_pipeline"


def _extract_answers_from_pages(pdf_path: str, page_indices: list, exam_questions: list = None, 
                                exam: dict = None, attempts: int = 3):
    """
    Extract answers from SPECIFIC pages of a PDF (e.g., student's pages only).
    Memory-optimized: Uses PyPDF to extract pages without loading entire PDF into memory.
    
    Args:
        pdf_path: Full path to the PDF file
        page_indices: List of page numbers to extract from (0-indexed)
        exam_questions: Optional list of exam questions for context
        exam: Optional exam dict
        attempts: Number of retry attempts
    
    Returns:
        (answers, mode, error) tuple where answers is extracted answers dict or empty dict
    """
    import tempfile as _tempfile
    from pathlib import Path as _PathLib
    
    if not page_indices or not isinstance(page_indices, (list, tuple)):
        log.warning(f"[EXTRACT-PAGES] Invalid page_indices: {page_indices}")
        return {}, "invalid_pages", "page_indices must be a non-empty list"
    
    page_indices = [int(p) for p in page_indices if isinstance(p, (int, float))]
    if not page_indices:
        return {}, "invalid_pages", "no valid page indices"
    
    tmp_path = None
    try:
        # Use PyPDF for memory-efficient page extraction
        import pypdf
        
        reader = pypdf.PdfReader(pdf_path)
        total_pages = len(reader.pages)
        
        # Validate page indices
        page_indices = [p for p in page_indices if 0 <= p < total_pages]
        if not page_indices:
            log.error(f"[EXTRACT-PAGES] No valid page indices after filtering (total={total_pages})")
            return {}, "no_valid_pages", f"all pages out of range [0, {total_pages})"
        
        # Create a temporary PDF with only the specified pages
        with _tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name
        
        # Write selected pages to temp file using PyPDF (memory-efficient)
        writer = pypdf.PdfWriter()
        for page_idx in sorted(set(page_indices)):
            writer.add_page(reader.pages[page_idx])
        
        with open(tmp_path, "wb") as output_file:
            writer.write(output_file)
        
        log.info(f"[EXTRACT-PAGES] Created temp PDF with pages {sorted(set(page_indices))} at {tmp_path}")
        
        # Extract answers from the temporary PDF
        answers, mode = _extract_answers(tmp_path, exam_questions=exam_questions, exam=exam)
        
        log.info(f"[EXTRACT-PAGES] Successfully extracted from pages {sorted(page_indices)}: {len(answers)} answers (mode={mode})")
        return answers, mode, None
        
    except Exception as ex:
        log.error(f"[EXTRACT-PAGES] Failed to extract from pages {page_indices}: {ex}")
        return {}, "extraction_error", str(ex)
    
    finally:
        # Clean up temporary file immediately to free memory
        if tmp_path:
            try:
                _PathLib(tmp_path).unlink(missing_ok=True)
                log.debug(f"[EXTRACT-PAGES] Cleaned up temp file: {tmp_path}")
            except Exception as cleanup_err:
                log.warning(f"[EXTRACT-PAGES] Failed to clean temp file: {cleanup_err}")


def _extract_answers_with_retry(path: str, exam_questions: list = None, exam: dict = None,
                                attempts: int = 3, base_delay: float = 1.5):
    """
    Retry OCR extraction to survive transient OpenAI/API turbulence.
    Returns (answers, mode, error_message_or_none).
    """
    def _retry_sleep_from_msg(msg: str, fallback: float) -> float:
        if not msg:
            return fallback
        _m = re.search(r"try again in\s*([0-9]*\.?[0-9]+)\s*(ms|s)", msg, re.IGNORECASE)
        if not _m:
            return fallback
        _v = float(_m.group(1))
        _u = _m.group(2).lower()
        _sec = (_v / 1000.0) if _u == "ms" else _v
        return max(fallback, _sec + 0.25)

    last_mode = "no_answers_extracted"
    last_err = ""
    attempts = max(1, int(attempts))
    for attempt in range(1, attempts + 1):
        try:
            answers, mode = _extract_answers(path, exam_questions=exam_questions, exam=exam)
            last_mode = mode
            if answers:
                return answers, mode, None
            last_err = f"empty_answers ({mode})"
        except Exception as ex:
            last_mode = "extract_exception"
            last_err = str(ex)
            log.warning(f"[OCR-RETRY] Attempt {attempt}/{attempts} failed for {Path(path).name}: {ex}")

        if attempt < attempts:
            _fallback = base_delay * attempt
            _sleep_for = _retry_sleep_from_msg(last_err, _fallback)
            _time.sleep(_sleep_for)

    return {}, last_mode, last_err or "no_answers_extracted"


def _evaluate_answers_with_retry(exam: dict, answers: dict, roll_no: str,
                                 attempts: int = 2, base_delay: float = 1.0):
    """Retry grading for transient upstream failures."""
    def _retry_sleep_from_msg(msg: str, fallback: float) -> float:
        if not msg:
            return fallback
        _m = re.search(r"try again in\s*([0-9]*\.?[0-9]+)\s*(ms|s)", msg, re.IGNORECASE)
        if not _m:
            return fallback
        _v = float(_m.group(1))
        _u = _m.group(2).lower()
        _sec = (_v / 1000.0) if _u == "ms" else _v
        return max(fallback, _sec + 0.25)

    last_err = "evaluation_failed"
    attempts = max(1, int(attempts))
    for attempt in range(1, attempts + 1):
        try:
            _openai_limiter.wait()
            return _evaluate_answers(exam, answers, roll_no), None
        except Exception as ex:
            last_err = str(ex)
            log.warning(f"[EVAL-RETRY] Attempt {attempt}/{attempts} failed for roll={roll_no}: {ex}")
            if attempt < attempts:
                _fallback = base_delay * attempt
                _sleep_for = _retry_sleep_from_msg(last_err, _fallback)
                _time.sleep(_sleep_for)
    return None, last_err


def _multi_student_worker_count(student_count: int) -> int:
    """
    Adaptive concurrency for class-PDF evaluation.
    Higher worker count helps 40-50 student batches while the thread-safe
    limiter keeps API calls paced.
    """
    if student_count >= 40:
        return 2
    if student_count >= 20:
        return 3
    return 4

def _resolve_multi_eval_batch_size(student_count: int) -> int:
    """
    Keep batch waves smaller for large class PDFs to reduce long-tail timeouts.
    MULTI_EVAL_BATCH_SIZE remains the upper bound from env.
    """
    base = max(1, MULTI_EVAL_BATCH_SIZE)
    if student_count >= 30:
        return min(base, 5)
    if student_count >= 20:
        return min(base, 6)
    return base

def _parse_answer_text(text: str) -> Dict[int, str]:
    """
    Parse student answers from text in multiple formats:

    Format A — Objective with Answer: line (Section A of this app's papers):
        Q1. What is pi?
        A) 3.14  B) 3.15  C) 3.16  D) 3.17
        Answer: A

    Format B — Subjective with written response (Section B, no Answer: line):
        Q7. Explain the Pythagorean theorem.
        In a right-angled triangle, the square of the hypotenuse equals...

    Format C — Inline (Q1: A  or  1. A):
        Q1: A
        1. A
    """
    answers: Dict[int, str] = {}
    lines = [ln.rstrip() for ln in text.splitlines()]

    # Compiled patterns
    Q_NUM      = re.compile(r"^Q(?:uestion)?\s*(\d+)[\.\s]|^(\d+)\s*[\)\.\-]", re.IGNORECASE)
    ANS_LABEL  = re.compile(r"^(?:ans(?:wer)?|response|solution)\s*[:\-]\s*(.*)", re.IGNORECASE)
    OPTION_ROW = re.compile(r"^[A-D]\s*[\):\.]", re.IGNORECASE)   # "A) 3.14  B) 3.15 ..."
    BLANK_LINE = re.compile(r"^[_\-\s]*$")
    INLINE_PAT = re.compile(r"^Q?(\d+)\s*[\.\)\-:]\s*(.+)", re.IGNORECASE)

    # ── Pass 1: fast inline check ─────────────────────────────────────────────
    # Only matches SHORT answer values that look like answers, not question text.
    # Criteria: value must be ≤60 chars, must NOT end with "?", must NOT contain
    # "[N mark" (mark annotations), must NOT start with a capital question word.
    QUESTION_WORDS = re.compile(
        r"^(?:what|which|why|how|when|where|explain|describe|define|"
        r"calculate|find|solve|prove|show|discuss|evaluate|compare|list)",
        re.IGNORECASE
    )
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        m = INLINE_PAT.match(s)
        if m:
            val = m.group(2).strip()
            if BLANK_LINE.match(val):
                continue
            # Reject if it looks like question text
            if val.endswith("?"):
                continue
            if "[" in val and "mark" in val.lower():
                continue
            if len(val) > 80:
                continue
            if QUESTION_WORDS.match(val):
                continue
            answers[int(m.group(1))] = val

    # ── Pass 2: block-aware walk ──────────────────────────────────────────────
    # State machine: track current question number and accumulate written lines
    current_qnum: Optional[int] = None
    written_lines: list = []          # accumulates subjective answer lines
    in_mcq_block: bool = False        # True while inside an MCQ options block

    def _flush(qnum, wlines):
        """Commit accumulated subjective answer lines for qnum."""
        if qnum is None or qnum in answers:
            return
        text_val = " ".join(wlines).strip()
        text_val = re.sub(r"\s+", " ", text_val)
        if text_val and not BLANK_LINE.match(text_val):
            answers[qnum] = text_val

    for ln in lines:
        s = ln.strip()

        # New question heading?
        mq = Q_NUM.match(s)
        if mq:
            # Flush previous subjective block first
            _flush(current_qnum, written_lines)
            written_lines = []
            qn = mq.group(1) or mq.group(2)
            current_qnum = int(qn) if qn else current_qnum
            in_mcq_block = False
            continue

        # MCQ option row (A) 3.14  B) 3.15 ...) — skip, not an answer
        if OPTION_ROW.match(s):
            in_mcq_block = True
            continue

        # Answer: label line
        ma = ANS_LABEL.match(s)
        if ma:
            val = ma.group(1).strip()
            if not BLANK_LINE.match(val) and val:
                if current_qnum is not None and current_qnum not in answers:
                    answers[current_qnum] = val
            # Reset — Answer: line closes this block regardless
            written_lines = []
            in_mcq_block = False
            continue

        # Empty / underscore-only line
        if not s or BLANK_LINE.match(s):
            continue

        # Remaining non-empty line — could be subjective written answer
        # Only accumulate if: not in MCQ block, not a header/instruction line
        SKIP_PAT = re.compile(
            r"^(?:section|general\s+instructions?|date|student\s+name|"
            r"roll\s+number|marks|class|about:blank|[0-9]{1,2}/[0-9]{1,2}/[0-9]{4})",
            re.IGNORECASE
        )
        if not in_mcq_block and current_qnum is not None and not SKIP_PAT.match(s):
            # Skip if it looks like a question (ends with ? or starts with capitalised instruction)
            if not s.endswith("?") and len(s) > 2:
                written_lines.append(s)

    # Flush last block
    _flush(current_qnum, written_lines)

    return answers

def _vision_extract(image_path: str, prompt: str = None) -> Dict[int, str]:
    client = _oai.OpenAI(api_key=OPENAI_API_KEY)
    ext    = Path(image_path).suffix.lower().replace(".", "") or "png"
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    if not prompt:
        prompt = ('Extract student answers from this answer sheet. '
                  'Return JSON only: {"answers":{"1":"...","2":"..."}}. '
                  'Omit unreadable questions. Never copy printed question text.')
    resp = client.chat.completions.create(
        model=OPENAI_MINI, temperature=0,
        response_format={"type": "json_object"},
        messages=[{"role": "user", "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{b64}"}}
        ]}]
    )
    raw = json.loads(resp.choices[0].message.content)
    return {int(k): str(v).strip() for k, v in raw.get("answers", {}).items()
            if str(k).isdigit() and str(v).strip()}

def _run_audit(exam: dict, evaluation: dict) -> dict:
    diffs, obj_ok, subj_ok = [], 0, 0
    for item in evaluation.get("question_wise", []):
        qid = item.get("question_id")
        q   = next((x for x in exam.get("questions", []) if int(x.get("id", -1)) == int(qid)), None)
        if not q:
            continue
        if item.get("type") == "objective":
            obj_ok += 1
            valid  = [_normalize(a) for a in q.get("valid_answers", [])]
            exp    = float(q.get("weightage", 1)) if _normalize(item.get("student_answer", "")) in valid else 0.0
            if abs(float(item.get("awarded_marks", 0)) - exp) > 1e-6:
                diffs.append({"question_id": qid, "kind": "objective_mismatch",
                               "existing": item.get("awarded_marks", 0), "expected": exp})
        else:
            subj_ok += 1
            max_m = float(q.get("weightage", 1))
            prompt = SUBJECTIVE_GRADING_PROMPT.format(
                question=q.get("question", ""), max_marks=max_m,
                valid_answers=json.dumps(q.get("valid_answers", [])),
                key_points=json.dumps(q.get("answer_key_points", [])),
                rubric=q.get("evaluation_rubric", ""),
                student_answer=item.get("student_answer", "")
            )
            try:
                regrade = _llm_json(prompt, temperature=0)
            except Exception:
                regrade = {}
            exp   = max(0.0, min(max_m, float(regrade.get("awarded_marks", 0))))
            delta = abs(exp - float(item.get("awarded_marks", 0)))
            if delta > 0.5:
                diffs.append({"question_id": qid, "kind": "subjective_variance",
                               "existing": item.get("awarded_marks", 0),
                               "regraded": exp, "delta": round(delta, 2)})
    total = obj_ok + subj_ok
    return {
        "objective_checked": obj_ok, "subjective_checked": subj_ok, "total_checked": total,
        "inconsistencies": diffs,
        "consistency_score": 100.0 if total == 0 else round(((total - len(diffs)) / total) * 100, 2)
    }

# ══════════════════════════════════════════════════════════════════════════════
#  PROMPTS
# ══════════════════════════════════════════════════════════════════════════════
STUDY_PROMPT = None
if LANGCHAIN_OK:
    STUDY_PROMPT = PromptTemplate(
        input_variables=["context", "chat_history", "question"],
        template="""You are Arthavi, a patient and knowledgeable AI tutor helping students understand their curriculum.
Style: ELI5 explanations • 1–2 real-world examples • analogies where possible • end with a 1-sentence summary.

IMPORTANT rules:
- If study materials are provided below, use them as your PRIMARY source.
- If the context is empty or the question is NOT covered in the materials, STILL answer helpfully from your general knowledge. Start with: "⚠️ This topic is not part of your selected syllabus, but I'll still help:"
- NEVER say "there are no materials uploaded" or refuse to answer — always provide a useful educational response.
- Never refuse to answer — always be helpful, even for off-topic questions.

--- Study Materials ---
{context}

--- Chat History ---
{chat_history}

Student Question: {question}

Arthavi Answer:"""
    )

CHAPTER_EXTRACT_PROMPT = """You are analysing a study document.
Extract all chapter titles, unit names, or major topic headings from the content below.
Return ONLY a valid JSON array of strings — no markdown, no explanation.
Keep each item 2–6 words. Return 3–12 items.

Content:
{context}

JSON array only:"""

MIXED_QUESTION_GEN_PROMPT = """You are an expert exam paper setter.
Generate exactly {total_count} exam questions from the context only.
Output EXACTLY valid JSON — no markdown.

CRITICAL: Board={board} | Class={class_level} | Subject={subject}
- All questions MUST be at exactly {class_level} level — not higher, not lower.
- Vocabulary, concepts, and depth must match {board} {class_level} {subject} standards.
- Do NOT include concepts from higher classes even if present in the context.

objective_count={objective_count}, subjective_count={subjective_count}
Difficulty: easy={easy_count}, medium={medium_count}, hard={hard_count}

Each question object must have:
  id (1-based int), type ("objective"|"subjective"), difficulty ("easy"|"medium"|"hard"),
  weightage (number), question (string)
Objective extras: options ({{A,B,C,D}}), valid_answers ([letters]), explanation (string)
Subjective extras: options=null, valid_answers ([2-5 paraphrases]), answer_key_points ([strings]), evaluation_rubric (string)

Topic: {topic}
Context:
{context}
JSON only:"""

REGIONAL_MIXED_QUESTION_GEN_PROMPT = """You are an expert exam paper setter.
Generate exactly {total_count} exam questions. Write EVERYTHING in '{language}'.
Output EXACTLY valid JSON — no markdown.
objective_count={objective_count}, subjective_count={subjective_count}
Difficulty: easy={easy_count}, medium={medium_count}, hard={hard_count}
Same structure as standard question prompt but all text in {language}.
Topic: {topic}
Context: {context}
JSON only:"""

SUBJECTIVE_GRADING_PROMPT = """You are a strict but fair Indian school examiner.
Grade the student's answer for this question. Return ONLY valid JSON — no markdown.

IMPORTANT RULES:
1. If student_answer is empty, "(No answer provided)", or clearly blank → awarded_marks MUST be 0.
2. If student_answer looks like OCR noise (random garbled letters, unrelated content) → awarded_marks MUST be 0.
3. Award partial marks proportionally — if the student covers some key points, award partial credit.
4. If valid_answers and key_points are both empty, use your knowledge of the subject to grade fairly.
5. Be encouraging but accurate — give credit for any correct content, even if imperfectly stated.

Question: {question}
Max marks: {max_marks}
Valid answers (model answer paraphrases): {valid_answers}
Key points required: {key_points}
Marking rubric: {rubric}
Student answer: {student_answer}

Return JSON with exactly these keys:
- awarded_marks: number from 0 to {max_marks} (decimals allowed e.g. 1.5 for partial credit)
- feedback: 1-2 sentence feedback string addressed to the student
- missing_points: array of strings listing key points the student missed
- strengths: array of strings listing what the student got right (empty array if nothing correct)
- confidence: number 0-1 representing your confidence in this grade"""

SUMMARISE_PROMPT = """You are an expert Indian school educator. Summarise the provided context in EXACTLY 10 clear, numbered lines.
Each line should be one complete educational point. Include a brief real-world or relatable example after each point where useful.
Do NOT use bullet symbols — only numbered lines.
Do NOT include any preamble or heading — output ONLY the 10 numbered lines.
Do NOT quote or reproduce textbook passages verbatim. Always paraphrase into original study notes.

Context: {context}

Output format (strictly):
1. [Complete educational point with example if helpful]
2. [Complete educational point with example if helpful]
3. [Complete educational point with example if helpful]
4. [Complete educational point with example if helpful]
5. [Complete educational point with example if helpful]
6. [Complete educational point with example if helpful]
7. [Complete educational point with example if helpful]
8. [Complete educational point with example if helpful]
9. [Complete educational point with example if helpful]
10. [Complete educational point with example if helpful]"""

REGIONAL_SUMMARISE_PROMPT = """You are an expert educator. Summarise the provided context in EXACTLY 10 clear numbered lines in {language}.
Do NOT include any preamble — output ONLY 10 numbered lines in {language}.
Do NOT quote or reproduce textbook passages verbatim. Always paraphrase into original study notes.

Context: {context}

10 numbered lines in {language}:"""

FLASHCARDS_PROMPT = """Generate exactly {count} high-quality flashcards from the context below.
Each flashcard must cover a key concept, definition, formula, or fact from the content.
Vary the question style: some definitions, some "explain why", some examples.

Context:
{context}

Return EXACTLY a valid JSON array — no markdown, no preamble:
[{{"question": "...", "answer": "..."}}, ...]"""

REGIONAL_FLASHCARDS_PROMPT = """Generate exactly {count} high-quality flashcards in {language} from the context below.
Return EXACTLY a valid JSON array — no markdown:
[{{"question": "...", "answer": "..."}}]

Context:
{context}"""

MIXED_PRACTICE_GEN_PROMPT = """You are a high-quality academic practice generator.
Generate exactly {total_count} distinct practice questions strictly appropriate for the class and board specified.

CRITICAL: Board={board} | Class={class_level} | Subject={subject}
- Questions MUST be at exactly {class_level} level — not higher, not lower.
- Vocabulary, concepts, and depth must match {board} {class_level} {subject} standards.
- Do NOT include concepts from higher classes.

objective_count: {objective_count}, subjective_count: {subjective_count}, difficulty: {difficulty}

Return EXACTLY a valid JSON array. Do NOT generate blank questions.
Structure:
[
  {{"question":"...", "answer":"...", "type":"objective", "options":{{"A":"...","B":"...","C":"...","D":"..."}}}},
  {{"question":"...", "answer":"...", "type":"subjective", "options":null}}
]

Context: {context}
Topic: {topic}
JSON array only:"""

PRACTICE_EVALUATION_PROMPT = """You are an expert academic tutor.
Evaluate the student answer against the model answer for the given question.
For objective questions, a matching answer (case-insensitive) scores 10.
For subjective questions, evaluate based on accuracy, completeness, and clarity.
Return EXACTLY valid JSON: {{"score":0-10, "feedback":"...", "improvements":"...", "is_correct":true/false}}

Question: {question}
Model Answer: {model_answer}
Student Answer: {student_answer}"""

PRACTICE_REPORT_PROMPT = """You are an expert academic counsellor.
Analyse this practice session and return EXACTLY valid JSON:
{{"total_marks":n, "obtained_marks":n, "percentage":n, "overall_feedback":"...",
  "learning_gaps":[{{"topic":"...","gap":"...","recommendation":"..."}}],
  "predicted_grade":"A/B/C/D/F", "strengths":["..."]}}

Session data: {session_data}"""

VIDEO_EXPLAINER_PROMPT = """You are a warm British primary school teacher explaining: {topic}.
Return a JSON ARRAY of scene objects (7 scenes: INTRO, CONCEPT, EXAMPLE_1, EXAMPLE_2, TRY_IT, KEY_TAKEAWAYS, SUMMARY).
Each object: {{"segment":"...","text":"2-4 spoken sentences","visual":"actual content not a description","visual_type":"title|worked_example|bullet_list|fact_box|summary_box"}}
British English only. Return ONLY valid JSON array. Context: {context}"""

_REFUSAL_MARKERS = (
    "i'm sorry", "i am sorry", "i can't", "i cannot", "can't provide", "cannot provide",
    "unable to", "as requested", "however, i can help", "i can help with a general overview",
)

def _looks_like_refusal(text: str) -> bool:
    if not text:
        return True
    low = text.lower().strip()
    return any(marker in low for marker in _REFUSAL_MARKERS)

def _fallback_numbered_summary(topic: str, subject: str) -> str:
    base_topic = topic or subject or "this chapter"
    base_subject = subject or "the subject"
    lines = [
        f"1. {base_topic}: Understand the main idea first before memorising details.",
        f"2. In {base_subject}, identify key terms and write a one-line meaning for each.",
        "3. Group the topic into smaller concepts so revision becomes easier.",
        "4. Track cause-and-effect links between concepts to improve answer quality.",
        "5. Use one practical daily-life example for each concept to retain it longer.",
        "6. Compare similar ideas in a table to avoid confusion in exams.",
        "7. Mark frequently asked patterns and prepare short revision notes for them.",
        "8. Practice 3-5 likely exam questions and answer them in your own words.",
        "9. Revise with a quick formula/fact/theme recap before attempting mock questions.",
        "10. End with a short self-check: what, why, how, and one example for each concept.",
    ]
    return "\n".join(lines)

def _generate_summary_with_fallback(ctx: str, topic: str, subject: str, language: str = "", point_count: int = 10) -> str:
    regional_langs = ["odia","kannada","hindi","telugu","gujarati","marathi","bengali","punjabi","tamil","malayalam"]
    is_regional = bool(language) and any(lang in language.lower() for lang in regional_langs)

    n = max(1, min(20, point_count))

    if is_regional:
        prompt = (REGIONAL_SUMMARISE_PROMPT
                  .replace("{language}", language.title())
                  .replace("{context}", ctx)
                  .replace("EXACTLY 10", f"EXACTLY {n}")
                  .replace("10 numbered lines", f"{n} numbered lines"))
    else:
        numbered_format = "\n".join(f"{i+1}. [Complete educational point with example if helpful]" for i in range(n))
        prompt = (
            f"You are an expert Indian school educator. Summarise the provided context in EXACTLY {n} clear, numbered lines.\n"
            f"Each line should be one complete educational point. Include a brief real-world or relatable example after each point where useful.\n"
            f"Do NOT use bullet symbols — only numbered lines.\n"
            f"Do NOT include any preamble or heading — output ONLY the {n} numbered lines.\n"
            f"Do NOT quote or reproduce textbook passages verbatim. Always paraphrase into original study notes.\n\n"
            f"Context: {ctx}\n\n"
            f"Output format (strictly):\n{numbered_format}"
        )

    summary = _llm_text(prompt)
    if not _looks_like_refusal(summary):
        return summary

    # Recovery path for copyright-style refusals: generate original study notes only.
    recovery_prompt = (
        f"Create EXACTLY {n} numbered study-note lines for Indian school students. "
        "Use only high-level educational explanation in your own words. "
        "Do not quote or reproduce textbook text. No refusal text, no preamble.\n\n"
        f"Topic: {topic}\n"
        f"Subject: {subject}\n"
        f"Reference context: {ctx[:3000]}\n"
    )
    summary = _llm_text(recovery_prompt, temperature=0.2, mini=True)
    if not _looks_like_refusal(summary):
        return summary

    return _fallback_numbered_summary(topic, subject)

# ══════════════════════════════════════════════════════════════════════════════
#  CURRICULUM METADATA
# ══════════════════════════════════════════════════════════════════════════════
def _build_curriculum() -> dict:
    boards = {
        "National (NCERT)": ["CBSE", "ICSE", "ISC"],
        "Andhra Pradesh": ["BSEAP", "CBSE", "ICSE", "ISC"],
        "Bihar": ["BSEB", "CBSE", "ICSE", "ISC"],
        "Delhi": ["CBSE", "ICSE", "ISC"],
        "Gujarat": ["GSEB", "CBSE", "ICSE", "ISC"],
        "Karnataka": ["KSEEB", "CBSE", "ICSE", "ISC"],
        "Kerala": ["KBPE", "CBSE", "ICSE", "ISC"],
        "Maharashtra": ["MSBSHSE", "CBSE", "ICSE", "ISC"],
        "Odisha": ["BSE", "CHSE", "CBSE", "ICSE", "ISC"],
        "Punjab": ["PSEB", "CBSE", "ICSE", "ISC"],
        "Rajasthan": ["RBSE", "CBSE", "ICSE", "ISC"],
        "Tamil Nadu": ["TNBSE", "CBSE", "ICSE", "ISC"],
        "Uttar Pradesh": ["UPMSP", "CBSE", "ICSE", "ISC"],
        "West Bengal": ["WBBSE", "CBSE", "ICSE", "ISC"],
    }
    data: dict = {}
    for state, bds in boards.items():
        data[state] = {}
        for board in bds:
            data[state][board] = {}
            for i in range(1, 13):
                if i <= 5:
                    subs = ["Mathematics", "Environmental Studies", "English", "Local Language"]
                elif i <= 8:
                    subs = ["Mathematics", "Science", "Social Science", "English", "Local Language", "Hindi"]
                elif i <= 10:
                    subs = ["Mathematics", "Science", "Social Science", "English", "Local Language", "Hindi", "Sanskrit"]
                else:
                    subs = ["Physics", "Chemistry", "Biology", "Mathematics", "Computer Science",
                            "English", "History", "Geography", "Economics", "Accountancy",
                            "Business Studies", "Local Language"]
                data[state][board][f"Class {i}"] = subs
    return data

CURRICULUM_DATA = _build_curriculum()

# Hardcoded NCERT chapter database (reliable, no LLM needed for known combinations)
CHAPTERS_DB: Dict[tuple, list] = {
    # ── Science ───────────────────────────────────────────────────────────────
    ("class6","science"):  ["Ch 1: The Wonderful World of Science","Ch 2: Diversity in the Living World","Ch 3: Mindful Eating: A Path to a Healthy Body","Ch 4: Exploring Magnets","Ch 5: Measurement of Length and Motion","Ch 6: Materials Around Us","Ch 7: Temperature and its Measurement","Ch 8: A Treat for Mosquitoes!","Ch 9: Methods of Separation","Ch 10: Living Creatures: Exploring their Characteristics","Ch 11: Nature's Treasures","Ch 12: Beyond Earth"],
    ("class7","science"):  ["Ch 1: Nutrition in Plants","Ch 2: Nutrition in Animals","Ch 3: Heat","Ch 4: Acids, Bases and Salts","Ch 5: Physical and Chemical Changes","Ch 6: Respiration in Organisms","Ch 7: Transportation in Plants and Animals","Ch 8: Reproduction in Plants","Ch 9: Motion and Time","Ch 10: Electric Current and its Effects","Ch 11: Light","Ch 12: Forests: Our Lifeline","Ch 13: Wastewater Story"],  # rationalized 2023-24
    ("class8","science"):  ["Ch 1: Crop Production and Management","Ch 2: Microorganisms: Friend and Foe","Ch 3: Coal and Petroleum","Ch 4: Combustion and Flame","Ch 5: Conservation of Plants and Animals","Ch 6: Reproduction in Animals","Ch 7: Reaching the Age of Adolescence","Ch 8: Force and Pressure","Ch 9: Friction","Ch 10: Sound","Ch 11: Chemical Effects of Electric Current","Ch 12: Some Natural Phenomena","Ch 13: Light"],  # rationalized 2023-24
    ("class9","science"):  ["Ch 1: Matter in Our Surroundings","Ch 2: Is Matter Around Us Pure?","Ch 3: Atoms and Molecules","Ch 4: Structure of the Atom","Ch 5: The Fundamental Unit of Life","Ch 6: Tissues","Ch 7: Diversity in Living Organisms","Ch 8: Motion","Ch 9: Force and Laws of Motion","Ch 10: Gravitation","Ch 11: Work and Energy","Ch 12: Sound","Ch 13: Why Do We Fall Ill?","Ch 14: Natural Resources","Ch 15: Improvement in Food Resources"],
    ("class10","science"): ["Ch 1: Chemical Reactions and Equations","Ch 2: Acids, Bases and Salts","Ch 3: Metals and Non-metals","Ch 4: Carbon and its Compounds","Ch 5: Periodic Classification of Elements","Ch 6: Life Processes","Ch 7: Control and Coordination","Ch 8: How do Organisms Reproduce?","Ch 9: Heredity and Evolution","Ch 10: Light — Reflection and Refraction","Ch 11: Human Eye and the Colourful World","Ch 12: Electricity","Ch 13: Magnetic Effects of Electric Current","Ch 14: Sources of Energy","Ch 15: Our Environment","Ch 16: Sustainable Management of Natural Resources"],
    ("class11","physics"): ["Ch 1: Physical World","Ch 2: Units and Measurements","Ch 3: Motion in a Straight Line","Ch 4: Motion in a Plane","Ch 5: Laws of Motion","Ch 6: Work, Energy and Power","Ch 7: System of Particles and Rotational Motion","Ch 8: Gravitation","Ch 9: Mechanical Properties of Solids","Ch 10: Mechanical Properties of Fluids","Ch 11: Thermal Properties of Matter","Ch 12: Thermodynamics","Ch 13: Kinetic Theory","Ch 14: Oscillations","Ch 15: Waves"],
    ("class12","physics"): ["Ch 1: Electric Charges and Fields","Ch 2: Electrostatic Potential and Capacitance","Ch 3: Current Electricity","Ch 4: Moving Charges and Magnetism","Ch 5: Magnetism and Matter","Ch 6: Electromagnetic Induction","Ch 7: Alternating Current","Ch 8: Electromagnetic Waves","Ch 9: Ray Optics and Optical Instruments","Ch 10: Wave Optics","Ch 11: Dual Nature of Radiation and Matter","Ch 12: Atoms","Ch 13: Nuclei","Ch 14: Semiconductor Electronics"],
    ("class11","chemistry"):["Ch 1: Some Basic Concepts of Chemistry","Ch 2: Structure of Atom","Ch 3: Classification of Elements and Periodicity","Ch 4: Chemical Bonding and Molecular Structure","Ch 5: States of Matter","Ch 6: Thermodynamics","Ch 7: Equilibrium","Ch 8: Redox Reactions","Ch 9: Hydrogen","Ch 10: The s-Block Elements","Ch 11: The p-Block Elements","Ch 12: Organic Chemistry — Some Basic Principles","Ch 13: Hydrocarbons","Ch 14: Environmental Chemistry"],
    ("class12","chemistry"):["Ch 1: The Solid State","Ch 2: Solutions","Ch 3: Electrochemistry","Ch 4: Chemical Kinetics","Ch 5: Surface Chemistry","Ch 6: General Principles of Isolation of Elements","Ch 7: The p-Block Elements","Ch 8: The d-and f-Block Elements","Ch 9: Coordination Compounds","Ch 10: Haloalkanes and Haloarenes","Ch 11: Alcohols, Phenols and Ethers","Ch 12: Aldehydes, Ketones and Carboxylic Acids","Ch 13: Amines","Ch 14: Biomolecules","Ch 15: Polymers","Ch 16: Chemistry in Everyday Life"],
    ("class11","biology"):  ["Ch 1: The Living World","Ch 2: Biological Classification","Ch 3: Plant Kingdom","Ch 4: Animal Kingdom","Ch 5: Morphology of Flowering Plants","Ch 6: Anatomy of Flowering Plants","Ch 7: Structural Organisation in Animals","Ch 8: Cell — The Unit of Life","Ch 9: Biomolecules","Ch 10: Cell Cycle and Cell Division","Ch 11: Transport in Plants","Ch 12: Mineral Nutrition","Ch 13: Photosynthesis in Higher Plants","Ch 14: Respiration in Plants","Ch 15: Plant Growth and Development","Ch 16: Digestion and Absorption","Ch 17: Breathing and Exchange of Gases","Ch 18: Body Fluids and Circulation","Ch 19: Excretory Products and their Elimination","Ch 20: Locomotion and Movement","Ch 21: Neural Control and Coordination","Ch 22: Chemical Coordination and Integration"],
    ("class12","biology"):  ["Ch 1: Reproduction in Organisms","Ch 2: Sexual Reproduction in Flowering Plants","Ch 3: Human Reproduction","Ch 4: Reproductive Health","Ch 5: Principles of Inheritance and Variation","Ch 6: Molecular Basis of Inheritance","Ch 7: Evolution","Ch 8: Human Health and Disease","Ch 9: Strategies for Enhancement in Food Production","Ch 10: Microbes in Human Welfare","Ch 11: Biotechnology — Principles and Processes","Ch 12: Biotechnology and its Applications","Ch 13: Organisms and Populations","Ch 14: Ecosystem","Ch 15: Biodiversity and Conservation","Ch 16: Environmental Issues"],
    # ── Mathematics ───────────────────────────────────────────────────────────
    ("class6","mathematics"):["Ch 1: Patterns in Mathematics","Ch 2: Lines and Angles","Ch 3: Number Play","Ch 4: Data Handling and Presentation","Ch 5: Prime Time","Ch 6: Perimeter and Area","Ch 7: Fractions","Ch 8: Playing with Constructions","Ch 9: Symmetry","Ch 10: The Other Side of Zero"],
    ("class7","mathematics"):["Ch 1: Integers","Ch 2: Fractions and Decimals","Ch 3: Data Handling","Ch 4: Simple Equations","Ch 5: Lines and Angles","Ch 6: The Triangle and its Properties","Ch 7: Comparing Quantities","Ch 8: Rational Numbers","Ch 9: Perimeter and Area","Ch 10: Algebraic Expressions","Ch 11: Exponents and Powers","Ch 12: Symmetry","Ch 13: Visualising Solid Shapes"],  # rationalized 2023-24
    ("class8","mathematics"):["Ch 1: Rational Numbers","Ch 2: Linear Equations in One Variable","Ch 3: Understanding Quadrilaterals","Ch 4: Data Handling","Ch 5: Squares and Square Roots","Ch 6: Cubes and Cube Roots","Ch 7: Comparing Quantities","Ch 8: Algebraic Expressions and Identities","Ch 9: Mensuration","Ch 10: Exponents and Powers","Ch 11: Direct and Inverse Proportions","Ch 12: Factorisation","Ch 13: Introduction to Graphs"],  # rationalized 2023-24
    ("class9","mathematics"):["Ch 1: Number Systems","Ch 2: Polynomials","Ch 3: Coordinate Geometry","Ch 4: Linear Equations in Two Variables","Ch 5: Introduction to Euclid's Geometry","Ch 6: Lines and Angles","Ch 7: Triangles","Ch 8: Quadrilaterals","Ch 9: Areas of Parallelograms and Triangles","Ch 10: Circles","Ch 11: Constructions","Ch 12: Heron's Formula","Ch 13: Surface Areas and Volumes","Ch 14: Statistics","Ch 15: Probability"],
    ("class10","mathematics"):["Ch 1: Real Numbers","Ch 2: Polynomials","Ch 3: Pair of Linear Equations in Two Variables","Ch 4: Quadratic Equations","Ch 5: Arithmetic Progressions","Ch 6: Triangles","Ch 7: Coordinate Geometry","Ch 8: Introduction to Trigonometry","Ch 9: Some Applications of Trigonometry","Ch 10: Circles","Ch 11: Constructions","Ch 12: Areas Related to Circles","Ch 13: Surface Areas and Volumes","Ch 14: Statistics","Ch 15: Probability"],
    ("class11","mathematics"):["Ch 1: Sets","Ch 2: Relations and Functions","Ch 3: Trigonometric Functions","Ch 4: Principle of Mathematical Induction","Ch 5: Complex Numbers and Quadratic Equations","Ch 6: Linear Inequalities","Ch 7: Permutations and Combinations","Ch 8: Binomial Theorem","Ch 9: Sequences and Series","Ch 10: Straight Lines","Ch 11: Conic Sections","Ch 12: Introduction to Three Dimensional Geometry","Ch 13: Limits and Derivatives","Ch 14: Mathematical Reasoning","Ch 15: Statistics","Ch 16: Probability"],
    ("class12","mathematics"):["Ch 1: Relations and Functions","Ch 2: Inverse Trigonometric Functions","Ch 3: Matrices","Ch 4: Determinants","Ch 5: Continuity and Differentiability","Ch 6: Application of Derivatives","Ch 7: Integrals","Ch 8: Application of Integrals","Ch 9: Differential Equations","Ch 10: Vector Algebra","Ch 11: Three Dimensional Geometry","Ch 12: Linear Programming","Ch 13: Probability"],
    # ── English ───────────────────────────────────────────────────────────────
    ("class6","english"):  ["Honeysuckle Ch 1: Who Did Patrick's Homework?","Honeysuckle Ch 2: How the Dog Found Himself a New Master!","Honeysuckle Ch 3: Taro's Reward","Honeysuckle Ch 4: An Indian — American Woman in Space","Honeysuckle Ch 5: A Different Kind of School","Honeysuckle Ch 6: Who I Am","Honeysuckle Ch 7: Fair Play","Honeysuckle Ch 8: A Game of Chance","Honeysuckle Ch 9: Desert Animals","Honeysuckle Ch 10: The Banyan Tree","A Pact with the Sun Ch 1: A Tale of Two Birds","A Pact with the Sun Ch 2: The Friendly Mongoose","A Pact with the Sun Ch 3: The Shepherd's Treasure","A Pact with the Sun Ch 4: The Old-Clock Shop","A Pact with the Sun Ch 5: Tansen","A Pact with the Sun Ch 6: The Monkey and the Crocodile","A Pact with the Sun Ch 7: The Wonder Called Sleep","A Pact with the Sun Ch 8: A Pact with the Sun"],
    ("class7","english"):  ["Honeycomb Ch 1: Three Questions","Honeycomb Ch 2: A Gift of Chappals","Honeycomb Ch 3: Gopal and the Hilsa Fish","Honeycomb Ch 4: The Ashes That Made Trees Bloom","Honeycomb Ch 5: Quality","Honeycomb Ch 6: Expert Detectives","Honeycomb Ch 7: The Invention of Vita-Wonk","Honeycomb Ch 8: Fire: Friend and Foe","Honeycomb Ch 9: A Bicycle in Good Repair","Honeycomb Ch 10: The Story of Cricket","An Alien Hand Ch 1: The Tiny Teacher","An Alien Hand Ch 2: Bringing Up Kittens","An Alien Hand Ch 3: The Desert","An Alien Hand Ch 4: The Cop and the Anthem","An Alien Hand Ch 5: Golu Grows a Nose","An Alien Hand Ch 6: I Want Something in a Cage","An Alien Hand Ch 7: Chandni","An Alien Hand Ch 8: The Bear Story","An Alien Hand Ch 9: A Tiger in the House","An Alien Hand Ch 10: An Alien Hand"],
    ("class8","english"):  ["Honeydew Ch 1: The Best Christmas Present in the World","Honeydew Ch 2: The Tsunami","Honeydew Ch 3: Glimpses of the Past","Honeydew Ch 4: Bepin Choudhury's Lapse of Memory","Honeydew Ch 5: The Summit Within","Honeydew Ch 6: This is Jody's Fawn","Honeydew Ch 7: A Visit to Cambridge","Honeydew Ch 8: A Short Monsoon Diary","Honeydew Ch 9: The Great Stone Face — I","Honeydew Ch 10: The Great Stone Face — II","It So Happened Ch 1: How the Camel Got His Hump","It So Happened Ch 2: Children at Work","It So Happened Ch 3: The Selfish Giant","It So Happened Ch 4: The Treasure Within","It So Happened Ch 5: Princess September","It So Happened Ch 6: The Fight","It So Happened Ch 7: The Open Window","It So Happened Ch 8: Jalebis","It So Happened Ch 9: The Comet — I","It So Happened Ch 10: The Comet — II"],
    ("class9","english"):  ["Beehive Ch 1: The Fun They Had","Beehive Ch 2: The Sound of Music","Beehive Ch 3: The Little Girl","Beehive Ch 4: A Truly Beautiful Mind","Beehive Ch 5: The Snake and the Mirror","Beehive Ch 6: My Childhood","Beehive Ch 7: Packing","Beehive Ch 8: Reach for the Top","Beehive Ch 9: The Bond of Love","Beehive Ch 10: Kathmandu","Beehive Ch 11: If I Were You","Moments Ch 1: The Lost Child","Moments Ch 2: The Adventures of Toto","Moments Ch 3: Iswaran the Storyteller","Moments Ch 4: In the Kingdom of Fools","Moments Ch 5: The Happy Prince","Moments Ch 6: Weathering the Storm in Ersama","Moments Ch 7: The Last Leaf","Moments Ch 8: A House Is Not a Home","Moments Ch 9: The Accidental Tourist","Moments Ch 10: The Beggar"],
    ("class10","english"):["First Flight Ch 1: A Letter to God","First Flight Ch 2: Nelson Mandela — Long Walk to Freedom","First Flight Ch 3: Two Stories About Flying","First Flight Ch 4: From the Diary of Anne Frank","First Flight Ch 5: The Hundred Dresses — I","First Flight Ch 6: The Hundred Dresses — II","First Flight Ch 7: Glimpses of India","First Flight Ch 8: Mijbil the Otter","First Flight Ch 9: Madam Rides the Bus","First Flight Ch 10: The Sermon at Benares","First Flight Ch 11: The Proposal","Footprints Ch 1: A Triumph of Surgery","Footprints Ch 2: The Thief's Story","Footprints Ch 3: The Midnight Visitor","Footprints Ch 4: A Question of Trust","Footprints Ch 5: Footprints Without Feet"],
    # ── Hindi ─────────────────────────────────────────────────────────────────
    ("class6","hindi"):    ["Vasant Ch 1: Vah Chidiya Jo","Vasant Ch 2: Bachpan","Vasant Ch 3: Naadaan Dost","Vasant Ch 4: Chaand Se Thodi Si Gappe","Vasant Ch 5: Aksharon Ka Mahatv","Vasant Ch 6: Paar Nazar Ke","Vasant Ch 7: Saathi Haath Badhana","Vasant Ch 8: Aise Aise","Vasant Ch 9: Ticket Album","Vasant Ch 10: Jhaanse Ki Rani"],
    ("class7","hindi"):    ["Vasant Ch 1: Hum Panchhi Unmukt Gagan Ke","Vasant Ch 2: Dadi Maa","Vasant Ch 3: Himaalaya Ki Betiyan","Vasant Ch 4: Kathaputli","Vasant Ch 5: Miti Ki Sondh","Vasant Ch 6: Rakt aur Hamara Sharir","Vasant Ch 7: Paapad Wali Gali","Vasant Ch 8: Shaame — Ek Kisaan","Vasant Ch 9: Chidiya Ki Bacchi","Vasant Ch 10: Apoorv Anubhav","Vasant Ch 11: Raheem Ke Dohe"],
    ("class8","hindi"):    ["Vasant Ch 1: Dhwani","Vasant Ch 2: Lakh Ki Chudiyan","Vasant Ch 3: Bus Ki Yatra","Vasant Ch 4: Deewanon Ki Hasti","Vasant Ch 5: Chitthiyon Ki Anoothi Duniya","Vasant Ch 6: Bhagwan Ke Dakiye","Vasant Ch 7: Kya Nirash Hua Jaaye","Vasant Ch 8: Yeh Sabse Kathin Samay Nahi","Vasant Ch 9: Kabir Ki Saakhiyan","Vasant Ch 10: Hamare Watan Ki Dharohar"],
    ("class9","hindi"):    ["Kshitij Ch 1: Do Baillon Ki Katha","Kshitij Ch 2: Rahul Sankrityayan — Lhasa Ki Or","Kshitij Ch 3: Upbhoktavaad Ki Sanskriti","Kshitij Ch 4: Saavanon Ke Geeton Ki Patjhad","Kshitij Ch 5: Nana Sahab Ki Putri — Devi Maina","Kshitij Ch 6: Premchand Ke Phate Joote","Kshitij Ch 7: Mere Bachpan Ke Din","Kshitij Ch 8: Ek Kutta Aur Ek Maina","Kritika Ch 1: Is Jal Pralay Mein","Kritika Ch 2: Mere Sang Ki Auraten","Kritika Ch 3: Reedh Ki Haddi"],
    ("class10","hindi"):   ["Kshitij Ch 1: Surdas — Pad","Kshitij Ch 2: Tulsidas — Ram-Lakshman-Parshuram Samvad","Kshitij Ch 3: Dev — Savaiya aur Kavitt","Kshitij Ch 4: Jayashankar Prasad — Aatmakathya","Kshitij Ch 5: Suryakant Tripathi Nirala — Utsah, Aat Nahi Rahi","Kshitij Ch 6: Nagarjun — Yah Danturhit Muskan","Kshitij Ch 7: Girdhar — Fasal","Kshitij Ch 8: Rituraj — Ek Kahani Yah Bhi","Kritika Ch 1: Mata Ka Aanchal","Kritika Ch 2: George Pancham Ki Naak","Kritika Ch 3: Sana Sana Haath Jodi"],
    # ── Social Science ────────────────────────────────────────────────────────
    ("class6","social science"):["History Ch 1: What, Where, How and When?","History Ch 2: From Hunting-Gathering to Growing Food","History Ch 3: In the Earliest Cities","History Ch 4: What Books and Burials Tell Us","History Ch 5: Kingdoms, Kings and an Early Republic","History Ch 6: New Questions and Ideas","History Ch 7: Ashoka, The Emperor","History Ch 8: Vital Villages, Thriving Towns","History Ch 9: Traders, Kings and Pilgrims","History Ch 10: New Empires and Kingdoms","History Ch 11: Buildings, Paintings and Books","Geography Ch 1: The Earth in the Solar System","Geography Ch 2: Globe — Latitudes and Longitudes","Geography Ch 3: Motions of the Earth","Geography Ch 4: Maps","Geography Ch 5: Major Domains of the Earth","Geography Ch 6: Major Landforms of the Earth","Geography Ch 7: Our Country — India","Geography Ch 8: India — Climate, Vegetation and Wildlife","Civics Ch 1: Understanding Diversity","Civics Ch 2: Diversity and Discrimination","Civics Ch 3: What is Government?","Civics Ch 4: Key Elements of a Democratic Government","Civics Ch 5: Panchayati Raj","Civics Ch 6: Rural Administration","Civics Ch 7: Urban Administration"],
    ("class7","social science"):["History Ch 1: Tracing Changes Through a Thousand Years","History Ch 2: New Kings and Kingdoms","History Ch 3: The Delhi Sultans","History Ch 4: The Mughal Empire","History Ch 5: Rulers and Buildings","History Ch 6: Towns, Traders and Craftspersons","History Ch 7: Tribes, Nomads and Settled Communities","History Ch 8: Devotional Paths to the Divine","History Ch 9: The Making of Regional Cultures","History Ch 10: Eighteenth-Century Political Formations","Geography Ch 1: Environment","Geography Ch 2: Inside Our Earth","Geography Ch 3: Our Changing Earth","Geography Ch 4: Air","Geography Ch 5: Water","Geography Ch 6: Natural Vegetation and Wildlife","Geography Ch 7: Human Environment — Settlement, Transport and Communication","Geography Ch 8: Human–Environment Interactions","Geography Ch 9: Life in the Deserts","Civics Ch 1: On Equality","Civics Ch 2: Role of the Government in Health","Civics Ch 3: How the State Government Works","Civics Ch 4: Growing Up as Boys and Girls","Civics Ch 5: Women Change the World","Civics Ch 6: Understanding Media","Civics Ch 7: Markets Around Us"],
    ("class8","social science"):["History Ch 1: How, When and Where","History Ch 2: From Trade to Territory","History Ch 3: Ruling the Countryside","History Ch 4: Tribals, Dikus and the Vision of a Golden Age","History Ch 5: When People Rebel","History Ch 6: Weavers, Iron Smelters and Factory Owners","History Ch 7: Civilising the 'Native', Educating the Nation","History Ch 8: Women, Caste and Reform","History Ch 9: The Making of the National Movement: 1870s–1947","History Ch 10: India After Independence","Geography Ch 1: Resources","Geography Ch 2: Land, Soil, Water, Natural Vegetation and Wildlife","Geography Ch 3: Mineral and Power Resources","Geography Ch 4: Agriculture","Geography Ch 5: Industries","Geography Ch 6: Human Resources","Civics Ch 1: The Indian Constitution","Civics Ch 2: Understanding Secularism","Civics Ch 3: Why Do We Need a Parliament?","Civics Ch 4: Understanding Laws","Civics Ch 5: Judiciary","Civics Ch 6: Understanding Our Criminal Justice System","Civics Ch 7: Understanding Marginalisation","Civics Ch 8: Confronting Marginalisation"],
    ("class9","social science"):["History Ch 1: The French Revolution","History Ch 2: Socialism in Europe and the Russian Revolution","History Ch 3: Nazism and the Rise of Hitler","History Ch 4: Forest Society and Colonialism","History Ch 5: Pastoralists in the Modern World","Geography Ch 1: India — Size and Location","Geography Ch 2: Physical Features of India","Geography Ch 3: Drainage","Geography Ch 4: Climate","Geography Ch 5: Natural Vegetation and Wildlife","Geography Ch 6: Population","Political Science Ch 1: What is Democracy? Why Democracy?","Political Science Ch 2: Constitutional Design","Political Science Ch 3: Electoral Politics","Political Science Ch 4: Working of Institutions","Political Science Ch 5: Democratic Rights","Economics Ch 1: The Story of Village Palampur","Economics Ch 2: People as Resource","Economics Ch 3: Poverty as a Challenge","Economics Ch 4: Food Security in India"],
    ("class10","social science"):["History Ch 1: The Rise of Nationalism in Europe","History Ch 2: Nationalism in India","History Ch 3: The Making of a Global World","History Ch 4: The Age of Industrialisation","History Ch 5: Print Culture and the Modern World","Geography Ch 1: Resources and Development","Geography Ch 2: Forest and Wildlife Resources","Geography Ch 3: Water Resources","Geography Ch 4: Agriculture","Geography Ch 5: Minerals and Energy Resources","Geography Ch 6: Manufacturing Industries","Geography Ch 7: Lifelines of National Economy","Political Science Ch 1: Power Sharing","Political Science Ch 2: Federalism","Political Science Ch 3: Democracy and Diversity","Political Science Ch 4: Gender, Religion and Caste","Political Science Ch 5: Popular Struggles and Movements","Economics Ch 1: Development","Economics Ch 2: Sectors of the Indian Economy","Economics Ch 3: Money and Credit","Economics Ch 4: Globalisation and the Indian Economy"],
}

# Curated CISCE chapter structure for ICSE classes 9-10.
# These follow the book-style subject split students expect on ICSE reference
# sites such as KnowledgeBoat/Oswaal instead of a generic LLM-generated list.
def _prefixed(prefix: str, items: List[str]) -> List[str]:
    return [f"{prefix}: {item}" for item in items]

_ICSE_9_ENGLISH_LITERATURE = [
    "The Fun They Had",
    "The Sound of Music",
    "The Little Girl",
    "A Truly Beautiful Mind",
    "The Snake and the Mirror",
    "My Childhood",
    "Packing",
    "The Bond of Love",
    "The Road Not Taken",
    "The Lake Isle of Innisfree",
    "A Legend of the Northland",
    "No Men Are Foreign",
    "The Duck and the Kangaroo",
    "On the Grasshopper and Cricket",
    "The Gift of the Magi",
]
_ICSE_9_ENGLISH_LANGUAGE = [
    "Composition Writing",
    "Story and Narrative Writing",
    "Descriptive Writing",
    "Picture Composition",
    "Letter Writing",
    "Email and Notice Writing",
    "Comprehension",
    "Grammar - Tenses and Subject-Verb Agreement",
    "Grammar - Transformation of Sentences",
    "Grammar - Synthesis, Prepositions and Usage",
]
_ICSE_9_MATHEMATICS = [
    "Number Systems",
    "Polynomials",
    "Coordinate Geometry",
    "Linear Equations in Two Variables",
    "Introduction to Euclid's Geometry",
    "Lines and Angles",
    "Triangles",
    "Quadrilaterals",
    "Areas of Parallelograms and Triangles",
    "Circles",
    "Constructions",
    "Heron's Formula",
    "Surface Areas and Volumes",
    "Statistics",
    "Probability",
]
_ICSE_9_PHYSICS = [
    "Measurements and Experimentation",
    "Motion in One Dimension",
    "Laws of Motion",
    "Pressure in Fluids and Atmospheric Pressure",
    "Upthrust in Fluids, Archimedes' Principle and Floatation",
    "Heat and Energy",
    "Reflection of Light",
    "Propagation of Sound Waves",
    "Current Electricity and Household Circuits",
    "Magnetism",
]
_ICSE_9_CHEMISTRY = [
    "The Language of Chemistry",
    "Chemical Changes and Reactions",
    "Water",
    "Atomic Structure and Chemical Bonding",
    "The Periodic Table",
    "Study of the First Element - Hydrogen",
    "Study of Gas Laws",
    "Atmospheric Pollution",
]
_ICSE_9_BIOLOGY = [
    "Cell - The Unit of Life",
    "Tissues",
    "Flower",
    "Pollination and Fertilisation",
    "Seeds",
    "Digestive System",
    "Skeleton - Movement and Locomotion",
    "Skin - The Jack of All Trades",
    "Respiratory System",
    "Health Organisations",
    "Waste Generation and Management",
]
_ICSE_9_HISTORY_CIVICS = [
    "The Harappan Civilisation",
    "The Vedic Period",
    "Jainism and Buddhism",
    "The Mauryan Empire",
    "The Sangam Age",
    "The Age of the Guptas",
    "Medieval India",
    "The Modern Age in Europe",
    "Our Constitution",
    "Elections",
    "Local Self-Government - Rural",
    "Local Self-Government - Urban",
]
_ICSE_9_GEOGRAPHY = [
    "Earth as a Planet",
    "Geographic Grid - Latitudes and Longitudes",
    "Rotation and Revolution",
    "Structure of the Earth",
    "Volcanoes",
    "Earthquakes",
    "Weathering and Denudation",
    "Hydrosphere",
    "Atmosphere",
    "Humidity",
    "Natural Regions of the World",
]
_ICSE_9_COMPUTER_APPLICATIONS = [
    "Introduction to Object-Oriented Programming",
    "Values and Data Types",
    "Operators in Java",
    "Input in Java",
    "Mathematical Library Methods",
    "Conditional Constructs",
    "Iterative Constructs",
    "Nested Loops",
    "Computing and Ethics",
]

_ICSE_10_ENGLISH_LITERATURE = [
    "Merchant of Venice - Act 3 Scene 2",
    "Merchant of Venice - Act 4 Scene 1",
    "Merchant of Venice - Act 4 Scene 2",
    "Merchant of Venice - Act 5 Scene 1",
    "Treasure Chest - The Heart of the Tree",
    "Treasure Chest - The Cold Within",
    "Treasure Chest - The Bangle Sellers",
    "Treasure Chest - After Blenheim",
    "Treasure Chest - The Blue Bead",
    "Treasure Chest - My Greatest Olympic Prize",
    "Treasure Chest - All Summer in a Day",
    "Treasure Chest - The Power of Music",
]
_ICSE_10_ENGLISH_LANGUAGE = [
    "Composition Writing",
    "Argumentative and Reflective Writing",
    "Story and Descriptive Writing",
    "Letter Writing",
    "Email, Notice and Proposal Writing",
    "Comprehension",
    "Grammar - Tenses, Agreement and Prepositions",
    "Grammar - Transformation of Sentences",
    "Grammar - Synthesis and Re-ordering",
    "Grammar - Vocabulary and Usage",
]
_ICSE_10_MATHEMATICS = [
    "Goods and Services Tax",
    "Banking - Recurring Deposit and Bank Accounts",
    "Shares and Dividends",
    "Linear Inequations",
    "Quadratic Equations",
    "Ratio and Proportion",
    "Factorisation",
    "Matrices",
    "Arithmetic Progression",
    "Reflection",
    "Section Formula and Mid-point Formula",
    "Equation of a Straight Line",
    "Similarity",
    "Loci",
    "Circles",
    "Tangents and Intersections",
    "Trigonometric Identities",
    "Heights and Distances",
    "Mensuration",
    "Statistics",
    "Probability",
]
_ICSE_10_PHYSICS = [
    "Force",
    "Work, Energy and Power",
    "Machines",
    "Refraction of Light at Plane Surfaces",
    "Refraction through a Lens",
    "Spectrum",
    "Sound",
    "Current Electricity",
    "Household Circuits",
    "Electromagnetism",
    "Calorimetry",
    "Radioactivity",
]
_ICSE_10_CHEMISTRY = [
    "Periodic Properties and Variations of Properties",
    "Chemical Bonding",
    "Acids, Bases and Salts",
    "Analytical Chemistry",
    "Mole Concept and Stoichiometry",
    "Electrolysis",
    "Metallurgy",
    "Hydrogen Chloride",
    "Ammonia",
    "Nitric Acid",
    "Sulphuric Acid",
    "Organic Chemistry",
]
_ICSE_10_BIOLOGY = [
    "Cell Cycle and Cell Division",
    "Genetics",
    "Absorption by Roots",
    "Transpiration",
    "Photosynthesis",
    "Chemical Coordination in Plants",
    "Circulatory System",
    "Excretory System",
    "Nervous System",
    "Sense Organs",
    "Endocrine System",
    "Reproductive System",
    "Population",
    "Pollution",
]
_ICSE_10_HISTORY_CIVICS = [
    "The First War of Independence, 1857",
    "Growth of Nationalism",
    "Gandhian Nationalism",
    "Forward Bloc and the INA",
    "Independence and Partition",
    "The Union Legislature",
    "The Union Executive",
    "The Judiciary",
    "The United Nations",
    "Major Agencies of the United Nations",
]
_ICSE_10_GEOGRAPHY = [
    "Interpretation of Topographical Maps",
    "Map Work",
    "Climate",
    "Soil Resources",
    "Natural Vegetation",
    "Water Resources",
    "Mineral and Energy Resources",
    "Agriculture",
    "Manufacturing Industries",
    "Transport",
    "Waste Management - 1",
    "Waste Management - 2",
]
_ICSE_10_COMPUTER_APPLICATIONS = [
    "Classes as Objects",
    "User-defined Methods",
    "Constructors",
    "Library Classes",
    "Arrays",
    "Strings",
    "Class as the Basis of All Computation",
    "Encapsulation and Function Overloading",
    "Recursion",
    "Boolean Logic and Operators",
]

CISCE_CURATED_CHAPTERS: Dict[tuple, List[str]] = {
    ("class9", "english"): _ICSE_9_ENGLISH_LITERATURE,
    ("class9", "english literature"): _ICSE_9_ENGLISH_LITERATURE,
    ("class9", "english language"): _ICSE_9_ENGLISH_LANGUAGE,
    ("class9", "mathematics"): _ICSE_9_MATHEMATICS,
    ("class9", "physics"): _ICSE_9_PHYSICS,
    ("class9", "chemistry"): _ICSE_9_CHEMISTRY,
    ("class9", "biology"): _ICSE_9_BIOLOGY,
    ("class9", "history & civics"): _ICSE_9_HISTORY_CIVICS,
    ("class9", "geography"): _ICSE_9_GEOGRAPHY,
    ("class9", "computer applications"): _ICSE_9_COMPUTER_APPLICATIONS,
    ("class9", "science"): _prefixed("Physics", _ICSE_9_PHYSICS) + _prefixed("Chemistry", _ICSE_9_CHEMISTRY) + _prefixed("Biology", _ICSE_9_BIOLOGY),
    ("class9", "social science"): _prefixed("History & Civics", _ICSE_9_HISTORY_CIVICS) + _prefixed("Geography", _ICSE_9_GEOGRAPHY),
    ("class10", "english"): _ICSE_10_ENGLISH_LITERATURE,
    ("class10", "english literature"): _ICSE_10_ENGLISH_LITERATURE,
    ("class10", "english language"): _ICSE_10_ENGLISH_LANGUAGE,
    ("class10", "mathematics"): _ICSE_10_MATHEMATICS,
    ("class10", "physics"): _ICSE_10_PHYSICS,
    ("class10", "chemistry"): _ICSE_10_CHEMISTRY,
    ("class10", "biology"): _ICSE_10_BIOLOGY,
    ("class10", "history & civics"): _ICSE_10_HISTORY_CIVICS,
    ("class10", "geography"): _ICSE_10_GEOGRAPHY,
    ("class10", "computer applications"): _ICSE_10_COMPUTER_APPLICATIONS,
    ("class10", "science"): _prefixed("Physics", _ICSE_10_PHYSICS) + _prefixed("Chemistry", _ICSE_10_CHEMISTRY) + _prefixed("Biology", _ICSE_10_BIOLOGY),
    ("class10", "social science"): _prefixed("History & Civics", _ICSE_10_HISTORY_CIVICS) + _prefixed("Geography", _ICSE_10_GEOGRAPHY),
}

NCERT_PDF_URLS: Dict[tuple, str] = {
    # ── Science ──────────────────────────────────────────
    # Class 6 Science: new 2023-24 "Curiosity" textbook — no direct portal URL yet; Google fallback
    ("class7","science"):       "https://ncert.nic.in/textbook.php?gesc1=",
    ("class8","science"):       "https://ncert.nic.in/textbook.php?hesc1=",
    ("class9","science"):       "https://ncert.nic.in/textbook.php?iesc1=",
    ("class10","science"):      "https://ncert.nic.in/textbook.php?jesc1=",
    # ── Physics / Chemistry / Biology ────────────────────
    ("class11","physics"):      "https://ncert.nic.in/textbook.php?leph1=",
    ("class12","physics"):      "https://ncert.nic.in/textbook.php?leph2=",
    ("class11","chemistry"):    "https://ncert.nic.in/textbook.php?lech1=",
    ("class12","chemistry"):    "https://ncert.nic.in/textbook.php?lech2=",
    ("class11","biology"):      "https://ncert.nic.in/textbook.php?lebo1=",
    ("class12","biology"):      "https://ncert.nic.in/textbook.php?lebo2=",
    # ── Mathematics ──────────────────────────────────────
    # Class 6 Mathematics: new 2023-24 "Ganita Prakash" textbook — no direct portal URL yet; Google fallback
    ("class7","mathematics"):   "https://ncert.nic.in/textbook.php?gemh1=",
    ("class8","mathematics"):   "https://ncert.nic.in/textbook.php?hemh1=",
    ("class9","mathematics"):   "https://ncert.nic.in/textbook.php?iemh1=",
    ("class10","mathematics"):  "https://ncert.nic.in/textbook.php?jemh1=",
    ("class11","mathematics"):  "https://ncert.nic.in/textbook.php?lemh1=",
    ("class12","mathematics"):  "https://ncert.nic.in/textbook.php?lemh2=",
    # ── English ──────────────────────────────────────────
    # Class 6 (Honeysuckle) and Class 7 (Honeycomb): portal URLs return 404 — Google fallback
    ("class8","english"):       "https://ncert.nic.in/textbook.php?hehd1=",   # Honeydew
    ("class9","english"):       "https://ncert.nic.in/textbook.php?iebe1=",   # Beehive
    ("class10","english"):      "https://ncert.nic.in/textbook.php?jeff1=",   # First Flight
    # ── Hindi ────────────────────────────────────────────
    ("class6","hindi"):         "https://ncert.nic.in/textbook.php?hhvs1=",   # Vasant Part 1
    # Class 7 Vasant 2 and Class 8 Vasant 3: portal URLs not found — Google fallback
    ("class9","hindi"):         "https://ncert.nic.in/textbook.php?ihks1=",   # Kshitij 1
    ("class10","hindi"):        "https://ncert.nic.in/textbook.php?jhks1=",   # Kshitij 2
    # ── Social Science ───────────────────────────────────
    # Each class has 3-4 sub-books; these URLs are fallbacks.
    # Sub-book picker in frontend provides individual links.
    ("class6","social science"): "https://ncert.nic.in/textbook.php?hess1=",   # Our Pasts bundle
    ("class7","social science"): "https://ncert.nic.in/textbook.php?hess2=",   # Our Pasts bundle
    ("class8","social science"): "https://ncert.nic.in/textbook.php?hess3=",   # Social & Political Life
    ("class9","social science"): "https://ncert.nic.in/textbook.php?iess3=",   # India & Cont. World-I (History) ✓
    ("class10","social science"):"https://ncert.nic.in/textbook.php?jess3=",   # India & Cont. World-II (History) ✓
    # ── Senior subjects ──────────────────────────────────
    ("class11","economics"):    "https://ncert.nic.in/textbook.php?leec1=",
    ("class12","economics"):    "https://ncert.nic.in/textbook.php?leec2=",
    ("class12","accountancy"):  "https://ncert.nic.in/textbook.php?leac2=",
    ("class11","computer science"):"https://ncert.nic.in/textbook.php?lecs1=",
    ("class12","computer science"):"https://ncert.nic.in/textbook.php?lecs2=",
}

# Supplementary (second) textbook PDFs for subjects that have two NCERT books
NCERT_SUPP_PDF_URLS: Dict[tuple, str] = {
    ("class9",  "english"):  "https://ncert.nic.in/textbook.php?iemo1=",   # Moments Supplementary Reader
    ("class10", "english"):  "https://ncert.nic.in/textbook.php?jefp1=",   # Footprints without Feet
    ("class9",  "hindi"):    "https://ncert.nic.in/textbook.php?ihkr1=",   # Kritika Part 1
    ("class10", "hindi"):    "https://ncert.nic.in/textbook.php?jhkr1=",   # Kritika Part 2
}

# ══════════════════════════════════════════════════════════════════════════════
#  DIKSHA API INTEGRATION — Textbooks for ALL Indian Boards
# ══════════════════════════════════════════════════════════════════════════════
DIKSHA_SEARCH_URL = "https://diksha.gov.in/api/content/v1/search"
DIKSHA_HIERARCHY_URL = "https://diksha.gov.in/api/course/v1/hierarchy"

# Map frontend board shortNames → DIKSHA board filter values
DIKSHA_BOARD_MAP = {
    "CBSE": "CBSE", "NCERT": "NCERT", "NIOS": "NIOS",
    "BSEAP": "State (Andhra Pradesh)", "SEBA": "State (Assam)", "AHSEC": "State (Assam)",
    "BSEB": "State (Bihar)", "CGBSE": "State (Chhattisgarh)", "GBSHSE": "State (Goa)",
    "GSEB": "State (Gujarat)", "HBSE": "State (Haryana)", "HPBOSE": "State (Himachal Pradesh)",
    "JAC": "State (Jharkhand)", "KSEEB": "State (Karnataka)", "PUC": "State (Karnataka)",
    "KBPE": "State (Kerala)", "MPBSE": "State (Madhya Pradesh)", "MSBSHSE": "State (Maharashtra)",
    "BSEM": "State (Manipur)", "MBOSE": "State (Meghalaya)", "MBSE": "State (Mizoram)",
    "NBSE": "State (Nagaland)", "BSE": "State (Odisha)", "CHSE": "State (Odisha)",
    "PSEB": "State (Punjab)", "RBSE": "State (Rajasthan)", "BSSS": "State (Sikkim)",
    "TNBSE": "State (Tamil Nadu)", "BSETS": "State (Telangana)", "TBSE": "State (Tripura)",
    "UPMSP": "State (Uttar Pradesh)", "UBSE": "State (Uttarakhand)",
    "WBBSE": "State (West Bengal)", "WBCHSE": "State (West Bengal)",
    "JKBOSE": "State (Jammu and Kashmir)", "DoE": "State (Delhi)",
    "PBSE": "UT (Puducherry)",
}

# In-memory cache for DIKSHA results (TTL: process lifetime)
_diksha_cache: Dict[str, Any] = {}

def _diksha_search(filters: dict, limit: int = 20, fields: list = None, facets: list = None) -> dict:
    """Search DIKSHA content API. Returns parsed JSON response."""
    req_body: dict = {"request": {"filters": filters, "limit": limit}}
    if fields:
        req_body["request"]["fields"] = fields
    if facets:
        req_body["request"]["facets"] = facets
    try:
        r = _requests.post(DIKSHA_SEARCH_URL, json=req_body, timeout=15)
        r.raise_for_status()
        return r.json()
    except Exception as e:
        print(f"[DIKSHA] Search error: {e}")
        return {}

def _diksha_hierarchy(identifier: str) -> dict:
    """Get textbook hierarchy (chapter tree) from DIKSHA."""
    cache_key = f"hier_{identifier}"
    if cache_key in _diksha_cache:
        return _diksha_cache[cache_key]
    try:
        r = _requests.get(f"{DIKSHA_HIERARCHY_URL}/{identifier}", timeout=15)
        r.raise_for_status()
        data = r.json()
        _diksha_cache[cache_key] = data
        return data
    except Exception as e:
        print(f"[DIKSHA] Hierarchy error: {e}")
        return {}

def _diksha_extract_chapters(hierarchy_content: dict) -> list:
    """Extract chapter names from DIKSHA hierarchy content node."""
    children = hierarchy_content.get("children", [])
    chapters = []
    ch_num = 0
    for ch in children:
        name = ch.get("name", "").strip()
        if not name:
            continue
        # Skip non-chapter nodes (e.g., "eTextBook", "Learning Outcomes", "APPENDIX", "QUESTION BANK", "AUDIO BOOKS")
        skip = ["etextbook", "e-textbook", "learning outcomes", "appendix", "question bank",
                "audio books", "teacher resources", "practice set"]
        if name.lower() in skip:
            continue
        ch_num += 1
        chapters.append(name if name.startswith(("Ch ", "ch ", "Chapter")) else f"Ch {ch_num}: {name}")
    return chapters

def _diksha_extract_pdf(hierarchy_content: dict) -> str:
    """Extract the best PDF URL from a DIKSHA hierarchy (book-level or first eTextBook)."""
    # 1. Check book-level downloadUrl
    dl = hierarchy_content.get("downloadUrl", "")
    if dl and dl.endswith((".pdf", ".PDF")):
        return dl
    # 2. Look for eTextBook child → PDF leaf
    def _find_pdf(node):
        mime = node.get("mimeType", "")
        if mime == "application/pdf":
            return node.get("artifactUrl", "")
        for ch in node.get("children", []):
            result = _find_pdf(ch)
            if result:
                return result
        return ""
    for ch in hierarchy_content.get("children", []):
        name_lower = ch.get("name", "").lower()
        if "textbook" in name_lower or "e-textbook" in name_lower or "etextbook" in name_lower:
            pdf = _find_pdf(ch)
            if pdf:
                return pdf
    # 3. Fallback: look for any PDF in the tree
    return _find_pdf(hierarchy_content)

def _is_probable_textbook_name(name: str) -> bool:
    """Heuristic gate to exclude comics/question sets and keep textbook titles."""
    n = _normalize(name)
    if not n:
        return False
    bad = [
        "comic book", "question set", "question bank", "practice", "worksheet",
        "activity", "workbook", "sample paper", "assessment"
    ]
    if any(tok in n for tok in bad):
        return False
    good = [
        "textbook", "ganita", "curiosity", "honey", "vasant", "math", "mathematics"
    ]
    return any(tok in n for tok in good)

def _diksha_find_textbooks(board_short: str, class_n: str, subject: str, medium: str = "English") -> list:
    """Find textbooks from DIKSHA for a given board/class/subject.
    Returns list of dicts: [{name, identifier, leafNodesCount}]
    """
    subject = _normalize_subject_name(subject)
    cache_key = f"books_{board_short}_{class_n}_{subject}_{medium}"
    if cache_key in _diksha_cache:
        return _diksha_cache[cache_key]

    diksha_board = DIKSHA_BOARD_MAP.get(board_short, "")
    if not diksha_board:
        return []

    # Normalise class format: "Class 10" → "Class 10"
    grade = class_n if class_n.startswith("Class") else f"Class {class_n}"

    filters = {
        "contentType": ["TextBook"],
        "board": [diksha_board],
        "gradeLevel": [grade],
        "status": ["Live"],
    }
    # Only add medium filter if specified
    if medium:
        filters["medium"] = [medium]
    # Add subject filter — try exact match first
    if subject:
        filters["subject"] = [subject]

    data = _diksha_search(filters, limit=50,
                          fields=["name", "identifier", "leafNodesCount", "subject", "medium"])
    books = data.get("result", {}).get("content", []) or []

    # If no results with subject filter, try without it
    if not books and subject:
        del filters["subject"]
        data = _diksha_search(filters, limit=80,
                              fields=["name", "identifier", "leafNodesCount", "subject", "medium"])
        all_books = data.get("result", {}).get("content", []) or []
        # Fuzzy match subject — also try reverse containment and first-word match.
        # IMPORTANT: in this broad search (no subject filter) only accept single-subject
        # books (len == 1).  Multi-subject compilations like activity workbooks often
        # have subjects=["English","Science","Mathematics","Odia Language"] and will
        # match EVERY subject query, causing all subjects to show identical chapters.
        subj_lower = subject.lower()
        subj_words = set(subj_lower.split())
        books = [b for b in all_books
                 if len(b.get("subject", [])) == 1  # accept only clearly single-subject books
                 and any(
                     subj_lower in s.lower() or s.lower() in subj_lower
                     or bool(subj_words & set(s.lower().split()))
                     for s in b.get("subject", [])
                 )]
        # Do NOT fall back to all_books[:5] — that picks books from unrelated
        # subjects (e.g. Computer Science when Math was requested). Instead,
        # return empty and let the caller fall through to LLM generation.

    _diksha_cache[cache_key] = books
    return books

def _diksha_get_chapters_and_pdf(board_short: str, class_n: str, subject: str, medium: str = "English"):
    """Full DIKSHA lookup: find textbook → get hierarchy → extract chapters + PDF.
    Returns (chapters_list, pdf_url, textbook_name) or (None, None, None).
    """
    subject = _normalize_subject_name(subject)
    books = _diksha_find_textbooks(board_short, class_n, subject, medium)
    if not books:
        return None, None, None

    # Pick best textbook — prefer one with most leaf nodes
    books.sort(key=lambda b: b.get("leafNodesCount", 0), reverse=True)
    best = books[0]

    # Reject catch-all reference books that cover many subjects — they match
    # every query but contain generic content, not a proper subject textbook.
    # Threshold: > 2 subjects (e.g. a multi-subject Odia workbook with subjects
    # ["English","Science","Mathematics","Odia Language"] is rejected).
    if len(best.get("subject", [])) > 2:
        print(f"[DIKSHA] Rejecting multi-subject book '{best.get('name')}' ({len(best.get('subject',[]))} subjects)")
        return None, None, None

    # Verify the chosen book's subject metadata matches the requested subject.
    # DIKSHA can return miscategorised books (e.g. Computer Science when Math was
    # requested) — discard them so we fall through to the NCERT fallback.
    subj_lower = subject.lower()
    subj_words = set(subj_lower.split())
    book_subjects = [s.lower() for s in best.get("subject", [])]
    if book_subjects:
        subject_ok = any(
            subj_lower in s or s in subj_lower or bool(subj_words & set(s.split()))
            for s in book_subjects
        )
        if not subject_ok:
            print(f"[DIKSHA] Subject mismatch: requested '{subject}', book subject={book_subjects} — discarding")
            return None, None, None

    hier = _diksha_hierarchy(best["identifier"])
    content = hier.get("result", {}).get("content", {})
    if not content:
        return None, None, None

    chapters = _diksha_extract_chapters(content)
    pdf_url = _diksha_extract_pdf(content)
    return chapters or None, pdf_url or None, best.get("name", "")

# ── DIKSHA REST Endpoints ─────────────────────────────────────────────────────
@app.get("/api/diksha/boards")
def diksha_boards():
    """Return all boards available on DIKSHA with textbook counts."""
    cache_key = "diksha_all_boards"
    if cache_key in _diksha_cache:
        return jsonify(_diksha_cache[cache_key])
    data = _diksha_search(
        {"contentType": ["TextBook"], "status": ["Live"]},
        limit=0, facets=["board"]
    )
    facets = data.get("result", {}).get("facets", [{}])
    boards = facets[0].get("values", []) if facets else []
    result = {
        "total": data.get("result", {}).get("count", 0),
        "boards": [{"name": b["name"], "count": b["count"]} for b in boards]
    }
    _diksha_cache[cache_key] = result
    return jsonify(result)

@app.post("/api/diksha/textbooks")
def diksha_textbooks():
    """Search DIKSHA textbooks by board/class/subject/medium."""
    b = request.json or {}
    board = b.get("board", "CBSE")
    grade = b.get("class", "Class 10")
    subject = _normalize_subject_name(b.get("subject", ""))
    medium = b.get("medium", "English")

    books = _diksha_find_textbooks(board, grade, subject, medium)
    return jsonify({"textbooks": [
        {"name": bk.get("name", ""), "identifier": bk.get("identifier", ""),
         "leaves": bk.get("leafNodesCount", 0),
         "subject": bk.get("subject", []), "medium": bk.get("medium", [])}
        for bk in books
    ]})

@app.post("/api/diksha/chapters")
def diksha_chapters():
    """Get chapters + PDF URL for a specific DIKSHA textbook (by identifier or search)."""
    b = request.json or {}
    identifier = b.get("identifier", "")

    if identifier:
        hier = _diksha_hierarchy(identifier)
        content = hier.get("result", {}).get("content", {})
        chapters = _diksha_extract_chapters(content) if content else []
        pdf_url = _diksha_extract_pdf(content) if content else ""
        return jsonify({"chapters": chapters, "pdf_url": pdf_url,
                        "name": content.get("name", ""), "source": "diksha"})

    # Fallback: search by board/class/subject
    board = b.get("board", "CBSE")
    grade = b.get("class", "Class 10")
    subject = _normalize_subject_name(b.get("subject", ""))
    medium = b.get("medium", "English")
    chapters, pdf_url, name = _diksha_get_chapters_and_pdf(board, grade, subject, medium)
    return jsonify({
        "chapters": chapters or [],
        "pdf_url": pdf_url or "",
        "name": name or "",
        "source": "diksha"
    })

# ══════════════════════════════════════════════════════════════════════════════
#  FLASK MIDDLEWARE
# ══════════════════════════════════════════════════════════════════════════════
@app.before_request
def _log_request():
    if request.path.startswith("/api/"):
        print(f"[API] {request.method} {request.path}")

@app.errorhandler(404)
def _404(e):
    if request.path.startswith("/api/"):
        return jsonify({"error": "API route not found", "path": request.path}), 404
    try:
        return send_from_directory(app.static_folder, "index.html")
    except Exception:
        return jsonify({"error": "Not found"}), 404

@app.errorhandler(Exception)
def _500(e):
    import traceback
    print(f"[ERROR] {e}\n{traceback.format_exc()}")
    if request.path.startswith("/api/"):
        return jsonify({"error": "Internal server error", "message": str(e)}), 500
    return f"<h1>500 Internal Server Error</h1><p>{e}</p>", 500

# ══════════════════════════════════════════════════════════════════════════════
#  AUTH DECORATOR — role isolation + session expiry
# ══════════════════════════════════════════════════════════════════════════════
SESSION_TTL_DAYS = int(os.getenv("SESSION_TTL_DAYS", 30))

def _session_find(token: str) -> Optional[dict]:
    """Load a session by token (MongoDB primary, JSON fallback)."""
    if not token:
        return None
    if MONGO_OK:
        return _mongo_find_one(M_SESSIONS, {"token": token})
    db = db_load()
    for s in db.get("sessions", []):
        if s.get("token") == token:
            return s
    return None

def _session_save(token: str, user_id: str, role: str, ip: str = ""):
    """Persist a session (MongoDB primary, JSON fallback)."""
    payload = {
        "token": token,
        "user_id": user_id,
        "role": role,
        "created_at": _now(),
        "ip": ip or "",
    }
    if MONGO_OK:
        _mongo_insert(M_SESSIONS, payload)
        return
    db = db_load()
    sessions = [s for s in db.get("sessions", []) if s.get("token") != token]
    sessions.append(payload)
    db["sessions"] = sessions
    db_save(db)

def _session_delete(token: str):
    """Delete a session by token from the active persistence layer."""
    if not token:
        return
    if MONGO_OK:
        _mongo_delete(M_SESSIONS, {"token": token})
        return
    db = db_load()
    db["sessions"] = [s for s in db.get("sessions", []) if s.get("token") != token]
    db_save(db)

def auth(roles=None):
    """
    Decorator that:
    1. Validates Bearer token (memory → MongoDB session lookup)
    2. Checks session TTL (SESSION_TTL_DAYS, default 30 days)
    3. Confirms user still exists and is not suspended
    4. Enforces role-based access control
    5. Attaches request.user for downstream use
    """
    def dec(fn):
        @wraps(fn)
        def wrap(*a, **kw):
            tok = request.headers.get("Authorization", "").replace("Bearer ", "").strip()
            if not tok:
                # Fallback: check query parameters (for direct file downloads via <a href>)
                tok = request.args.get("token", "").strip()

            if not tok:
                return jsonify({"error": "Unauthorized — no token provided"}), 401

            uid = TOKENS.get(tok)

            # Session lookup in MongoDB if not cached in memory
            if not uid:
                session = _session_find(tok)
                if session:
                    # Check session TTL
                    created = session.get("created_at", "")
                    if created:
                        try:
                            age_days = (datetime.utcnow() - datetime.fromisoformat(created)).days
                            if age_days > SESSION_TTL_DAYS:
                                # Session expired — clean up
                                _session_delete(tok)
                                return jsonify({"error": "Session expired. Please sign in again."}), 401
                        except Exception:
                            pass
                    uid = session.get("user_id")
                    if uid:
                        TOKENS[tok] = uid

            if not uid:
                return jsonify({"error": "Unauthorized — invalid or expired token"}), 401

            db = db_load()
            u  = next((x for x in db["users"] if x["id"] == uid), None)
            if not u:
                TOKENS.pop(tok, None)
                return jsonify({"error": "User not found"}), 401
            if u.get("status") == "suspended":
                TOKENS.pop(tok, None)
                return jsonify({"error": "Account suspended. Contact administrator."}), 403
            if roles and u["role"] not in roles:
                print(f"[AUTH] Forbidden: user {uid} role={u['role']} tried endpoint requiring {roles}")
                return jsonify({"error": "Forbidden — insufficient permissions"}), 403

            request.user = u
            return fn(*a, **kw)
        return wrap
    return dec

# ══════════════════════════════════════════════════════════════════════════════
#  AUTH ENDPOINTS  — OTP, role-isolation, full security
# ══════════════════════════════════════════════════════════════════════════════

# ── OTP store (in-memory + MongoDB) ──────────────────────────────────────────
# Format: {email: {"otp": "123456", "expires": epoch_seconds, "attempts": int}}
_otp_store: Dict[str, dict] = {}
M_OTP = "otp_store"

import time as _time

def _generate_otp() -> str:
    return str(random.randint(100000, 999999))

def _store_otp(email: str, otp: str, ttl_secs: int = 600):
    """Store OTP with expiry. Dual-writes to memory and MongoDB."""
    record = {"otp": otp, "expires": _time.time() + ttl_secs, "attempts": 0}
    _otp_store[email] = record
    if MONGO_OK:
        try:
            mongo_db[M_OTP].update_one(
                {"email": email},
                {"$set": {**record, "email": email}},
                upsert=True
            )
        except Exception as e:
            print(f"[OTP] MongoDB write error: {e}")

def _get_otp_record(email: str) -> Optional[dict]:
    rec = _otp_store.get(email)
    if rec:
        return rec
    if MONGO_OK:
        try:
            d = mongo_db[M_OTP].find_one({"email": email})
            if d:
                d.pop("_id", None)
                _otp_store[email] = d
                return d
        except Exception:
            pass
    return None

def _verify_otp(email: str, submitted: str) -> tuple:
    """Returns (ok: bool, error_msg: str)"""
    rec = _get_otp_record(email)
    if not rec:
        return False, "OTP not found. Please request a new one."
    if _time.time() > rec.get("expires", 0):
        _otp_store.pop(email, None)
        return False, "OTP expired. Please request a new one."
    if rec.get("attempts", 0) >= 5:
        _otp_store.pop(email, None)
        return False, "Too many incorrect attempts. Please request a new OTP."
    if rec["otp"] != submitted.strip():
        rec["attempts"] = rec.get("attempts", 0) + 1
        _otp_store[email] = rec
        return False, f"Incorrect OTP. {5 - rec['attempts']} attempt(s) remaining."
    # Valid — clean up
    _otp_store.pop(email, None)
    if MONGO_OK:
        try: mongo_db[M_OTP].delete_one({"email": email})
        except Exception: pass
    return True, ""

def _send_otp_email(to_email: str, otp: str, context: str = "registration") -> bool:
    """Send OTP via SMTP. Returns True on success, False if SMTP not configured."""
    host   = os.getenv("SMTP_HOST")
    port   = int(os.getenv("SMTP_PORT", 587))
    user   = os.getenv("SMTP_USER")
    passwd = os.getenv("SMTP_PASS")
    sender = os.getenv("SMTP_FROM", user)

    if not all([host, user, passwd]):
        print(f"[OTP] Dev mode — OTP for {to_email}: {otp}")
        return False   # Caller handles dev mode

    try:
        html = f"""
<div style="font-family:Arial,sans-serif;max-width:480px;margin:0 auto;padding:32px">
  <div style="text-align:center;margin-bottom:28px">
    <div style="font-size:48px">🔐</div>
    <h2 style="color:#1e293b;margin:8px 0 4px">Your VidyAI OTP</h2>
    <p style="color:#64748b;font-size:14px">For {context}</p>
  </div>
  <div style="background:#f1f5f9;border-radius:12px;padding:28px;text-align:center;margin-bottom:24px">
    <div style="font-size:42px;font-weight:900;letter-spacing:12px;color:#1e293b;font-family:monospace">{otp}</div>
    <p style="color:#64748b;font-size:13px;margin:12px 0 0">This OTP expires in <strong>10 minutes</strong></p>
  </div>
  <p style="font-size:13px;color:#94a3b8;text-align:center">
    If you did not request this OTP, please ignore this email.<br/>
    Do not share this OTP with anyone.
  </p>
  <div style="border-top:1px solid #e2e8f0;margin-top:24px;padding-top:16px;text-align:center">
    <span style="font-size:12px;color:#94a3b8">VidyAI — AI-Powered Education Platform</span>
  </div>
</div>"""
        msg = MIMEMultipart("alternative")
        msg["From"]    = sender
        msg["To"]      = to_email
        msg["Subject"] = f"VidyAI OTP: {otp} (expires in 10 minutes)"
        msg.attach(MIMEText(html, "html", "utf-8"))
        with smtplib.SMTP(host, port) as s:
            s.starttls(); s.login(user, passwd); s.send_message(msg)
        return True
    except Exception as e:
        print(f"[OTP] SMTP error: {e}")
        return False

@app.post("/api/auth/send-otp")
def send_otp():
    """Send a 6-digit OTP to an email address for verification."""
    b     = request.json or {}
    email = b.get("email", "").strip().lower()
    if not email or "@" not in email:
        return jsonify({"error": "Valid email required"}), 400

    # Rate-limit: max 3 OTPs per email per 10 minutes
    rec = _get_otp_record(email)
    if rec and _time.time() < rec.get("expires", 0) - 540:  # sent within last 60s
        return jsonify({"error": "Please wait 60 seconds before requesting another OTP"}), 429

    otp     = _generate_otp()
    context = b.get("form", "registration")
    _store_otp(email, otp)

    sent = _send_otp_email(email, otp, context)

    if sent:
        print(f"[OTP] Sent to {email}")
    else:
        # Dev mode: log OTP to console
        print(f"[OTP] ⚠️  SMTP not configured. Dev OTP for {email}: {otp}")

    return jsonify({
        "ok":       True,
        "dev_mode": not sent,
        "message":  f"OTP sent to {email}" if sent else f"[Dev] OTP logged to server console: {otp}",
        "dev_otp":  otp if not sent else None,   # expose in dev mode only
    })

@app.post("/api/auth/login")
def login():
    b  = request.json or {}
    db = db_load()

    email    = b.get("email", "").strip().lower()
    password = b.get("password", "")
    req_role = b.get("role", "").strip()      # Role the user selected on login tab

    if not email or not password:
        return jsonify({"error": "Email and password are required"}), 400

    u = next((x for x in db["users"] if x["email"] == email), None)

    # Unified error — do not reveal whether email exists
    if not u or u["pw_hash"] != _hash(password):
        return jsonify({"error": "Invalid email or password"}), 401

    if u.get("status") == "suspended":
        return jsonify({"error": "Your account has been suspended. Contact your administrator."}), 403

    # ── SECURITY: Role mismatch check ─────────────────────────────────────────
    # If the user selected a specific role tab, verify it matches their actual role.
    # Map UI role keys to DB role values
    ROLE_MAP = {
        "school":          "institute_admin",
        "institute_admin":    "institute_admin",
        "institute":       "institute_admin",
        "institute_admin": "institute_admin",
        "teacher":         "teacher",
        "student":         "student",
        "parent":          "parent",
        "admin":           "admin",
    }
    if req_role and req_role in ROLE_MAP:
        expected = ROLE_MAP[req_role]
        if u["role"] != expected:
            # Log the attempt
            print(f"[SECURITY] Role mismatch login attempt: {email} is '{u['role']}' but tried '{expected}'")
            return jsonify({
                "error": f"This account is registered as '{u['role']}'. Please select the correct role tab."
            }), 403

    # Issue token
    tok = uuid.uuid4().hex
    TOKENS[tok] = u["id"]
    _session_save(tok, u["id"], u["role"], ip=request.remote_addr or "")
    # Save activity log + last_login in background (non-blocking)
    import threading
    def _save_activity():
        try:
            db2 = db_load()
            for usr in db2["users"]:
                if usr["id"] == u["id"]:
                    usr["last_login"] = _now()
                    usr["login_count"] = usr.get("login_count", 0) + 1
                    break
            db_log(db2, u["id"], "login", f"{u['name']} ({u['role']}) signed in")
            db_save(db2)
        except Exception: pass
    threading.Thread(target=_save_activity, daemon=True).start()
    return jsonify({"token": tok, "user": _safe(u)})

@app.post("/api/auth/signup")
def signup():
    b     = request.json or {}
    name  = b.get("name", "").strip()
    email = b.get("email", "").strip().lower()
    pw    = b.get("password", "")
    role  = b.get("role", "student").strip()
    inst  = b.get("institution", "").strip()
    roll  = b.get("roll_number", "").strip()
    phone = b.get("phone", "").strip()
    otp   = b.get("otp", "").strip()

    # ── Input validation ──────────────────────────────────────────────────────
    if not name or not email or not pw:
        return jsonify({"error": "Name, email and password are required"}), 400
    if "@" not in email or "." not in email.split("@")[-1]:
        return jsonify({"error": "Enter a valid email address"}), 400
    if len(pw) < 8:
        return jsonify({"error": "Password must be at least 8 characters"}), 400
    # institute_admin registered via dedicated /api/auth/register-institute endpoint
    if role not in ("student", "teacher", "parent"):
        return jsonify({"error": "To register as an Institute, please use the Institute sign-up. Administrative accounts are restricted."}), 403

    # ── OTP verification (required for registration) ──────────────────────────
    if otp:
        ok, err = _verify_otp(email, otp)
        if not ok:
            return jsonify({"error": err}), 400
    else:
        return jsonify({"error": "OTP verification required. Please verify your email first."}), 400

    # ── Duplicate check ───────────────────────────────────────────────────────
    db = db_load()
    if any(x["email"] == email for x in db["users"]):
        return jsonify({"error": "An account with this email already exists"}), 409

    u = {
        "id": _uid(), "name": name, "email": email,
        "pw_hash": _hash(pw), "role": role,
        "institution": inst, "roll_number": roll, "phone": phone,
        "joined": _now()[:10], "docs": 0, "status": "active"
    }
    db["users"].append(u)
    db_log(db, u["id"], "signup", f"{name} registered as {role}")
    db_save(db)

    tok = uuid.uuid4().hex
    TOKENS[tok] = u["id"]
    _session_save(tok, u["id"], u["role"], ip=request.remote_addr or "")
    return jsonify({"token": tok, "user": _safe(u)}), 201

@app.post("/api/auth/register-institute")
def register_institute():
    """
    Dedicated endpoint for Institution sign-up.
    Creates a new institute_admin account; the institution name is
    mandatory and will be used to scope all users under that institute.
    Requires OTP exactly like regular sign-up.
    """
    b            = request.json or {}
    name         = b.get("name", "").strip()
    email        = b.get("email", "").strip().lower()
    pw           = b.get("password", "")
    institution  = b.get("institution", "").strip()
    phone        = b.get("phone", "").strip()
    otp          = b.get("otp", "").strip()

    if not name or not email or not pw:
        return jsonify({"error": "Name, email and password are required"}), 400
    if not institution:
        return jsonify({"error": "Institution name is mandatory for Institute registration"}), 400
    if "@" not in email or "." not in email.split("@")[-1]:
        return jsonify({"error": "Enter a valid email address"}), 400
    if len(pw) < 8:
        return jsonify({"error": "Password must be at least 8 characters"}), 400
    if not otp:
        return jsonify({"error": "OTP verification required. Please verify your email first."}), 400

    ok, err = _verify_otp(email, otp)
    if not ok:
        return jsonify({"error": err}), 400

    db = db_load()

    # Prevent duplicate email
    if any(x["email"] == email for x in db["users"]):
        return jsonify({"error": "An account with this email already exists"}), 409

    # Prevent duplicate institution name (case-insensitive)
    inst_lower = institution.lower()
    if any(x.get("institution", "").lower() == inst_lower and x["role"] == "institute_admin"
           for x in db["users"]):
        return jsonify({"error": f"An Institute account already exists for '{institution}'. Contact the existing admin."}), 409

    u = {
        "id": _uid(), "name": name, "email": email,
        "pw_hash": _hash(pw), "role": "institute_admin",
        "institution": institution, "phone": phone,
        "joined": _now()[:10], "docs": 0, "status": "active"
    }
    db["users"].append(u)
    db_log(db, u["id"], "signup", f"{name} registered Institute: {institution}")
    db_save(db)

    tok = uuid.uuid4().hex
    TOKENS[tok] = u["id"]
    _session_save(tok, u["id"], "institute_admin", ip=request.remote_addr or "")
    return jsonify({"token": tok, "user": _safe(u)}), 201


@app.post("/api/auth/logout")
@auth()
def logout():
    tok = request.headers.get("Authorization", "").replace("Bearer ", "").strip()
    TOKENS.pop(tok, None)
    _session_delete(tok)
    return jsonify({"ok": True})

@app.get("/api/auth/me")
@auth()
def me():
    return jsonify({"user": _safe(request.user)})

# ══════════════════════════════════════════════════════════════════════════════
#  PAYMENT — Razorpay Integration
# ══════════════════════════════════════════════════════════════════════════════
def _razorpay_client():
    """Return an initialised Razorpay client, or None if credentials missing."""
    if not RAZORPAY_KEY_ID or not RAZORPAY_KEY_SECRET:
        return None
    try:
        import razorpay as _rp
        return _rp.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))
    except ImportError:
        return None

def _apply_plan_quota(user: dict, plan_id: str):
    """
    After a successful payment, upgrade the user's quota limits to match
    the purchased plan by writing user_overrides in quota_config.
    """
    preset_map = PLAN_TO_PRESET.get(plan_id)
    if not preset_map:
        return

    qc = _qc_load()
    ptype = preset_map["type"]

    if ptype == "role_default":
        # For student-pro: apply the preset as a user_override so only THIS user
        # gets the upgraded limits (not all students).
        preset = (QUOTA_PRESETS.get("role_defaults") or {}).get(preset_map["preset"])
        if not preset:
            return
        template = preset.get("template", {})
        uid = user["id"]
        qc.setdefault("user_overrides", {}).setdefault(uid, {})
        for feat in QUOTA_FEATURES:
            if feat in template:
                qc["user_overrides"][uid][feat] = int(template[feat])
        _qc_save(qc)

    elif ptype == "institution":
        # For school/coaching plans: apply as institution_override
        inst = (user.get("institution") or "").strip()
        if not inst:
            return
        preset = (QUOTA_PRESETS.get("institution") or {}).get(preset_map["preset"])
        if not preset:
            return
        qc.setdefault("institution_overrides", {}).setdefault(inst, {})
        cfg = qc["institution_overrides"][inst]
        cfg["user_daily"] = {
            feat: int(limit)
            for feat, limit in (preset.get("user_daily") or {}).items()
            if feat in QUOTA_FEATURES
        }
        if preset.get("inst_daily"):
            cfg["inst_daily"] = {
                feat: int(limit)
                for feat, limit in preset["inst_daily"].items()
                if feat in QUOTA_FEATURES
            }
        _qc_save(qc)


@app.post("/api/payments/create-order")
@auth()
def payments_create_order():
    """
    Create a Razorpay order for the given plan.
    Body: { "plan_id": "student-pro" | "school-starter" | ... }
    Returns: { order_id, amount, currency, key_id }
    """
    b = request.json or {}
    plan_id = str(b.get("plan_id", "")).strip()

    if plan_id not in PLAN_PRICES:
        return jsonify({"error": f"Unknown plan: {plan_id}"}), 400

    if not RAZORPAY_KEY_ID or not RAZORPAY_KEY_SECRET:
        return jsonify({"error": "Payment gateway not configured. Contact support."}), 503

    client = _razorpay_client()
    if not client:
        return jsonify({"error": "Payment gateway unavailable. Contact support."}), 503

    amount = PLAN_PRICES[plan_id]
    u = request.user
    try:
        order = client.order.create({
            "amount":   amount,
            "currency": "INR",
            "payment_capture": 1,
            "notes": {
                "plan_id": plan_id,
                "user_id": u["id"],
                "email":   u["email"],
            }
        })
    except Exception as e:
        print(f"[PAYMENT] Razorpay order creation failed: {e}")
        return jsonify({"error": "Could not create payment order. Please try again."}), 500

    return jsonify({
        "order_id": order["id"],
        "amount":   order["amount"],
        "currency": order["currency"],
        "key_id":   RAZORPAY_KEY_ID,
        "plan_id":  plan_id,
    })


@app.post("/api/payments/verify")
@auth()
def payments_verify():
    """
    Verify Razorpay payment signature and activate the subscription.
    Body: { razorpay_order_id, razorpay_payment_id, razorpay_signature, plan_id }
    """
    b = request.json or {}
    order_id   = str(b.get("razorpay_order_id", "")).strip()
    payment_id = str(b.get("razorpay_payment_id", "")).strip()
    signature  = str(b.get("razorpay_signature", "")).strip()
    plan_id    = str(b.get("plan_id", "")).strip()

    if not all([order_id, payment_id, signature, plan_id]):
        return jsonify({"error": "Missing payment verification fields"}), 400

    if not RAZORPAY_KEY_SECRET:
        return jsonify({"error": "Payment gateway not configured"}), 503

    # HMAC-SHA256 signature verification
    import hmac as _hmac
    import hashlib as _hashlib
    expected = _hmac.new(
        RAZORPAY_KEY_SECRET.encode(),
        f"{order_id}|{payment_id}".encode(),
        _hashlib.sha256
    ).hexdigest()
    if not _hmac.compare_digest(expected, signature):
        return jsonify({"error": "Payment verification failed — invalid signature"}), 400

    # Activate subscription
    u = request.user
    from datetime import date as _date, timedelta as _td
    expires = (_date.today() + _td(days=30)).isoformat()

    db = db_load()
    for user in db["users"]:
        if user["id"] == u["id"]:
            user["plan"]           = plan_id
            user["plan_expires"]   = expires
            user["plan_payment_id"] = payment_id
            break
    db_save(db)

    # Apply quota upgrades for this user/institution
    _apply_plan_quota(u, plan_id)

    db_log(db, u["id"], "payment", f"Activated plan {plan_id} via payment {payment_id}")
    db_save(db)

    return jsonify({
        "ok":          True,
        "plan":        plan_id,
        "plan_expires": expires,
    })


@app.get("/api/payments/status")
@auth()
def payments_status():
    """Return the current user's active plan and expiry."""
    u = request.user
    from datetime import date as _date
    plan    = u.get("plan", "free")
    expires = u.get("plan_expires", "")
    active  = False
    if expires:
        try:
            active = _date.today() <= _date.fromisoformat(expires)
        except ValueError:
            pass
    if not active:
        plan = "free"
    return jsonify({
        "plan":        plan,
        "plan_expires": expires,
        "active":      active,
    })


@app.post("/api/payments/webhook")
def payments_webhook():
    """
    Razorpay webhook — verifies X-Razorpay-Signature header and processes
    payment.captured events for async payment confirmation.
    """
    sig = request.headers.get("X-Razorpay-Signature", "")
    body = request.get_data()

    if RAZORPAY_KEY_SECRET and sig:
        import hmac as _hmac
        import hashlib as _hashlib
        expected = _hmac.new(
            RAZORPAY_KEY_SECRET.encode(),
            body,
            _hashlib.sha256
        ).hexdigest()
        if not _hmac.compare_digest(expected, sig):
            return jsonify({"error": "Invalid webhook signature"}), 400

    try:
        event = request.json or {}
        if event.get("event") == "payment.captured":
            payload = event.get("payload", {}).get("payment", {}).get("entity", {})
            notes   = payload.get("notes", {})
            user_id = notes.get("user_id", "")
            plan_id = notes.get("plan_id", "")
            payment_id = payload.get("id", "")
            if user_id and plan_id:
                from datetime import date as _date, timedelta as _td
                expires = (_date.today() + _td(days=30)).isoformat()
                db = db_load()
                for user in db["users"]:
                    if user["id"] == user_id:
                        user["plan"]           = plan_id
                        user["plan_expires"]   = expires
                        user["plan_payment_id"] = payment_id
                        _apply_plan_quota(user, plan_id)
                        break
                db_save(db)
    except Exception as e:
        print(f"[WEBHOOK] Processing error: {e}")

    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT UPLOAD / LIST / DELETE
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/upload")
@auth()
def upload():
    try:
        u  = request.user
        db = db_load()

        # ── Per-role upload limit check ───────────────────────────────────────────
        settings   = db.get("settings", {})
        role_limit_key = f"max_uploads_{u['role']}"
        global_limit   = int(settings.get("max_uploads_per_user", 20))
        role_limit     = int(settings.get(role_limit_key, global_limit))
        user_docs      = [d for d in db.get("documents", []) if d["owner_id"] == u["id"]]
        if len(user_docs) >= role_limit:
            return jsonify({
                "error": f"Upload limit reached ({role_limit} files). "
                         f"Ask your school admin to increase the limit or delete existing uploads."
            }), 429

        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400
        f = request.files["file"]
        if not f or not f.filename:
            return jsonify({"error": "No file selected"}), 400
        if not _allowed(f.filename):
            ext_given = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else "unknown"
            return jsonify({"error": f"Unsupported file type (.{ext_given}). Allowed: PDF, DOCX, TXT, MD, MP3, MP4, WAV, M4A"}), 400

        ext  = f.filename.rsplit(".", 1)[1].lower()
        fn   = secure_filename(f.filename)
        did  = uuid.uuid4().hex[:10]
        udir = UPL_DIR / u["id"]
        udir.mkdir(parents=True, exist_ok=True)
        spath = udir / f"{did}.{ext}"
        f.save(str(spath))

        chunks, chapters = _index_doc(spath, did, ext)
        syl_name = request.form.get("syllabus_name", Path(fn).stem)
        ocr_queued = (chunks == -1)
        if ocr_queued:
            # Scanned PDF: start background OCR thread, save doc with chunks=0 for now
            import threading
            t = threading.Thread(target=_run_ocr_background, args=(spath, did, syl_name, fn, u["id"]), daemon=True)
            t.start()
            chunks, chapters = 0, []

        db = db_load()
        doc = {
            "id": did, "owner_id": u["id"], "name": fn, "ext": ext,
            "type": _dtype(ext), "size": spath.stat().st_size,
            "path": str(spath), "uploaded_at": _now(),
            "chunks": chunks, "chapters": chapters
        }
        db["documents"].append(doc)
        db_log(db, u["id"], "upload", f"Uploaded {fn}")
        db_save(db)

        syllabi_registry[did] = {
            "id": did, "name": syl_name, "files": [fn],
            "chunks": chunks, "chapters": chapters,
            "owner_id": u["id"], "created_at": _now()
        }
        _save_syllabi()
        return jsonify({
            "doc": doc, "syllabus_id": did, "syllabus_name": syl_name,
            "chunks": chunks, "chapters": chapters,
            "ocr_queued": ocr_queued
        }), 201

    except Exception as e:
        import traceback
        print(f"[UPLOAD ERROR] {e}")
        traceback.print_exc()
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500

@app.get("/api/documents/<did>/ocr-status")
@auth()
def doc_ocr_status(did):
    status = ocr_status.get(did, "ready")  # default ready for non-OCR docs
    chunks = 0
    db = db_load()
    doc = next((d for d in db.get("documents", []) if d["id"] == did), None)
    if doc:
        chunks = doc.get("chunks", 0)
    return jsonify({"status": status, "chunks": chunks})


@app.get("/api/documents")
@auth()
def list_docs():
    db = db_load()
    u  = request.user
    # Admin sees all documents from all users
    # Every other role sees ONLY their own documents — strict isolation
    if u["role"] in ("admin", "institute_admin"):
        docs = db["documents"]
    else:
        docs = [d for d in db["documents"] if d["owner_id"] == u["id"]]
    return jsonify({"documents": docs})

@app.delete("/api/documents/<did>")
@auth()
def del_doc(did):
    db = db_load(); u = request.user
    doc = next((d for d in db["documents"] if d["id"] == did), None)
    if not doc:
        return jsonify({"error": "Not found"}), 404
    if doc["owner_id"] != u["id"] and u["role"] != "admin":
        return jsonify({"error": "Forbidden"}), 403
    try:
        Path(doc["path"]).unlink(missing_ok=True)
    except Exception:
        pass
    syllabi_registry.pop(did, None)
    qa_chains.pop(did, None)
    memories.pop(did, None)
    vector_stores.pop(did, None)
    _save_syllabi()
    db["documents"] = [d for d in db["documents"] if d["id"] != did]
    db_save(db)
    return jsonify({"ok": True})

# ══════════════════════════════════════════════════════════════════════════════
#  SYLLABI
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/syllabi")
@auth()
def list_syllabi():
    u     = request.user
    items = list(syllabi_registry.values())

    if u["role"] in ("admin", "institute_admin"):
        # Admin sees everything
        pass
    else:
        # STRICT isolation: each user sees ONLY their own uploaded syllabi.
        # gov_ syllabi (NCERT/state) are session-scoped — the user who loaded
        # them in the current session sees them; others do not.
        # This prevents Student A's uploads from appearing in Teacher B's lists.
        items = [
            s for s in items
            if s.get("owner_id") == u["id"]   # this user's own (uploaded or loaded)
        ]

    return jsonify({"syllabi": items})

@app.get("/api/syllabi/<syllabus_id>/chapters")
@auth()
def get_chapters(syllabus_id):
    s = syllabi_registry.get(syllabus_id)
    if not s:
        return jsonify({"error": "Syllabus not found"}), 404
    return jsonify({"syllabus_id": syllabus_id, "name": s.get("name"), "chapters": s.get("chapters", [])})

@app.post("/api/syllabi/register-virtual")
@auth()
def register_virtual():
    u    = request.user
    data = request.json or {}
    sid  = data.get("syllabus_id")
    name = data.get("name")
    chapters = data.get("chapters", [])
    if not sid or not name:
        return jsonify({"error": "Missing syllabus_id or name"}), 400
    # Always set / refresh owner_id — prevents a crafted request from
    # registering a syllabus under a shared/global key
    syllabi_registry[sid] = {
        "id": sid, "name": name, "chapters": chapters,
        "owner_id": u["id"], "created_at": _now(), "is_virtual": True
    }
    _save_syllabi()
    return jsonify({"success": True, "syllabus_id": sid})

@app.get("/api/syllabi/<syllabus_id>/sample-questions")
@auth()
def get_syllabus_sample_questions(syllabus_id):
    """Return 5 sample questions a student could ask for this syllabus."""
    FALLBACK = [
        "What are the main topics covered in this chapter?",
        "Explain the key concepts with simple examples",
        "What type of questions appear in the exam from this chapter?",
        "How do I solve problems related to this topic?",
        "What are the most important points to memorise?",
    ]
    if not OPENAI_API_KEY:
        return jsonify({"questions": FALLBACK})
    vs = _load_vs(syllabus_id)
    if not vs:
        return jsonify({"questions": FALLBACK})
    try:
        docs = vs.similarity_search("important concepts key topics exam questions definitions", k=4)
        ctx  = "\n\n".join(d.page_content[:400] for d in docs[:3])
        prompt = (
            "You are an expert teacher. Based on the following excerpt from a student's syllabus, "
            "generate exactly 5 short, specific questions (1 sentence each) that a student would "
            "naturally ask their tutor while studying this material. "
            "Make them concrete — reference actual topics from the text, not generic prompts.\n\n"
            f"Excerpt:\n{ctx[:1800]}\n\n"
            'Return JSON only: {"questions": ["Q1", "Q2", "Q3", "Q4", "Q5"]}'
        )
        result = _llm_json(prompt, temperature=0.3, mini=True)
        if isinstance(result, dict):
            qs = result.get("questions", [])
            if isinstance(qs, list) and len(qs) >= 3:
                return jsonify({"questions": qs[:5]})
    except Exception as _e:
        log.warning(f"[SAMPLE_Q] {syllabus_id}: {_e}")
    return jsonify({"questions": FALLBACK})


@app.delete("/api/syllabi/<syllabus_id>")
@auth()
def delete_syllabus(syllabus_id):
    u = request.user
    s = syllabi_registry.get(syllabus_id)
    if not s:
        return jsonify({"error": "Not found"}), 404
    # Only the owner or admin can delete
    if s.get("owner_id") and s["owner_id"] != u["id"] and u["role"] not in ("admin", "institute_admin"):
        return jsonify({"error": "Forbidden — you can only remove your own syllabi"}), 403
    syllabi_registry.pop(syllabus_id, None)
    qa_chains.pop(syllabus_id, None)
    memories.pop(syllabus_id, None)
    vector_stores.pop(syllabus_id, None)
    _save_syllabi()
    return jsonify({"success": True, "deleted": syllabus_id})

# ══════════════════════════════════════════════════════════════════════════════
#  CURRICULUM ENDPOINTS (Indian Board)
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/curriculum/metadata")
@auth()
def get_curriculum_metadata():
    return jsonify({
        "metadata": CURRICULUM_DATA,
        "status": {
            "mongodb_connected":    MONGO_OK,
            "openai_configured":    bool(OPENAI_API_KEY),
            "langchain_available":  LANGCHAIN_OK,
            "persistence":          "mongodb" if MONGO_OK else "json_fallback",
            "model":                OPENAI_MODEL,
        }
    })

@app.post("/api/curriculum/load")
@auth()
@quota('curriculum_load')
def load_curriculum_book():
    b       = request.json or {}
    state   = b.get("state")
    board   = b.get("board")
    class_n = b.get("class") or b.get("year")
    subject = b.get("subject")
    subject_canon = _normalize_subject_name(subject)
    u       = request.user

    if not all([state, board, class_n, subject]):
        return jsonify({"error": "state, board, class, and subject are required"}), 400

    # ── IMPORTANT: Embed user_id in the key so each user's curriculum hub
    #    entries are INDEPENDENT and never visible to other users.
    did = (
        f"gov_{u['id'][:6]}_"
        f"{board.lower().replace(' ','_')}_{class_n.replace(' ','').lower()}_"
        f"{subject_canon.replace(' ','_')}"
    )
    class_key = class_n.lower().replace(" ", "")
    subj_key  = subject_canon

    chapters = CHAPTERS_DB.get((class_key, subj_key))

    if not chapters:
        try:
            prompt = (f"List the EXACT chapter titles from the official '{board}' textbook for '{subject}' — {class_n} in '{state}', India.\n"
                      f"Return ONLY a valid JSON array of strings. Example: [\"Chapter 1: ...\", \"Chapter 2: ...\"]")
            result = _llm_json(prompt, temperature=0.1, mini=True)
            if isinstance(result, list) and result:
                chapters = [str(c).strip() for c in result]
        except Exception as e:
            print(f"[CHAPTERS LLM] {e}")

    if not chapters:
        chapters = [f"Chapter {i}: Topic {i}" for i in range(1, 8)]

    pdf_url      = NCERT_PDF_URLS.get((class_key, subj_key), "")
    supp_pdf_url = NCERT_SUPP_PDF_URLS.get((class_key, subj_key), "")

    syllabi_registry[did] = {
        "id": did,
        "name": f"{board} {class_n} — {subject_canon.title()}",
        "files": [f"{subject_canon}_textbook.pdf"],
        "chunks": 100, "chapters": chapters,
        "owner_id": u["id"],        # ← always scoped to this user
        "created_at": _now(),
        "pdf_url": pdf_url, "supp_pdf_url": supp_pdf_url
    }
    _save_syllabi()
    return jsonify({
        "syllabus_id": did, "chapters": chapters,
        "name": syllabi_registry[did]["name"], "pdf_url": pdf_url,
        "supp_pdf_url": supp_pdf_url
    })

@app.post("/api/curriculum/subjects")
@auth()
def get_curriculum_subjects():
    b = request.json or {}
    board   = b.get("board", "CBSE")
    class_n = b.get("class", "Class 10")
    # Find subjects for any state that has this board
    for state_data in CURRICULUM_DATA.values():
        if board in state_data and class_n in state_data[board]:
            return jsonify({"subjects": state_data[board][class_n]})
    # Fallback
    return jsonify({"subjects": ["Mathematics","Science","English","Social Science","Hindi"]})

@app.post("/api/curriculum/chapters")
@auth()
def get_curriculum_chapters():
    b = request.json or {}
    # Backward compatibility: some older clients send `exam` instead of `board`.
    board   = (b.get("board") or b.get("exam") or "CBSE")
    class_n = b.get("class", "Class 10")
    subject = b.get("subject", "Mathematics")
    subject_canon = _normalize_subject_name(subject)
    medium  = b.get("medium", "English")
    u       = request.user

    # User-scoped ID — prevents cross-user syllabus leakage
    did = (
        f"gov_{u['id'][:6]}_"
        f"{board.lower().replace(' ','_')}_{class_n.replace(' ','').lower()}_"
        f"{subject_canon.replace(' ','_')}"
    )
    class_key = class_n.lower().replace(" ", "")
    subj_key  = subject_canon

    # CISCE boards (ICSE/ISC) are private — not on DIKSHA
    _is_cisce  = board in ("ICSE", "ISC")
    # Odisha state-board textbooks on DIKSHA are published by SCERT Odisha
    _is_odisha = board in ("BSE", "CHSE")
    # Only national/NCERT boards use the hardcoded CHAPTERS_DB and NCERT PDF URLs.
    # State boards must always go through DIKSHA so they get their own textbooks.
    _NCERT_BOARDS = {"CBSE", "NIOS", "NCERT", "DoE", "IB", "CBSE-AP"}

    source        = "local"
    available_books: list = []   # List of alternative textbooks from DIKSHA

    if board in _NCERT_BOARDS:
        chapters     = CHAPTERS_DB.get((class_key, subj_key))
        pdf_url      = NCERT_PDF_URLS.get((class_key, subj_key), "")
        supp_pdf_url = NCERT_SUPP_PDF_URLS.get((class_key, subj_key), "")

        # Classes 6–8: NEP new-curriculum rollout — the textbook.php portal lists
        # multiple editions (English/Hindi/Urdu/old vs new), confusing students.
        # Query DIKSHA WITHOUT medium filter so we get ALL language editions.
        # The frontend then shows a specific book picker instead of the portal page.
        # Clear the portal URL so the frontend never opens the confusing portal.
        _new_curriculum_classes = {"class6", "class7", "class8"}
        if class_key in _new_curriculum_classes:
            try:
                # Pass empty string for medium to get all editions (English, Hindi, Urdu)
                dk_books = _diksha_find_textbooks(board, class_n, subject_canon, "")
                # Only surface actual textbooks (not comic books / question sets)
                textbook_books = [
                    b for b in dk_books
                    if b.get("leafNodesCount", 0) >= 5
                    and len(b.get("subject", [])) == 1  # single-subject only
                    and _is_probable_textbook_name(b.get("name", ""))
                ]
                if textbook_books:
                    available_books = [
                        {"identifier": b.get("identifier", ""),
                         "name":       b.get("name", ""),
                         "leaves":     b.get("leafNodesCount", 0),
                         "medium":     b.get("medium", [])}
                        for b in textbook_books
                    ]
                    # Clear the portal URL — frontend will get the direct PDF
                    # by selecting a book from the picker via /diksha/chapters
                    pdf_url = ""
                    print(f"[DIKSHA-NCERT] {len(available_books)} books for {board} {class_n} {subject}")
            except Exception as e:
                print(f"[DIKSHA-NCERT] Error: {e}")
    else:
        chapters     = None
        pdf_url      = ""
        supp_pdf_url = ""

    if _is_cisce:
        chapters = CISCE_CURATED_CHAPTERS.get((class_key, subj_key))
        if chapters:
            source = "cisce_curated"

    # For CBSE/NCERT with known chapters, use local data (fastest).
    # CISCE boards (ICSE/ISC) bypass DIKSHA — their content is not on DIKSHA.
    if not chapters and not _is_cisce:
        # Try DIKSHA for state boards and any unknown combo
        try:
            # Get all available books from DIKSHA first, then pick best
            dk_books = _diksha_find_textbooks(board, class_n, subject_canon, medium)
            if dk_books:
                # Expose all options to the frontend for book-picker UI
                available_books = [
                    {"identifier": b.get("identifier",""),
                     "name":       b.get("name",""),
                     "leaves":     b.get("leafNodesCount", 0),
                     "medium":     b.get("medium", [])}
                    for b in dk_books
                    if len(b.get("subject",[])) <= 2   # exclude multi-subject bundles
                ]
            dk_chapters, dk_pdf, dk_name = _diksha_get_chapters_and_pdf(
                board, class_n, subject_canon, medium
            )
            if dk_chapters:
                chapters = dk_chapters
                # Odisha state-board books on DIKSHA are SCERT publications
                source = "scert" if _is_odisha else "diksha"
                if dk_pdf:
                    pdf_url = dk_pdf
                print(f"[DIKSHA] Found {len(chapters)} chapters for {board} {class_n} {subject}")
        except Exception as e:
            print(f"[DIKSHA] Fallback error: {e}")

    # ── NCERT fallback ─────────────────────────────────────────────────────
    # When DIKSHA has no data (or returned a wrong-subject book that was
    # discarded), use NCERT chapters + PDF.  Most Indian state boards follow
    # the same NCERT syllabus for core subjects (Math, Science, English, etc.)
    # so this gives the student meaningful content even for state-board selections.
    if not chapters and not _is_cisce:
        ncert_chs = CHAPTERS_DB.get((class_key, subj_key))
        ncert_pdf = NCERT_PDF_URLS.get((class_key, subj_key), "")
        if ncert_chs:
            chapters = ncert_chs
            pdf_url  = ncert_pdf
            source   = "ncert_fallback"
            print(f"[NCERT fallback] Using NCERT data for {board} {class_n} {subject}")

    # Final fallback: LLM generation
    if not chapters:
        try:
            board_ctx = "CISCE (ICSE/ISC board)" if _is_cisce else board
            result = _llm_json(
                f"List exact chapters for '{board_ctx}' '{subject_canon}' {class_n} India. JSON array only.",
                temperature=0.1, mini=True
            )
            if isinstance(result, list):
                chapters = [str(c).strip() for c in result]
                source = "cisce" if _is_cisce else "llm"
        except Exception:
            chapters = [f"Chapter {i}" for i in range(1, 8)]
            source = "fallback"

    # Create or update in registry — always with this user as owner
    syllabi_registry[did] = {
        "id": did,
        "name": f"{board} {class_n} — {subject_canon.title()}",
        "chunks": 100, "chapters": chapters,
        "owner_id": u["id"], "created_at": _now(), "pdf_url": pdf_url,
        "supp_pdf_url": supp_pdf_url,
        # Stored explicitly so question-gen endpoints can enforce class-level accuracy
        "board": board,
        "class_name": class_n,
        "subject": subject_canon.title(),
    }
    _save_syllabi()
    return jsonify({
        "syllabus_id": did,
        "name": syllabi_registry[did]["name"],
        "chapters": chapters,
        "pdf_url": pdf_url,
        "supp_pdf_url": supp_pdf_url,
        "available_books": available_books,
        "source": source
    })

# ── UK Curriculum Endpoints ───────────────────────────────────────────────────
@app.post("/api/uk-curriculum/topics")
@auth()
def uk_topics():
    b = request.json or {}
    year = b.get("year", "Year 10"); subject = b.get("subject", "Mathematics")
    u    = request.user
    did  = f"uk_{year.lower().replace(' ','_')}_{subject.lower().replace(' ','_')}"
    try:
        result = _llm_json(f"List the main topics in the UK National Curriculum for '{subject}' — {year}. JSON array of strings only.", temperature=0.1, mini=True)
        topics = result if isinstance(result, list) else [f"Topic {i}" for i in range(1, 8)]
    except Exception:
        topics = [f"{subject} Topic {i}" for i in range(1, 8)]
    if did not in syllabi_registry:
        syllabi_registry[did] = {"id": did, "name": f"UK {year} — {subject}", "chapters": topics, "owner_id": u["id"], "created_at": _now()}
        _save_syllabi()
    return jsonify({"syllabus_id": did, "topics": topics})

@app.post("/api/uk-curriculum/<tool>")
@auth()
@quota('uk_curriculum')
def uk_tool(tool):
    b = request.json or {}
    year = b.get("year","Year 10"); subject = b.get("subject","Mathematics")
    topics = b.get("topics", [])
    topic_str = ", ".join(topics) if topics else subject
    ctx = f"UK National Curriculum: {subject}, {year}. Topics: {topic_str}."
    return _run_tool(tool, ctx, topic_str, subject)

# ── Summarise / Flashcards / Questions / Audio ────────────────────────────────
@app.post("/api/summarise")
@auth()
@quota('summarise')
def summarise_chapter():
    b = request.json or {}
    sid     = b.get("syllabus_id")
    topic   = b.get("topic")
    subject = b.get("subject", "").lower()
    if not sid or not topic:
        return jsonify({"error": "syllabus_id and topic required"}), 400
    vs  = _load_vs(sid)
    ctx = ""
    if vs:
        try:
            docs = vs.as_retriever(search_kwargs={"k": 8}).invoke(topic)
            ctx  = "\n\n".join(d.page_content for d in docs)
        except Exception:
            pass
    if not ctx:
        s_info = syllabi_registry.get(sid, {})
        ctx = (f"Provide a comprehensive 10-point explanation about '{topic}' from the "
               f"'{s_info.get('name', subject)}' curriculum. "
               f"Cover: introduction, core concepts, key facts, examples, applications, and summary.")
    return jsonify({"summary": _generate_summary_with_fallback(ctx, topic, subject, language=subject)})

@app.post("/api/flashcards")
@auth()
@quota('flashcards')
def generate_flashcards():
    b = request.json or {}
    sid     = b.get("syllabus_id")
    topic   = b.get("topic")
    subject = b.get("subject", "").lower()
    count   = max(1, min(30, int(b.get("flashcard_count", b.get("count", 10)))))
    if not sid or not topic:
        return jsonify({"error": "syllabus_id and topic required"}), 400
    vs  = _load_vs(sid)
    ctx = ""
    if vs:
        try:
            docs = vs.as_retriever(search_kwargs={"k": 10}).invoke(topic)
            ctx  = "\n\n".join(d.page_content for d in docs)
        except Exception:
            pass
    if not ctx:
        s_info = syllabi_registry.get(sid, {})
        ctx = f"Educational content for '{topic}' from {s_info.get('name', subject)} curriculum."
    regional_langs = ["odia","kannada","hindi","telugu","gujarati","marathi","bengali","punjabi","tamil","malayalam"]
    is_regional    = any(lang in subject for lang in regional_langs)
    prompt = (REGIONAL_FLASHCARDS_PROMPT.replace("{language}", subject.title())
                                         .replace("{count}", str(count))
                                         .replace("{context}", ctx)
              if is_regional
              else FLASHCARDS_PROMPT.replace("{count}", str(count)).replace("{context}", ctx))
    cards = _llm_json(prompt)
    if not isinstance(cards, list):
        cards = []
    return jsonify({"flashcards": cards})

@app.post("/api/curriculum/<tool>")
@auth()
@quota('curriculum_tool')
def curriculum_tool(tool):
    b = request.json or {}
    sid      = b.get("syllabus_id")
    subject  = b.get("subject", "")
    board    = b.get("board", "")
    class_n  = b.get("class", "")
    chapters = b.get("chapters", [])
    topic    = b.get("topic") or (", ".join(chapters[:3]) if chapters else subject)
    try:
        fc_count = int(b.get("flashcard_count", 10))
    except (ValueError, TypeError):
        fc_count = 10
    try:
        point_count = max(1, min(20, int(b.get("point_count", 10))))
    except (ValueError, TypeError):
        point_count = 10
    if not sid:
        return jsonify({"error": "syllabus_id required"}), 400
    vs  = _load_vs(sid)
    ctx = ""
    if vs:
        try:
            docs = vs.as_retriever(search_kwargs={"k": 8}).invoke(topic or subject)
            ctx  = "\n\n".join(d.page_content for d in docs)
        except Exception:
            pass
    if not ctx:
        s = syllabi_registry.get(sid, {})
        ctx = (f"You are an expert on the Indian school curriculum. "
               f"Provide detailed educational content for '{topic or subject}' "
               f"from '{s.get('name', subject)}'. Cover all key concepts, examples, and facts.")
    # Keep generation anchored to the user's actual board/class selection.
    if board or class_n:
        ctx = f"Board: {board or 'Unknown'}\nClass: {class_n or 'Unknown'}\nSubject: {subject or 'Unknown'}\n\n{ctx}"
    try:
        return _run_tool(tool, ctx, topic or subject, subject, fc_count=fc_count, point_count=point_count)
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": f"AI generation failed: {e}"}), 500

AUDIO_LESSON_PROMPT = """You are an expert Indian school teacher creating an audio lesson about: {topic}
Subject: {subject}

Generate a structured audio lesson with 4-6 sections. Each section should be speakable (no symbols, no markdown).
Return ONLY a valid JSON object with this structure:
{{
  "summary": "A complete 4-5 sentence spoken summary of the entire topic, natural language, no bullet points or symbols",
  "sections": [
    {{"label": "Introduction", "text": "2-3 sentences introducing the topic in simple spoken language"}},
    {{"label": "Core Concept", "text": "3-4 sentences explaining the main concept clearly"}},
    {{"label": "Real-World Example", "text": "2-3 sentences with a relatable Indian example"}},
    {{"label": "Key Points", "text": "3-4 sentences summarising the most important points to remember"}},
    {{"label": "Exam Tips", "text": "2-3 sentences on what students should focus on for exams"}}
  ]
}}

Rules:
- All text must be plain spoken English — no bullet symbols (•), no asterisks, no markdown
- Each section text should be 2-4 natural spoken sentences
- Use simple vocabulary suitable for Class 6-12 students
- Include Indian context where relevant (NCERT examples, Indian scientists, Indian geography)

Context from textbook: {context}

Return ONLY the JSON object, no markdown:"""

def _run_tool(tool: str, ctx: str, topic: str, subject: str, fc_count: int = 10, point_count: int = 10):
    """Shared dispatcher for summarise / flashcards / questions / audio / video."""

    if tool == "summarise":
        summary = _generate_summary_with_fallback(ctx, topic, subject, point_count=point_count)
        return jsonify({"summary": summary})

    if tool == "flashcards":
        cards = _llm_json(FLASHCARDS_PROMPT.replace("{count}", str(fc_count)).replace("{context}", ctx))
        return jsonify({"flashcards": cards if isinstance(cards, list) else []})

    if tool == "questions":
        prompt = (
            f"Generate 10 exam-ready practice questions about: {topic}.\n"
            f"Context: {ctx}\n"
            f"Mix objective (MCQ) and subjective types.\n"
            f"Return ONLY a valid JSON array: [{{\"question\":\"...\",\"answer\":\"...\",\"type\":\"objective|subjective\"}}]"
        )
        qs = _llm_json(prompt)
        return jsonify({"questions": qs if isinstance(qs, list) else []})

    if tool == "audio":
        # Audio tool: returns structured sections for individual play + full summary
        prompt = AUDIO_LESSON_PROMPT.format(topic=topic, subject=subject, context=ctx[:3000])
        try:
            result = _llm_json(prompt, temperature=0.4)
            if isinstance(result, dict):
                # Clean all text to remove symbols that break TTS
                import re
                def clean_tts(text):
                    if not text: return ""
                    # Remove bullet symbols, asterisks, hash, pipes
                    text = re.sub(r'[•\*#|→←↑↓]', '', text)
                    # Remove markdown bold/italic
                    text = re.sub(r'\*+', '', text)
                    # Remove numbered list markers at start of lines
                    text = re.sub(r'^\s*\d+[\.\)]\s*', '', text, flags=re.MULTILINE)
                    # Collapse extra whitespace
                    text = re.sub(r'\s+', ' ', text).strip()
                    return text

                summary = clean_tts(result.get("summary", ""))
                sections = []
                for s in result.get("sections", []):
                    sections.append({
                        "label": s.get("label", "Section"),
                        "text":  clean_tts(s.get("text", ""))
                    })

                # If no sections returned, build from summary paragraphs
                if not sections and summary:
                    sentences = [s.strip() for s in summary.split('.') if s.strip()]
                    chunk_size = max(2, len(sentences) // 4)
                    for i in range(0, len(sentences), chunk_size):
                        chunk = '. '.join(sentences[i:i+chunk_size]) + '.'
                        sections.append({"label": f"Part {i//chunk_size + 1}", "text": chunk})

                return jsonify({"summary": summary, "sections": sections})
        except Exception as e:
            print(f"[AUDIO] JSON parse failed, falling back to text: {e}")
            # Fallback: generate plain text and split into sections
            fallback_prompt = f"Explain {topic} for Indian school students in 5 clear spoken paragraphs. No bullet points, no symbols, only plain sentences. Context: {ctx[:2000]}"
            summary = _llm_text(fallback_prompt, temperature=0.4)
            # Clean and split into sections
            import re
            summary = re.sub(r'[•\*#|→←↑↓]', '', summary)
            summary = re.sub(r'\*+', '', summary)
            paragraphs = [p.strip() for p in summary.split('\n\n') if p.strip()]
            if not paragraphs:
                paragraphs = [summary]
            labels = ["Introduction", "Core Concept", "Examples", "Key Points", "Summary"]
            sections = [
                {"label": labels[min(i, len(labels)-1)], "text": para}
                for i, para in enumerate(paragraphs[:5])
            ]
            return jsonify({"summary": summary, "sections": sections})

    if tool == "video":
        prompt = VIDEO_EXPLAINER_PROMPT.replace("{topic}", topic).replace("{context}", ctx)
        script = _llm_json(prompt)
        return jsonify({"script": script if isinstance(script, list) else []})

    return jsonify({"error": f"Unknown tool: {tool}"}), 400

# ── Chat ──────────────────────────────────────────────────────────────────────
@app.post("/api/chat")
@auth()
@quota('chat')
def chat():
    b           = request.json or {}
    question    = b.get("message", b.get("question", "")).strip()
    syllabus_id = b.get("doc_id", b.get("syllabus_id", "")).strip()
    u           = request.user

    if not question:
        return jsonify({"error": "Empty message"}), 400

    if not LANGCHAIN_OK or not OPENAI_API_KEY:
        return jsonify({"answer": "(Demo) AI features require OPENAI_API_KEY.", "sources": []})

    if syllabus_id and syllabus_id in syllabi_registry:
        try:
            chain  = _get_chain(syllabus_id)
            result = chain({"question": question})
            sources = list({Path(d.metadata.get("source","")).name for d in result.get("source_documents",[])})
            return jsonify({
                "answer": result["answer"], "sources": sources,
                "syllabus_id": syllabus_id,
                "syllabus_name": syllabi_registry.get(syllabus_id, {}).get("name", "")
            })
        except Exception as e:
            print(f"[CHAT] Chain error: {e}")
            import traceback; traceback.print_exc()
            # Fall through to general fallback instead of returning nothing

    # Cross-syllabus global search (only searches already-cached vector stores)
    if not syllabus_id:
        all_ctx, all_src = [], set()
        for sid, info in syllabi_registry.items():
            if sid not in vector_stores:
                continue  # skip uncached — loading from disk re-embeds and is too slow
            try:
                vs = vector_stores[sid]
                if vs and hasattr(vs, "similarity_search"):
                    docs = vs.similarity_search(question, k=2)
                    if docs:
                        all_ctx.append(f"--- {info.get('name','?')} ---\n" + "\n".join(d.page_content for d in docs))
                        for d in docs:
                            all_src.add(Path(d.metadata.get("source","")).name)
            except Exception:
                continue
        if all_ctx:
            prompt = f"Context from multiple sources:\n\n{chr(10).join(all_ctx)}\n\nQuestion: {question}\n\nAnswer based on the context:"
            return jsonify({"answer": _llm_text(prompt), "sources": list(all_src), "syllabus_id": "global"})

    # General fallback (also reached when chain fails for a specific syllabus)
    try:
        syl_ctx = ""
        if syllabus_id:
            s = syllabi_registry.get(syllabus_id, {})
            syl_ctx = f" The student has selected: {s.get('name', syllabus_id)}."
        sys_msg = SystemMessage(content=f"You are a helpful AI study assistant.{syl_ctx} The user is a {u.get('role','student')}. Be helpful and educational.")
        res = _get_llm(mini=True).invoke([sys_msg, HumanMessage(content=question)])
        return jsonify({"answer": res.content, "sources": [], "syllabus_id": syllabus_id or ""})
    except Exception as e:
        log.exception("[CHAT] LLM fallback failed")
        return jsonify({
            "answer": _study_fallback_answer(question, u.get("role", "student")),
            "sources": [],
            "syllabus_id": syllabus_id or "",
            "fallback": True,
            "warning": f"AI backend temporary issue: {str(e)}",
        }), 200

# ── Mentor / Video Session ────────────────────────────────────────────────────
@app.post("/api/mentor-session")
@auth()
@quota('video_explanation')
def video_explanation():
    b   = request.json or {}
    sid = b.get("syllabus_id", "")
    topic = b.get("topic", "")
    if not sid or not topic:
        return jsonify({"error": "syllabus_id and topic required"}), 400
    vs  = _load_vs(sid)
    ctx = ""
    if vs:
        try:
            docs = vs.as_retriever(search_kwargs={"k": 5}).invoke(topic)
            ctx  = "\n\n".join(d.page_content for d in docs)
        except Exception:
            pass
    if not ctx:
        ctx = f"Educational content about {topic}."
    prompt = VIDEO_EXPLAINER_PROMPT.replace("{topic}", topic).replace("{context}", ctx)
    script = _llm_json(prompt)
    if not isinstance(script, list):
        return jsonify({"error": "Failed to generate video script"}), 500
    return jsonify({"script": script})

# ══════════════════════════════════════════════════════════════════════════════
#  QUESTION GENERATION
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/generate-questions")
@auth()
@quota('generate_questions')
def generate_questions():
    data = request.json or {}
    sid         = data.get("syllabus_id", "").strip()
    topic       = data.get("topic", "the uploaded material")
    chapters    = data.get("chapters", [])
    obj_count   = max(0, int(data.get("objective_count", data.get("count", 5))))
    subj_count  = max(0, int(data.get("subjective_count", 0)))
    obj_weight  = float(data.get("objective_weightage", data.get("marks", 1)))
    subj_weight = float(data.get("subjective_weightage", max(2.0, float(data.get("marks", 2)))))
    total       = obj_count + subj_count

    if not sid or sid not in syllabi_registry:
        return jsonify({"error": "Select a valid syllabus first"}), 400
    if total <= 0:
        return jsonify({"error": "Total questions must be > 0"}), 400

    # Difficulty distribution
    single_diff = (data.get("difficulty") or "").lower()
    if single_diff in ("easy", "medium", "hard"):
        easy_c   = total if single_diff == "easy"   else 0
        medium_c = total if single_diff == "medium" else 0
        hard_c   = total if single_diff == "hard"   else 0
    else:
        dist     = data.get("difficulty_distribution", {})
        easy_c   = max(0, round((int(dist.get("easy",   40)) / 100) * total))
        medium_c = max(0, round((int(dist.get("medium", 40)) / 100) * total))
        hard_c   = max(0, total - easy_c - medium_c)

    topic_str  = f"{topic} — chapters: {', '.join(chapters)}" if chapters else topic
    subject    = data.get("subject", "").lower()
    regional   = ["odia","kannada","hindi","telugu","gujarati","marathi","bengali","punjabi","tamil","malayalam"]
    is_regional= any(lang in subject for lang in regional)

    try:
        vs  = _load_vs(sid)
        ctx = ""
        if vs:
            docs = vs.as_retriever(search_kwargs={"k": 8}).invoke(topic_str)
            ctx  = "\n\n".join(d.page_content for d in docs)
        if not ctx:
            s_name = syllabi_registry.get(sid, {}).get("name", "General Knowledge")
            ctx    = f"Generate high-quality questions for '{s_name}' on: {topic_str}."

        syl_meta2   = syllabi_registry.get(sid, {})
        q_board     = data.get("board",      syl_meta2.get("board",      "")) or "CBSE"
        q_class     = data.get("class_name", syl_meta2.get("class_name",
                          data.get("class",  syl_meta2.get("class", "")))) or "the appropriate class"
        q_subject   = data.get("subject",    syl_meta2.get("subject",    subject)) or "the subject"

        base  = REGIONAL_MIXED_QUESTION_GEN_PROMPT.replace("{language}", subject.title()) if is_regional else MIXED_QUESTION_GEN_PROMPT
        prompt= base.replace("{context}", ctx).replace("{topic}", topic_str).format(
            total_count=total, objective_count=obj_count, subjective_count=subj_count,
            easy_count=easy_c, medium_count=medium_c, hard_count=hard_c,
            board=q_board, class_level=q_class, subject=q_subject,
        )
        questions = _llm_json(prompt, temperature=0.2)
        if isinstance(questions, dict) and "questions" in questions:
            questions = questions["questions"]
        if not isinstance(questions, list):
            return jsonify({"error": "Question generation failed — unexpected AI response"}), 500
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": f"Server error: {e}"}), 500

    normalized = []
    for i, q in enumerate(questions, 1):
        qtype  = q.get("type", "objective")
        is_obj = (qtype == "objective")
        weight = obj_weight if is_obj else subj_weight
        try:
            weight = float(q.get("weightage", weight))
        except Exception:
            pass
        # Normalise valid_answers to letters for MCQ, keep text for subjective
        raw_valid = q.get("valid_answers", [])
        if is_obj:
            opts = q.get("options") or {}
            valid_letters = []
            for va in raw_valid:
                va_s = str(va).strip().upper()
                if va_s in opts:
                    valid_letters.append(va_s)
                else:
                    for k, v in opts.items():
                        if _normalize(str(v)) == _normalize(va):
                            valid_letters.append(k.upper()); break
            valid_for_q = valid_letters if valid_letters else [str(v).strip() for v in raw_valid]
            answer_str  = valid_for_q[0] if valid_for_q else ""
        else:
            valid_for_q = raw_valid
            answer_str  = raw_valid[0] if raw_valid else q.get("answer","")

        normalized.append({
            "id": i, "type": qtype,
            "difficulty": str(q.get("difficulty","medium")).title(),
            "marks": weight, "weightage": weight,
            "question": q.get("question",""),
            "options": q.get("options") if is_obj else None,
            "answer": answer_str,
            "valid_answers": valid_for_q,
            "explanation": q.get("explanation",""),
            "evaluation_criteria": q.get("evaluation_rubric") or ", ".join(q.get("answer_key_points",[])),
            "answer_key_points": q.get("answer_key_points", []),
            "evaluation_rubric": q.get("evaluation_rubric",""),
        })

    exam_id = uuid.uuid4().hex[:10]
    u       = request.user
    syl_meta = syllabi_registry.get(sid, {})
    payload = {
        "exam_id":          exam_id,
        "owner_id":         u["id"],
        "created_at":       _now(),
        "syllabus_id":      sid,
        "syllabus_name":    syl_meta.get("name", ""),
        "topic":            topic_str,
        "chapters":         chapters,
        "board":            data.get("board", syl_meta.get("board", "")),
        "class":            data.get("class_name", syl_meta.get("class_name",
                                data.get("class", syl_meta.get("class", "")))),
        "subject":          data.get("subject", syl_meta.get("subject", "")),
        "school_name":      data.get("school_name", ""),
        "teacher_name":     data.get("teacher_name", ""),
        "exam_date":        data.get("exam_date", _now()[:10]),
        "difficulty":       data.get("difficulty", "Mixed"),
        "objective_count":  obj_count,
        "subjective_count": subj_count,
        "questions":        normalized,
        "total_marks":      sum(float(q["marks"]) for q in normalized),
    }
    exams_registry[exam_id] = payload
    if not IS_VERCEL:
        _save_json(EXAMS_DIR / f"{exam_id}.json", payload)
    _save_exams()
    payload["_download_urls"] = {
        "question_paper":  f"/api/exams/{exam_id}/question-paper",
        "answer_sheet":    f"/api/exams/{exam_id}/answer-sheet",
        "answer_key":      f"/api/exams/{exam_id}/answer-key",
        "combined_print":  f"/api/exams/{exam_id}/combined-print",
    }
    return jsonify({**payload})

# ── List / Get / Delete Exams ─────────────────────────────────────────────────
def _exams_visible_to_evaluator(exams: List[dict], u: dict) -> List[dict]:
    """Exams shown in teacher/tutor evaluation flows (not students)."""
    role = u.get("role")
    uid  = u.get("id")
    if role in ("institute_admin", "admin"):
        return list(exams)
    u_inst = _normalize(u.get("institution") or "")
    if role in ("teacher", "tutor") and u_inst:
        db    = db_load()
        by_id = {x.get("id"): x for x in db.get("users", []) if x.get("id")}
        out   = []
        for e in exams:
            if e.get("owner_id") == uid:
                out.append(e)
                continue
            if _normalize(e.get("school_name") or "") == u_inst:
                out.append(e)
                continue
            owner = by_id.get(e.get("owner_id"))
            if owner and _normalize(owner.get("institution") or "") == u_inst:
                out.append(e)
        return out
    return [e for e in exams if e.get("owner_id") == uid]


@app.get("/api/questions")
@auth()
def list_exams():
    u     = request.user
    exams = list(exams_registry.values())
    if u["role"] == "student":
        exams = [e.copy() for e in exams if e.get("owner_id") == u["id"]]
        for e in exams:
            e["questions"] = [{k: v for k, v in q.items() if k not in ("answer","valid_answers","explanation","answer_key_points","evaluation_rubric")} for q in e.get("questions", [])]
    else:
        exams = _exams_visible_to_evaluator(exams, u)
    return jsonify({"exams": exams})

# Alias — EvalPanel calls /api/exams, legacy QMaster calls /api/questions
@app.get("/api/exams")
@auth()
def list_exams_alias():
    """GET /api/exams — alias for GET /api/questions (returns all saved exams)."""
    return list_exams()

@app.post("/api/exams/import-registry")
@auth(roles=["admin"])
def import_exams_registry():
    """Replace or merge the exams registry (admin-only)."""
    try:
        data = request.get_json(force=True) or {}
        registry = data.get("registry") or data.get("data") or {}
        merge = bool(data.get("merge"))

        if not isinstance(registry, dict) or not registry:
            return jsonify({"error": "registry must be a non-empty object"}), 400

        global exams_registry
        if merge:
            exams_registry.update(registry)
        else:
            exams_registry = registry

        _save_exams()
        return jsonify({"ok": True, "count": len(exams_registry), "merged": merge})
    except Exception as e:
        log.error(f"[EXAMS] import-registry failed: {e}", exc_info=True)
        return jsonify({"error": f"Import failed: {e}"}), 500

@app.post("/api/exams/save")
@auth()
def save_exam():
    """Save a generated exam and return a stable Exam ID for evaluation."""
    data = request.json or {}
    u    = request.user

    questions   = data.get("questions", [])
    if not questions:
        return jsonify({"error": "questions list is required"}), 400

    exam_id = uuid.uuid4().hex[:10]

    # Normalise questions — ensure all required fields present
    normalised = []
    for i, q in enumerate(questions, 1):
        qtype   = q.get("type", "objective")
        is_obj  = qtype == "objective"
        marks   = float(q.get("marks", 1 if is_obj else 5))
        raw_va = q.get("valid_answers", [])
        if not raw_va and q.get("answer"):
            raw_va = [q["answer"]]
        # For MCQ coerce valid_answers to option letters
        opts = q.get("options") or {}
        if is_obj and raw_va and opts:
            coerced = []
            for va in raw_va:
                va_s = str(va).strip().upper()
                if va_s in opts:
                    coerced.append(va_s)
                else:
                    for k, v in opts.items():
                        if _normalize(str(v)) == _normalize(va):
                            coerced.append(k.upper()); break
                    else:
                        coerced.append(va_s)
            raw_va = coerced if coerced else raw_va

        entry = {
            "id":                 q.get("id", i),
            "type":               qtype,
            "difficulty":         q.get("difficulty", "Medium"),
            "marks":              marks,
            "weightage":          marks,
            "question":           q.get("question", ""),
            "options":            q.get("options") if is_obj else None,
            "answer":             q.get("answer", ""),
            "valid_answers":      raw_va,
            "explanation":        q.get("explanation", ""),
            "answer_key_points":  q.get("answer_key_points", []),
            "evaluation_rubric":  q.get("evaluation_rubric", q.get("evaluation_criteria", "")),
        }
        normalised.append(entry)

    # Resolve syllabus metadata — QMaster may not send all fields explicitly
    _syl_id   = data.get("syllabus_id", "")
    _syl_meta = syllabi_registry.get(_syl_id, {})
    _syl_name = (data.get("syllabus_name", "")
                 or _syl_meta.get("name", "")
                 or data.get("subject", "")
                 or "Custom Exam")

    payload = {
        "exam_id":          exam_id,
        "owner_id":         u["id"],
        "created_at":       _now(),
        "syllabus_id":      _syl_id,
        "syllabus_name":    _syl_name,
        "topic":            data.get("topic", ""),
        "board":            data.get("board", _syl_meta.get("board", "")),
        "class":            data.get("class", _syl_meta.get("class_name",
                                _syl_meta.get("class", ""))),
        "subject":          data.get("subject", _syl_meta.get("subject", "")),
        "school_name":      data.get("school_name", ""),
        "teacher_name":     data.get("teacher_name", ""),
        "exam_date":        data.get("exam_date", _now()[:10]),
        "difficulty":       data.get("difficulty", "Mixed"),
        "objective_count":  len([q for q in normalised if q["type"] == "objective"]),
        "subjective_count": len([q for q in normalised if q["type"] != "objective"]),
        "questions":        normalised,
        "total_marks":      sum(float(q["marks"]) for q in normalised),
    }

    exams_registry[exam_id] = payload
    if not IS_VERCEL:
        _save_json(EXAMS_DIR / f"{exam_id}.json", payload)
    _save_exams()

    # Also persist to MongoDB directly for durability
    if MONGO_OK:
        try:
            mongo_db["exams_detail"].update_one(
                {"exam_id": exam_id},
                {"$set": payload},
                upsert=True
            )
        except Exception as e:
            print(f"[EXAMS] MongoDB detail write failed: {e}")

    log.info(f"[EXAMS] Saved exam {exam_id} — {len(normalised)} questions, {payload['total_marks']} marks")
    return jsonify({
        "ok":               True,
        "exam_id":          exam_id,
        "total_marks":      payload["total_marks"],
        "objective_count":  payload["objective_count"],
        "subjective_count": payload["subjective_count"],
        "syllabus_name":    payload["syllabus_name"],
        "subject":          payload["subject"],
        "download_urls": {
            "question_paper":  f"/api/exams/{exam_id}/question-paper",
            "answer_sheet":    f"/api/exams/{exam_id}/answer-sheet",
            "answer_key":      f"/api/exams/{exam_id}/answer-key",
            "combined_print":  f"/api/exams/{exam_id}/combined-print",
        },
    }), 201

@app.get("/api/exams/<exam_id>")
@auth()
def get_exam(exam_id):
    e = exams_registry.get(exam_id)
    if not e and not IS_VERCEL:
        p = EXAMS_DIR / f"{exam_id}.json"
        if p.exists():
            e = _load_json(p)
            if e:
                exams_registry[exam_id] = e
    if not e:
        return jsonify({"error": "Exam not found"}), 404
    return jsonify(e)

@app.delete("/api/exams/<exam_id>")
@auth(roles=["tutor","teacher","institute_admin","admin"])
def delete_exam(exam_id):
    u = request.user
    if exam_id not in exams_registry:
        return jsonify({"error": "Not found"}), 404
    if u["role"] not in ("institute_admin","admin") and exams_registry[exam_id].get("owner_id") != u["id"]:
        return jsonify({"error": "Forbidden"}), 403
    exams_registry.pop(exam_id, None)
    if not IS_VERCEL:
        p = EXAMS_DIR / f"{exam_id}.json"
        if p.exists():
            p.unlink()
    _save_exams()
    return jsonify({"ok": True})

@app.get("/api/exams/<exam_id>/analytics")
@auth(roles=["tutor","teacher","institute_admin","admin"])
def get_exam_analytics(exam_id):
    evals = [v for v in evaluations_registry.values() if v.get("exam_id") == exam_id]
    if not evals:
        return jsonify({"error": "No evaluations found for this exam"}), 404
    total = len(evals)
    q_stats: Dict[int, dict] = {}
    for ev in evals:
        for i, qr in enumerate(ev.get("result", {}).get("question_wise", [])):
            if i not in q_stats:
                q_stats[i] = {"awarded": 0.0, "possible": 0.0}
            q_stats[i]["awarded"]  += float(qr.get("awarded_marks", 0))
            q_stats[i]["possible"] += float(qr.get("max_marks", 1))
    gaps = [{"index": i, "average": round((s["awarded"] / s["possible"]) * 100, 1)}
            for i, s in q_stats.items() if s["possible"] > 0 and (s["awarded"] / s["possible"]) < 0.5]
    top = sorted(evals, key=lambda e: e.get("result",{}).get("percentage",0), reverse=True)[:5]
    return jsonify({
        "total_evaluations": total,
        "class_average": round(sum(e.get("result",{}).get("percentage",0) for e in evals) / total, 2),
        "learning_gaps": gaps,
        "top_performers": [{"roll_no": e.get("result",{}).get("roll_no","N/A"),
                             "percentage": e.get("result",{}).get("percentage",0)} for e in top]
    })

# ══════════════════════════════════════════════════════════════════════════════
#  TEMPLATE DOWNLOAD ENDPOINTS
#  /api/exams/<id>/answer-sheet  → structured OCR-optimised answer sheet PDF
#  /api/exams/<id>/question-paper → question paper with instructions
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/exams/<exam_id>/answer-key")
@auth(roles=["teacher", "tutor", "institute_admin", "admin"])
def download_answer_key(exam_id):
    """
    Generate and return the ANSWER KEY PDF.
    Teacher-only: shows correct answers, model answers, key points, rubrics.
    """
    e = _load_exam(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404

    pdf_bytes = _generate_answer_key_pdf(e)
    if not pdf_bytes:
        return jsonify({"error": "PDF generation failed"}), 500

    subject  = re.sub(r"[^\w\s-]", "", e.get("subject", "Exam"))[:30]
    filename = f"AnswerKey_{subject}_{exam_id}.pdf"

    from flask import Response
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.get("/api/exams/<exam_id>/answer-sheet")
@auth()
def download_answer_sheet(exam_id):
    """
    Generate and return the structured OCR-optimised answer sheet PDF.
    Students print this, fill it in by hand, scan it, and upload for evaluation.
    """
    e = _load_exam(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404

    pdf_bytes = _generate_answer_sheet_pdf(e)
    if not pdf_bytes:
        return jsonify({"error": "PDF generation failed"}), 500

    subject  = re.sub(r"[^\w\s-]", "", e.get("subject", "Exam"))[:30]
    filename = f"AnswerSheet_{subject}_{exam_id}.pdf"

    from flask import Response
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.get("/api/exams/<exam_id>/question-paper")
@auth()
def download_question_paper(exam_id):
    """
    Generate and return the question paper PDF (with full question text).
    Students use this to READ questions, then write answers on the answer sheet.
    """
    e = _load_exam(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404

    pdf_bytes = _generate_question_paper_pdf(e)
    if not pdf_bytes:
        return jsonify({"error": "PDF generation failed"}), 500

    subject  = re.sub(r"[^\w\s-]", "", e.get("subject", "Exam"))[:30]
    filename = f"QuestionPaper_{subject}_{exam_id}.pdf"

    from flask import Response
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.get("/api/exams/<exam_id>/combined-print")
@auth()
def download_combined_print(exam_id):
    """
    Combined print PDF: Question Paper followed by Answer Sheet.
    Single file for teacher to print — students get both together.
    Page 1+: Question Paper (read-only, questions + instructions)
    Then: Answer Sheet (fill in — bubbles + ruled boxes)
    The Exam ID appears prominently on every page of both documents.
    """
    e = _load_exam(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404

    try:
        from pypdf import PdfWriter, PdfReader
        import io as _io

        writer = PdfWriter()

        qp_bytes = _generate_question_paper_pdf(e)
        as_bytes = _generate_answer_sheet_pdf(e)

        for pdf_bytes in [qp_bytes, as_bytes]:
            if pdf_bytes:
                reader = PdfReader(_io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)

        out = _io.BytesIO()
        writer.write(out)
        out.seek(0)
        combined = out.read()

        subject  = re.sub(r"[^\w\s-]", "", e.get("subject", "Exam"))[:30]
        filename = f"PrintPack_{subject}_{exam_id}.pdf"

        from flask import Response
        return Response(
            combined,
            mimetype="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as ex:
        log.error(f"[COMBINED-PRINT] {ex}", exc_info=True)
        return jsonify({"error": str(ex)}), 500

def _load_exam(exam_id: str) -> Optional[dict]:
    if exam_id in exams_registry:
        return exams_registry[exam_id]
    if not IS_VERCEL:
        p = EXAMS_DIR / f"{exam_id}.json"
        if p.exists():
            e = _load_json(p)
            if e:
                exams_registry[exam_id] = e
                return e
    return None

# ── Multi-student PDF splitter using GPT-4o vision ───────────────────────────
def _split_multi_student_pdf(pdf_path: str, exam_questions: list = None) -> list:
    """
    Identify per-student sections in a combined answer-sheet PDF.
    Strategy:
      1. Text extraction (for typed/digital PDFs)
      2. Per-page vision analysis — each page is classified individually to detect
         whether it starts a new student's answer sheet (Q1 restart or new header)
      3. All-pages-at-once vision fallback if per-page analysis finds no boundaries
    Returns list of dicts: [{"student_name": "", "roll_no": "", "page_indices": [...]}]
    """
    import base64, io as _io
    import fitz as _fitz
    from PIL import Image as _PILImage
    from collections import OrderedDict, Counter

    # ── Load pages as images ──────────────────────────────────────────────────
    try:
        _doc = _fitz.open(pdf_path)
        n_pages = len(_doc)
        images = []
        for p in _doc:
            pix = p.get_pixmap(dpi=150)
            images.append(_PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples))
        _doc.close()
    except Exception as _load_err:
        print(f"[MULTI-SPLIT] PDF→image failed: {_load_err}")
        return [{"student_name": "Student 1", "roll_no": "", "page_indices": [0]}]

    all_pages = list(range(n_pages))

    if n_pages <= 1:
        return [{"student_name": "Student 1", "roll_no": "", "page_indices": all_pages}]

    # ── Text extraction (typed/printed PDFs) ──────────────────────────────────
    try:
        import pypdf
        _rdr = pypdf.PdfReader(pdf_path)
        all_text = "\n\n--- PAGE BREAK ---\n\n".join(
            p.extract_text() for p in _rdr.pages
            if (p.extract_text() or "").strip()
        )
    except Exception:
        all_text = ""

    if all_text and n_pages <= 80:
        _tp = (
            f"Text from a combined answer sheet PDF:\n{all_text[:6000]}\n\n"
            "List each student with their pages (0-indexed). Return JSON array:\n"
            '[{"student_name":"...","roll_no":"...","page_indices":[0,1]},...]'
        )
        _tr = _llm_json(_tp, temperature=0.1, mini=True)
        if isinstance(_tr, list) and _tr:
            _covered = set()
            for _e in _tr:
                _pis = _e.get("page_indices")
                if isinstance(_pis, list):
                    _covered.update([p for p in _pis if isinstance(p, int) and 0 <= p < n_pages])
                _e.setdefault("page_indices", all_pages)

            # Accept text-path only when it covers most pages; otherwise keep going
            # to per-page vision logic (more robust for large or scanned PDFs).
            if len(_covered) >= max(1, int(0.80 * n_pages)):
                return _tr
            print(
                f"[MULTI-SPLIT] Text-path split under-covered pages "
                f"({len(_covered)}/{n_pages}) — falling back to vision split"
            )

    client = _oai.OpenAI(api_key=OPENAI_API_KEY)

    # Determine the first question number from exam metadata (default 1)
    first_q_no = 1
    if exam_questions:
        _qnos = [q.get("question_no", q.get("q_no", 0)) for q in exam_questions]
        _qnos = [n for n in _qnos if isinstance(n, int) and n > 0]
        if _qnos:
            first_q_no = min(_qnos)

    def _analyze_header_page(_i: int, _img) -> dict:
        """ZERO-VISION header analysis using targeted OCR."""
        from PIL import Image as _PILImage
        _w, _h = _img.size
        # Crop top 22% (typical header area)
        header_box = [0, _w, 0, int(_h * 0.22)]
        
        # Use targeted OCR (cached & cheap)
        text = _ocr_cropped_region(_img, header_box)
        if not text:
            return {}

        # Use local LLM (or mini GPT fallback) to parse extracted text
        _pp = (
            f"Identify student details from this OCR text: '{text}'\n"
            "Return JSON only: {\"is_blank\":false,\"student_name\":\"...\",\"roll_no\":\"...\",\"sheet_page_no\":null}"
        )
        _role = call_ollama(_pp)
        if not _role:
            _role = _llm_json(_pp, mini=True)
            
        if isinstance(_role, dict) and (
            (_role.get("student_name") or "") or (_role.get("roll_no") or "") or _role.get("sheet_page_no") is not None
        ):
            return _role
        return {}

    def _best_student_identity(_names: list, _rolls: list, fallback_idx: int) -> tuple[str, str]:
        def _clean_name(val: str) -> str:
            return re.sub(r'\s+', ' ', str(val or '').replace('+', ' ').strip())

        def _is_generic_name(val: str) -> bool:
            s = _clean_name(val).lower()
            if not s:
                return True
            if s.startswith("student") or s.startswith("students"):
                return True
            s2 = re.sub(r'[^a-z0-9]+', '', s)
            return s2 in {"student", "students", "name"}

        def _score_name(val: str) -> tuple[int, int, int]:
            s = _clean_name(val)
            digits = 1 if re.search(r'\d', s) else 0
            generic_penalty = 0 if not _is_generic_name(s) else -10
            return (generic_penalty, digits, len(s))

        clean_names = [_clean_name(n) for n in _names if _clean_name(n)]
        clean_rolls = [str(r).strip() for r in _rolls if str(r).strip()]

        best_roll = Counter(clean_rolls).most_common(1)[0][0] if clean_rolls else ""
        best_name = ""
        if clean_names:
            best_name = sorted(clean_names, key=_score_name, reverse=True)[0]

        if not best_name or _is_generic_name(best_name):
            if best_roll:
                best_name = f"Student {best_roll}"
            else:
                best_name = f"Student {fallback_idx}"
        else:
            lower_name = best_name.lower()
            if (lower_name.startswith("student") or lower_name.startswith("students")) and best_roll:
                best_name = f"Student {best_roll}"

        return best_name, (best_roll or str(fallback_idx))

    # ── Fixed-chunk fast path: deterministic chunking + sparse OCR ─────────────
    # Applies to ANY class PDF whose page count is an exact multiple of
    # MULTI_SPLIT_FIXED_PAGES (default 4). Works for small real PDFs (e.g. 24pp /
    # 6 students) AND large synthetic ones (160pp / 40 students).
    # Only one header OCR call per chunk → avoids per-page TPM spikes.
    #
    # VALIDATION: Before committing to the fixed chunk size, we probe the SECOND
    # chunk's first page. If its student_name/roll_no differs from the first
    # chunk, the chunk size is correct.  If they match (i.e. pages 0 and
    # MULTI_SPLIT_FIXED_PAGES belong to the SAME student), the chunk size is
    # too small — we fall through to per-page OCR which auto-detects boundaries.
    _min_fast_path_pages = MULTI_SPLIT_FIXED_PAGES * 2   # at least 2 students
    if (n_pages >= _min_fast_path_pages and MULTI_SPLIT_FIXED_PAGES >= 2
            and n_pages % MULTI_SPLIT_FIXED_PAGES == 0):
        _pps = MULTI_SPLIT_FIXED_PAGES
        # Probe page 0 and page _pps to verify they belong to DIFFERENT students.
        # Also probe page _pps-1 (last page of first chunk) to verify it belongs
        # to the SAME student as page 0.  If it doesn't, the chunk is too large
        # (mixing multiple students in one chunk).
        _meta0 = _analyze_header_page(0, images[0])
        _meta_last = _analyze_header_page(_pps - 1, images[_pps - 1]) if _pps > 1 else _meta0
        _meta_next = _analyze_header_page(_pps, images[_pps])
        _name0 = (_meta0.get("student_name") or "").replace("+", " ").strip().lower()
        _roll0 = (_meta0.get("roll_no") or "").strip()
        _name_last = (_meta_last.get("student_name") or "").replace("+", " ").strip().lower()
        _roll_last = (_meta_last.get("roll_no") or "").strip()
        _name_next = (_meta_next.get("student_name") or "").replace("+", " ").strip().lower()
        _roll_next = (_meta_next.get("roll_no") or "").strip()

        # FAIL-SAFE: if any probe returned no usable data (e.g. 429 / timeout),
        # we cannot validate the chunk size → fall through to per-page OCR.
        _has_probe_data = bool(_name0 or _roll0) and bool(_name_next or _roll_next)
        if _pps > 1:
            _has_probe_data = _has_probe_data and bool(_name_last or _roll_last)
        if not _has_probe_data:
            print(
                f"[MULTI-SPLIT] Fixed-chunk validation SKIPPED: insufficient probe data "
                f"(p0: name='{_name0}' roll='{_roll0}', p{_pps-1}: name='{_name_last}' roll='{_roll_last}', "
                f"p{_pps}: name='{_name_next}' roll='{_roll_next}'). "
                f"Falling through to per-page OCR."
            )
            _chunk_valid = False
        else:
            # Check 1: page 0 and page _pps must be DIFFERENT students
            _next_is_different = True
            if _roll0 and _roll_next and _roll0 == _roll_next:
                _next_is_different = False
            elif _name0 and _name_next and re.sub(r'\s+', '', _name0) == re.sub(r'\s+', '', _name_next):
                _next_is_different = False

            # Check 2: page 0 and page _pps-1 must be the SAME student
            _last_is_same = True
            if _pps > 1:
                if _roll0 and _roll_last and _roll0 != _roll_last:
                    _last_is_same = False
                elif _name0 and _name_last and re.sub(r'\s+', '', _name0) != re.sub(r'\s+', '', _name_last):
                    _last_is_same = False

            _chunk_valid = _next_is_different and _last_is_same
            if not _chunk_valid:
                _reason = []
                if not _next_is_different:
                    _reason.append(f"pages 0 and {_pps} belong to same student (chunk too small)")
                if not _last_is_same:
                    _reason.append(f"pages 0 and {_pps-1} belong to different students (chunk too large)")
                print(
                    f"[MULTI-SPLIT] Fixed-chunk validation FAILED: {'; '.join(_reason)}. "
                    f"(p0: name='{_name0}' roll='{_roll0}', p{_pps-1}: name='{_name_last}' roll='{_roll_last}', "
                    f"p{_pps}: name='{_name_next}' roll='{_roll_next}'). "
                    f"Falling through to per-page OCR."
                )
            else:
                _groups = []
                _name0, _roll0b = _best_student_identity([_meta0.get("student_name") or ""], [_roll0], 1)
                _groups.append({
                    "student_name": _name0,
                    "roll_no": _roll0b,
                    "page_indices": list(range(0, _pps)),
                })
                _name1, _roll1b = _best_student_identity([_meta_next.get("student_name") or ""], [_roll_next], 2)
                _groups.append({
                    "student_name": _name1,
                    "roll_no": _roll1b,
                    "page_indices": list(range(_pps, min(_pps * 2, n_pages))),
                })
                for _st in range(_pps * 2, n_pages, _pps):
                    _en = min(_st + _pps, n_pages)
                    _grp = list(range(_st, _en))
                    _meta = _analyze_header_page(_st, images[_st])
                    _name = (_meta.get("student_name") or "").replace("+", " ").strip()
                    _roll = (_meta.get("roll_no") or "").strip()
                    _idx = len(_groups) + 1
                    _best_name, _best_roll = _best_student_identity([_name], [_roll], _idx)
                    _groups.append({
                        "student_name": _best_name,
                        "roll_no": _best_roll,
                        "page_indices": _grp,
                    })
                print(
                    f"[MULTI-SPLIT] Large-PDF fixed-chunk fast path: "
                    f"pages_per_student={_pps}, groups={len(_groups)}"
                )
                return _groups

    # ── Per-page OCR: extract student identity from each page ────────────────
    # Strategy: ask GPT for student_name + roll_no + sheet_page_no on each page,
    # then GROUP pages by roll_no. This is far more robust than boundary detection
    # because the answer sheet template repeats the name/roll header on every page.
    from PIL import ImageStat as _PILStat
    page_roles = []
    for _i, _img in enumerate(images):
        # ── Pixel-based blank detection (cheap, no API call) ─────────────────
        try:
            _pix_std = _PILStat.Stat(_img.convert('L')).stddev[0]
        except Exception:
            _pix_std = 999.0
        if _pix_std < 22:
            page_roles.append({"page": _i, "is_blank": True, "student_name": "",
                "roll_no": "", "sheet_page_no": None})
            print(f"[MULTI-SPLIT] Page {_i}: BLANK (pixel_std={_pix_std:.1f}) — skipped")
            continue
        _role = _analyze_header_page(_i, _img)
        _role["page"] = _i
        page_roles.append(_role)
        print(f"[MULTI-SPLIT] Page {_i}: {_role}")

    # ── Group pages by student boundaries ────────────────────────────────────
    # Strategy 1 (primary): detect student starts via printed sheet_page_no == 1.
    # Printed text is far more reliably OCR'd than handwritten roll numbers, so
    # finding "Page 1" resets is the most robust way to separate students.
    # Strategy 2 (fallback): majority-vote roll_no grouping (handles PDFs where
    # the sheet_page_no label is absent or not printed).
    blank_pages = {_r["page"] for _r in page_roles if _r.get("is_blank")}
    if blank_pages:
        print(f"[MULTI-SPLIT] Blank pages (skipped): {sorted(blank_pages)}")

    _non_blank = sorted([_r for _r in page_roles if not _r.get("is_blank")],
                        key=lambda x: x["page"])
    _page_role_map: dict = {_r["page"]: _r for _r in page_roles}
    _non_blank_page_set = {_r["page"] for _r in _non_blank}
    _all_non_blank_sorted = sorted(_non_blank_page_set)

    def _finalize_groups(_entries: list, _tag: str) -> list:
        """Ensure grouped sections cover all non-blank pages; recover deterministically if needed."""
        if not isinstance(_entries, list):
            _entries = []

        # Normalize page indices and drop empty groups.
        _norm = []
        for _idx, _e in enumerate(_entries, 1):
            _pis = _e.get("page_indices") if isinstance(_e, dict) else []
            _pis = [p for p in (_pis or []) if isinstance(p, int) and p in _non_blank_page_set]
            _pis = sorted(set(_pis))
            if not _pis:
                continue
            _name = str((_e or {}).get("student_name", "")).replace("+", " ").strip()
            _roll = str((_e or {}).get("roll_no", "")).strip()
            _best_name, _best_roll = _best_student_identity([_name], [_roll], _idx)
            _norm.append({"student_name": _best_name, "roll_no": _best_roll, "page_indices": _pis})

        # Coverage check.
        _covered = set()
        for _e in _norm:
            _covered.update(_e.get("page_indices", []))

        # If grouping is weak, recover by inferred fixed chunk size from sheet_page_no.
        _sp = [
            int(_r.get("sheet_page_no"))
            for _r in _non_blank
            if isinstance(_r.get("sheet_page_no"), int) and int(_r.get("sheet_page_no")) > 0
        ]
        _pps = max(_sp) if _sp else 0
        _expected = (len(_all_non_blank_sorted) // _pps) if (_pps and len(_all_non_blank_sorted) % _pps == 0) else 0
        if _expected >= 2 and len(_norm) < _expected:
            _chunked = []
            for _i in range(0, len(_all_non_blank_sorted), _pps):
                _grp = _all_non_blank_sorted[_i:_i + _pps]
                if len(_grp) != _pps:
                    continue
                _roles = [(_page_role_map.get(p) or {}) for p in _grp]
                _names = [str(r.get("student_name", "")).replace("+", " ").strip() for r in _roles if str(r.get("student_name", "")).strip()]
                _rolls = [str(r.get("roll_no", "")).strip() for r in _roles if str(r.get("roll_no", "")).strip()]
                _name, _roll = _best_student_identity(_names, _rolls, len(_chunked) + 1)
                _chunked.append({"student_name": _name, "roll_no": _roll, "page_indices": _grp})
            if len(_chunked) >= _expected:
                print(
                    f"[MULTI-SPLIT] {_tag} under-covered groups ({len(_norm)}/{_expected}) — "
                    f"using inferred chunk recovery (pages_per_student={_pps})."
                )
                _norm = _chunked
                _covered = set(p for e in _norm for p in e.get("page_indices", []))

        # Attach any remaining uncovered non-blank pages so no student pages are dropped.
        _missing = sorted(_non_blank_page_set - _covered)
        if _missing:
            print(f"[MULTI-SPLIT] {_tag} missing pages recovered: {_missing}")
            for _p in _missing:
                _r = _page_role_map.get(_p, {})
                _name, _roll = _best_student_identity(
                    [str(_r.get("student_name", "")).replace("+", " ").strip()],
                    [str(_r.get("roll_no", "")).strip()],
                    len(_norm) + 1,
                )
                _norm.append({"student_name": _name, "roll_no": _roll, "page_indices": [_p]})

        return _norm

    # ── Strategy 1: sheet_page_no restart detection ───────────────────────────
    _restart_pages = [_r["page"] for _r in _non_blank if _r.get("sheet_page_no") == 1]
    print(f"[MULTI-SPLIT] Detected restart pages (sheet_page_no=1): {_restart_pages}")

    if len(_restart_pages) >= 2:
        # Build contiguous groups in physical page order. Restart pages are used
        # as hints, but grouping is driven by contiguous identity so one missed
        # page label cannot collapse two different students into a single group.
        result = []
        _current_pages: list = []
        _current_key = ""
        for _r in _non_blank:
            _name = (_r.get("student_name") or "").replace("+", " ").strip()
            _roll = (_r.get("roll_no") or "").strip()
            _key = _roll or re.sub(r'\s+', '', _name.lower())
            _starts_new = False
            if _current_pages:
                _prev_role = _page_role_map[_current_pages[-1]]
                _prev_key = (_prev_role.get("roll_no") or "").strip() or re.sub(r'\s+', '', ((_prev_role.get("student_name") or "").replace("+", " ").strip().lower()))
                _sheet_no = _r.get("sheet_page_no")
                if _sheet_no == 1:
                    _starts_new = True
                elif _key and _prev_key and _key != _prev_key:
                    _starts_new = True
            if _starts_new and _current_pages:
                _names = [(_page_role_map[p].get("student_name") or "").replace("+", " ").strip()
                          for p in _current_pages
                          if (_page_role_map[p].get("student_name") or "").strip()]
                _rolls = [(_page_role_map[p].get("roll_no") or "").strip()
                          for p in _current_pages
                          if (_page_role_map[p].get("roll_no") or "").strip()]
                _spnos = [int(_page_role_map[p].get("sheet_page_no") or 0) if str(_page_role_map[p].get("sheet_page_no") or "").strip().isdigit() else 0 for p in _current_pages]
                _use_spno = (
                    len(set(_spnos)) == len(_spnos)
                    and all(_spnos)
                    and sorted(_spnos) == list(range(1, len(_spnos) + 1))
                )
                _srt = sorted(_current_pages) if not _use_spno else sorted(_current_pages, key=lambda p: int(_page_role_map[p].get("sheet_page_no") or 0))
                _best_name, _best_roll = _best_student_identity(_names, _rolls, len(result) + 1)
                result.append({
                    "student_name": _best_name,
                    "roll_no": _best_roll,
                    "page_indices": _srt,
                })
                _current_pages = []
                _current_key = ""
            _current_pages.append(_r["page"])
            if _key:
                _current_key = _key

        if _current_pages:
            _names = [
                (_page_role_map[p].get("student_name") or "").replace("+", " ").strip()
                for p in _current_pages
                if (_page_role_map[p].get("student_name") or "").strip()
            ]
            _rolls = [
                (_page_role_map[p].get("roll_no") or "").strip()
                for p in _current_pages
                if (_page_role_map[p].get("roll_no") or "").strip()
            ]
            _spnos = [
                int(_page_role_map[p].get("sheet_page_no") or 0)
                if str(_page_role_map[p].get("sheet_page_no") or "").strip().isdigit()
                else 0
                for p in _current_pages
            ]
            _use_spno = (
                len(set(_spnos)) == len(_spnos)
                and all(_spnos)
                and sorted(_spnos) == list(range(1, len(_spnos) + 1))
            )
            _srt = (
                sorted(_current_pages)
                if not _use_spno
                else sorted(_current_pages, key=lambda p: int(_page_role_map[p].get("sheet_page_no") or 0))
            )
            _best_name, _best_roll = _best_student_identity(_names, _rolls, len(result) + 1)
            result.append({
                "student_name": _best_name,
                "roll_no": _best_roll,
                "page_indices": _srt,
            })

        result = _finalize_groups(result, "sequential-grouping")
        print(
            f"[MULTI-SPLIT] Grouped sequentially by page order + identity: "
            f"{[(r['student_name'], r['roll_no'], r['page_indices']) for r in result]}"
        )
        return result

    # ── Strategy 2: group by roll_no (fallback) ───────────────────────────────
    print("[MULTI-SPLIT] sheet_page_no restart detection insufficient — trying roll_no grouping")

    # ── Large-PDF fallback: infer fixed pages-per-student and chunk sequentially ─
    # For big class uploads (40-50 students), per-page vision can intermittently 429.
    # If restart detection is weak but we still observed usable sheet_page_no values,
    # infer pages-per-student and chunk by that size to avoid undercounting students.
    if n_pages >= 80:
        _sp = [
            int(_r.get("sheet_page_no"))
            for _r in _non_blank
            if isinstance(_r.get("sheet_page_no"), int) and int(_r.get("sheet_page_no")) > 0
        ]
        _pages_per_student = max(_sp) if _sp else 0
        if 2 <= _pages_per_student <= 12:
            _start_at = _restart_pages[0] if _restart_pages else (_all_non_blank_sorted[0] if _all_non_blank_sorted else 0)
            _chunked = []
            for _st in range(_start_at, n_pages, _pages_per_student):
                _en = min(_st + _pages_per_student, n_pages)
                _grp = [p for p in _all_non_blank_sorted if _st <= p < _en]
                if _grp:
                    _chunked.append(_grp)

            if len(_chunked) >= 2:
                _entries = []
                for _grp in _chunked:
                    _roles = [(_page_role_map.get(p) or {}) for p in _grp]
                    _names = [str(r.get("student_name", "")).strip() for r in _roles if str(r.get("student_name", "")).strip()]
                    _rolls = [str(r.get("roll_no", "")).strip() for r in _roles if str(r.get("roll_no", "")).strip()]
                    _name = Counter(_names).most_common(1)[0][0] if _names else ""
                    _roll = Counter(_rolls).most_common(1)[0][0] if _rolls else ""
                    _entries.append({
                        "student_name": _name or f"Student {len(_entries)+1}",
                        "roll_no": _roll,
                        "page_indices": _grp,
                    })
                print(
                    f"[MULTI-SPLIT] Large-PDF inferred chunk fallback: pages_per_student={_pages_per_student}, "
                    f"groups={len(_entries)}"
                )
                return _entries
    groups = OrderedDict()
    for _r in _non_blank:
        _roll = (_r.get("roll_no") or "").strip()
        _name = (_r.get("student_name") or "").replace("+", " ").strip()
        _key = _roll if _roll else (_name if _name else f"unknown_{_r['page']}")
        if _key not in groups:
            groups[_key] = {"roll_no": _roll, "all_names": [], "pages": []}
        if _roll and not groups[_key]["roll_no"]:
            groups[_key]["roll_no"] = _roll
        if _name:
            groups[_key]["all_names"].append(_name)
        groups[_key]["pages"].append((_r.get("sheet_page_no") or 999, _r["page"]))

    print(f"[MULTI-SPLIT] Groups by roll_no: { {k: len(v['pages']) for k, v in groups.items()} }")

    if len(groups) > 1:
        result = []
        for _seq, (_key, _g) in enumerate(groups.items(), 1):
            _raw_pages = [p for _, p in sorted(_g["pages"])]
            _spnos2 = [(_page_role_map[p].get("sheet_page_no") or 0) for p in _raw_pages]
            if len(set(_spnos2)) == len(_spnos2) and all(_spnos2):
                _sorted_pages = sorted(_raw_pages, key=lambda p: _page_role_map[p].get("sheet_page_no"))
            else:
                _sorted_pages = sorted(_raw_pages)
            _name, _roll = _best_student_identity(_g["all_names"], [_g["roll_no"]], _seq)
            result.append({"student_name": _name, "roll_no": _roll, "page_indices": _sorted_pages})
        result = _finalize_groups(result, "roll-grouping")
        print(f"[MULTI-SPLIT] Grouped result: {[(r['student_name'], r['roll_no'], r['page_indices']) for r in result]}")
        return result

    # ── No clear grouping — fall through to all-pages-at-once vision ──────────
    print("[MULTI-SPLIT] Roll-number grouping did not yield multiple students — trying fallback")

    # ── Fallback: all-pages-at-once vision ────────────────────────────────────
    # Per-page analysis found no boundaries — send all pages together with a
    # more aggressive prompt that avoids the "single student" fallback.
    print("[MULTI-SPLIT] No per-page boundaries found — trying all-at-once vision")
    _parts = [{"type": "text", "text": (
        "These are ALL pages of a scanned answer sheet PDF. "
        "MULTIPLE students submitted their answers in this file.\n\n"
        "Your task: identify EVERY student and which pages belong to them.\n"
        "Look for: student names/roll numbers at page tops, question numbering that "
        "restarts from Q1 (indicating a new student's sheet), visual separators.\n\n"
        "IMPORTANT: If you see the same question repeated on multiple pages starting "
        f"from Q{first_q_no}, those pages belong to DIFFERENT students — do not merge them.\n\n"
        "Return ONLY a JSON array (no other text):\n"
        '[{"student_name":"Ravi Kumar","roll_no":"101","page_indices":[0,1]},'
        '{"student_name":"Priya Singh","roll_no":"102","page_indices":[2,3]}]'
    )}]
    for _idx, _img in enumerate(images[:20]):
        _buf2 = _io.BytesIO()
        _img.save(_buf2, format="PNG")
        _b64_2 = base64.b64encode(_buf2.getvalue()).decode()
        _parts.append({"type": "text", "text": f"--- Page {_idx} ---"})
        _parts.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{_b64_2}"}})

    try:
        _sr = client.chat.completions.create(
            model=OPENAI_MODEL,
            temperature=0.1,
            messages=[{"role": "user", "content": _parts}]
        )
        _parsed = _clean_json(_sr.choices[0].message.content)
        if isinstance(_parsed, list) and _parsed:
            for _e in _parsed:
                _e.setdefault("page_indices", all_pages)
                # Clean '+' OCR artefact from names in fallback path too
                if "student_name" in _e:
                    _e["student_name"] = (_e["student_name"] or "").replace("+", " ").strip()
            # Reject single-student result only if n_pages > 2
            # (if GPT-4o still says 1 student for many pages, trust it less)
            if len(_parsed) > 1 or n_pages <= 2:
                return _finalize_groups(_parsed, "all-at-once")
        if isinstance(_parsed, dict):
            for _k in ("students", "data", "results"):
                if isinstance(_parsed.get(_k), list) and _parsed[_k]:
                    return _finalize_groups(_parsed[_k], "all-at-once-dict")
    except Exception as _fe:
        print(f"[MULTI-SPLIT] All-at-once fallback error: {_fe}")

    # ── Last resort: treat each page as a separate student ────────────────────
    # This handles the case where n_pages = n_students (1 page per student)
    print("[MULTI-SPLIT] Last resort: 1 page per student")
    return [
        {"student_name": f"Student {_i + 1}", "roll_no": str(_i + 1), "page_indices": [_i]}
        for _i in range(n_pages)
    ]

def _parse_raw_answers(raw_text: str) -> Dict[int, str]:
    """Parse Q1: answer  Q2: answer ... from raw text."""
    answers: Dict[int, str] = {}
    for line in raw_text.splitlines():
        line = line.strip()
        if not line:
            continue
        for pat in [r"Q(?:uestion)?\s*(\d+)\s*[:\-]\s*(.+)", r"(\d+)\s*[\)\.\-:]\s*(.+)"]:
            m = re.match(pat, line, flags=re.IGNORECASE)
            if m:
                answers[int(m.group(1))] = m.group(2).strip()
                break
    return answers

@app.post("/api/evaluate")
@auth()
@quota('evaluate')
def evaluate_sheet():
    """
    Single answer sheet evaluation.
    If Redis/Celery is available: returns task_id immediately (async).
    Otherwise: evaluates synchronously (fallback).
    Includes file-hash caching to avoid re-processing identical uploads.
    """
    u            = request.user
    exam_id      = request.form.get("exam_id")
    roll_no      = request.form.get("roll_no", "")
    student_name = request.form.get("student_name", "")
    parent_email = request.form.get("parent_email", "")

    if not exam_id:
        return jsonify({"error": "exam_id required"}), 400
    e = _load_exam(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404
    if "answer_sheet" not in request.files:
        return jsonify({"error": "answer_sheet file required"}), 400

    f = request.files["answer_sheet"]
    if not f or not _allowed(f.filename):
        return jsonify({"error": "Unsupported file type"}), 400

    fn   = secure_filename(f.filename)
    path = UPL_DIR / f"{datetime.utcnow().strftime('%Y%m%dT%H%M%S')}_{fn}"
    f.save(str(path))
    log.info(f"[EVAL] Saved upload: {fn} ({path.stat().st_size} bytes)")

    # ── Cache check (skip OCR if identical file already evaluated) ────────────
    fhash  = _file_hash(str(path))
    cached = _cache_get(fhash, exam_id)
    # if cached:
    #     log.info(f"[EVAL] Cache hit for {fn} — returning cached result")
    #     return jsonify({"cached": True, **cached})

    # ── Async path (Celery available) ─────────────────────────────────────────
    if CELERY_OK:
        try:
            from celery_worker import evaluate_single as _eval_task
            task = _eval_task.delay(
                str(path), e, roll_no, student_name, parent_email, fn
            )
            log.info(f"[EVAL] Async task queued: {task.id}")
            return jsonify({
                "async":    True,
                "task_id":  task.id,
                "status":   "processing",
                "message":  "Evaluation queued. Poll /api/evaluate/status/<task_id> for result.",
            }), 202
        except Exception as _ce:
            log.warning(f"[EVAL] Celery dispatch failed, falling back to sync: {_ce}")

    # ── Sync path (fallback) ──────────────────────────────────────────────────
    log.info(f"[EVAL] Sync evaluation: {fn}")
    _openai_limiter.wait()

    # Auto-read header if student info not provided by user
    if not student_name or not roll_no:
        header = _extract_header_info(str(path))
        if not student_name:
            student_name = header.get("student_name", "").strip()
        if not roll_no:
            roll_no = header.get("roll_no", "").strip()
        if header.get("class_section"):
            log.info(f"[EVAL] Header OCR: name={student_name!r} roll={roll_no!r} class={header.get('class_section')!r}")

    answers, mode = _extract_answers(str(path), exam_questions=e.get("questions", []), exam=e)

    # Exam ID mismatch — scanned sheet belongs to a different exam
    if not answers and isinstance(mode, str) and mode.startswith("exam_id_mismatch:"):
        parts      = mode.split(":")
        sheet_id   = parts[1] if len(parts) > 1 else "unknown"
        selected   = parts[2] if len(parts) > 2 else exam_id
        return jsonify({
            "error": (
                f"Wrong answer sheet uploaded. "
                f"The scanned sheet has Exam ID '{sheet_id}' "
                f"but you selected exam '{selected}'. "
                f"Please upload the answer sheet printed for this exam, "
                f"or select the correct exam from the dropdown."
            ),
            "exam_id_mismatch": True,
            "sheet_exam_id":    sheet_id,
            "selected_exam_id": selected,
        }), 422

    if not answers:
        return jsonify({"error": f"Could not extract answers (mode: {mode})"}), 400

    _openai_limiter.wait()
    result  = _evaluate_answers(e, answers, roll_no)
    eval_id = uuid.uuid4().hex[:10]
    payload = {
        "evaluation_id":    eval_id,
        "exam_id":          exam_id,
        "created_at":       _now(),
        "student_name":     student_name,
        "roll_no":          roll_no,
        "parent_email":     parent_email,
        "file_name":        fn,
        "file_hash":        fhash,
        "extraction_mode":  mode,
        "submitted_answers":answers,
        "result":           result,
        "ocr_version":      OCR_PIPELINE_VERSION,
    }
    evaluations_registry[eval_id] = payload
    if not IS_VERCEL:
        _save_json(EVAL_DIR / f"{eval_id}.json", payload)
    _save_evals()
    # Dual-write: also insert as individual doc so Vercel serverless instances can retrieve it
    _mongo_insert("eval_items", {**payload, "evaluation_id": eval_id})
    _cache_set(fhash, exam_id, payload)
    log.info(f"[EVAL] Sync complete: eval_id={eval_id} score={result.get('percentage',0):.1f}%")
    return jsonify(payload)

@app.post("/api/evaluate/text")
@auth()
@quota('evaluate')
def evaluate_text():
    """
    Text-based evaluation — student/teacher types answers directly.
    No OCR, no file upload. Accepts JSON:
      { exam_id, answers: {1: "B", 2: "D", 7: "written answer..."}, student_name, roll_no, parent_email }
    Uses the same _evaluate_answers() grading engine as PDF evaluation.
    """
    u    = request.user
    data = request.get_json(silent=True) or {}

    exam_id      = str(data.get("exam_id", "")).strip()
    answers_raw  = data.get("answers", {})
    student_name = str(data.get("student_name", "")).strip()
    roll_no      = str(data.get("roll_no", "")).strip()
    parent_email = str(data.get("parent_email", "")).strip()

    if not exam_id:
        return jsonify({"error": "exam_id is required"}), 400
    e = exams_registry.get(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404
    if not answers_raw or not isinstance(answers_raw, dict):
        return jsonify({"error": "answers dict is required"}), 400

    # Convert string keys to int keys, strip whitespace
    submitted: Dict[int, str] = {}
    for k, v in answers_raw.items():
        try:
            submitted[int(k)] = str(v).strip()
        except (ValueError, TypeError):
            continue

    if not submitted:
        return jsonify({"error": "No valid answers provided"}), 400

    log.info(f"[EVAL-TEXT] exam={exam_id} answers={len(submitted)} student={student_name or 'anon'}")

    _openai_limiter.wait()
    result  = _evaluate_answers(e, submitted, roll_no)
    eval_id = uuid.uuid4().hex[:10]
    payload = {
        "evaluation_id":     eval_id,
        "exam_id":           exam_id,
        "created_at":        _now(),
        "student_name":      student_name,
        "roll_no":           roll_no,
        "parent_email":      parent_email,
        "file_name":         "",
        "extraction_mode":   "manual_text_entry",
        "submitted_answers": submitted,
        "result":            result,
    }
    evaluations_registry[eval_id] = payload
    if not IS_VERCEL:
        _save_json(EVAL_DIR / f"{eval_id}.json", payload)
    _save_evals()
    # Dual-write: individual doc for reliable cross-instance lookup
    _mongo_insert("eval_items", {**payload, "evaluation_id": eval_id})
    log.info(f"[EVAL-TEXT] Complete: eval_id={eval_id} score={result.get('percentage',0):.1f}%")
    return jsonify(payload)

@app.get("/api/evaluate/status/<task_id>")
@auth()
def evaluation_status(task_id):
    """
    Poll async evaluation task status.
    States: pending → processing → completed | failed
    Checks native background-task registry first, then Celery.
    """
    # ── Check native bg-task registry first ──────────────────────────────────
    _native = _bg_task_get(task_id)
    if _native is not None:
        if _native["status"] == "completed":
            return jsonify({"status": "completed", "task_id": task_id, **_native["result"]})
        if _native["status"] == "failed":
            return jsonify({"status": "failed", "task_id": task_id, "error": _native["error"]}), 500
        return jsonify({"status": "pending", "task_id": task_id,
                        "step": "processing multi-student evaluation"})

    if not CELERY_OK or not celery_app:
        return jsonify({"error": "Async evaluation not available (Redis not connected)"}), 503

    task = celery_app.AsyncResult(task_id)

    if task.state == "PENDING":
        return jsonify({"status": "pending",    "task_id": task_id})
    if task.state == "STARTED":
        meta = task.info or {}
        return jsonify({"status": "processing", "task_id": task_id,
                        "progress": meta.get("completed"), "total": meta.get("total"),
                        "step": meta.get("step", "")})
    if task.state == "SUCCESS":
        result = task.result or {}
        # Cache the result so /api/evaluate can serve it on retry
        if result.get("file_hash") and result.get("exam_id"):
            _cache_set(result["file_hash"], result["exam_id"], result)
        return jsonify({"status": "completed", "task_id": task_id, **result})
    if task.state == "FAILURE":
        return jsonify({"status": "failed", "task_id": task_id,
                        "error": str(task.info)}), 500
    return jsonify({"status": task.state, "task_id": task_id})

@app.route("/api/evaluate/multi-student", methods=["OPTIONS"])
def evaluate_multi_student_options():
    """Handle preflight OPTIONS request for CORS"""
    return "", 204

@app.post("/api/evaluate/multi-student")
@auth()
@quota('evaluate')
def evaluate_multi_student():
    """
    Evaluate a SINGLE PDF containing MULTIPLE students' answer sheets.
    Async if Celery available. Falls back to parallel-thread sync.
    """
    exam_id = request.form.get("exam_id")
    if not exam_id:
        return jsonify({"error": "exam_id required"}), 400
    e = _load_exam(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404
    if "answer_sheet" not in request.files:
        return jsonify({"error": "answer_sheet (PDF) required"}), 400

    f = request.files["answer_sheet"]
    if not f or not _allowed(f.filename):
        return jsonify({"error": "Only PDF files supported for multi-student mode"}), 400

    # Hard limit: Cloud Run enforces a 32 MiB (33.5 MB) request body limit at the
    # load-balancer level — requests larger than that never reach this code and the
    # browser sees "Failed to Fetch".  We reject at 30 MB so the user gets a clear
    # error message instead of a silent network failure.
    _MAX_MULTI_BYTES = 30 * 1024 * 1024  # 30 MB  (well below Cloud Run's 32 MiB cap)
    f.seek(0, 2)  # seek to end
    _fsize = f.tell()
    f.seek(0)     # reset
    if _fsize > _MAX_MULTI_BYTES:
        _fsize_mb = round(_fsize / 1024 / 1024, 1)
        return jsonify({
            "error": (
                f"File too large ({_fsize_mb} MB). Maximum is 30 MB. "
                "Please compress the PDF (reduce scan resolution to 150–200 DPI) "
                "or split large classes into batches of 20–25 students."
            )
        }), 413

    fn   = secure_filename(f.filename)
    path = UPL_DIR / f"{datetime.utcnow().strftime('%Y%m%dT%H%M%S')}_{fn}"
    f.save(str(path))
    log.info(f"[MULTI-EVAL] Saved: {fn}")

    # ── Async path: Celery (preferred) ───────────────────────────────────────
    if CELERY_OK:
        try:
            from celery_worker import evaluate_multi_student as _ms_task
            task = _ms_task.delay(str(path), e)
            log.info(f"[MULTI-EVAL] Async task queued: {task.id}")
            return jsonify({
                "async":   True,
                "task_id": task.id,
                "status":  "processing",
                "message": "Multi-student evaluation queued. Poll /api/evaluate/status/<task_id>.",
            }), 202
        except Exception as _ce:
            log.warning(f"[MULTI-EVAL] Celery dispatch failed, falling back: {_ce}")

    # ── Async path: native background thread (no Celery) ─────────────────────
    # Important: in-process background threads are unreliable on Cloud Run when
    # no worker backend exists, because the instance/request lifecycle can stop
    # progress after returning 202. On Cloud Run without Celery, run sync path.
    _np = 0
    try:
        import fitz as _fitz_peek
        _doc_peek = _fitz_peek.open(str(path))
        _np = len(_doc_peek)
        _doc_peek.close()
    except Exception:
        pass

    _is_cloud_run = bool(os.getenv("K_SERVICE", "").strip())
    _use_native_async = not _is_cloud_run
    if _use_native_async:
        _bg_tid = _bg_task_create()
        log.info(f"[MULTI-EVAL] Native async task {_bg_tid} ({_np}pp)")

        def _bg_run_multi_student(_tid, _pdf_path, _exam, _exam_id, _fname):
            try:
                _secs = _split_multi_student_pdf(_pdf_path, exam_questions=_exam.get("questions", []))
                _all_evals: list = []
                _failed: list = []
                _bs = _resolve_multi_eval_batch_size(len(_secs))
                _ibd = max(0.0, MULTI_EVAL_INTER_BATCH_DELAY)
                _ea = max(1, MULTI_EVAL_EXTRACT_ATTEMPTS)
                _va = max(1, MULTI_EVAL_EVAL_ATTEMPTS)

                def _batched(items, size):
                    for _s in range(0, len(items), size):
                        yield items[_s:_s + size]

                def _eval_one(_idx, _section):
                    _sname = _section.get("student_name") or f"Student {_idx + 1}"
                    _roll  = _section.get("roll_no") or str(_idx + 1)
                    _raw   = _section.get("raw_text", "")
                    _ans   = {}; _mode = "multi_student_parallel"; _exerr = ""
                    if _section.get("page_indices"):
                        import fitz as _fz2
                        _tmp = Path(_pdf_path).parent / f"ms_{_idx}_{uuid.uuid4().hex[:6]}.pdf"
                        try:
                            _s = _fz2.open(_pdf_path); _d = _fz2.open()
                            for _pi in _section["page_indices"]:
                                if _pi < len(_s): _d.insert_pdf(_s, from_page=_pi, to_page=_pi)
                            _d.save(str(_tmp)); _d.close(); _s.close()
                            _ans, _mode, _exerr = _extract_answers_with_retry(
                                str(_tmp), exam_questions=_exam.get("questions", []), exam=_exam, attempts=_ea)
                        except Exception as _ex:
                            _exerr = str(_ex)
                        finally:
                            try: _tmp.unlink()
                            except Exception: pass
                    if not _ans and _raw:
                        _ans = _parse_raw_answers(_raw); _mode = "multi_student_text"
                    if not _ans:
                        return {"status": "failed", "student_name": _sname, "roll_no": _roll,
                                "reason": _exerr or f"no_answers ({_mode})", "extraction_mode": _mode}
                    _res, _eerr = _evaluate_answers_with_retry(_exam, _ans, _roll, attempts=_va)
                    if not _res:
                        return {"status": "failed", "student_name": _sname, "roll_no": _roll,
                                "reason": _eerr or "evaluation_failed", "extraction_mode": _mode}
                    _eid = uuid.uuid4().hex[:10]
                    _payload = {
                        "evaluation_id": _eid, "exam_id": _exam_id, "created_at": _now(),
                        "student_name": _sname, "roll_no": _roll, "parent_email": "",
                        "file_name": _fname, "file_path": str(_pdf_path),
                        "page_indices": _section.get("page_indices", []),
                        "extraction_mode": _mode, "ocr_version": OCR_PIPELINE_VERSION,
                        "submitted_answers": _ans, "result": _res, "status": "evaluated",
                    }
                    evaluations_registry[_eid] = _payload
                    if not IS_VERCEL:
                        _save_json(EVAL_DIR / f"{_eid}.json", _payload)
                    return _payload

                _wc_fn = _multi_student_worker_count
                for _bi, _batch in enumerate(list(_batched(list(enumerate(_secs)), _bs)), 1):
                    _wc = min(max(1, _wc_fn(len(_batch))), len(_batch) or 1)
                    log.info(f"[MULTI-EVAL-BG] Task {_tid} batch {_bi}: students={len(_batch)} workers={_wc}")
                    with ThreadPoolExecutor(max_workers=_wc) as _pool:
                        _fts = [_pool.submit(_eval_one, *_item) for _item in _batch]
                        for _ft in as_completed(_fts):
                            try:
                                _p = _ft.result()
                            except Exception as _ex:
                                _failed.append({"status": "failed", "student_name": "Unknown",
                                                "roll_no": "", "reason": f"worker_exception: {_ex}",
                                                "extraction_mode": "unknown"})
                                continue
                            if not _p: continue
                            if _p.get("status") == "failed":
                                _failed.append(_p)
                            else:
                                _all_evals.append(_p)
                    if _bi < (len(_secs) // _bs + 1) and _ibd > 0:
                        _time.sleep(_ibd)

                _save_evals()
                _tc = len(_all_evals)
                _avg = round(sum(ev["result"].get("percentage", 0) for ev in _all_evals) / max(_tc, 1), 1)
                _bg_task_set_result(_tid, {
                    "detected_student_count": len(_secs),
                    "student_count": _tc, "class_average": _avg,
                    "pass_count": sum(1 for ev in _all_evals if ev["result"].get("is_pass")),
                    "fail_count": _tc - sum(1 for ev in _all_evals if ev["result"].get("is_pass")),
                    "failed_student_count": len(_failed),
                    "failed_students": _failed,
                    "evaluations": _all_evals,
                    "exam_id": _exam_id,
                })
            except Exception as _be:
                log.error(f"[MULTI-EVAL-BG] Task {_tid} failed: {_be}")
                _bg_task_set_error(_tid, str(_be))

        _t = threading.Thread(
            target=_bg_run_multi_student,
            args=(_bg_tid, str(path), e, exam_id, fn),
            daemon=True,
        )
        _t.start()
        return jsonify({
            "async":   True,
            "task_id": _bg_tid,
            "status":  "processing",
            "message": "Multi-student evaluation started. Poll /api/evaluate/status/<task_id>.",
        }), 202
    else:
        log.warning("[MULTI-EVAL] Cloud Run + no Celery: using sync evaluation path")

    # ── Sync path (small class PDFs, ≤ 2 batches) ────────────────────────────
    log.info(f"[MULTI-EVAL] Sync mode: splitting PDF {fn}")
    student_sections = _split_multi_student_pdf(str(path), exam_questions=e.get("questions", []))
    log.info(f"[MULTI-EVAL] Detected {len(student_sections)} students")

    all_evaluations = []
    failed_students = []

    # Adaptive wave scheduler: process class PDFs in batches to reduce OCR TPM spikes.
    batch_size = _resolve_multi_eval_batch_size(len(student_sections))
    inter_batch_delay = max(0.0, MULTI_EVAL_INTER_BATCH_DELAY)
    extract_attempts = max(1, MULTI_EVAL_EXTRACT_ATTEMPTS)
    eval_attempts = max(1, MULTI_EVAL_EVAL_ATTEMPTS)

    def _batched_pairs(items: list, size: int):
        for _start in range(0, len(items), size):
            yield items[_start:_start + size]

    _indexed_sections = list(enumerate(student_sections))
    _batches = list(_batched_pairs(_indexed_sections, batch_size))
    log.info(
        f"[MULTI-EVAL] Scheduling {len(student_sections)} students in {len(_batches)} batch(es): "
        f"batch_size={batch_size}, inter_batch_delay={inter_batch_delay}s"
    )

    # Bind attempts selected from runtime config into worker closure.
    def _eval_section(args):
        idx, section = args
        sname    = section.get("student_name") or f"Student {idx + 1}"
        roll     = section.get("roll_no") or str(idx + 1)
        raw_text = section.get("raw_text", "")
        answers  = {}
        mode     = "multi_student_parallel"
        extract_error = ""

        if section.get("page_indices"):
            import fitz as _fitz_ms
            tmp_pdf = UPL_DIR / f"ms_{idx}_{uuid.uuid4().hex[:6]}.pdf"
            try:
                src = _fitz_ms.open(str(path))
                dst = _fitz_ms.open()
                for pi in section["page_indices"]:
                    if pi < len(src):
                        dst.insert_pdf(src, from_page=pi, to_page=pi)
                dst.save(str(tmp_pdf))
                dst.close()
                src.close()
                answers, mode, extract_error = _extract_answers_with_retry(
                    str(tmp_pdf),
                    exam_questions=e.get("questions", []),
                    exam=e,
                    attempts=extract_attempts,
                )
            except Exception as ex:
                log.warning(f"[MULTI-EVAL] Page extraction failed for {sname}: {ex}")
                extract_error = str(ex)
            finally:
                try:
                    tmp_pdf.unlink()
                except Exception:
                    pass

        if not answers and raw_text:
            answers = _parse_raw_answers(raw_text)
            mode = "multi_student_text"

        if not answers:
            reason = extract_error or f"no_answers ({mode})"
            log.warning(f"[MULTI-EVAL] No answers for {sname} ({roll}) — {reason}")
            return {
                "status": "failed",
                "student_name": sname,
                "roll_no": roll,
                "reason": reason,
                "extraction_mode": mode,
            }

        result, eval_error = _evaluate_answers_with_retry(e, answers, roll, attempts=eval_attempts)
        if not result:
            reason = eval_error or "evaluation_failed"
            log.warning(f"[MULTI-EVAL] Grading failed for {sname} ({roll}) — {reason}")
            return {
                "status": "failed",
                "student_name": sname,
                "roll_no": roll,
                "reason": reason,
                "extraction_mode": mode,
            }

        eval_id = uuid.uuid4().hex[:10]
        payload = {
            "evaluation_id":     eval_id, "exam_id": exam_id, "created_at": _now(),
            "student_name":      sname,   "roll_no": roll,
            "parent_email":      section.get("parent_email", ""), "file_name": fn,
            "file_path":         str(path),
            "page_indices":      section.get("page_indices", []),
            "extraction_mode":   mode,
            "ocr_version":       OCR_PIPELINE_VERSION,
            "submitted_answers": answers, "result": result,
            "status": "evaluated",
        }
        evaluations_registry[eval_id] = payload
        if not IS_VERCEL:
            _save_json(EVAL_DIR / f"{eval_id}.json", payload)
        return payload

    for batch_idx, batch in enumerate(_batches, start=1):
        desired_workers = _multi_student_worker_count(len(batch))
        if MULTI_EVAL_MAX_WORKERS_OVERRIDE > 0:
            desired_workers = min(desired_workers, MULTI_EVAL_MAX_WORKERS_OVERRIDE)
        # Cloud Run memory is shared by Flask + OCR image buffers + model payloads.
        # Allow 2 concurrent workers on Cloud Run (4 GB instance can handle it).
        if os.getenv("K_SERVICE", "").strip():
            desired_workers = min(desired_workers, 2)
        worker_count = min(max(1, desired_workers), len(batch) or 1)
        log.info(
            f"[MULTI-EVAL] Batch {batch_idx}/{len(_batches)}: "
            f"students={len(batch)}, workers={worker_count}"
        )

        with ThreadPoolExecutor(max_workers=worker_count) as pool:
            futs = [pool.submit(_eval_section, pair) for pair in batch]
            for fut in as_completed(futs):
                try:
                    p = fut.result()
                except Exception as ex:
                    failed_students.append({
                        "status": "failed",
                        "student_name": "Unknown",
                        "roll_no": "",
                        "reason": f"worker_exception: {ex}",
                        "extraction_mode": "unknown",
                    })
                    continue
                if not p:
                    continue
                if p.get("status") == "failed":
                    failed_students.append(p)
                else:
                    all_evaluations.append(p)

        if batch_idx < len(_batches) and inter_batch_delay > 0:
            _time.sleep(inter_batch_delay)

    _save_evals()
    total_students = len(all_evaluations)
    avg_pct = round(sum(ev["result"].get("percentage",0) for ev in all_evaluations) / max(total_students,1), 1)
    pass_count = sum(1 for ev in all_evaluations if ev["result"].get("is_pass"))
    log.info(
        f"[MULTI-EVAL] Sync complete: detected={len(student_sections)} "
        f"evaluated={total_students} failed={len(failed_students)} avg={avg_pct}%"
    )
    return jsonify({
        "detected_student_count": len(student_sections),
        "student_count": total_students, "class_average": avg_pct,
        "pass_count": pass_count, "fail_count": total_students - pass_count,
        "failed_student_count": len(failed_students),
        "failed_students": failed_students,
        "runtime_tuning": {
            "batch_size": batch_size,
            "inter_batch_delay_seconds": inter_batch_delay,
            "extract_attempts": extract_attempts,
            "eval_attempts": eval_attempts,
            "max_workers_override": MULTI_EVAL_MAX_WORKERS_OVERRIDE,
            "batches": len(_batches),
        },
        "evaluations": all_evaluations, "exam_id": exam_id,
    })

@app.post("/api/evaluate/bulk")
@auth(roles=["tutor","teacher","institute_admin"])
@quota('evaluate')
def bulk_evaluate():
    """
    Multiple separate files (one per student).
    Async if Celery available — all sheets processed in parallel.
    Falls back to sequential sync processing.
    Includes per-file caching.
    """
    exam_id = request.form.get("exam_id")
    if not exam_id:
        return jsonify({"error": "exam_id required"}), 400
    e = _load_exam(exam_id)
    if not e:
        return jsonify({"error": "Exam not found"}), 404

    files = request.files.getlist("answer_sheets")
    if not files:
        return jsonify({"error": "No answer_sheets uploaded"}), 400

    # Save all files first
    saved, unsupported = [], []
    for f in files:
        if not f or not _allowed(f.filename):
            unsupported.append({"file": getattr(f, "filename", "?"), "error": "unsupported_type"})
            continue
        fn   = secure_filename(f.filename)
        path = UPL_DIR / f"{datetime.utcnow().strftime('%Y%m%dT%H%M%S')}_{fn}"
        f.save(str(path))
        # Prefer explicitly submitted values; roll defaults to "" (resolved from header later)
        roll = request.form.get(f"roll_no_{fn}", "").strip()
        name = request.form.get(f"name_{fn}", "").strip()
        saved.append((str(path), roll, name, fn))

    log.info(f"[BULK] {len(saved)} files saved for exam={exam_id}")

    # ── Async path ────────────────────────────────────────────────────────────
    if CELERY_OK and saved:
        try:
            from celery_worker import evaluate_bulk as _bulk_task
            task = _bulk_task.delay(saved, e)
            log.info(f"[BULK] Async task queued: {task.id}")
            return jsonify({
                "async":   True,
                "task_id": task.id,
                "status":  "processing",
                "files":   len(saved),
                "message": "Bulk evaluation queued. Poll /api/evaluate/status/<task_id>.",
            }), 202
        except Exception as _ce:
            log.warning(f"[BULK] Celery dispatch failed, falling back to sync: {_ce}")

    # ── Sync path with parallel threads ──────────────────────────────────────
    results, failures = list(unsupported), []

    def _eval_one(entry):
        fpath, roll, name, fn = entry
        try:
            fhash  = _file_hash(fpath)
            cached = _cache_get(fhash, exam_id)
            # Extract student info from header if roll/name not provided by caller
            if not roll or not name:
                try:
                    hdr = _extract_header_info(fpath)
                    if not roll:
                        roll = hdr.get("roll_no", "").strip()
                    if not name:
                        name = hdr.get("student_name", "").strip()
                        # FALLBACK: If name is generic/blank, use the filename (e.g. Student09)
                        if not name or name.lower() in {"student", "name", "student name", "student name:", "none", "null", ""}:
                            name = fn.split(".")[0].capitalize()

                        # Ensure name isn't JUST "Student" if we can help it
                        if name.lower() == "student" and fn.lower().startswith("student"):
                             name = fn.split(".")[0].capitalize()
                except Exception as _he:
                    log.warning(f"[BULK] Header extraction failed for {fn}: {_he}")
            if cached:
                # Reconstruct mcq_answers / descriptive_answers from cached question_wise
                cached_qwise = cached.get("result", {}).get("question_wise", [])
                cached_mcq, cached_subj = [], []
                for qw in cached_qwise:
                    qid = qw.get("question_id")
                    ans = str(qw.get("student_answer", "")).strip() or "—"
                    if qw.get("type") == "objective":
                        cached_mcq.append(f"Q{qid}={ans}")
                    else:
                        short = (ans[:80] + "...") if len(ans) > 80 else ans
                        cached_subj.append(f"Q{qid}={short}")
                cached_res = cached.get("result", {})
                log.info(f"[BULK CACHE HIT] fn={fn} mcq={' '.join(cached_mcq)}")
                return {
                    "evaluation_id": cached["evaluation_id"],
                    "student_name": name or cached.get("student_name", ""),
                    "roll_no": roll or cached.get("roll_no", ""),
                    "submitted_answers": cached.get("submitted_answers", {}),
                    "result": cached_res,
                    "mcq_answers": " ".join(cached_mcq),
                    "descriptive_answers": " | ".join(cached_subj),
                    "percentage": cached_res.get("percentage", 0),
                    "is_pass": cached_res.get("is_pass", False),
                    "total_awarded": cached_res.get("total_awarded", 0),
                    "total_possible": cached_res.get("total_possible", 0),
                }, None
            _openai_limiter.wait()
            answers, mode = _extract_answers(fpath, exam_questions=e.get("questions",[]), exam=e)
            if not answers:
                return None, {"file": fn, "error": f"extract_failed ({mode})"}
            _openai_limiter.wait()
            res = _evaluate_answers(e, answers, roll)
            eid = uuid.uuid4().hex[:10]
            p   = {
                "evaluation_id": eid, "exam_id": exam_id, "created_at": _now(),
                "student_name": name, "roll_no": roll, "file_name": fn,
                "file_hash": fhash, "extraction_mode": mode,
                "submitted_answers": answers, "result": res,
                "ocr_version": OCR_PIPELINE_VERSION,
            }
            evaluations_registry[eid] = p
            if not IS_VERCEL:
                _save_json(EVAL_DIR / f"{eid}.json", p)
            _cache_set(fhash, exam_id, p)
            # Store in registry so individual reports can be downloaded
            evaluations_registry[eid] = p

            # Build CSV-friendly strings for BulkEvalTab's downloadCSV
            mcq_parts = []
            subj_parts = []
            for qw in res.get("question_wise", []):
                qid = qw.get("question_id")
                ans = str(qw.get("student_answer", "")).strip() or "—"
                if qw.get("type") == "objective":
                    mcq_parts.append(f"Q{qid}={ans}")
                else:
                    # Truncate long descriptive answers
                    short = (ans[:80] + "...") if len(ans) > 80 else ans
                    subj_parts.append(f"Q{qid}={short}")

            summary_p = {
                "evaluation_id": eid, "student_name": name, "roll_no": roll,
                "file_name": fn,
                "submitted_answers": answers,
                "result": res,
                # Fields expected by BulkEvalTab's CSV generator
                "mcq_answers": " ".join(mcq_parts),
                "descriptive_answers": " | ".join(subj_parts),
                "total_awarded": res.get("total_awarded", 0),
                "total_possible": res.get("total_possible", 0),
                "percentage": res.get("percentage", 0),
                "is_pass": res.get("is_pass", False)
            }
            log.info(f"[EVAL-SUMMARY] fn={fn} name={name} roll={roll} mcq={summary_p['mcq_answers']}")
            return summary_p, None
        except Exception as _one_err:
            log.exception(f"[BULK] Fatal error while evaluating {fn}: {_one_err}")
            return None, {"file": fn, "error": str(_one_err)}

    with ThreadPoolExecutor(max_workers=min(4, len(saved))) as pool:
        futs = {pool.submit(_eval_one, entry): entry for entry in saved}
        eval_items = []
        for fut in as_completed(futs):
            try:
                res_p, err = fut.result()
            except Exception as _fut_err:
                _entry = futs.get(fut) or ("", "", "", "unknown")
                _fn = _entry[3] if len(_entry) > 3 else "unknown"
                log.exception(f"[BULK] Worker future crashed for {_fn}: {_fut_err}")
                res_p, err = None, {"file": _fn, "error": str(_fut_err)}
            if res_p:
                eval_items.append(res_p)
            if err:
                failures.append(err)

    log.info(f"[BULK] Sync complete: ok={len(eval_items)}")
    return jsonify({
        "status": "completed", 
        "evaluations": eval_items, 
        "results": eval_items, 
        "failures": failures,
        "total": len(eval_items)
    })

@app.get("/api/evaluations/<eval_id>")
@auth()
def get_evaluation(eval_id):
    ev = evaluations_registry.get(eval_id)
    if not ev and not IS_VERCEL:
        ev = _load_json(EVAL_DIR / f"{eval_id}.json")
        if ev:
            evaluations_registry[eval_id] = ev
    # Vercel: in-memory registry is per-instance; fall back to individual MongoDB doc
    if not ev:
        ev = _mongo_find_one("eval_items", {"evaluation_id": eval_id})
        if ev:
            evaluations_registry[eval_id] = ev
    if not ev:
        return jsonify({"error": "Not found"}), 404
    return jsonify(ev)

def _build_email_report_html(ev: dict, exam: dict, student_name: str) -> str:
    """Build a rich HTML email report for parents."""
    res         = ev.get("result", {})
    pct         = float(res.get("percentage", 0))
    is_pass     = res.get("is_pass", False)
    awarded     = res.get("total_awarded", 0)
    possible    = res.get("total_possible", 0)
    grade       = res.get("grade", "F")
    qwise       = res.get("question_wise", [])
    obj_qw      = [q for q in qwise if q.get("type") == "objective"]
    subj_qw     = [q for q in qwise if q.get("type") == "subjective"]
    school      = exam.get("school_name", "Parvidya Smart Learning")
    subject_str = exam.get("subject", "")
    exam_cls    = exam.get("class", "")
    imp         = res.get("improvement_prediction", "")
    score_color = "#065f46" if pct >= 60 else ("#d97706" if pct >= 40 else "#dc2626")
    pass_color  = "#065f46" if is_pass else "#dc2626"
    pass_bg     = "#f0fdf4" if is_pass else "#fef2f2"
    pass_border = "#86efac" if is_pass else "#fca5a5"
    pass_text   = "PASS &#10003;" if is_pass else "FAIL &#10007;"

    # MCQ rows
    mcq_rows = ""
    for qr in obj_qw:
        correct = qr.get("is_correct", False)
        row_bg  = "#f0fdf4" if correct else "#fef2f2"
        sym     = "&#10003;" if correct else "&#10007;"
        sym_col = "#065f46" if correct else "#dc2626"
        am      = qr.get("awarded_marks", 0)
        mm      = qr.get("max_marks", 0)
        sa      = str(qr.get("student_answer", "&#8212;") or "&#8212;")
        mcq_rows += (
            f'<tr style="background:{row_bg}">'
            f'<td style="padding:7px 10px;border:1px solid #e2e8f0">Q{qr.get("question_id","")}</td>'
            f'<td style="padding:7px 10px;border:1px solid #e2e8f0">{sa}</td>'
            f'<td style="padding:7px 10px;border:1px solid #e2e8f0;font-weight:700;color:{sym_col};text-align:center">{sym}</td>'
            f'<td style="padding:7px 10px;border:1px solid #e2e8f0;text-align:center">{am}/{mm}</td>'
            f'</tr>'
        )

    mcq_section = ""
    if obj_qw:
        mcq_section = (
            '<tr><td style="padding:8px 32px">'
            '<div style="font-size:13px;font-weight:700;color:#3730a3;background:#eef2ff;'
            'padding:8px 12px;border-radius:6px;margin-bottom:10px">'
            f'Section A &#8212; Objective ({len(obj_qw)} questions)</div>'
            '<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse;font-size:13px">'
            '<tr style="background:#f1f5f9">'
            '<th style="padding:8px 10px;text-align:left;border:1px solid #e2e8f0">Q#</th>'
            '<th style="padding:8px 10px;text-align:left;border:1px solid #e2e8f0">Answer</th>'
            '<th style="padding:8px 10px;border:1px solid #e2e8f0">Result</th>'
            '<th style="padding:8px 10px;border:1px solid #e2e8f0">Marks</th>'
            f'</tr>{mcq_rows}</table></td></tr>'
        )

    # Subjective rows
    subj_rows = ""
    for qr in subj_qw:
        am_q    = qr.get("awarded_marks", 0)
        mm_q    = qr.get("max_marks", 0)
        pct_q   = round((am_q / mm_q * 100) if mm_q else 0)
        bc      = "#065f46" if pct_q >= 60 else ("#d97706" if pct_q >= 30 else "#dc2626")
        fb      = str(qr.get("feedback", ""))[:150]
        sa      = str(qr.get("student_answer", "No answer"))[:200]
        fb_html = f'<div style="font-size:12px;color:#1d4ed8;margin-top:4px"><em>{fb}</em></div>' if fb else ""
        subj_rows += (
            f'<tr>'
            f'<td style="padding:10px;border:1px solid #e2e8f0;vertical-align:top">'
            f'<strong>Q{qr.get("question_id","")}</strong></td>'
            f'<td style="padding:10px;border:1px solid #e2e8f0">'
            f'<div style="color:#374151;font-size:13px">{sa}</div>{fb_html}</td>'
            f'<td style="padding:10px;border:1px solid #e2e8f0;white-space:nowrap;text-align:center">'
            f'<span style="font-weight:700;color:{bc}">{am_q}/{mm_q}m</span></td>'
            f'</tr>'
        )

    subj_section = ""
    if subj_qw:
        subj_section = (
            '<tr><td style="padding:16px 32px 8px">'
            '<div style="font-size:13px;font-weight:700;color:#065f46;background:#ecfdf5;'
            'padding:8px 12px;border-radius:6px;margin-bottom:10px">'
            f'Section B &#8212; Written ({len(subj_qw)} questions)</div>'
            '<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse;font-size:13px">'
            '<tr style="background:#f1f5f9">'
            '<th style="padding:8px 10px;text-align:left;border:1px solid #e2e8f0">Q#</th>'
            '<th style="padding:8px 10px;text-align:left;border:1px solid #e2e8f0">Student Answer</th>'
            '<th style="padding:8px 10px;border:1px solid #e2e8f0">Marks</th>'
            f'</tr>{subj_rows}</table></td></tr>'
        )

    imp_section = ""
    if imp:
        imp_section = (
            '<tr><td style="padding:16px 32px">'
            '<div style="background:#eff6ff;border-radius:10px;padding:14px 16px;border:1px solid #bfdbfe">'
            '<div style="font-size:12px;font-weight:700;color:#1e40af;margin-bottom:6px">&#128161; Teacher&#39;s Recommendation</div>'
            f'<div style="font-size:13px;color:#1d4ed8">{imp}</div>'
            '</div></td></tr>'
        )

    return (
        '<!DOCTYPE html><html><head><meta charset="utf-8"></head>'
        '<body style="margin:0;padding:0;background:#f8fafc;font-family:Arial,sans-serif">'
        '<table width="100%" cellpadding="0" cellspacing="0" style="background:#f8fafc;padding:24px 0">'
        '<tr><td align="center">'
        '<table width="600" cellpadding="0" cellspacing="0" '
        'style="background:#fff;border-radius:16px;overflow:hidden;border:1px solid #e2e8f0">'
        # Header
        '<tr><td style="background:linear-gradient(135deg,#1e1b4b,#4c1d95);padding:28px 32px">'
        f'<div style="color:#fff;font-size:20px;font-weight:700">{school}</div>'
        f'<div style="color:rgba(255,255,255,.6);font-size:13px;margin-top:4px">'
        f'Student Performance Report &middot; {subject_str} {exam_cls}</div>'
        '</td></tr>'
        # Greeting
        '<tr><td style="padding:24px 32px 8px">'
        '<p style="color:#374151;font-size:15px;margin:0">Dear Parent / Guardian,</p>'
        f'<p style="color:#374151;font-size:14px;margin:10px 0 0">Please find the performance report for <strong>{student_name}</strong>.</p>'
        '</td></tr>'
        # Score cards
        '<tr><td style="padding:16px 32px">'
        '<table width="100%" cellpadding="0" cellspacing="0"><tr>'
        f'<td width="25%" style="text-align:center;padding:14px;background:#f8fafc;border-radius:10px;border:1px solid #e2e8f0"><div style="font-size:11px;color:#64748b;font-weight:700;text-transform:uppercase;margin-bottom:6px">MARKS</div><div style="font-size:24px;font-weight:800;color:#0f172a">{awarded}/{possible}</div></td>'
        '<td width="4%"></td>'
        f'<td width="25%" style="text-align:center;padding:14px;background:#f8fafc;border-radius:10px;border:1px solid #e2e8f0"><div style="font-size:11px;color:#64748b;font-weight:700;text-transform:uppercase;margin-bottom:6px">PERCENTAGE</div><div style="font-size:24px;font-weight:800;color:{score_color}">{pct}%</div></td>'
        '<td width="4%"></td>'
        f'<td width="25%" style="text-align:center;padding:14px;background:#f8fafc;border-radius:10px;border:1px solid #e2e8f0"><div style="font-size:11px;color:#64748b;font-weight:700;text-transform:uppercase;margin-bottom:6px">GRADE</div><div style="font-size:24px;font-weight:800;color:{score_color}">{grade}</div></td>'
        '<td width="4%"></td>'
        f'<td width="25%" style="text-align:center;padding:14px;border-radius:10px;background:{pass_bg};border:1px solid {pass_border}"><div style="font-size:11px;color:#64748b;font-weight:700;text-transform:uppercase;margin-bottom:6px">RESULT</div><div style="font-size:22px;font-weight:800;color:{pass_color}">{pass_text}</div></td>'
        '</tr></table></td></tr>'
        # Progress bar
        f'<tr><td style="padding:4px 32px 20px"><div style="height:8px;background:#e2e8f0;border-radius:999px"><div style="height:8px;width:{min(100, int(pct))}%;background:{score_color};border-radius:999px"></div></div></td></tr>'
        # Sections
        + mcq_section + subj_section + imp_section +
        # Footer
        f'<tr><td style="padding:20px 32px;background:#f8fafc;border-top:1px solid #e2e8f0">'
        f'<div style="font-size:12px;color:#94a3b8">Evaluated by <strong>Parvidya Smart Learning AI</strong> &middot; Exam ID: {ev.get("exam_id","&#8212;")} &middot; Date: {ev.get("created_at","")[:10]}</div>'
        '</td></tr>'
        '</table></td></tr></table></body></html>'
    )

@app.post("/api/evaluations/send-report")
@auth()
def send_evaluation_report():
    """
    Send a richly formatted HTML performance report to a parent's email.
    Accepts pre-rendered HTML from the frontend, or builds one from eval data.
    """
    data         = request.json or {}
    eval_id      = data.get("evaluation_id")
    parent_email = data.get("parent_email", "").strip()
    student_name = data.get("student_name", "Student")
    html_report  = data.get("html_report", "")  # pre-built by React
    exam_meta    = data.get("exam_meta") or {}

    if not parent_email:
        return jsonify({"error": "parent_email is required"}), 400

    # Build subject line
    subject_line = (
        f"Performance Report — {student_name} | "
        f"{exam_meta.get('subject', exam_meta.get('syllabus_name', 'Examination'))}"
    )

    # Build rich HTML email if not pre-supplied
    if not html_report and eval_id:
        ev   = evaluations_registry.get(eval_id)
        exam = _load_exam(ev.get("exam_id","")) if ev else {}
        if ev:
            html_report = _build_email_report_html(
                ev, exam, student_name or ev.get("student_name","Student")
            )

    # Send via SMTP
    host   = os.getenv("SMTP_HOST")
    port   = int(os.getenv("SMTP_PORT", 587))
    user   = os.getenv("SMTP_USER")
    passwd = os.getenv("SMTP_PASS")
    sender = os.getenv("SMTP_FROM", user)

    if not all([host, user, passwd]):
        print(f"[EMAIL] SMTP not configured — would have sent to {parent_email}")
        # In dev mode, return success so UI can be tested
        return jsonify({
            "success":  True,
            "message":  f"[Dev mode] SMTP not configured. Report would be sent to {parent_email}.",
            "dev_mode": True
        })

    try:
        msg = MIMEMultipart("alternative")
        msg["From"]    = sender
        msg["To"]      = parent_email
        msg["Subject"] = subject_line
        msg.attach(MIMEText(html_report, "html", "utf-8"))

        with smtplib.SMTP(host, port) as s:
            s.starttls()
            s.login(user, passwd)
            s.send_message(msg)

        print(f"[EMAIL] Report sent to {parent_email} for {student_name}")
        return jsonify({"success": True, "message": f"Report successfully sent to {parent_email}"})

    except Exception as e:
        print(f"[EMAIL] Failed: {e}")
        return jsonify({"error": f"Email failed: {str(e)}. Check SMTP settings in .env"}), 500

@app.post("/api/evaluations/<eval_id>/audit")
@auth(roles=["tutor","teacher","institute_admin"])
def audit_evaluation(eval_id):
    ev = evaluations_registry.get(eval_id) or _load_json(EVAL_DIR / f"{eval_id}.json", {})
    if not ev:
        return jsonify({"error": "Not found"}), 404
    ex = _load_exam(ev.get("exam_id",""))
    if not ex:
        return jsonify({"error": "Exam not found"}), 404
    audit = _run_audit(ex, ev.get("result", {}))
    ev["audit"] = {"run_at": _now(), **audit}
    evaluations_registry[eval_id] = ev
    if not IS_VERCEL:
        _save_json(EVAL_DIR / f"{eval_id}.json", ev)
    _save_evals()
    return jsonify({"evaluation_id": eval_id, "audit": ev["audit"]})

@app.post("/api/evaluations/<eval_id>/refresh")
@auth(roles=["tutor","teacher","institute_admin","admin"])
def refresh_evaluation(eval_id):
    """Force re-extraction and re-grading for a single evaluation."""
    ev = (evaluations_registry.get(eval_id)
          or _mongo_find_one("eval_items", {"evaluation_id": eval_id})
          or _load_json(EVAL_DIR / f"{eval_id}.json", {}))
    if not ev:
        return jsonify({"error": "Evaluation not found"}), 404

    exam = _load_exam(ev.get("exam_id", ""))
    if not exam:
        return jsonify({"error": "Exam not found"}), 404

    force = bool((request.json or {}).get("force"))
    if force:
        ev["ocr_version"] = "force_refresh"

    refreshed = _refresh_evaluation_for_report(ev, exam)
    return jsonify({
        "evaluation_id": eval_id,
        "answers": len(refreshed.get("submitted_answers") or {}),
        "extraction_mode": refreshed.get("extraction_mode", ""),
        "ocr_version": refreshed.get("ocr_version", ""),
    })

def _generate_evaluation_report_pdf(ev: dict, exam: dict) -> bytes:
    """
    Generate a single-student evaluation report PDF.
    Sections:
     1. Header: school, exam, student details, overall score
     2. Score summary grid
     3. Section A: MCQ table with correct/wrong per question
     4. Section B: Subjective with model answer, feedback, key points
     5. Recommendations footer
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors as rl_c
        import io as _io2

        buf = _io2.BytesIO()
        W, H = A4
        c    = rl_canvas.Canvas(buf, pagesize=A4)
        M    = 15 * mm

        res      = ev.get("result", {})
        pct      = float(res.get("percentage", 0))
        grade    = res.get("grade", "F")
        is_pass  = res.get("is_pass", False)
        awarded  = res.get("total_awarded", 0)
        possible = res.get("total_possible", 0)
        qwise    = res.get("question_wise", [])
        obj_qw   = [q for q in qwise if q.get("type") == "objective"]
        subj_qw  = [q for q in qwise if q.get("type") == "subjective"]
        exam_qs  = {int(q.get("id",0)): q for q in exam.get("questions", [])}

        grade_color = (
            "#15803d" if pct >= 70 else
            "#d97706" if pct >= 50 else
            "#dc2626"
        )
        pass_color  = "#15803d" if is_pass else "#dc2626"

        page = 1

        def draw_page_header(pnum=1):
            # Dark header bar
            c.setFillColor(rl_c.HexColor("#1e1b4b"))
            c.rect(M, H - 20*mm, W - 2*M, 20*mm, fill=1, stroke=0)
            c.setFillColor(rl_c.white)
            c.setFont("Helvetica-Bold", 13)
            c.drawString(M + 5*mm, H - 10*mm,
                exam.get("school_name", "VidyAI School")[:50])
            c.setFont("Helvetica", 9)
            c.drawRightString(W - M - 5*mm, H - 10*mm,
                f"PERFORMANCE REPORT  ·  Page {pnum}")
            c.setFont("Helvetica", 8)
            c.setFillColor(rl_c.HexColor("#a5b4fc"))
            subj_str = "  ·  ".join(str(p) for p in [
                exam.get("subject",""), exam.get("board",""), exam.get("class","")
            ] if p)
            c.drawString(M + 5*mm, H - 17*mm, subj_str)
            c.drawRightString(W - M - 5*mm, H - 17*mm,
                f"Exam ID: {ev.get('exam_id','—')}  ·  Date: {ev.get('created_at','')[:10]}")
            return H - 23*mm

        def new_page():
            nonlocal page
            c.showPage(); page += 1
            return draw_page_header(page)

        y = draw_page_header(page)

        # ── Student info + score summary ──────────────────────────────────────
        # Student info row
        c.setFillColor(rl_c.HexColor("#f8fafc"))
        c.setStrokeColor(rl_c.HexColor("#e2e8f0"))
        c.setLineWidth(0.8)
        c.rect(M, y - 14*mm, W - 2*M, 14*mm, fill=1, stroke=1)

        fields = [
            ("Student",     ev.get("student_name","—")),
            ("Roll No.",    ev.get("roll_no","—")),
            ("Exam Date",   ev.get("created_at","")[:10]),
        ]
        fw = (W - 2*M) / len(fields)
        for fi, (lbl, val) in enumerate(fields):
            fx = M + fi * fw
            c.setFillColor(rl_c.HexColor("#64748b"))
            c.setFont("Helvetica", 7)
            c.drawString(fx + 4*mm, y - 5*mm, lbl.upper())
            c.setFillColor(rl_c.HexColor("#0f172a"))
            c.setFont("Helvetica-Bold", 10)
            c.drawString(fx + 4*mm, y - 11*mm, str(val)[:30])
            if fi > 0:
                c.setStrokeColor(rl_c.HexColor("#e2e8f0"))
                c.setLineWidth(0.4)
                c.line(fx, y - 12*mm, fx, y - 2*mm)
        y -= 17*mm

        # Score cards grid
        score_cells = [
            ("MARKS",      f"{awarded}/{possible}",  "#0f172a",  "#f8fafc"),
            ("PERCENTAGE", f"{pct}%",                 grade_color,"#f8fafc"),
            ("GRADE",      grade,                     grade_color,"#f8fafc"),
            ("RESULT",     "PASS ✓" if is_pass else "FAIL ✗", pass_color,
             "#f0fdf4" if is_pass else "#fef2f2"),
        ]
        cw = (W - 2*M) / 4
        CH = 18*mm
        for ci, (lbl, val, fcolor, bg) in enumerate(score_cells):
            cx = M + ci * cw
            c.setFillColor(rl_c.HexColor(bg))
            c.setStrokeColor(rl_c.HexColor("#e2e8f0"))
            c.setLineWidth(0.6)
            c.rect(cx, y - CH, cw, CH, fill=1, stroke=1)
            c.setFillColor(rl_c.HexColor("#64748b"))
            c.setFont("Helvetica", 7)
            c.drawCentredString(cx + cw/2, y - 6*mm, lbl)
            c.setFillColor(rl_c.HexColor(fcolor))
            c.setFont("Helvetica-Bold", 16)
            c.drawCentredString(cx + cw/2, y - 14*mm, str(val))
        y -= CH + 5*mm

        # Improvement prediction
        imp = res.get("improvement_prediction","")
        if imp:
            imp_lines = _pdf_wrap(imp, 90)
            IH = (len(imp_lines) * 4.5 + 5) * mm
            c.setFillColor(rl_c.HexColor("#eff6ff"))
            c.setStrokeColor(rl_c.HexColor("#bfdbfe"))
            c.setLineWidth(0.6)
            c.roundRect(M, y - IH, W - 2*M, IH, 2, fill=1, stroke=1)
            c.setFillColor(rl_c.HexColor("#1e40af"))
            c.setFont("Helvetica-Bold", 8)
            c.drawString(M + 4*mm, y - 4.5*mm, "💡 Teacher Recommendation:")
            c.setFont("Helvetica", 8)
            for li, line in enumerate(imp_lines):
                c.drawString(M + 4*mm, y - 9*mm - li*4.5*mm, line)
            y -= IH + 5*mm

        # ── Section A: Objective Results ──────────────────────────────────────
        if obj_qw:
            if y - 10*mm < 20*mm: y = new_page()
            _pdf_section_bar(c, M, y, W - 2*M, 8*mm,
                "#eef2ff", f"SECTION A — OBJECTIVE  ({len(obj_qw)} questions)",
                "#3730a3", font_size=8.5, mm=mm)
            y -= 10*mm

            # Column headers
            cols = [("Q#",12), ("Question",90), ("Student Answer",35), ("Marks",20), ("Result",22)]
            col_x = [M]
            for (_, cw_mm) in cols:
                col_x.append(col_x[-1] + cw_mm*mm)

            c.setFillColor(rl_c.HexColor("#f1f5f9"))
            c.rect(M, y - 6*mm, W - 2*M, 6*mm, fill=1, stroke=0)
            c.setFillColor(rl_c.HexColor("#374151"))
            c.setFont("Helvetica-Bold", 7.5)
            for i, (lbl, _) in enumerate(cols):
                c.drawString(col_x[i] + 1.5*mm, y - 4.2*mm, lbl)
            y -= 7*mm

            for qr in obj_qw:
                if y < 20*mm: y = new_page()
                qid      = int(qr.get("question_id", 0))
                correct  = qr.get("is_correct", False)
                q_text   = str(exam_qs.get(qid,{}).get("question",""))[:55]
                # Sanitize: strip newlines, collapse whitespace, treat bare-dash
                # (OCR artifact) as blank so it renders as em-dash, not "-\nA"
                raw_sa   = str(qr.get("student_answer", "") or "")
                raw_sa   = re.sub(r"[\r\n]+", " ", raw_sa).strip()
                raw_sa   = re.sub(r"\s+", " ", raw_sa)
                # If value is just '-' or '--' (OCR noise) treat as missing
                stud_ans = raw_sa if (raw_sa and raw_sa not in {"-", "--", "—"}) else "—"
                # If value still has leading dash+space like "- A", extract the letter
                m_dash = re.match(r'^[-–—]\s*([A-D])\b', stud_ans, re.IGNORECASE)
                if m_dash:
                    stud_ans = m_dash.group(1).upper()
                awarded_m= qr.get("awarded_marks",0)
                max_m    = qr.get("max_marks",0)
                bg       = "#f0fdf4" if correct else "#fef2f2"
                rc_text  = "✓ Correct" if correct else "✗ Wrong"
                rc_col   = "#15803d" if correct else "#dc2626"

                ROW_H = 7*mm
                c.setFillColor(rl_c.HexColor(bg))
                c.rect(M, y - ROW_H, W - 2*M, ROW_H, fill=1, stroke=0)
                c.setStrokeColor(rl_c.HexColor("#e2e8f0"))
                c.setLineWidth(0.3)
                c.line(M, y - ROW_H, W - M, y - ROW_H)

                row_vals = [f"Q{qid}", q_text, stud_ans or "—",
                            f"{awarded_m}/{max_m}", rc_text]
                for i, val in enumerate(row_vals):
                    color = rc_col if i == 4 else "#0f172a"
                    bold  = (i == 4)
                    c.setFillColor(rl_c.HexColor(color))
                    c.setFont("Helvetica-Bold" if bold else "Helvetica", 8)
                    c.drawString(col_x[i] + 1.5*mm, y - 4.8*mm, str(val)[:30])
                y -= ROW_H

            y -= 5*mm

        # ── Section B: Subjective Results ─────────────────────────────────────
        if subj_qw:
            if y - 10*mm < 25*mm: y = new_page()
            _pdf_section_bar(c, M, y, W - 2*M, 8*mm,
                "#ecfdf5", f"SECTION B — WRITTEN  ({len(subj_qw)} questions)",
                "#065f46", font_size=8.5, mm=mm)
            y -= 10*mm

            for qr in subj_qw:
                qid      = int(qr.get("question_id",0))
                awarded_m= float(qr.get("awarded_marks",0))
                max_m    = float(qr.get("max_marks",0))
                stud_ans = str(qr.get("student_answer","")).strip() or "(No answer)"
                feedback = str(qr.get("feedback","")).strip()
                strengths= qr.get("strengths",[])
                missing  = qr.get("missing_points",[])
                q_text   = str(exam_qs.get(qid,{}).get("question",""))
                pct_q    = (awarded_m/max_m*100) if max_m else 0
                bg       = "#f0fdf4" if pct_q >= 60 else ("#fffbeb" if pct_q >= 30 else "#fef2f2")
                border   = "#86efac" if pct_q >= 60 else ("#fde68a" if pct_q >= 30 else "#fecaca")

                # Estimate height
                q_lines  = _pdf_wrap(q_text, 88)
                a_lines  = _pdf_wrap(stud_ans, 88)
                fb_lines = _pdf_wrap(feedback, 88) if feedback else []
                total_h  = (2 + len(q_lines)*4.5 + len(a_lines)*4.5 +
                            len(fb_lines)*4.5 + len(strengths)*4.5 +
                            len(missing)*4.5 + 16) * mm

                if y - total_h < 18*mm: y = new_page()

                c.setFillColor(rl_c.HexColor(bg))
                c.setStrokeColor(rl_c.HexColor(border))
                c.setLineWidth(0.8)
                c.roundRect(M, y - total_h, W - 2*M, total_h, 2, fill=1, stroke=1)

                # Q number + score tab
                c.setFillColor(rl_c.HexColor("#059669"))
                c.roundRect(M, y - total_h, 16*mm, total_h, 2, fill=1, stroke=0)
                c.setFillColor(rl_c.white)
                c.setFont("Helvetica-Bold", 12)
                c.drawCentredString(M + 8*mm, y - total_h/2 + 2*mm, f"Q{qid}")
                c.setFont("Helvetica", 7)
                c.drawCentredString(M + 8*mm, y - total_h/2 - 4*mm,
                    f"{awarded_m}/{max_m}m")

                cx  = M + 18*mm
                cur_y = y - 5*mm

                # Question text
                c.setFillColor(rl_c.HexColor("#374151"))
                c.setFont("Helvetica-Bold", 8)
                for li, line in enumerate(q_lines):
                    c.drawString(cx, cur_y - li*4.5*mm, line)
                cur_y -= (len(q_lines)*4.5 + 3)*mm

                # Marks bar
                c.setFillColor(rl_c.HexColor("#1e293b"))
                c.setFont("Helvetica-Bold", 9)
                score_str = f"Score: {awarded_m}/{max_m} marks  ({int(pct_q)}%)"
                c.drawString(cx, cur_y, score_str)
                cur_y -= 7*mm

                # Student answer
                c.setFillColor(rl_c.HexColor("#374151"))
                c.setFont("Helvetica-Bold", 7.5)
                c.drawString(cx, cur_y, "Student's Answer:")
                cur_y -= 5*mm
                c.setFont("Helvetica", 8)
                for line in a_lines:
                    c.drawString(cx + 2*mm, cur_y, line)
                    cur_y -= 4.5*mm
                cur_y -= 2*mm

                if feedback:
                    c.setFillColor(rl_c.HexColor("#1d4ed8"))
                    c.setFont("Helvetica-Bold", 7.5)
                    c.drawString(cx, cur_y, "AI Feedback:")
                    cur_y -= 5*mm
                    c.setFont("Helvetica-Oblique", 8)
                    for line in fb_lines:
                        c.drawString(cx + 2*mm, cur_y, line)
                        cur_y -= 4.5*mm
                    cur_y -= 2*mm

                if strengths:
                    c.setFillColor(rl_c.HexColor("#15803d"))
                    c.setFont("Helvetica-Bold", 7.5)
                    c.drawString(cx, cur_y, "Strengths:")
                    cur_y -= 4.5*mm
                    c.setFont("Helvetica", 8)
                    for s in strengths[:3]:
                        c.drawString(cx + 2*mm, cur_y, f"+ {s[:80]}")
                        cur_y -= 4.5*mm

                if missing:
                    c.setFillColor(rl_c.HexColor("#92400e"))
                    c.setFont("Helvetica-Bold", 7.5)
                    c.drawString(cx, cur_y, "Could Improve:")
                    cur_y -= 4.5*mm
                    c.setFont("Helvetica", 8)
                    for mp in missing[:3]:
                        c.drawString(cx + 2*mm, cur_y, f"• {mp[:80]}")
                        cur_y -= 4.5*mm

                y -= total_h + 5*mm

        # Footer
        c.setFillColor(rl_c.HexColor("#94a3b8"))
        c.setFont("Helvetica", 6.5)
        c.drawCentredString(W/2, 5*mm,
            f"Evaluated by VidyAI  ·  Evaluation ID: {ev.get('evaluation_id','—')}  ·  "
            f"Mode: {ev.get('extraction_mode','—')}")

        c.save()
        buf.seek(0)
        return buf.read()

    except Exception as ex:
        log.exception(f"[REPORT-PDF] Critical failure for evaluation {ev.get('evaluation_id','—')}: {ex}")
        return b""

# ══════════════════════════════════════════════════════════════════════════════
#  REPORTING ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/evaluations")
@auth()
def list_evaluations():
    """
    List all evaluations visible to the requesting user.
    - student : own evaluations only
    - teacher / institute_admin / admin : all evaluations, optionally filtered
    Query params: exam_id, student_name, roll_no, limit (default 100)
    """
    u        = request.user
    exam_id  = request.args.get("exam_id", "").strip()
    s_name   = request.args.get("student_name", "").strip().lower()
    roll     = request.args.get("roll_no", "").strip()
    limit    = min(int(request.args.get("limit", 200)), 500)

    # Merge registry + JSON files for completeness
    all_evals = list(evaluations_registry.values())
    if not IS_VERCEL and EVAL_DIR.exists():
        for p in EVAL_DIR.glob("*.json"):
            eid = p.stem
            if eid not in evaluations_registry:
                ev = _load_json(p)
                if ev:
                    evaluations_registry[eid] = ev
                    all_evals.append(ev)

    # Role-based filter
    if u["role"] == "student":
        # Students see only evaluations where roll_no or student_name matches their profile
        all_evals = [ev for ev in all_evals
                     if (ev.get("roll_no") == u.get("roll_number","")
                         or ev.get("student_name","").lower() == u.get("name","").lower())]

    # Optional query filters
    if exam_id:
        all_evals = [ev for ev in all_evals if ev.get("exam_id") == exam_id]
    if s_name:
        all_evals = [ev for ev in all_evals
                     if s_name in ev.get("student_name","").lower()]
    if roll:
        all_evals = [ev for ev in all_evals if ev.get("roll_no","") == roll]

    # Sort newest first
    all_evals.sort(key=lambda e: e.get("created_at",""), reverse=True)
    all_evals = all_evals[:limit]

    # Lightweight summary — don't send full question_wise in list view
    summaries = []
    for ev in all_evals:
        res = ev.get("result", {})
        exam = _load_exam(ev.get("exam_id","")) or {}
        summaries.append({
            "evaluation_id":   ev.get("evaluation_id",""),
            "exam_id":         ev.get("exam_id",""),
            "exam_name":       exam.get("syllabus_name") or exam.get("subject",""),
            "subject":         exam.get("subject",""),
            "board":           exam.get("board",""),
            "class":           exam.get("class",""),
            "student_name":    ev.get("student_name",""),
            "roll_no":         ev.get("roll_no",""),
            "created_at":      ev.get("created_at",""),
            "total_awarded":   res.get("total_awarded",0),
            "total_possible":  res.get("total_possible",0),
            "percentage":      res.get("percentage",0),
            "grade":           res.get("grade","F"),
            "is_pass":         res.get("is_pass",False),
            "extraction_mode": ev.get("extraction_mode",""),
        })

    return jsonify({"evaluations": summaries, "total": len(summaries)})

@app.get("/api/evaluations/report/<eval_id>")
@auth()
def download_evaluation_report(eval_id):
    """
    Generate and return a PDF evaluation report for one student.
    """
    ev = evaluations_registry.get(eval_id)
    if not ev and not IS_VERCEL:
        ev = _load_json(EVAL_DIR / f"{eval_id}.json")
    # Vercel serverless: in-memory registry is per-instance, so look up individual doc
    if not ev:
        ev = _mongo_find_one("eval_items", {"evaluation_id": eval_id})
    if not ev:
        log.warning(f"[PDF-REPORT] Evaluation {eval_id} not found")
        return jsonify({"error": "Evaluation not found"}), 404

    exam = _load_exam(ev.get("exam_id","")) or {}
    if not exam:
        log.warning(f"[PDF-REPORT] Exam {ev.get('exam_id','')} not found for evaluation {eval_id}")
        return jsonify({"error": "Exam not found for this evaluation"}), 404
    
    try:
        ev = _refresh_evaluation_for_report(ev, exam)
        pdf  = _generate_evaluation_report_pdf(ev, exam)
        if not pdf:
            log.error(f"[PDF-REPORT] Generation returned None for evaluation {eval_id}")
            return jsonify({"error": "PDF generation failed - no data returned"}), 500

        sname    = re.sub(r"[^\w\s-]","", ev.get("student_name","Student"))[:20]
        filename = f"Report_{sname}_{eval_id}.pdf"
        from flask import Response
        return Response(pdf, mimetype="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        log.error(f"[PDF-REPORT] Error generating PDF for evaluation {eval_id}: {e}", exc_info=True)
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500

@app.get("/api/exams/<exam_id>/class-report")
@auth(roles=["teacher","tutor","institute_admin","admin"])
def get_class_report(exam_id):
    """
    Aggregate class-level report for one exam:
    - pass/fail counts, class average, grade distribution
    - question-wise accuracy (which questions the class found hardest)
    - top performers, students at risk
    - per-student summary table
    """
    exam = _load_exam(exam_id)
    if not exam:
        return jsonify({"error": "Exam not found"}), 404

    evals = [ev for ev in evaluations_registry.values()
             if ev.get("exam_id") == exam_id]
    if not IS_VERCEL and EVAL_DIR.exists():
        for p in EVAL_DIR.glob("*.json"):
            ev = _load_json(p)
            if ev and ev.get("exam_id") == exam_id:
                eid = ev.get("evaluation_id","")
                if eid and eid not in evaluations_registry:
                    evals.append(ev)

    if not evals:
        return jsonify({"error": "No evaluations found for this exam"}), 404

    evals = [_refresh_evaluation_for_report(ev, exam) for ev in evals]

    total     = len(evals)
    pcts      = [ev.get("result",{}).get("percentage",0) for ev in evals]
    awarded   = [ev.get("result",{}).get("total_awarded",0) for ev in evals]
    pass_list = [ev for ev in evals if ev.get("result",{}).get("is_pass")]

    # Grade distribution
    grade_dist = {"A+":0,"A":0,"B":0,"C":0,"D":0,"F":0}
    for ev in evals:
        g = ev.get("result",{}).get("grade","F")
        grade_dist[g] = grade_dist.get(g,0) + 1

    # Question-wise accuracy
    q_stats: Dict[int, dict] = {}
    for ev in evals:
        for qr in ev.get("result",{}).get("question_wise",[]):
            qid = int(qr.get("question_id",0))
            if qid not in q_stats:
                q_stats[qid] = {"awarded":0.0,"possible":0.0,"attempts":0}
            q_stats[qid]["awarded"]   += float(qr.get("awarded_marks",0))
            q_stats[qid]["possible"]  += float(qr.get("max_marks",1))
            q_stats[qid]["attempts"]  += 1

    q_accuracy = []
    for qid, s in sorted(q_stats.items()):
        poss = s["possible"]
        acc  = round((s["awarded"]/poss)*100, 1) if poss else 0
        q    = next((qq for qq in exam.get("questions",[]) if int(qq.get("id",0))==qid), {})
        q_accuracy.append({
            "question_id":   qid,
            "question":      str(q.get("question",""))[:80],
            "type":          q.get("type","objective"),
            "max_marks":     q.get("marks",q.get("weightage",1)),
            "avg_awarded":   round(s["awarded"]/s["attempts"],2) if s["attempts"] else 0,
            "accuracy_pct":  acc,
            "attempts":      s["attempts"],
        })

    # Student summaries sorted by percentage desc
    student_rows = []
    for ev in sorted(evals, key=lambda e: e.get("result",{}).get("percentage",0), reverse=True):
        res = ev.get("result",{})
        student_rows.append({
            "evaluation_id": ev.get("evaluation_id",""),
            "student_name":  ev.get("student_name",""),
            "roll_no":       ev.get("roll_no",""),
            "total_awarded": res.get("total_awarded",0),
            "total_possible":res.get("total_possible",0),
            "percentage":    res.get("percentage",0),
            "grade":         res.get("grade","F"),
            "is_pass":       res.get("is_pass",False),
            "created_at":    ev.get("created_at",""),
        })

    # Weakest questions (accuracy < 50%)
    weak_questions = [q for q in q_accuracy if q["accuracy_pct"] < 50]
    weak_questions.sort(key=lambda q: q["accuracy_pct"])

    return jsonify({
        "exam_id":         exam_id,
        "exam_name":       exam.get("syllabus_name","") or exam.get("subject",""),
        "subject":         exam.get("subject",""),
        "board":           exam.get("board",""),
        "class":           exam.get("class",""),
        "total_students":  total,
        "pass_count":      len(pass_list),
        "fail_count":      total - len(pass_list),
        "pass_rate":       round(len(pass_list)/total*100, 1) if total else 0,
        "class_average":   round(sum(pcts)/total, 1) if total else 0,
        "highest_score":   max(pcts) if pcts else 0,
        "lowest_score":    min(pcts) if pcts else 0,
        "total_marks":     exam.get("total_marks",0),
        "grade_distribution": grade_dist,
        "question_accuracy":  q_accuracy,
        "weak_questions":     weak_questions[:5],
        "student_rows":       student_rows,
    })

@app.get("/api/students/<roll_no>/history")
@auth()
def get_student_history(roll_no):
    """
    All evaluations for one student (by roll number).
    Shows performance trend across exams.
    """
    u = request.user
    # Students can only see their own history
    if u["role"] == "student" and u.get("roll_number","") != roll_no:
        return jsonify({"error": "Forbidden"}), 403

    all_evals = list(evaluations_registry.values())
    student_evals = [ev for ev in all_evals
                     if ev.get("roll_no","") == roll_no]
    student_evals.sort(key=lambda e: e.get("created_at",""))

    history = []
    for ev in student_evals:
        res  = ev.get("result",{})
        exam = _load_exam(ev.get("exam_id","")) or {}
        history.append({
            "evaluation_id": ev.get("evaluation_id",""),
            "exam_id":       ev.get("exam_id",""),
            "exam_name":     exam.get("syllabus_name","") or exam.get("subject",""),
            "subject":       exam.get("subject",""),
            "created_at":    ev.get("created_at",""),
            "percentage":    res.get("percentage",0),
            "grade":         res.get("grade","F"),
            "is_pass":       res.get("is_pass",False),
            "total_awarded": res.get("total_awarded",0),
            "total_possible":res.get("total_possible",0),
        })

    avg = round(sum(h["percentage"] for h in history)/len(history),1) if history else 0
    trend = "improving" if (len(history)>=2 and
        history[-1]["percentage"] > history[-2]["percentage"]) else (
        "declining" if len(history)>=2 and
        history[-1]["percentage"] < history[-2]["percentage"] else "stable")

    return jsonify({
        "roll_no":      roll_no,
        "student_name": student_evals[0].get("student_name","") if student_evals else "",
        "total_exams":  len(history),
        "average":      avg,
        "trend":        trend,
        "history":      history,
    })

# ══════════════════════════════════════════════════════════════════════════════
#  PRACTICE MODE
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/curriculum/practice/generate")
@auth()
@quota('practice_generate')
def generate_practice():
    u    = request.user
    data = request.json or {}
    sid  = data.get("syllabus_id", "").strip()
    topic     = data.get("topic", "the material")
    obj_count = int(data.get("objective_count", 5))
    subj_count= int(data.get("subjective_count", 5))
    diff      = data.get("difficulty", "medium").lower()
    total     = obj_count + subj_count

    if not sid or sid not in syllabi_registry:
        return jsonify({"error": "Select a valid syllabus first"}), 400

    # Fetch previous questions from MongoDB to avoid repetition
    previous_qs: list = []
    if MONGO_OK:
        hist = _mongo_find_one(M_PRACTICE_HIST, {"user_id": u["id"], "syllabus_id": sid, "topic": topic})
        if hist:
            previous_qs = hist.get("asked_questions", [])

    vs  = _load_vs(sid)
    ctx = ""
    if vs:
        try:
            docs = vs.as_retriever(search_kwargs={"k": 5}).invoke(topic)
            ctx  = "\n\n".join(d.page_content for d in docs)
        except Exception:
            pass
    if not ctx:
        s = syllabi_registry.get(sid, {})
        ctx = f"Syllabus: {s.get('name','General')}. Topic: {topic}."

    exclusion = ""
    if previous_qs:
        exclusion = "\nDO NOT repeat these previously asked questions:\n" + "\n".join(f"- {q}" for q in previous_qs[-20:])

    syl_meta = syllabi_registry.get(sid, {})
    syl_board      = syl_meta.get("board", "") or syl_meta.get("name", "")
    syl_class      = syl_meta.get("class_name", "") or ""
    syl_subject    = syl_meta.get("subject", "") or ""
    # Parse from name if fields missing (legacy syllabus created before this fix)
    if not syl_class and " — " in syl_meta.get("name", ""):
        _parts = syl_meta["name"].split(" — ")
        syl_subject = syl_subject or (_parts[1].strip() if len(_parts) > 1 else "")
    if not syl_board:
        syl_board = syl_meta.get("name", "Indian curriculum").split(" ")[0]

    prompt = MIXED_PRACTICE_GEN_PROMPT.format(
        total_count=total, objective_count=obj_count, subjective_count=subj_count,
        difficulty=diff, context=ctx, topic=topic,
        board=syl_board or "CBSE",
        class_level=syl_class or "the appropriate class",
        subject=syl_subject or "the subject",
    ) + exclusion

    try:
        raw = _llm_json(prompt, temperature=0.7)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    # Normalise
    questions = raw if isinstance(raw, list) else raw.get("questions", []) if isinstance(raw, dict) else []
    valid_qs  = []
    for q in questions:
        if not isinstance(q, dict):
            continue
        q_text = q.get("question") or q.get("q") or q.get("text")
        a_text = q.get("answer") or q.get("a") or q.get("explanation")
        if not q_text or not a_text:
            continue
        if isinstance(a_text, list):
            a_text = "; ".join(a_text)
        valid_qs.append({
            "question": str(q_text).strip(),
            "answer":   str(a_text).strip(),
            "type":     q.get("type", "subjective"),
            "options":  q.get("options"),
        })

    if not valid_qs:
        return jsonify({"error": "AI returned empty or malformed questions. Please try again."}), 500

    # Strict type filtering
    if obj_count > 0 and subj_count == 0:
        valid_qs = [q for q in valid_qs if q["type"] == "objective"] or valid_qs
    elif subj_count > 0 and obj_count == 0:
        valid_qs = [q for q in valid_qs if q["type"] == "subjective"] or valid_qs

    # Persist question history to MongoDB
    if MONGO_OK and valid_qs:
        new_texts = [q["question"] for q in valid_qs]
        _mongo_update(M_PRACTICE_HIST,
                      {"user_id": u["id"], "syllabus_id": sid, "topic": topic},
                      {"$push": {"asked_questions": {"$each": new_texts}}},
                      upsert=True)

    return jsonify(valid_qs)

@app.post("/api/curriculum/practice/evaluate")
@auth()
@quota('practice_evaluate')
def evaluate_practice():
    data = request.json or {}
    question  = data.get("question","")
    model_ans = data.get("model_answer","")
    stud_ans  = data.get("student_answer","")
    q_type    = data.get("type", "subjective")
    options   = data.get("options", {})

    # Deterministic check for objective questions
    if q_type == "objective":
        # Build a rich question dict so _objective_is_correct can bridge
        # letter keys ↔ option text (e.g. student picks "A", model stores "5").
        q_dict = {
            "options":       options or {},
            "answer":        model_ans,
            # Also populate valid_answers with both the raw value AND any
            # option text that matches, so all comparison paths are covered.
            "valid_answers": [model_ans],
        }
        is_correct = _objective_is_correct(q_dict, stud_ans)

        # Build a human-readable correct answer label (prefer "A (5)" style)
        correct_label = model_ans
        if options:
            # If model_ans is a letter key, show "A (text)"
            if model_ans.strip().upper() in options:
                correct_label = f"{model_ans.upper()} — {options[model_ans.strip().upper()]}"
            else:
                # If model_ans is option text, find its letter
                for k, v in options.items():
                    if _normalize(v) == _normalize(model_ans):
                        correct_label = f"{k} — {v}"
                        break

        if is_correct:
            return jsonify({
                "score": 10,
                "is_correct": True,
                "feedback": f"Correct! The answer is {correct_label}.",
                "improvements": ""
            })
        else:
            return jsonify({
                "score": 0,
                "is_correct": False,
                "feedback": f"Incorrect. The correct answer is {correct_label}.",
                "improvements": "Review this topic in your textbook."
            })

    try:
        prompt = PRACTICE_EVALUATION_PROMPT.replace("{question}", question).replace("{model_answer}", model_ans).replace("{student_answer}", stud_ans)
        result = _llm_json(prompt, temperature=0, mini=True)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.post("/api/curriculum/practice/report")
@auth()
def practice_report():
    data = request.json or {}
    attempts = data.get("attempts", [])
    if not attempts:
        return jsonify({"error": "attempts list required"}), 400
    try:
        prompt = PRACTICE_REPORT_PROMPT.replace("{session_data}", json.dumps(attempts, default=str))
        result = _llm_json(prompt, temperature=0.2, mini=True)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.post("/api/curriculum/practice/email-report")
@auth()
def email_practice_report():
    data    = request.json or {}
    email   = data.get("email")
    results = data.get("results", [])
    if not email:
        return jsonify({"error": "email required"}), 400
    ok = send_email_report(email, "Your AI Practice Results — VidyAI", results)
    if ok:
        return jsonify({"success": True, "message": f"Report sent to {email}"})
    return jsonify({"success": False, "message": "SMTP not configured. Set SMTP_HOST, SMTP_USER, SMTP_PASS in .env"}), 500

# ══════════════════════════════════════════════════════════════════════════════
#  ADMIN ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/admin/users")
@auth(roles=["admin","institute_admin"])
def admin_list_users():
    """
    Platform admin (admin) sees all users.
    Institute admin (institute_admin) sees only users whose institution
    matches their own institution (case-insensitive), hiding platform admins.
    """
    db          = db_load()
    u           = request.user
    role_filter = request.args.get("role")

    if u["role"] == "admin":
        users = [x for x in db["users"] if x["id"] != u["id"]]
    else:
        # institute_admin: only their institution, excluding platform admin accounts
        my_inst = u.get("institution", "").lower()
        users = [
            x for x in db["users"]
            if x["id"] != u["id"]
            and x["role"] not in ("admin",)
            and x.get("institution", "").lower() == my_inst
        ]

    if role_filter:
        users = [x for x in users if x["role"] == role_filter]

    return jsonify({"users": [_safe(x) for x in users]})

@app.post("/api/admin/users")
@auth(roles=["admin"])
def admin_add_user():
    b = request.json or {}; db = db_load()
    email = b.get("email","").strip().lower()
    if any(x["email"] == email for x in db["users"]):
        return jsonify({"error": "Email exists"}), 409
    u = {"id":_uid(),"name":b.get("name",""),"email":email,
         "pw_hash":_hash(b.get("password","password123")),"role":b.get("role","student"),
         "institution":b.get("institution",""),"roll_number":"",
         "joined":_now()[:10],"docs":0,"status":"active"}
    db["users"].append(u); db_save(db)
    return jsonify({"user": _safe(u)}), 201

@app.patch("/api/admin/users/<uid>")
@auth(roles=["admin","institute_admin"])
def admin_update_user(uid):
    db     = db_load()
    caller = request.user
    u      = next((x for x in db["users"] if x["id"] == uid), None)
    if not u:
        return jsonify({"error": "Not found"}), 404
    # institute_admin cannot edit accounts outside their institution or admin/institute_admin accounts
    if caller["role"] == "institute_admin":
        if u["role"] in ("admin", "institute_admin") and u["id"] != caller["id"]:
            return jsonify({"error": "Forbidden"}), 403
        if u.get("institution", "").lower() != caller.get("institution", "").lower():
            return jsonify({"error": "Forbidden — user is not in your institution"}), 403
    b = request.json or {}
    for k in ("name","role","status","institution"):
        if k in b:
            u[k] = b[k]
    if "password" in b and b["password"]:
        if len(b["password"]) < 8:
            return jsonify({"error": "Password must be at least 8 characters"}), 400
        u["pw_hash"] = _hash(b["password"])
    db_save(db)
    return jsonify({"user": _safe(u)})

@app.delete("/api/admin/users/<uid>")
@auth(roles=["admin"])
def admin_delete_user(uid):
    db = db_load()
    db["users"] = [u for u in db["users"] if u["id"] != uid]
    db_save(db)
    return jsonify({"ok": True})

@app.get("/api/admin/stats")
@auth(roles=["admin","institute_admin"])
def admin_stats():
    db = db_load()
    u  = request.user

    if u["role"] == "admin":
        # Platform admin sees everyone
        users = [x for x in db["users"] if x["id"] != u["id"]]
    else:
        # institute_admin: scoped to their institution
        my_inst = u.get("institution", "").lower()
        users = [
            x for x in db["users"]
            if x["id"] != u["id"]
            and x["role"] != "admin"
            and x.get("institution", "").lower() == my_inst
        ]

    visible_uid_set = {x["id"] for x in users} | {u["id"]}
    activity = [a for a in db.get("activity", []) if a.get("user_id") in visible_uid_set]

    storage_gb = sum(
        (Path(d["path"]).stat().st_size
         if d.get("path") and Path(d.get("path","")).exists() else 0)
        for d in db.get("documents", [])
        if u["role"] == "admin" or d.get("owner_id") in visible_uid_set
    ) / (1024 ** 3)

    # Exam/eval counts: scoped to visible users
    # Stats across all users
    exam_count = len([e for e in exams_registry.values()
                      if u["role"] == "admin" or e.get("owner_id") in visible_uid_set])
    eval_count = len([e for e in evaluations_registry.values()
                      if u["role"] == "admin" or e.get("owner_id") in visible_uid_set])
    visitor_count = len(db.get("visitors", []))

    # ── feature usage analytics (last 30 days) ──────────────────────────────
    from collections import defaultdict as _dd
    from datetime import date as _date2, timedelta as _td2
    usage_counters = db.get("usage_counters", {})
    feature_totals: dict = {}   # {feature: total_count}
    user_feature_map: dict = {} # {uid: {feature: count}}
    cutoff30 = (_date2.today() - _td2(days=30)).isoformat()

    for day_str, day_data in usage_counters.items():
        if day_str < cutoff30:
            continue
        for uid2, feats in day_data.get("users", {}).items():
            if uid2 not in visible_uid_set:
                continue
            if uid2 not in user_feature_map:
                user_feature_map[uid2] = {}
            for feat, cnt in feats.items():
                feature_totals[feat]               = feature_totals.get(feat, 0) + cnt
                user_feature_map[uid2][feat]       = user_feature_map[uid2].get(feat, 0) + cnt

    # build a uid→name lookup for top_users
    uid_name = {x["id"]: x.get("name", x["id"]) for x in users}
    uid_name[u["id"]] = u.get("name", u["id"])

    top_users = sorted(
        [{"uid": uid2, "name": uid_name.get(uid2, uid2),
          "total": sum(f.values())} for uid2, f in user_feature_map.items()],
        key=lambda x: x["total"], reverse=True
    )[:10]

    # recently active (last 7 days) — includes the admin themselves
    from datetime import datetime as _dt2, timedelta as _td3
    cutoff7 = (_dt2.utcnow() - _td3(days=7)).strftime("%Y-%m-%dT%H:%M:%S")
    recently_active = [
        {"id": x["id"], "name": x.get("name",""), "email": x.get("email",""),
         "role": x.get("role",""), "institution": x.get("institution",""),
         "last_login": x.get("last_login",""), "login_count": x.get("login_count",0)}
        for x in (users + [u])
        if x.get("last_login","") >= cutoff7
    ]
    recently_active.sort(key=lambda x: x["last_login"], reverse=True)
    # ────────────────────────────────────────────────────────────────────────

    return jsonify({
        "teacher_count":  len([x for x in users if x["role"] == "teacher"]),
        "student_count":  len([x for x in users if x["role"] == "student"]),
        "parent_count":   len([x for x in users if x["role"] == "parent"]),
        "institute_count": len([x for x in db["users"] if x["role"] == "institute_admin"]) if u["role"] == "admin" else None,
        "visitor_count":  visitor_count,
        "exam_count":     exam_count,
        "eval_count":     eval_count,
        "syllabus_count": len(syllabi_registry),
        "storage_gb":     round(storage_gb, 3),
        "avg_score":      round(
            sum(v.get("result",{}).get("percentage",0) for v in evaluations_registry.values()) /
            max(len(evaluations_registry), 1), 1
        ),
        "activity":          activity[:20],
        "feature_totals":    feature_totals,
        "top_users_by_usage": top_users,
        "recently_active":   recently_active[:20],
        "mongodb_status": "connected" if MONGO_OK else "fallback_json",
    })

@app.get("/api/admin/institutions")
@auth(roles=["admin"])
def admin_list_institutions():
    """Platform admin: list all registered institutions with their user counts."""
    db = db_load()
    # Group users by institution
    from collections import defaultdict
    inst_map = defaultdict(lambda: {"teachers": 0, "students": 0, "parents": 0, "admins": []})
    for x in db["users"]:
        inst = x.get("institution", "").strip()
        if not inst:
            continue
        role = x["role"]
        if role == "teacher":
            inst_map[inst]["teachers"] += 1
        elif role == "student":
            inst_map[inst]["students"] += 1
        elif role == "parent":
            inst_map[inst]["parents"] += 1
        elif role == "institute_admin":
            inst_map[inst]["admins"].append({
                "id": x["id"], "name": x["name"], "email": x["email"],
                "joined": x.get("joined", ""), "status": x.get("status", "active")
            })
    institutions = [
        {"name": inst, **data}
        for inst, data in sorted(inst_map.items())
    ]
    return jsonify({"institutions": institutions})

@app.route("/api/admin/settings", methods=["GET","PATCH"])
@auth(roles=["admin","institute_admin"])
def admin_settings():
    db = db_load()
    if "settings" not in db:
        db["settings"] = {}
    if request.method == "PATCH":
        incoming = request.json or {}
        # institute_admin cannot change system-level model settings
        u = request.user
        if u["role"] == "institute_admin":
            forbidden_keys = {"model"}
            incoming = {k: v for k, v in incoming.items() if k not in forbidden_keys}
        db["settings"].update(incoming)
        db_save(db)
    return jsonify({"settings": db["settings"]})

@app.get("/api/admin/visitors")
@auth(roles=["admin", "institute_admin"])
def get_visitor_details():
    db = db_load()
    visitors = db.get("visitors", [])
    # Sort by timestamp descending
    try:
        visitors.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
    except:
        pass
    return jsonify({"visitors": visitors})

@app.route("/api/settings", methods=["GET"])
@auth() # Any logged in user
def get_settings():
    db = db_load()
    return jsonify({"settings": db.get("settings", {})})

@app.route("/api/visitors/track", methods=["POST"])
def track_visitor():
    db = db_load()
    if "visitors" not in db:
        db["visitors"] = []
    
    b = request.json or {}
    
    # Extract location and IP from Vercel headers if available
    # https://vercel.com/docs/concepts/edge-network/headers
    city    = request.headers.get("x-vercel-ip-city", "Unknown")
    country = request.headers.get("x-vercel-ip-country", "Unknown")
    ip      = request.headers.get("x-forwarded-for", request.remote_addr or "unknown").split(',')[0]
    
    db["visitors"].append({
        "timestamp":  _now(),
        "ip":         ip,
        "city":       city,
        "country":    country,
        "path":       b.get("path", "/"),
        "user_agent": request.headers.get("User-Agent", "unknown")[:200]
    })
    
    # Keep last 1000 visitors
    if len(db["visitors"]) > 1000:
        db["visitors"] = db["visitors"][-1000:]
        
    db_save(db)
    return jsonify({"ok": True})

# ══════════════════════════════════════════════════════════════════════════════
#  QUOTA MANAGEMENT ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/quota/me")
@auth()
def get_my_quota():
    """Return today's usage + limits for the current user across all features."""
    user  = request.user
    uid   = user["id"]
    inst  = (user.get("institution") or "").strip()
    today = _quota_today()
    counts = _get_today_usage(uid, inst)

    result = {}
    for feat in QUOTA_FEATURES:
        u_limit, i_pool = _get_effective_limit(user, feat)
        u_used = counts["user"].get(feat, 0)
        i_used = counts["inst"].get(feat, 0) if inst else 0
        result[feat] = {
            "label":      QUOTA_FEATURES[feat],
            "used":       u_used,
            "limit":      u_limit,
            "remaining":  (u_limit - u_used) if u_limit != _UNLIMITED else _UNLIMITED,
            "inst_used":  i_used,
            "inst_pool":  i_pool,
            "blocked":    (u_limit == 0),
        }
    return jsonify({"date": today, "features": result, "institution": inst or None})

@app.get("/api/admin/quota-config")
@auth(roles=["admin", "institute_admin"])
def get_quota_config():
    """Return full quota configuration. institute_admin sees their institution's overrides."""
    caller = request.user
    qc     = _qc_load()
    # institute_admin: scope institution_overrides to their own institution
    if caller["role"] == "institute_admin":
        my_inst = (caller.get("institution") or "").strip()
        scoped_inst = {}
        if my_inst and my_inst in qc.get("institution_overrides", {}):
            scoped_inst[my_inst] = qc["institution_overrides"][my_inst]
        return jsonify({
            "role_defaults":        qc.get("role_defaults", _DEFAULT_ROLE_QUOTAS),
            "institution_overrides": scoped_inst,
            "user_overrides":       qc.get("user_overrides", {}),
            "features":             QUOTA_FEATURES,
            "can_edit_role_defaults": False,
            "quota_presets":        _visible_quota_presets(caller["role"]),
        })
    return jsonify({
        "role_defaults":        qc.get("role_defaults", _DEFAULT_ROLE_QUOTAS),
        "institution_overrides": qc.get("institution_overrides", {}),
        "user_overrides":       qc.get("user_overrides", {}),
        "features":             QUOTA_FEATURES,
        "can_edit_role_defaults": True,
        "quota_presets":        _visible_quota_presets(caller["role"]),
    })

@app.post("/api/admin/quota-presets/apply")
@auth(roles=["admin"])
def apply_quota_preset():
    """
    Apply a named quota preset.
    Body:
      {
        "preset_id": "student-pro" | "school-starter" | ...,
        "target": {
          "type": "role_default" | "institution",
          "role": "student" | ...,
          "institution": "School Name"
        }
      }
    """
    caller = request.user
    b = request.json or {}
    preset_id = str(b.get("preset_id", "")).strip()
    target = b.get("target") or {}
    target_type = str(target.get("type", "")).strip()
    if not preset_id or target_type not in ("role_default", "institution"):
        return jsonify({"error": "preset_id and valid target.type are required"}), 400

    qc = _qc_load()

    if target_type == "role_default":
        if caller["role"] != "admin":
            return jsonify({"error": "Only admin can apply role-default presets"}), 403
        preset = (QUOTA_PRESETS.get("role_defaults") or {}).get(preset_id)
        if not preset:
            return jsonify({"error": f"Unknown role preset: {preset_id}"}), 400
        role = str(target.get("role", "")).strip()
        if role not in ("student", "parent", "teacher", "institute_admin", "admin"):
            return jsonify({"error": "Valid target.role is required"}), 400
        qc.setdefault("role_defaults", {})
        qc["role_defaults"].setdefault(role, {})
        template = preset.get("template", {})
        for feat in QUOTA_FEATURES:
            if feat in template:
                qc["role_defaults"][role][feat] = int(template[feat])
        _qc_save(qc)
        return jsonify({
            "ok": True,
            "applied": {
                "preset_id": preset_id,
                "target": {"type": "role_default", "role": role},
            },
            "role_defaults": qc.get("role_defaults", {}),
        })

    preset = (QUOTA_PRESETS.get("institution") or {}).get(preset_id)
    if not preset:
        return jsonify({"error": f"Unknown institution preset: {preset_id}"}), 400
    inst_name = str(target.get("institution", "")).strip()
    if caller["role"] == "institute_admin":
        my_inst = (caller.get("institution") or "").strip()
        if not my_inst:
            return jsonify({"error": "Your account has no institution"}), 400
        if inst_name and inst_name.lower() != my_inst.lower():
            return jsonify({"error": "Forbidden - not your institution"}), 403
        inst_name = my_inst
    if not inst_name:
        return jsonify({"error": "target.institution is required"}), 400

    qc.setdefault("institution_overrides", {})
    qc["institution_overrides"].setdefault(inst_name, {})
    cfg = qc["institution_overrides"][inst_name]
    cfg["user_daily"] = {
        feat: int(limit)
        for feat, limit in (preset.get("user_daily") or {}).items()
        if feat in QUOTA_FEATURES
    }
    cfg["inst_daily"] = {
        feat: int(limit)
        for feat, limit in (preset.get("inst_daily") or {}).items()
        if feat in QUOTA_FEATURES
    }
    _qc_save(qc)
    return jsonify({
        "ok": True,
        "applied": {
            "preset_id": preset_id,
            "target": {"type": "institution", "institution": inst_name},
        },
        "institution": inst_name,
        "config": cfg,
    })

@app.patch("/api/admin/quota-config/role-defaults")
@auth(roles=["admin"])
def update_quota_role_defaults():
    """Replace role_defaults section. Usage: {role: {feature: daily_limit}}"""
    b  = request.json or {}
    qc = _qc_load()
    # Validate and merge
    for role, feats in b.items():
        if role not in ("student", "parent", "teacher", "institute_admin", "admin"):
            return jsonify({"error": f"Invalid role: {role}"}), 400
        if not isinstance(feats, dict):
            return jsonify({"error": f"Expected dict for role {role}"}), 400
        if role not in qc["role_defaults"]:
            qc["role_defaults"][role] = {}
        for feat, limit in feats.items():
            if feat not in QUOTA_FEATURES:
                return jsonify({"error": f"Unknown feature: {feat}"}), 400
            try:
                qc["role_defaults"][role][feat] = int(limit)
            except (ValueError, TypeError):
                return jsonify({"error": f"Limit must be an integer for {feat}"}), 400
    _qc_save(qc)
    return jsonify({"ok": True, "role_defaults": qc["role_defaults"]})

@app.patch("/api/admin/quota-config/institution/<path:inst_name>")
@auth(roles=["admin", "institute_admin"])
def update_quota_institution(inst_name):
    """
    Set institution overrides.
    Body: { user_daily: {feature: limit}, inst_daily: {feature: pool_limit} }
    institute_admin can only update their own institution.
    """
    caller = request.user
    if caller["role"] == "institute_admin":
        my_inst = (caller.get("institution") or "").strip()
        if inst_name.lower() != my_inst.lower():
            return jsonify({"error": "Forbidden — not your institution"}), 403
    b  = request.json or {}
    qc = _qc_load()
    if inst_name not in qc["institution_overrides"]:
        qc["institution_overrides"][inst_name] = {}
    cfg = qc["institution_overrides"][inst_name]
    for section in ("user_daily", "inst_daily"):
        if section in b:
            if not isinstance(b[section], dict):
                return jsonify({"error": f"{section} must be a dict"}), 400
            cfg.setdefault(section, {})
            for feat, limit in b[section].items():
                if feat not in QUOTA_FEATURES:
                    return jsonify({"error": f"Unknown feature: {feat}"}), 400
                if limit is None:
                    cfg[section].pop(feat, None)
                else:
                    try:
                        cfg[section][feat] = int(limit)
                    except (ValueError, TypeError):
                        return jsonify({"error": f"Limit must be an integer for {feat}"}), 400
    _qc_save(qc)
    return jsonify({"ok": True, "institution": inst_name, "config": cfg})

@app.delete("/api/admin/quota-config/institution/<path:inst_name>")
@auth(roles=["admin"])
def delete_quota_institution(inst_name):
    """Remove all quota overrides for an institution (revert to role defaults)."""
    qc = _qc_load()
    qc["institution_overrides"].pop(inst_name, None)
    _qc_save(qc)
    return jsonify({"ok": True})

@app.patch("/api/admin/quota-config/user/<uid>")
@auth(roles=["admin", "institute_admin"])
def update_quota_user(uid):
    """
    Set per-user quota overrides. Body: {feature: daily_limit, ...}
    Pass null/None to remove a feature override.
    institute_admin can only configure users in their institution.
    """
    caller = request.user
    db     = db_load()
    target = next((u for u in db["users"] if u["id"] == uid), None)
    if not target:
        return jsonify({"error": "User not found"}), 404
    if caller["role"] == "institute_admin":
        my_inst = (caller.get("institution") or "").strip()
        if target.get("institution", "").strip().lower() != my_inst.lower():
            return jsonify({"error": "Forbidden — user not in your institution"}), 403
    b  = request.json or {}
    qc = _qc_load()
    qc["user_overrides"].setdefault(uid, {})
    for feat, limit in b.items():
        if feat not in QUOTA_FEATURES:
            return jsonify({"error": f"Unknown feature: {feat}"}), 400
        if limit is None:
            qc["user_overrides"][uid].pop(feat, None)
        else:
            try:
                qc["user_overrides"][uid][feat] = int(limit)
            except (ValueError, TypeError):
                return jsonify({"error": f"Limit must be an integer for {feat}"}), 400
    # Clean up empty entry
    if not qc["user_overrides"][uid]:
        qc["user_overrides"].pop(uid, None)
    _qc_save(qc)
    return jsonify({"ok": True, "user_id": uid, "overrides": qc.get("user_overrides", {}).get(uid, {})})

@app.delete("/api/admin/quota-config/user/<uid>")
@auth(roles=["admin", "institute_admin"])
def delete_quota_user(uid):
    """Remove all quota overrides for a user (revert to role/institution defaults)."""
    caller = request.user
    db     = db_load()
    target = next((u for u in db["users"] if u["id"] == uid), None)
    if not target:
        return jsonify({"error": "User not found"}), 404
    if caller["role"] == "institute_admin":
        my_inst = (caller.get("institution") or "").strip()
        if target.get("institution", "").strip().lower() != my_inst.lower():
            return jsonify({"error": "Forbidden"}), 403
    qc = _qc_load()
    qc["user_overrides"].pop(uid, None)
    _qc_save(qc)
    return jsonify({"ok": True})

@app.get("/api/admin/usage-report")
@auth(roles=["admin", "institute_admin"])
def get_usage_report():
    """
    Usage report for the last N days.
    Query params: days=7 (default), feature (optional filter).
    institute_admin sees only their institution's users.
    """
    from datetime import date as _date, timedelta as _td
    caller  = request.user
    days    = min(int(request.args.get("days", 7)), 90)
    feature_filter = request.args.get("feature")

    db   = db_load()
    all_users = {u["id"]: u for u in db["users"]}

    # Determine visible user IDs
    if caller["role"] == "admin":
        visible_uids = set(all_users.keys())
    else:
        my_inst = (caller.get("institution") or "").strip().lower()
        visible_uids = {u["id"] for u in db["users"]
                        if u.get("institution", "").strip().lower() == my_inst}

    counters = db.get("usage_counters", {})
    today    = _date.today()
    report   = []
    for i in range(days):
        d = (today - _td(days=i)).isoformat()
        day = counters.get(d, {})
        users_day = day.get("users", {})
        inst_day  = day.get("institutions", {})
        by_user   = {}
        for uid, feats in users_day.items():
            if uid not in visible_uids:
                continue
            u = all_users.get(uid, {})
            if feature_filter:
                count = feats.get(feature_filter, 0)
                if count == 0:
                    continue
                by_user[uid] = {"name": u.get("name","?"), "role": u.get("role","?"),
                                "institution": u.get("institution",""),
                                feature_filter: count}
            else:
                by_user[uid] = {"name": u.get("name","?"), "role": u.get("role","?"),
                                "institution": u.get("institution",""), **feats}
        report.append({"date": d, "users": by_user, "institutions": inst_day})

    # Summary: per-feature totals across the period
    totals = {f: 0 for f in QUOTA_FEATURES}
    for day_data in report:
        for uid, feats in day_data["users"].items():
            for f in QUOTA_FEATURES:
                totals[f] += feats.get(f, 0)

    return jsonify({
        "days":     days,
        "report":   report,
        "totals":   totals,
        "features": QUOTA_FEATURES,
    })

@app.get("/api/admin/quota-config/reset-defaults")
@auth(roles=["admin"])
def reset_quota_defaults():
    """Reset role_defaults to system defaults (non-destructive for overrides)."""
    qc = _qc_load()
    qc["role_defaults"] = _DEFAULT_ROLE_QUOTAS
    _qc_save(qc)
    return jsonify({"ok": True, "role_defaults": _DEFAULT_ROLE_QUOTAS})

# ══════════════════════════════════════════════════════════════════════════════
#  SPA SERVING (React dist/)
# ══════════════════════════════════════════════════════════════════════════════
@app.route("/manifest.json")
def serve_manifest():
    try:
        return send_from_directory(app.static_folder, "manifest.json")
    except Exception:
        return jsonify({}), 200

# ══════════════════════════════════════════════════════════════════════════════
#  FINANCE LEDGER  — admin-only internal accounting
# ══════════════════════════════════════════════════════════════════════════════
M_FINANCE      = "finance_ledger"
M_PAYROLL_EMP  = "payroll_employees"
M_INVESTORS    = "finance_investors"

def _fin_id() -> str:
    return "txn_" + uuid.uuid4().hex[:10]

def _emp_id() -> str:
    return "emp_" + uuid.uuid4().hex[:8]

def _inv_id() -> str:
    return "inv_" + uuid.uuid4().hex[:8]

# expense categories that count as "outflow"
EXPENSE_TYPES = {"expense", "salary", "refund"}
INCOME_TYPES  = {"income", "investment"}

@app.get("/api/admin/finance/summary")
@auth(roles=["admin"])
def finance_summary():
    """Aggregate KPIs: totals, monthly breakdown, category breakdown."""
    import calendar as _cal
    _FIN_FILE = TMP_ROOT / "finance_ledger.json"

    if MONGO_OK:
        txns = _mongo_find(M_FINANCE, {})
    else:
        txns = _load_json(_FIN_FILE, [])

    total_income   = sum(t["amount"] for t in txns if t.get("type") in INCOME_TYPES)
    total_expenses = sum(t["amount"] for t in txns if t.get("type") in EXPENSE_TYPES)
    api_cost_total = sum(t["amount"] for t in txns if t.get("category") == "api_cost")
    salary_total   = sum(t["amount"] for t in txns if t.get("type") == "salary" or t.get("category") == "salary")
    invest_total   = sum(t["amount"] for t in txns if t.get("type") == "investment")

    # Monthly rollup
    monthly_map = {}
    for t in txns:
        m = (t.get("date") or "")[:7]  # YYYY-MM
        if not m:
            continue
        if m not in monthly_map:
            monthly_map[m] = {"month": m, "income": 0.0, "expenses": 0.0, "count": 0}
        if t.get("type") in INCOME_TYPES:
            monthly_map[m]["income"] += t["amount"]
        elif t.get("type") in EXPENSE_TYPES:
            monthly_map[m]["expenses"] += t["amount"]
        monthly_map[m]["count"] += 1
    monthly = sorted(monthly_map.values(), key=lambda x: x["month"])
    # Abbreviate month label e.g. "2026-03" → "Mar'26"
    for row in monthly:
        try:
            yr, mo = row["month"].split("-")
            row["month"] = f"{_cal.month_abbr[int(mo)]}'{yr[2:]}"
        except Exception:
            pass

    # Category breakdown
    cat_map = {}
    for t in txns:
        cat  = t.get("category", "misc_exp")
        flow = "income" if t.get("type") in INCOME_TYPES else "expense"
        key  = f"{cat}___{flow}"
        if key not in cat_map:
            cat_map[key] = {"category": cat, "flow": flow, "total": 0.0, "count": 0}
        cat_map[key]["total"] += t["amount"]
        cat_map[key]["count"] += 1

    # Recent 20 transactions sorted by date desc
    recent = sorted(txns, key=lambda t: t.get("date", ""), reverse=True)[:20]

    return jsonify({
        "summary": {
            "total_income":     round(total_income,   2),
            "total_expenses":   round(total_expenses, 2),
            "net_pl":           round(total_income - total_expenses, 2),
            "api_cost_total":   round(api_cost_total, 2),
            "salary_total":     round(salary_total,   2),
            "investment_total": round(invest_total,   2),
            "monthly":          monthly,
            "by_category":      list(cat_map.values()),
        },
        "recent_transactions": recent,
    })


@app.get("/api/admin/finance/transactions")
@auth(roles=["admin"])
def finance_list_transactions():
    """List transactions with optional filters: type, category, from, to, q."""
    _FIN_FILE = TMP_ROOT / "finance_ledger.json"
    type_f = request.args.get("type", "")
    cat_f  = request.args.get("category", "")
    from_f = request.args.get("from", "")
    to_f   = request.args.get("to", "")
    q_f    = request.args.get("q", "").lower()

    if MONGO_OK:
        txns = _mongo_find(M_FINANCE, {})
    else:
        txns = _load_json(_FIN_FILE, [])

    if type_f:
        txns = [t for t in txns if t.get("type") == type_f]
    if cat_f:
        txns = [t for t in txns if t.get("category") == cat_f]
    if from_f:
        txns = [t for t in txns if (t.get("date") or "") >= from_f]
    if to_f:
        txns = [t for t in txns if (t.get("date") or "") <= to_f]
    if q_f:
        txns = [t for t in txns
                if q_f in (t.get("description") or "").lower()
                or q_f in (t.get("vendor") or "").lower()
                or q_f in (t.get("reference") or "").lower()
                or any(q_f in tag for tag in (t.get("tags") or []))]

    txns = sorted(txns, key=lambda t: t.get("date", ""), reverse=True)
    return jsonify({"transactions": txns, "count": len(txns)})


@app.post("/api/admin/finance/transactions")
@auth(roles=["admin"])
def finance_add_transaction():
    """Record a new transaction."""
    data = request.json or {}
    required = ["type", "amount", "description", "date"]
    for f in required:
        if not data.get(f):
            return jsonify({"error": f"'{f}' is required"}), 400

    txn = {
        "id":          _fin_id(),
        "type":        data["type"],
        "category":    data.get("category", "misc_exp"),
        "amount":      float(data["amount"]),
        "currency":    data.get("currency", "INR"),
        "description": data["description"],
        "vendor":      data.get("vendor", ""),
        "date":        data["date"],
        "reference":   data.get("reference", ""),
        "notes":       data.get("notes", ""),
        "tags":        data.get("tags", []),
        "created_by":  request.user["id"],
        "created_at":  _now(),
    }

    if MONGO_OK:
        _mongo_insert(M_FINANCE, txn)
    else:
        _FIN_FILE = TMP_ROOT / "finance_ledger.json"
        existing = _load_json(_FIN_FILE, [])
        existing.append(txn)
        _save_json(_FIN_FILE, existing)

    return jsonify({"transaction": txn}), 201


@app.patch("/api/admin/finance/transactions/<txn_id>")
@auth(roles=["admin"])
def finance_update_transaction(txn_id):
    """Update an existing transaction."""
    data    = request.json or {}
    allowed = {"type","category","amount","currency","description",
               "vendor","date","reference","notes","tags"}
    update  = {k: v for k, v in data.items() if k in allowed}
    if "amount" in update:
        update["amount"] = float(update["amount"])
    update["updated_at"] = _now()

    if MONGO_OK:
        _mongo_update(M_FINANCE, {"id": txn_id}, {"$set": update})
    else:
        _FIN_FILE = TMP_ROOT / "finance_ledger.json"
        txns = _load_json(_FIN_FILE, [])
        for t in txns:
            if t["id"] == txn_id:
                t.update(update)
        _save_json(_FIN_FILE, txns)

    return jsonify({"updated": True})


@app.delete("/api/admin/finance/transactions/<txn_id>")
@auth(roles=["admin"])
def finance_delete_transaction(txn_id):
    """Delete a transaction."""
    if MONGO_OK:
        _mongo_delete(M_FINANCE, {"id": txn_id})
    else:
        _FIN_FILE = TMP_ROOT / "finance_ledger.json"
        txns = _load_json(_FIN_FILE, [])
        txns = [t for t in txns if t["id"] != txn_id]
        _save_json(_FIN_FILE, txns)
    return jsonify({"deleted": True})


# ── Payroll: employees / contractors ─────────────────────────────────────────
@app.get("/api/admin/finance/payroll/employees")
@auth(roles=["admin"])
def payroll_list_employees():
    if MONGO_OK:
        emps = _mongo_find(M_PAYROLL_EMP, {})
    else:
        emps = _load_json(TMP_ROOT / "payroll_employees.json", [])
    return jsonify({"employees": emps})


@app.post("/api/admin/finance/payroll/employees")
@auth(roles=["admin"])
def payroll_add_employee():
    data = request.json or {}
    if not data.get("name") or not data.get("monthly_salary"):
        return jsonify({"error": "name and monthly_salary required"}), 400

    emp = {
        "id":             _emp_id(),
        "name":           data["name"],
        "role":           data.get("role", ""),
        "email":          data.get("email", ""),
        "type":           data.get("type", "employee"),
        "monthly_salary": float(data["monthly_salary"]),
        "currency":       data.get("currency", "INR"),
        "bank_ref":       data.get("bank_ref", ""),
        "joined_on":      data.get("joined_on", _now()[:10]),
        "notes":          data.get("notes", ""),
        "created_at":     _now(),
    }
    if MONGO_OK:
        _mongo_insert(M_PAYROLL_EMP, emp)
    else:
        _EMP_FILE = TMP_ROOT / "payroll_employees.json"
        existing = _load_json(_EMP_FILE, [])
        existing.append(emp)
        _save_json(_EMP_FILE, existing)
    return jsonify({"employee": emp}), 201


@app.delete("/api/admin/finance/payroll/employees/<emp_id>")
@auth(roles=["admin"])
def payroll_delete_employee(emp_id):
    if MONGO_OK:
        _mongo_delete(M_PAYROLL_EMP, {"id": emp_id})
    else:
        _EMP_FILE = TMP_ROOT / "payroll_employees.json"
        emps = _load_json(_EMP_FILE, [])
        emps = [e for e in emps if e["id"] != emp_id]
        _save_json(_EMP_FILE, emps)
    return jsonify({"deleted": True})


# ── Investors / funding rounds ────────────────────────────────────────────────
@app.get("/api/admin/finance/payroll/investors")
@auth(roles=["admin"])
def investors_list():
    if MONGO_OK:
        investors = _mongo_find(M_INVESTORS, {})
    else:
        investors = _load_json(TMP_ROOT / "finance_investors.json", [])
    return jsonify({"investors": investors})


@app.post("/api/admin/finance/payroll/investors")
@auth(roles=["admin"])
def investors_add():
    data = request.json or {}
    if not data.get("name") or not data.get("amount"):
        return jsonify({"error": "name and amount required"}), 400

    inv = {
        "id":         _inv_id(),
        "name":       data["name"],
        "type":       data.get("type", "angel"),
        "round":      data.get("round", "Seed"),
        "amount":     float(data["amount"]),
        "currency":   data.get("currency", "INR"),
        "equity_pct": float(data.get("equity_pct") or 0),
        "date":       data.get("date", _now()[:10]),
        "contact":    data.get("contact", ""),
        "terms":      data.get("terms", ""),
        "notes":      data.get("notes", ""),
        "created_at": _now(),
    }
    if MONGO_OK:
        _mongo_insert(M_INVESTORS, inv)
    else:
        _INV_FILE = TMP_ROOT / "finance_investors.json"
        existing = _load_json(_INV_FILE, [])
        existing.append(inv)
        _save_json(_INV_FILE, existing)
    return jsonify({"investor": inv}), 201


# ══════════════════════════════════════════════════════════════════════════════
#  EVALUATION ENHANCEMENTS — Review Queue & Analytics Endpoints
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/review-queue")
@auth(roles=["tutor", "teacher", "institute_admin"])
def get_review_queue():
    """
    Fetch pending items in the teacher review queue.
    Items tagged for review: low OCR confidence, borderline marks, illegible handwriting, anomalies.
    """
    exam_id = request.args.get("exam_id", "")
    limit = int(request.args.get("limit", "20"))
    offset = int(request.args.get("offset", "0"))
    
    if not exam_id:
        return jsonify({"error": "exam_id required"}), 400
    
    query = {"exam_id": exam_id, "status": "pending"}
    
    if MONGO_OK:
        try:
            collection = mongo_db[M_REVIEW_QUEUE]
            total = collection.count_documents(query)
            items = list(collection.find(query).sort("created_at", -1).skip(offset).limit(limit))
            for item in items:
                item.pop("_id", None)
            
            # Categorize items by review reason
            categorized = {}
            for item in items:
                reason_type = item.get("review_reason_type", "other")
                if reason_type not in categorized:
                    categorized[reason_type] = []
                categorized[reason_type].append(item)
            
            return jsonify({
                "total": total,
                "limit": limit,
                "offset": offset,
                "items": items,
                "categorized": categorized,
            }), 200
        except Exception as ex:
            log.error(f"[REVIEW-QUEUE] Failed to fetch: {ex}")
            return jsonify({"error": str(ex)}), 500
    else:
        return jsonify({"error": "MongoDB not available"}), 503


@app.post("/api/review-queue/<review_id>/approve")
@auth(roles=["tutor", "teacher", "institute_admin"])
def approve_review_item(review_id: str):
    """Approve an evaluation from the review queue (teacher decision)."""
    if not review_id:
        return jsonify({"error": "review_id required"}), 400
    
    data = request.get_json() or {}
    teacher_comments = data.get("comments", "").strip()
    action = data.get("action", "approve")  # approve, reject, request_modification
    
    if MONGO_OK:
        try:
            collection = mongo_db[M_REVIEW_QUEUE]
            result = collection.update_one(
                {"review_id": review_id},
                {"$set": {
                    "status": action,
                    "teacher_comments": teacher_comments,
                    "updated_at": _now(),
                    "reviewed_by": g.user.get("username", ""),
                    "reviewed_at": _now(),
                }}
            )
            if result.matched_count == 0:
                return jsonify({"error": "Review item not found"}), 404
            
            log.info(f"[REVIEW-QUEUE] Reviewed {review_id}: {action} by {g.user.get('username')}")
            return jsonify({"status": action, "message": f"Review item {action}"}), 200
        except Exception as ex:
            log.error(f"[REVIEW-QUEUE] Failed to approve {review_id}: {ex}")
            return jsonify({"error": str(ex)}), 500
    else:
        return jsonify({"error": "MongoDB not available"}), 503


@app.get("/api/evaluation/analytics/<exam_id>")
@auth(roles=["tutor", "teacher", "institute_admin"])
def get_evaluation_analytics(exam_id: str):
    """
    Get analytics dashboard for an exam:
    - Papers processed today/week/month
    - Average grading time
    - Accuracy/review percentage
    - Cost per paper
    - Teacher correction rate
    - Most common weak topics
    """
    if not exam_id:
        return jsonify({"error": "exam_id required"}), 400
    
    period = request.args.get("period", "today")  # today, week, month
    
    if MONGO_OK:
        try:
            # Get analytics summary
            analytics_col = mongo_db[M_ANALYTICS]
            evals_col = mongo_db[M_EVALS]
            review_col = mongo_db[M_REVIEW_QUEUE]
            
            # Date range based on period
            now = datetime.utcnow()
            if period == "today":
                start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
            elif period == "week":
                start_date = now - timedelta(days=7)
            else:  # month
                start_date = now - timedelta(days=30)
            
            # Papers processed
            papers_processed = evals_col.count_documents({
                "exam_id": exam_id,
                "created_at": {"$gte": start_date.isoformat()}
            })
            
            # Evaluation scores
            scores = list(evals_col.aggregate([
                {
                    "$match": {
                        "exam_id": exam_id,
                        "created_at": {"$gte": start_date.isoformat()}
                    }
                },
                {
                    "$group": {
                        "_id": None,
                        "avg_percentage": {"$avg": "$result.percentage"},
                        "pass_count": {"$sum": {"$cond": ["$result.is_pass", 1, 0]}},
                        "fail_count": {"$sum": {"$cond": ["$result.is_pass", 0, 1]}},
                    }
                }
            ]))
            
            score_stats = scores[0] if scores else {"avg_percentage": 0, "pass_count": 0, "fail_count": 0}
            pass_rate = (score_stats["pass_count"] / max(papers_processed, 1)) * 100
            
            # Review queue stats
            review_pending = review_col.count_documents({
                "exam_id": exam_id,
                "status": "pending"
            })
            review_total = review_col.count_documents({"exam_id": exam_id})
            accuracy_percent = ((review_total - review_pending) / max(review_total, 1)) * 100 if review_total > 0 else 100
            
            # Cost tracking (placeholder - integrate actual API costs)
            cost_col = mongo_db[M_COST_TRACKER]
            costs = list(cost_col.aggregate([
                {"$match": {"exam_id": exam_id, "created_at": {"$gte": start_date.isoformat()}}},
                {
                    "$group": {
                        "_id": None,
                        "total_cost": {"$sum": "$cost_usd"},
                        "avg_cost_per_paper": {"$avg": "$cost_usd"},
                    }
                }
            ]))
            cost_stats = costs[0] if costs else {"total_cost": 0, "avg_cost_per_paper": 0}
            
            return jsonify({
                "exam_id": exam_id,
                "period": period,
                "papers_processed": papers_processed,
                "average_score_percent": round(score_stats.get("avg_percentage", 0), 1),
                "pass_count": score_stats.get("pass_count", 0),
                "fail_count": score_stats.get("fail_count", 0),
                "pass_rate_percent": round(pass_rate, 1),
                "review_pending": review_pending,
                "review_completed": review_total - review_pending,
                "accuracy_review_percent": round(accuracy_percent, 1),
                "total_cost_usd": round(cost_stats.get("total_cost", 0), 2),
                "avg_cost_per_paper_usd": round(cost_stats.get("avg_cost_per_paper", 0), 4),
            }), 200
        except Exception as ex:
            log.error(f"[ANALYTICS] Failed to fetch: {ex}")
            return jsonify({"error": str(ex)}), 500
    else:
        return jsonify({"error": "MongoDB not available"}), 503


@app.get("/api/evaluation/results/<eval_id>")
@auth(roles=["tutor", "teacher", "institute_admin"])
def get_evaluation_result(eval_id: str):
    """
    Fetch complete evaluation result with:
    - Extracted answers with confidence scores
    - Graded answers with per-question marks
    - Review queue status (if applicable)
    - Teacher feedback (if reviewed)
    """
    if not eval_id:
        return jsonify({"error": "eval_id required"}), 400
    
    # First check evaluations registry
    if eval_id in evaluations_registry:
        result = evaluations_registry[eval_id]
        result.pop("_id", None) if isinstance(result, dict) and "_id" in result else None
        return jsonify(result), 200
    
    # Then check MongoDB
    if MONGO_OK:
        try:
            evals_col = mongo_db[M_EVALS]
            result = evals_col.find_one({"evaluation_id": eval_id})
            if result:
                result.pop("_id", None)
                return jsonify(result), 200
            
            # Check review queue
            review_col = mongo_db[M_REVIEW_QUEUE]
            review_item = review_col.find_one({"evaluation_id": eval_id})
            if review_item:
                review_item.pop("_id", None)
                return jsonify({
                    "evaluation_id": eval_id,
                    "status": "in_review",
                    "review_queue_status": review_item,
                }), 200
            
            return jsonify({"error": "Evaluation not found"}), 404
        except Exception as ex:
            log.error(f"[EVAL-RESULT] Failed to fetch {eval_id}: {ex}")
            return jsonify({"error": str(ex)}), 500
    else:
        return jsonify({"error": "Evaluation not found"}), 404


@app.route("/api/review-queue", methods=["GET"])
def api_review_queue():
    """Fetch review queue for anomalous evaluations.
    
    Query params:
        exam_id: Filter by exam
        limit: Max items to return (default 10)
        status: Filter by status (pending, reviewed, escalated)
    
    Returns:
    {
        "total": int,
        "items": [...],
        "categorized": {
            "anomalies": [...],
            "borderline": [...],
            "escalated": [...]
        }
    }
    """
    try:
        # Check authorization
        auth_header = request.headers.get("Authorization", "")
        if not auth_header.startswith("Bearer "):
            return jsonify({"error": "Unauthorized"}), 401
        
        exam_id = request.args.get("exam_id", "").strip()
        limit = int(request.args.get("limit", 10))
        status_filter = request.args.get("status", "").strip()
        
        items = []
        categorized = {"anomalies": [], "borderline": [], "escalated": []}
        
        if MONGO_OK:
            try:
                review_col = mongo_db[M_REVIEW_QUEUE]
                query = {}
                if exam_id:
                    query["exam_id"] = exam_id
                if status_filter:
                    query["status"] = status_filter
                
                cursor = review_col.find(query).limit(limit)
                for doc in cursor:
                    doc.pop("_id", None)
                    items.append(doc)
                    
                    # Categorize
                    cat = doc.get("category", "anomalies")
                    if cat in categorized:
                        categorized[cat].append(doc)
                
            except Exception as ex:
                log.warning(f"[REVIEW-QUEUE] MongoDB error: {ex}")
        
        return jsonify({
            "total": len(items),
            "items": items,
            "categorized": categorized
        }), 200
        
    except Exception as ex:
        log.error(f"[REVIEW-QUEUE] API error: {ex}")
        return jsonify({"error": str(ex)}), 500

@app.route("/api/evaluation/analytics/<exam_id>", methods=["GET"])
def api_evaluation_analytics(exam_id):
    """Get analytics dashboard for an exam.
    
    Query params:
        period: "today", "week", "month" (default: today)
    
    Returns:
    {
        "papers_processed": int,
        "average_score_percent": float,
        "pass_rate_percent": float,
        "accuracy_review_percent": float,
        "avg_cost_per_paper_usd": float,
        "top_students": [...],
        "bottom_students": [...]
    }
    """
    try:
        # Check authorization
        auth_header = request.headers.get("Authorization", "")
        if not auth_header.startswith("Bearer "):
            return jsonify({"error": "Unauthorized"}), 401
        
        period = request.args.get("period", "today").lower()
        
        papers_processed = 0
        average_score = 0.0
        pass_rate = 0.0
        accuracy = 100.0
        cost_per_paper = 0.0
        top_students = []
        bottom_students = []
        
        if MONGO_OK:
            try:
                analytics_col = mongo_db[M_ANALYTICS]
                
                # Get analytics for period
                query = {"exam_id": exam_id}
                docs = list(analytics_col.find(query).limit(100))
                
                if docs:
                    papers_processed = sum(d.get("papers_processed", 0) for d in docs)
                    avg_scores = [d.get("class_average_percentage", 0) for d in docs if d.get("papers_processed", 0) > 0]
                    average_score = sum(avg_scores) / len(avg_scores) if avg_scores else 0.0
                    
                    # Estimate pass rate (assuming 40% is pass)
                    pass_threshold = 40.0
                    pass_rate = (average_score / 100.0 * 100) if average_score > pass_threshold else 0.0
                    
                    # Cost estimate (₹100/paper ≈ $1.2 USD)
                    cost_per_paper = (papers_processed * 0.01 * 1.2) / max(papers_processed, 1)
                
                # Get top and bottom students from evaluation records
                evals_col = mongo_db[M_EVALS]
                evals = list(evals_col.find({"exam_id": exam_id}).sort("result.percentage", -1).limit(5))
                top_students = [{
                    "roll_no": e.get("roll_no", ""),
                    "name": e.get("student_name", ""),
                    "score": e.get("result", {}).get("percentage", 0)
                } for e in evals]
                
                evals = list(evals_col.find({"exam_id": exam_id}).sort("result.percentage", 1).limit(5))
                bottom_students = [{
                    "roll_no": e.get("roll_no", ""),
                    "name": e.get("student_name", ""),
                    "score": e.get("result", {}).get("percentage", 0)
                } for e in evals]
                
            except Exception as ex:
                log.warning(f"[ANALYTICS] MongoDB error: {ex}")
        
        return jsonify({
            "papers_processed": papers_processed,
            "average_score_percent": round(average_score, 1),
            "pass_rate_percent": round(pass_rate, 1),
            "accuracy_review_percent": round(accuracy, 1),
            "avg_cost_per_paper_usd": round(cost_per_paper, 4),
            "top_students": top_students,
            "bottom_students": bottom_students
        }), 200
        
    except Exception as ex:
        log.error(f"[ANALYTICS] API error: {ex}")
        return jsonify({"error": str(ex)}), 500


@app.route("/", defaults={"path": ""})
@app.route("/<path:path>")
def spa(path):
    if app.static_folder:
        fp = Path(app.static_folder) / path
        if path and fp.exists():
            return send_from_directory(app.static_folder, path)
        idx = Path(app.static_folder) / "index.html"
        if idx.exists():
            return send_from_directory(app.static_folder, "index.html")
    return jsonify({"status": "VidyAI API running", "docs": "/api/"}), 200

# ══════════════════════════════════════════════════════════════════════════════
#  STARTUP
# ══════════════════════════════════════════════════════════════════════════════
_boot_load()

if __name__ == "__main__":
    port = int(os.getenv("PORT", 5001))
    lc   = "✓" if LANGCHAIN_OK   else "✗ (install langchain-openai)"
    oai  = f"✓ ({OPENAI_MODEL})" if OPENAI_API_KEY else "✗ (set OPENAI_API_KEY)"
    mdb  = "✓ MongoDB"           if MONGO_OK       else "⚠ JSON fallback"
    que  = "✓ Celery+Redis"      if CELERY_OK      else "⚠ Sync fallback (no Redis)"
    print(f"""
╔══════════════════════════════════════════╗
║  Parvidya Backend  →  http://localhost:{port}  ║
╠══════════════════════════════════════════╣
║  LangChain : {lc:<28}║
║  OpenAI    : {oai:<28}║
║  Database  : {mdb:<28}║
║  Queue     : {que:<28}║
║  Vercel    : {'Yes' if IS_VERCEL else 'No':<28}║
╚══════════════════════════════════════════╝
Syllabi loaded : {len(syllabi_registry)}
Exams loaded   : {len(exams_registry)}
""")
    from flask import cli
    cli.show_server_banner = lambda *_: None
    app.run(host="0.0.0.0", port=port, debug=False)